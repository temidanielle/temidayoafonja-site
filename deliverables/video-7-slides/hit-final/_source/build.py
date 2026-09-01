# -*- coding: utf-8 -*-
"""Build the Video 7 H.I.T. final recording and Shorts package.

Formatting helpers are the approved house system carried over unchanged
from the Videos 1-6 packages; only wording, metadata and the manifest
are Video 7 specific.
"""
import os, sys, shutil, zipfile, hashlib
sys.path.insert(0, "/tmp/v7hit")
from script_text import LINES, SPOKEN, MARKERS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0x9B,0x2C,0x10)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

ROOT="/tmp/v7hit/Video_7_HIT_FINAL"
LF=os.path.join(ROOT,"LONG_FORM"); SH=os.path.join(ROOT,"SHORTS")
shutil.rmtree(ROOT, ignore_errors=True)
os.makedirs(LF); os.makedirs(SH)

TITLE="How to Show Your Impact at Work When You Built It From Scratch"
THUMB="MAKE INVISIBLE WORK VISIBLE"
CTA="Keep the Proof"
CTA_SUB="Keep the Proof — A 60-Minute Career Evidence System"
CTA_URL="https://temidayoafonja.com/keep-the-proof"
NEXT="How to Explain Your Career Change"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO 7",size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t,before=20):
    return keep(P(d,t,size=14,bold=True,color=NAVY,before=before,after=8),True)
def H2(d,t,before=13):
    return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=before,after=5),True)

def compress(d, line_spacing=1.18, after_scale=0.80):
    """Tighten one document's vertical rhythm without touching type size,
    weight, colour or wording. Keeps editor briefs off a near-empty last page."""
    for p in d.paragraphs:
        pf=p.paragraph_format
        if pf.line_spacing and 1.0 < pf.line_spacing < 1.4 and pf.line_spacing > line_spacing:
            pf.line_spacing = line_spacing
        if pf.space_after is not None:
            pf.space_after = Pt(round(pf.space_after.pt*after_scale,1))
        if pf.space_before is not None and pf.space_before.pt:
            pf.space_before = Pt(round(pf.space_before.pt*after_scale,1))
    return d

def pairlist(d, items, indent="—  ", gap="        ", after=4, budget=78):
    """Lay short bullet items several to a line, packed by width.

    Every item is kept verbatim; nothing is removed, reworded or reordered,
    and neither type size nor leading changes. This reclaims the unused
    right-hand part of a short-bullet line, which is what keeps the Shorts
    editor brief to two pages. The width budget is measured against the
    inspection renderer's font, which is wider than Calibri, so Word fits at
    least as much per line and never less."""
    lines, cur = [], []
    for it in items:
        trial = cur + [it]
        width = sum(len(indent) + len(x) for x in trial) + len(gap) * (len(trial) - 1)
        if cur and width > budget:
            lines.append(cur); cur = [it]
        else:
            cur = trial
    if cur: lines.append(cur)
    for group in lines:
        keep(P(d, gap.join(indent + x for x in group), after=after))





# ------------------------------------------------- 1. teleprompter DOCX + TXT
TEL_DOCX="Video7TeleprompterScriptwithslidemarkers_HIT_v2.0.docx"
TEL_TXT ="Video7TeleprompterScriptwithslidemarkers_HIT_v2.0.txt"
RD_DOCX ="Video7ReadingScriptnomarkers_HIT_v2.0.docx"
RD_TXT  ="Video7ReadingScriptnomarkers_HIT_v2.0.txt"

d=newdoc(True)
head(d,TITLE,"Video 7  ·  Teleprompter script with slide markers",
     "Spoken script is the large text. A slide marker in a tinted band tells "
     "the editor which slide to bring up; it is not spoken.")
for line in LINES:
    if line.startswith("[SLIDE:"):
        p=P(d,"SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(),size=11,bold=True,
            color=NAVY,before=14,after=14,spacing=1.1)
        shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
    else:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,TEL_DOCX))
tel=[TITLE,"Video 7  ·  Teleprompter script with slide markers",""]
for line in LINES:
    if line.startswith("[SLIDE:"):
        tel += ["", "SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(), ""]
    else: tel += [line, ""]
open(os.path.join(LF,TEL_TXT),"w").write("\n".join(tel).strip()+"\n")

# ------------------------------------------------- 2. reading script DOCX+TXT
d=newdoc(True)
head(d,TITLE,"Video 7  ·  Reading script, no markers",
     "Spoken language only. No slide markers, no timestamps, no production "
     "directions.")
for line in SPOKEN:
    keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,RD_DOCX))
open(os.path.join(LF,RD_TXT),"w").write("\n\n".join(SPOKEN)+"\n")
print("long-form scripts written")

# --------------------------------------------------- 3. long-form editor brief
# SCRIPT MARKER -> DECK SLIDE -> ACTUAL DECK TITLE AFTER THE AUTHORISED
# CORRECTIONS. The mapping is by POSITION, not by name: several marker names
# do not read like the slide title they cue, which is exactly why this table
# exists. Deck titles below were read back out of the corrected PPTX files.
MAPPING=[
 (1,"Core Distinction",1,
  "BUILDING WHILE OPERATING / THE INFRASTRUCTURE WAS STILL MATURING."),
 (2,"The Invisible-Work Problem",2,
  "WHY THE WORK GOES QUIET / MOST WORKPLACE EVIDENCE IS COMPARATIVE."),
 (3,"Reconstruct the Before",3,"THE THREE MOVES"),
 (4,"Before Questions",4,"MOVE ONE / RECONSTRUCT THE STARTING CONDITION"),
 (5,"Show the Judgment",5,"WHAT WAS STILL MATURING?"),
 (6,"Built From Scratch Does Not Mean Built Alone",6,
  "MOVE TWO / THE OUTPUT IS THE SMALLEST PART."),
 (7,"Building While Operating",7,
  "JUDGMENT IS THE PART THAT DOES NOT TRANSFER INTO A SLIDE."),
 (8,"Keep the Proof",8,"MOVE THREE / WHAT IS DIFFERENT NOW?"),
 (9,"Proof Without Overclaiming",9,
  "A NUMBER YOU CANNOT DEFEND IS WORSE THAN NO NUMBER."),
 (10,"What Did the Building Form?",10,"THREE PARAGRAPHS"),
 (11,"Make Invisible Work Visible",11,
  "KEEP THE PROOF / A 60-MINUTE CAREER EVIDENCE SYSTEM"),
 (12,"Keep the Proof / Watch Next",12,
  "WATCH NEXT / HOW TO EXPLAIN YOUR CAREER CHANGE"),
]

d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 7",size=12,bold=True,color=GOLD,after=2,caps=True)
P(d,TITLE,size=20,bold=True,color=NAVY,after=6,spacing=1.1)
p=P(d,"This document is for the editor. It is NOT Temidayo's teleprompter and "
     "must not be placed on the recording screen.",size=11,italic=True,
     color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"Locked metadata",before=14)
for k,v in (("Title",TITLE),("Thumbnail",THUMB),("CTA",CTA),
            ("CTA URL",CTA_URL),("Watch next",NEXT),
            ("Capability Formation question",
             "Can the organization see and use capability it already has?"),
            ("Core distinction",
             "Sometimes the hardest work to prove is the work that made later "
             "work possible.")):
    keep(P(d,"%-32s %s"%(k+":",v),size=11,after=5))

H1(d,"Script marker → deck slide → deck title",before=14)
keep(P(d,"The twelve [SLIDE: ...] markers in the teleprompter map to the "
       "twelve deck slides BY POSITION, in order. Several marker names do not "
       "read like the slide they cue — marker 5 is named Show the Judgment but "
       "cues WHAT WAS STILL MATURING?, and marker 8 is named Keep the Proof "
       "but cues MOVE THREE. Follow the slide number in this table, not the "
       "marker name.",size=10.5,color=DIM,after=10,spacing=1.25))
# A real Word table: space-padded columns collapse in a proportional font,
# which is what made the first attempt run together into one line.
tbl=d.add_table(rows=1, cols=4)
tbl.style="Table Grid"
tbl.autofit=False
for cell,wdt,label in zip(tbl.rows[0].cells,(0.45,2.65,0.55,2.75),
                          ("MK","SCRIPT MARKER","SLIDE","DECK TITLE")):
    cell.width=Inches(wdt)
    cp=cell.paragraphs[0]; cp.paragraph_format.space_after=Pt(2)
    r=cp.add_run(label); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=NAVY
    shade(cp,BAND_NAVY)
for mk,name,sl,title in MAPPING:
    cells=tbl.add_row().cells
    for cell,wdt,val,bold in zip(cells,(0.45,2.65,0.55,2.75),
                                 (str(mk),name,str(sl),title),
                                 (False,False,False,True)):
        cell.width=Inches(wdt)
        cp=cell.paragraphs[0]
        cp.paragraph_format.space_after=Pt(2); cp.paragraph_format.line_spacing=1.1
        r=cp.add_run(val); r.font.size=Pt(8.5); r.bold=bold
        if bold: r.font.color.rgb=NAVY
P(d,"",size=6,after=2)
keep(P(d,"Deck titles above were read back out of the corrected PPTX files "
       "after the authorised Slide 1, 4, 5 and 12 changes were applied.",
       size=10.5,italic=True,color=DIM,before=6,after=8,spacing=1.25))

H1(d,"Authorised slide corrections — applied",before=14)
p=P(d,"AUTHORISED AND APPLIED. Slides 1, 4, 5 and 12 carried wording that "
     "contradicted the approved factual framing. Text only changed; design, "
     "typography, palette, layout and box positions are unchanged.",size=11,
     bold=True,color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
for sl,frm,to in [
 ("Slide 1 (upper)","YOU WERE NOT / IMPROVING SOMETHING.",
  "BUILDING WHILE OPERATING"),
 ("Slide 1 (lower)","YOU WERE THE BEFORE.",
  "THE INFRASTRUCTURE / WAS STILL MATURING."),
 ("Slide 4 (title)","DOCUMENT THE ABSENCE / YOU WALKED INTO.",
  "RECONSTRUCT THE STARTING CONDITION"),
 ("Slide 4 (sub-copy)","two paragraphs of absence framing",
  "What was incomplete, inconsistent or difficult before the work?"),
 ("Slide 5 (title)","WHAT DID NOT EXIST?","WHAT WAS STILL MATURING?"),
 ("Slide 5 (list)","eight absence claims",
  "the seven approved still-maturing categories, sentence case"),
 ("Slide 12 (title)","HOW TO EXPLAIN A / NONLINEAR CAREER / WITHOUT LOOKING / "
  "UNFOCUSED","HOW TO EXPLAIN YOUR / CAREER CHANGE")]:
    keep(P(d,sl,size=11,bold=True,color=NAVY,before=8,after=3))
    keep(P(d,"FROM:  %s"%frm,size=10.5,after=3))
    keep(P(d,"TO:    %s"%to,size=10.5,bold=True,after=4))
keep(P(d,"Slide 5's list dropped from eight items to seven. The item boxes "
       "were widened from 5.28in to 5.86in so the longer category names hold "
       "one line each at the existing 21pt DM Sans; the eighth slot and its "
       "divider rule were removed. Every line was measured against the real "
       "embedded fonts before the change was applied.",size=10.5,color=DIM,
       before=6,after=8,spacing=1.25))
keep(P(d,"Reveal frames corrected: 1, 2 (Slide 1); 8, 9 (Slide 4); 10, 11 "
       "(Slide 5); 24 (Slide 12). All seventeen other reveal frames are "
       "byte-identical. Main slides 2, 3, 6, 7, 8, 9, 10 and 11 are "
       "byte-identical.",bold=True,after=10,spacing=1.25))

H1(d,"First 30 seconds — H.I.T.",before=14)
P(d,"H = Hook. I = Interest. T = Trust. One audiovisual unit: immediate human "
    "recognition, meaningful visual interest, specific personal evidence and a "
    "clear viewer payoff by roughly 20 to 30 seconds. No exaggerated "
    "“nothing existed” language. No forced statistic. No heroic "
    "solo-builder framing.",after=14)

def beat(t,anchor,layer,body):
    H2(d,t,before=10)
    if anchor:
        p=P(d,"Spoken:  “%s”"%anchor,size=10.5,italic=True,color=DIM,after=8)
        shade(p,BAND_CREAM)
    if layer: keep(P(d,layer,size=11,bold=True,color=GOLD,after=6))
    for b in body: keep(P(d,b,after=5))

beat("0:00–0:06",
     "Some of the hardest work to prove is the work that made everything else "
     "possible.","H = HOOK",
     ["Visual: direct to camera.",
      "On-screen:  THE WORK BECAME INVISIBLE",
      "Do not use THERE WAS NOTHING HERE or any equivalent overstatement."])
beat("0:06–0:17",
     "In one senior people role, I joined a global remote-first company where "
     "parts of the change, onboarding, governance and organizational-"
     "effectiveness infrastructure were still incomplete.",
     "T = TRUST  /  I = INTEREST",
     ["Visual:","     CHANGE","     ONBOARDING","     GOVERNANCE",
      "     ORG EFFECTIVENESS","Do not show company name or logo."])
beat("0:17–0:24","I wasn’t building everything from zero. I was building while "
     "operating…",None,
     ["On-screen:  BUILDING WHILE OPERATING","This is the factual tension."])
beat("0:24–0:30",None,"PAYOFF",
     ["Progressively reveal:","     RECONSTRUCT THE BEFORE",
      "     SHOW THE JUDGMENT","     KEEP THE PROOF",
      "Then enter the existing deck."])

H1(d,"Factual visual boundary",before=14)
P(d,"Never visually imply:",after=5)
for x in ["nothing existed before Temidayo;","Temidayo built everything alone;",
 "one initiative caused all later outcomes;","the company was dysfunctional."]:
    keep(P(d,"—  "+x,after=4))
p=P(d,"Use:  incomplete  ·  still maturing  ·  needed stronger mechanisms  ·  "
     "building while operating",size=11,bold=True,color=GOLD,before=6,after=10,
     spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"Personal proof",before=14)
keep(P(d,"Approved visible proof:",after=6))
for x in ["People became a more regular input into executive planning within "
 "about 12 months;",
 "one measure of how well new hires felt integrated moved from 47 to 75 "
 "during the onboarding redesign Temidayo led with her team;",
 "post-acquisition integration completed within 90 days with no critical "
 "talent loss;",
 "a global CSR program launched and scaled across all regions."]:
    keep(P(d,"—  "+x,after=5,spacing=1.25))
keep(P(d,"Do not combine all outcomes into one causal result card. Keep them "
       "distinct.",bold=True,before=4,after=8))
keep(P(d,"Company name is not used publicly. The chapter stays unnamed.",
       bold=True,after=10))

H1(d,"CSR visual boundary",before=14)
P(d,"If illustrating the CSR program, do not use company branding and do not "
    "invent regions, participation metrics, donation values, volunteer hours "
    "or any social-impact result.",after=6,spacing=1.25)
p=P(d,"GLOBAL CSR PROGRAM   ·   LAUNCHED → SCALED ACROSS REGIONS  is "
     "sufficient.",size=11,bold=True,color=GOLD,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"Excluded metrics",before=14)
p=P(d,"Do not use an approximately 30% retention improvement or a "
     "more-than-$2M avoided-turnover figure anywhere: not in the script, "
     "Shorts, description, briefs, slides or metadata.",size=11,bold=True,
     color=RED,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"CTA and watch next",before=14)
keep(P(d,"One offer only: %s — %s"%(CTA_SUB,CTA_URL),after=5,spacing=1.25))
keep(P(d,"Do not add the Capability Formation Field Kit or the Career "
       "Decision Evidence Check.",bold=True,after=8))
keep(P(d,"Watch next: %s  (Video 4)"%NEXT,bold=True,after=5))
keep(P(d,"Do not leave Subscribe as the only end-screen element.",bold=True,
       color=RED,before=4,after=8))
compress(d, 1.15, 0.62)
d.save(os.path.join(LF,"Video_7_EDITOR_ONLY_HIT_Brief_v2.0.docx"))
print("editor brief written")

# The thirteen working chapter lines, defined once and reused by the publishing
# package, its reference section and the separate description-only document.
# Offsets are script-derived at 145 wpm from the canonical spoken text; the
# first twelve sit at the twelve slide markers, the thirteenth at the Watch
# Next hand-off. They are estimates and the editor must replace them.
CHAPTERS=[("00:00","When Valuable Work Becomes Invisible"),
 ("01:24","The Invisible-Work Problem"),
 ("02:15","Reconstruct the Before"),
 ("02:55","Questions That Reveal the Starting Condition"),
 ("03:45","Show the Judgment Behind the Build"),
 ("04:25","Built From Scratch Does Not Mean Built Alone"),
 ("05:17","What Building While Operating Taught Me"),
 ("06:52","Keep Evidence Before the Infrastructure Becomes Normal"),
 ("07:39","How to Describe Results Without Overclaiming"),
 ("08:40","What Did Building This Form in You?"),
 ("09:48","Make Invisible Work Visible"),
 ("10:25","Keep the Proof"),
 ("10:55","How to Explain Your Career Change")]
CHAPTER_LINES=["%s %s"%(t,c) for t,c in CHAPTERS]

PRIMARY="how to show your impact at work"
SUPPORTING=("how to show impact at work · how to prove your impact at work · "
 "career impact · foundational work · first in role · built from scratch · "
 "invisible work · career evidence · leadership impact · career portability")
TAGS=("how to show your impact at work, how to show impact at work, how to "
 "prove your impact at work, career impact, foundational work, first in role, "
 "built from scratch, invisible work, career evidence, leadership impact, "
 "career growth, career portability, Temidayo Afonja, Capability Formation")
DESC=[
 "Some of the most valuable work in a career is also the hardest to prove — "
 "especially when part of your job was building the infrastructure that later "
 "became “the way we work.”",
 "In this video, I show you how to make foundational work visible without "
 "pretending you built everything alone or overstating what existed before "
 "you arrived.",
 "✨ Reconstruct the before — What was incomplete, unclear or difficult "
 "before the work?",
 "✨ Show the judgment — What did you have to notice, decide or design before "
 "the mechanism could work?",
 "✨ Keep the proof — What changed, and what permitted evidence can you "
 "safely support?",
 "I also share evidence from one senior people chapter where I was building "
 "while operating: strengthening change, onboarding, governance and "
 "organizational-effectiveness mechanisms while the business continued moving.",
 "That chapter also included launching a global CSR program that scaled "
 "across all regions.",
 "The goal is not to say:",
 "“I built everything.”",
 "It is to answer:",
 "“What part of the infrastructure was I responsible for making possible — "
 "and what did building it teach me to do?”",
 "Sometimes the most transferable part of being first is what you had to "
 "learn before the role could work at all.","",
 "🧭 KEEP THE PROOF",
 "A 60-Minute Career Evidence System:",
 CTA_URL,"",
 "⏱️ CHAPTERS"]+CHAPTER_LINES+["",
 "▶️ WATCH NEXT", NEXT, "[ADD VIDEO 4 LINK]","",
 "🔗 CONNECT AND EXPLORE",
 "Website:","https://temidayoafonja.com",
 "LinkedIn:","https://www.linkedin.com/in/temidayo-afonja",
 "Substack:","https://temidayoafonja.substack.com","",
 "#CareerGrowth #Leadership #CareerImpact"]
PINNED=["What is something you helped build that now looks completely normal "
 "inside the organization?",
 "You do not need to name the company or share confidential details.",
 "Try answering:","What was incomplete before?",
 "What did you have to make possible?",
 "What evidence shows that the environment changed?",
 "If keeping that evidence is the difficult part, Keep the Proof is here:",
 CTA_URL]

EMOJI_NOTE=("The restrained emoji system is part of the approved standard: ✨ "
  "teaching points, 🧭 CTA/resource, ⏱️ chapters, ▶️ Watch Next, 🔗 Connect "
  "and Explore. Do not remove them and do not add more.")

def description_block(d, heading_before=14, upload_doc=False):
    """The copy-ready description, its end marker and the internal note.

    upload_doc=True moves the editorial emoji instruction ABOVE an explicit
    COPY-READY ... BEGIN marker, so nothing internal sits inside or immediately
    before the block a person selects when pasting into YouTube."""
    if upload_doc:
        H1(d,"INTERNAL NOTE — DO NOT PASTE INTO YOUTUBE",before=heading_before)
        p=P(d,EMOJI_NOTE,size=10.5,italic=True,color=RED,after=12,spacing=1.25)
        shade(p,BAND_CREAM); keep(p)
        p=keep(P(d,"COPY-READY YOUTUBE DESCRIPTION — BEGIN",size=11,bold=True,
                 color=NAVY,before=14,after=12,spacing=1.2))
        shade(p,BAND_NAVY)
    else:
        H1(d,"Description",before=heading_before)
        keep(P(d,EMOJI_NOTE,size=10.5,italic=True,color=DIM,after=10))
    for para in DESC:
        keep(P(d,para if para else " ",after=7 if para else 3))
    keep(P(d,"— END OF THE COPY-READY DESCRIPTION —",size=10,bold=True,
           color=DIM,before=14,after=12,spacing=1.2))
    H1(d,"Internal note — do not paste into YouTube",before=14)
    p=P(d,"WORKING ESTIMATES — EDITOR MUST REPLACE FROM FINAL CUT",size=11,
        bold=True,color=RED,after=6,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    p=P(d,"The chapter timestamps above are script-derived, not final. Replace "
        "every one of them using the finished cut before publication. Do not "
        "force the edit to match these estimates.",size=10.5,bold=True,
        italic=True,color=RED,after=10,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    H1(d,"Working chapters — reference copy",before=14)
    keep(P(d,"Identical to the thirteen chapter lines inside the description "
           "above.",size=10.5,italic=True,color=DIM,after=8))
    for line in CHAPTER_LINES: keep(P(d,line,size=11,after=4))


# ----------------------------------------------------- 4. publishing package
d=newdoc()
head(d,TITLE,"Video 7  ·  Publishing package",
     "Everything needed to upload. Working timestamps must be replaced with "
     "real ones from the finished edit.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14); P(d,PRIMARY,after=10)
H1(d,"Supporting search language",before=14); P(d,SUPPORTING,after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only. Do not put the full tag field in the "
       "public description.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
H1(d,"Watch next",before=14)
keep(P(d,"%s  (Video 4)"%NEXT,bold=True,after=5))
keep(P(d,"Slide 12 carries this title. The correction record is in the EDITOR "
       "ONLY brief.",size=10.5,color=DIM,after=8,spacing=1.25))
compress(d)
d.save(os.path.join(LF,"Video_7_Publishing_Package_HIT_v2.0.docx"))

# ------------------------------- 4b. separate description-only document
d=newdoc()
head(d,TITLE,"Video 7  ·  YouTube description",
     "Upload copy only. Everything below the end marker is internal and must "
     "not be pasted into YouTube.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14); P(d,PRIMARY,after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"Watch next",before=14); keep(P(d,"%s  (Video 4)"%NEXT,bold=True,after=8))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
compress(d)
DESC_DOC="/tmp/v7hit/Video_7_YouTube_Description_HIT.docx"
d.save(DESC_DOC)
print("publishing package and description-only document written")

# ---------------------------------------------------------------- 5. Shorts
from shorts_text import SHORTS
LABELS=["SHORT 1","SHORT 2","SHORT 3","SHORT 4"]
for (fn,role,hook,copy),label in zip(SHORTS,LABELS):
    d=newdoc(True)
    P(d,"VIDEO 7 SHORT",size=10,bold=True,color=GOLD,after=4,caps=True)
    P(d,label,size=20,bold=True,color=NAVY,after=8,spacing=1.1)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    keep(P(d,"Verbal hook:  “%s”"%hook,size=11,color=DIM,after=5))
    keep(P(d,"Related long-form:  %s"%TITLE,size=11,color=DIM,after=10))
    H1(d,"RECORDING COPY",before=12)
    for line in copy:
        keep(P(d,line,size=13.5,color=INK,after=10,spacing=1.5))
    d.save(os.path.join(SH,fn))

# ------------------------------------------------------ 6. Shorts editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 7 — FOUR STANDALONE SHORTS",size=18,bold=True,color=NAVY,
  after=8,spacing=1.1)
p=P(d,"This document is for the editor. It is separate from the four Short "
     "recording documents and must not be placed on Temidayo's recording "
     "screen.",size=11,italic=True,color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)
H1(d,"How these are produced",before=14)
keep(P(d,"These are separately recorded 9:16 Shorts. They are NOT excerpts cut "
       "from the long-form video.",bold=True,after=10))
P(d,"Each should have:",after=6)
pairlist(d,["an immediate verbal hook;","a matching on-screen hook;",
 "meaningful visual interest;","accurate mobile-safe captions;",
 "restrained editorial pacing;",
 "Video 7 added as the YouTube Related Video when available."])

def short(label,role,onscreen,body):
    H1(d,label,before=14)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    p=keep(P(d,"On-screen hook:  %s"%onscreen,size=11,bold=True,color=GOLD,after=8))
    shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))
    keep(P(d,"Related Video:  Video 7",size=10.5,color=DIM,before=4,after=6))

short("SHORT 1","Recognition","THE WORK BECAME INVISIBLE",
 ["Visual:  BEFORE  →  NOW NORMAL",
  "Avoid “nothing existed” in every caption and card.",
  "End:  RECONSTRUCT THE BEFORE"])
short("SHORT 2","Distinction","BUILT FROM SCRATCH ≠ BUILT ALONE",
 ["Visual:  LEAD  /  PARTNER  /  BUILD WITH OTHERS",
  "Do not imply solo ownership.",
  "End:  WHAT WAS I RESPONSIBLE FOR MAKING POSSIBLE?"])
short("SHORT 3","Personal proof","BUILDING WHILE OPERATING",
 ["Visual:","     CHANGE","     ONBOARDING","     GOVERNANCE",
  "     ORG EFFECTIVENESS","     GLOBAL CSR",
  "Do not use company branding.",
  "For CSR:  GLOBAL CSR PROGRAM  ·  LAUNCHED → SCALED ACROSS REGIONS",
  "No invented CSR metrics."])
short("SHORT 4","Practical action","RECONSTRUCT THE BEFORE",
 ["Reveal:","     WHAT COULD PEOPLE NOT DO RELIABLY?",
  "     WHERE WAS OWNERSHIP UNCLEAR?",
  "     WHAT DEPENDED ON INFORMAL EFFORT?",
  "     WHAT KEPT RECURRING?",
  "End:  MAKE FOUNDATIONAL WORK VISIBLE"])

H1(d,"All Shorts — visual boundaries",before=14)
P(d,"Do not use:",after=5)
pairlist(d,["stock “builder” imagery;","construction metaphors;",
 "toolbelt graphics;","company logos;","fake shock expressions;",
 "hyperactive zooms;","invented CSR metrics;","solo-ownership framing;",
 "“nothing existed” language;","trendy caption templates."],after=3)
compress(d, 1.18, 0.62)
d.save(os.path.join(SH,"Video_7_Shorts_EDITOR_ONLY_HIT_Brief.docx"))
print("shorts and shorts editor brief written")

# ---------------------------------------------------------------- 7. README
ZIPNAME="Video_7_HIT_FINAL_Recording_and_Shorts_Package.zip"
FILES=(["LONG_FORM/"+f for f in sorted(os.listdir(LF))]
      +["SHORTS/"+f for f in sorted(os.listdir(SH))])
R=["VIDEO 7 — H.I.T. FINAL RECORDING PACKAGE","",
 "Title:             %s"%TITLE,
 "Thumbnail:         %s"%THUMB,
 "CTA:               %s"%CTA,
 "CTA URL:           %s"%CTA_URL,
 "Watch next:        %s"%NEXT,"",
 "Long-form:         Revised under H.I.T.",
 "Personal proof:    Uses the unnamed global remote-first B2B SaaS chapter.",
 "Company name:      NOT USED PUBLICLY.","",
 "Approved proof:",
 "  - People became a more regular input into executive planning within about",
 "    12 months",
 "  - onboarding measure 47 -> 75",
 "  - post-acquisition integration completed within 90 days with no critical",
 "    talent loss",
 "  - global CSR program launched and scaled across all regions","",
 "Excluded:",
 "  - ~30% retention improvement",
 "  - >$2M avoided turnover","",
 "Shorts:            Four separately recorded vertical scripts.",
 "Editor",
 "instructions:      Separated from all recording copy.",
 "Description-only",
 "document:          Video_7_YouTube_Description_HIT.docx, created separately",
 "                   OUTSIDE this package ZIP.","",
 "Verified in the",
 "live files:        12 main slides, 24 reveal-build frames.","",
 "Slide corrections: Slides 1, 4, 5 and 12 carried wording that contradicted",
 "                   the approved factual framing and were corrected under",
 "                   explicit authorisation. Text only; design, typography,",
 "                   palette, layout and box positions unchanged.",
 "                   Reveal frames corrected: 1, 2, 8, 9, 10, 11, 24.",
 "                   Main slides 2, 3, 6, 7, 8, 9, 10, 11 and the other",
 "                   seventeen reveal frames are byte-identical.",
 "Thumbnail:         UNCHANGED.","",
 "-"*70,"","WHAT EACH FILE IS","",
 "LONG_FORM/","",
 "  "+TEL_DOCX,
 "  "+TEL_TXT,
 "      Temidayo's recording copy. Spoken script in large text; slide markers",
 "      in tinted bands. The markers are not spoken.","",
 "  "+RD_DOCX,
 "  "+RD_TXT,
 "      The same spoken words with the slide markers removed.","",
 "  Video_7_EDITOR_ONLY_HIT_Brief_v2.0.docx",
 "      For the editor. Locked metadata, the script-marker to deck-slide",
 "      mapping table, the authorised slide-correction record, the H.I.T.",
 "      first-30-second plan, the factual and CSR visual boundaries, the",
 "      excluded metrics and the CTA. Not for the teleprompter.","",
 "  Video_7_Publishing_Package_HIT_v2.0.docx",
 "      Title, thumbnail, search language, the copy-ready description with",
 "      its approved emoji system, working chapter estimates, pinned comment",
 "      and the tag field.","",
 "SHORTS/","",
 "  Four recording documents, one per Short. These contain Temidayo's",
 "  recording copy and no editor directions.","",
 "  Video_7_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "      For the editor. On-screen hooks and visual treatment for all four.","",
 "-"*70,"","ALL FILES IN THIS PACKAGE","",]
for f in FILES: R.append("  "+f)
R+=["  README_FINAL.txt","  SHA256SUMS.txt","",
 "  Video_7_YouTube_Description_HIT.docx is deliberately NOT in this ZIP.","",
 "-"*70,"","WORKING CHAPTER TIMESTAMPS","",
 "The thirteen chapter timestamps are WORKING ESTIMATES derived from the",
 "script at 145 words per minute. They were not measured from an edit. The",
 "editor must replace every one of them from the finished cut before",
 "publishing.","",
 "-"*70,"","FACT AND PROOF BOUNDARY","",
 "The chapter is not named publicly. Only the four approved outcomes are",
 "used, and they are kept distinct rather than combined into one causal",
 "result. No CSR participation rate, volunteer hour, donation amount,",
 "region count or social-impact result is invented. Neither the approximately",
 "30% retention improvement nor the more-than-$2M avoided-turnover figure",
 "appears anywhere in this package.","",
 "-"*70,"","CHECKSUMS","",
 "SHA256SUMS.txt covers the other 12 user-facing files in this package. It",
 "does not hash itself, and it carries no ZIP checksum. The archive's own",
 "SHA-256 is in the sibling file:",
 "  "+ZIPNAME+".sha256","",
 "-"*70,"","WHAT WAS NOT CHANGED","",
 "Every website file, every product, the approved thumbnail and every other",
 "video are unchanged. The only visual changes in this pass are the",
 "authorised Slide 1, 4, 5 and 12 text corrections and their matching reveal",
 "frames.",""]
open(os.path.join(ROOT,"README_FINAL.txt"),"w").write("\n".join(R))

# ------------------------------------------- 8. checksums and the master ZIP
MANIFEST=[
 "LONG_FORM/"+TEL_DOCX,
 "LONG_FORM/"+TEL_TXT,
 "LONG_FORM/"+RD_DOCX,
 "LONG_FORM/"+RD_TXT,
 "LONG_FORM/Video_7_EDITOR_ONLY_HIT_Brief_v2.0.docx",
 "LONG_FORM/Video_7_Publishing_Package_HIT_v2.0.docx",
 "SHORTS/Video_7_Short_1_Work_Became_Invisible.docx",
 "SHORTS/Video_7_Short_2_Built_From_Scratch_Not_Alone.docx",
 "SHORTS/Video_7_Short_3_Building_While_Operating.docx",
 "SHORTS/Video_7_Short_4_Reconstruct_The_Before.docx",
 "SHORTS/Video_7_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "README_FINAL.txt",
]
SUMS="SHA256SUMS.txt"
ZIP="/tmp/v7hit/"+ZIPNAME

def sha256(p):
    h=hashlib.sha256()
    with open(p,"rb") as fh:
        for b in iter(lambda: fh.read(1<<20), b""): h.update(b)
    return h.hexdigest()

for m in MANIFEST:
    assert os.path.isfile(os.path.join(ROOT,m)), "missing from build: "+m
on_disk=set()
for dp,dn,fn in os.walk(ROOT):
    dn[:]=[x for x in dn if x!="__pycache__"]
    for f in fn:
        if f.endswith(".pyc"): continue
        on_disk.add(os.path.relpath(os.path.join(dp,f),ROOT).replace(os.sep,"/"))
unexpected=sorted(on_disk-set(MANIFEST)-{SUMS})
assert not unexpected, "unexpected files in package directory: %r"%unexpected

L=["# VIDEO 7 - H.I.T. FINAL RECORDING PACKAGE",
   "# SHA-256 of the 12 user-facing files in this package.",
   "# SHA256SUMS.txt cannot hash itself. The master ZIP cannot contain its own",
   "# checksum either; it is published in the sibling file",
   "# "+ZIPNAME+".sha256",
   "# Video_7_YouTube_Description_HIT.docx sits outside this package; its",
   "# SHA-256 is reported in the delivery summary.",""]
for m in MANIFEST: L.append("%s  %s"%(sha256(os.path.join(ROOT,m)),m))
open(os.path.join(ROOT,SUMS),"w").write("\n".join(L)+"\n")

if os.path.exists(ZIP): os.remove(ZIP)
with zipfile.ZipFile(ZIP,"w",zipfile.ZIP_DEFLATED) as z:
    for m in MANIFEST+[SUMS]:
        z.write(os.path.join(ROOT,m), "Video_7_HIT_FINAL/"+m)
zsha=sha256(ZIP)
open(ZIP+".sha256","w").write("%s  %s\n"%(zsha,os.path.basename(ZIP)))

PROV="/tmp/v7hit/_source"
shutil.rmtree(PROV,ignore_errors=True); os.makedirs(PROV)
for f in ("script_text.py","shorts_text.py","build.py","qa.py","proxy.py",
          "drop_orphan.py","canonical_script.txt"):
    src="/tmp/v7hit/"+f
    if os.path.isfile(src): shutil.copy2(src, os.path.join(PROV,f))
CANON_SRC=("/root/.claude/uploads/f121668d-e262-5eb8-9b22-0eaa1006a361/"
           "5b35ceab-Video_7_Code_Prompt_HIT_Final.txt")
if os.path.isfile(CANON_SRC):
    shutil.copy2(CANON_SRC, os.path.join(PROV,"Video_7_Code_Prompt_HIT_Final.txt"))
print("ZIP sha256:",zsha)
print("description-only doc sha256:",sha256(DESC_DOC))
