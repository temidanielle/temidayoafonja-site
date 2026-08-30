# -*- coding: utf-8 -*-
"""Author the Video 7 production package DOCX.

This file is the single source of truth for Video 7. The teleprompter DOCX and
the clean recording TXT are generated from the DOCX it writes, never typed
separately, so the three can never drift.

Structure mirrors the approved Video 6 package section for section.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
OUT = os.path.join(ROOT,
                   "YouTube_Video_7_Production_Package_Impact_Without_Blueprint.docx")

NAVY = RGBColor(0x0F, 0x23, 0x46)
GOLD = RGBColor(0x8A, 0x6D, 0x1E)
DIM  = RGBColor(0x5A, 0x6B, 0x82)
INK  = RGBColor(0x1A, 0x1A, 0x1A)

TITLE = "How to Show Your Impact at Work When You Built It From Scratch"
KEYWORD = "how to show your impact at work"
CTA_NAME = "Keep the Proof — A 60-Minute Career Evidence System"
CTA_URL = "https://temidayoafonja.com/keep-the-proof"
THUMB = "MAKE INVISIBLE WORK VISIBLE"

# ---------------------------------------------------------------------------
# Section 4. The spoken script. Timed headers are production directions and are
# stripped from the clean TXT; every other line here is spoken verbatim.
# ---------------------------------------------------------------------------
SCRIPT = [
 ("0:00–0:35 | HOOK + PROMISE", [
  "Some of the most valuable work you have ever done may be the hardest work "
  "to describe. Not because it was small. Because it was first.",
  "If you built something where there had been nothing — a function, a "
  "process, a way of deciding, a set of relationships that had never existed "
  "in that shape before — then you already know the problem. The work was "
  "real. The record of it is thin.",
  "In the next ten minutes I want to give you a way to make that work legible "
  "again. Not louder. Legible.",
 ]),
 ("0:35–1:20 | EARLY PROOF", [
  "When I was first in roles, I sometimes had to create the language, the "
  "relationships and the decision clarity before the work itself could "
  "succeed.",
  "That is a strange position to sit in. On paper you have a title and a set "
  "of responsibilities. In practice, the first thing you are doing is "
  "building the conditions that make those responsibilities possible at all.",
  "Nobody schedules that. It rarely appears in a job description. And when "
  "the year closes and someone asks what you accomplished, the honest answer "
  "— I made this possible — can sound like nothing at all.",
  "It was not nothing. It was the part that had to happen first.",
 ]),
 ("1:20–2:25 | WHY FOUNDATIONAL WORK GOES QUIET", [
  "So let me be precise about why this work disappears, because the reason is "
  "not the one people usually give you.",
  "Most workplace evidence is comparative. It depends on a before and an "
  "after. A baseline moved. A number improved. A cycle time dropped. That is "
  "the shape almost every performance conversation expects.",
  "Foundational work breaks that shape, because you are the before. There was "
  "no owner, so there is no prior performance to compare against. There was "
  "no system, so there was nothing keeping score. You did not move a metric. "
  "You created the conditions under which a metric could exist.",
  "Which means the work is not invisible because you were too quiet about it. "
  "It is invisible because the instrument that would have recorded it did not "
  "exist yet.",
  "I want to be careful here, because that distinction matters. This is not a "
  "story about being overlooked. It is a description of how measurement "
  "works, and of what happens to anyone who arrives before the measuring "
  "does.",
  "That changes what you should do about it. The answer is not to speak up "
  "more. The answer is to reconstruct the record — and there are three moves. "
  "Reconstruct the before. Name what you built. Show what it returned.",
 ]),
 ("2:25–4:20 | MOVE ONE — RECONSTRUCT THE BEFORE", [
  "Start with the first move, which is the one most people skip and the one "
  "that does the most work.",
  "You are not describing what the place was like when you arrived. You are "
  "documenting the absence you walked into.",
  "So write down what did not exist. Was there an owner, or did the problem "
  "belong to everyone and therefore to no one? Was there a system, or a set "
  "of disconnected efforts that happened to overlap? Was there shared "
  "language for the problem, or did three teams describe it three different "
  "ways and assume they were agreeing? Was there a baseline anyone actually "
  "trusted? Was there a decision process, or did every case get escalated on "
  "its own merits? Were the relationships already in place across the "
  "functions you needed, or did you have to build each one? Was there a "
  "standard? Was there a repeatable method, or did the outcome depend "
  "entirely on who happened to be doing it that month?",
  "Be plain when you write this. Not: the process was immature. Instead: "
  "there was no process — requests arrived through three different channels "
  "and nobody reconciled them.",
  "The absence is your baseline. It is the only honest one available to you, "
  "and once it is written down, everything that follows has something to "
  "stand on.",
 ]),
 ("4:20–6:20 | MOVE TWO — NAME WHAT YOU BUILT", [
  "The second move is to name what the work actually required you to build.",
  "Most people answer this with the output. The tracker. The programme. The "
  "playbook. The team. The output is the smallest part of it, and it is also "
  "the part most easily dismissed as a document somebody made.",
  "Go underneath it. What definitions did you have to write before anyone "
  "could agree on what they were counting? What relationships did you have to "
  "build, and with whom, before the work could move at all? What decisions "
  "did you make with incomplete information — and what did you choose, and "
  "why? Whose alignment did you have to earn, and what did earning it "
  "actually take? And what did you put in place that other people could use "
  "afterwards, without you in the room?",
  "That last one carries the most weight. A capability is something the "
  "organisation can do now that it could not do before, whether or not you "
  "are still there.",
  "Then be specific about judgment, because judgment is the part that does "
  "not transfer into a slide, and so it is the part that goes unrecorded.",
  "If you narrowed the scope so the first version would actually land, that "
  "was judgment. If you sequenced the difficult conversation before the "
  "design work because you knew the design would not survive without it, that "
  "was judgment. Name the choice. Name the alternative you rejected. Name the "
  "reason.",
  "That is not self-promotion. It is a record of thinking, and it is the "
  "thing an experienced reader is actually looking for.",
 ]),
 ("6:20–8:30 | MOVE THREE — SHOW WHAT THE WORK RETURNED", [
  "The third move is to show what the work returned.",
  "This is where most people reach for a percentage they do not have, feel "
  "uneasy about it, and then say nothing at all.",
  "Return does not have to be a number. The question underneath it is "
  "simpler: what is different now, and how would somebody else be able to "
  "tell?",
  "Adoption is evidence. Who uses it, and how many of them. Continued use is "
  "stronger evidence, because it means the thing survived a quarter, survived "
  "a reorganisation, survived you moving on.",
  "Better decisions are evidence. Decisions that used to take three meetings "
  "now take one, or get made at the level where they belong instead of "
  "travelling upward. Reduced ambiguity is evidence. So is repeatability — "
  "the outcome no longer depends on who is doing it.",
  "Stakeholder recognition counts, and it is often already written down "
  "somewhere: the note, the review comment, the request to run it again for "
  "another team. A clean handoff is evidence, because somebody else can carry "
  "this now. Avoided risk is evidence, though it needs care, because you are "
  "describing something that did not happen.",
  "And a capability that remained after the original project ended is the "
  "strongest evidence there is.",
  "Some of those can be counted. Many cannot. Both are real, and the ones you "
  "cannot count are not lesser — they are simply harder to say well, which is "
  "why so few people say them at all.",
 ]),
 ("8:30–9:15 | EVIDENCE WITHOUT INVENTED NUMBERS", [
  "Which brings me to a boundary I want to be direct about.",
  "Do not manufacture precision. A number you cannot defend is worse than no "
  "number, because the first serious question dissolves it and takes the rest "
  "of your account down with it. Roughly forty per cent, estimated on a "
  "Tuesday, will not survive a director asking how you measured it.",
  "What you can say is this: it was used by every team in the region for two "
  "years after I left it. That is concrete, it is verifiable, and it does not "
  "pretend to be a calculation.",
  "And keep only the evidence you are permitted to keep. Outcomes, decisions, "
  "what you learned, non-confidential examples described in your own words. "
  "Not documents. Not client or customer data. Not anything proprietary or "
  "employer-owned. What you are building is a record of your judgment and "
  "your results — not a copy of your employer's files.",
 ]),
 ("9:15–9:55 | THE APPLICATION", [
  "So here is the whole thing in one pass, and it is short enough to do this "
  "week.",
  "Take one piece of work you built from nothing and write three paragraphs.",
  "Before: what did not exist. No owner, no system, no shared language, no "
  "baseline, no method.",
  "Build: what you created underneath the output — the definitions, the "
  "relationships, the decisions, the judgment, the capability that stayed.",
  "Return: what is different now, and how someone else could tell.",
  "That is three paragraphs. It is also the answer to a question you will be "
  "asked in an interview, in a promotion case, or by a new manager who was "
  "not there when any of it happened. Write it while you still remember the "
  "specifics, because the specifics are the first thing to go.",
 ]),
 ("9:55–10:25 | PRIMARY CTA", [
  "If you want a structured way to do this, I built Keep the Proof — a "
  "sixty-minute career evidence system. It takes you through reconstructing "
  "the before, naming what you built and recording what the work returned, in "
  "a form you can keep and reuse. You will find it at "
  "temidayoafonja.com/keep-the-proof.",
 ]),
 ("10:25–10:45 | WATCH NEXT", [
  "And if the record you are building spans work that does not sit in a "
  "straight line, watch how to explain a nonlinear career without looking "
  "unfocused, next.",
 ]),
]

SPOKEN = [p for _, ps in SCRIPT for p in ps]
WORDS = sum(len(p.split()) for p in SPOKEN)

# ---------------------------------------------------------------------------
SLIDE_COPY = [
 ("1", "No blueprint", "2",
  "YOU WERE NOT IMPROVING SOMETHING.  |  YOU WERE THE BEFORE."),
 ("2", "Why it goes quiet", "2",
  "WHY THE WORK GOES QUIET  |  MOST WORKPLACE EVIDENCE IS COMPARATIVE.  |  "
  "A baseline moved. A number improved. A cycle time dropped.  |  "
  "FOUNDATIONAL WORK HAS NO PRIOR STATE.  |  The instrument that would have "
  "recorded it did not exist yet."),
 ("3", "The three moves", "3",
  "THE THREE MOVES  |  1 RECONSTRUCT THE BEFORE  |  2 NAME WHAT YOU BUILT  |  "
  "3 SHOW WHAT IT RETURNED"),
 ("4", "Reconstruct the before", "2",
  "MOVE ONE  |  DOCUMENT THE ABSENCE YOU WALKED INTO.  |  Not the state of "
  "the world when you arrived.  |  What was missing from it."),
 ("5", "What did not exist", "2",
  "WHAT DID NOT EXIST?  |  No owner  |  No system  |  No shared language  |  "
  "No baseline anyone trusted  |  No decision process  |  No relationships "
  "across functions  |  No standard  |  No repeatable method"),
 ("6", "Name what you built", "2",
  "MOVE TWO  |  THE OUTPUT IS THE SMALLEST PART.  |  DEFINITIONS · "
  "RELATIONSHIPS · DECISIONS  |  ALIGNMENT · REPEATABLE CAPABILITY  |  What "
  "the organisation can do now that it could not do before."),
 ("7", "The judgment involved", "1",
  "JUDGMENT IS THE PART THAT DOES NOT TRANSFER INTO A SLIDE.  |  Name the "
  "choice. Name the alternative you rejected. Name the reason."),
 ("8", "What the work returned", "3",
  "MOVE THREE  |  WHAT IS DIFFERENT NOW?  |  ADOPTION · CONTINUED USE · "
  "BETTER DECISIONS  |  REDUCED AMBIGUITY · REPEATABILITY · RECOGNITION  |  "
  "A CLEAN HANDOFF · A CAPABILITY THAT REMAINED"),
 ("9", "Evidence without invented numbers", "2",
  "A NUMBER YOU CANNOT DEFEND IS WORSE THAN NO NUMBER.  |  “Used by every "
  "team in the region for two years after I left it.”  |  WHAT YOU MAY KEEP  "
  "|  Outcomes, decisions, what you learned, non-confidential examples in "
  "your own words. Not documents, not data, not anything employer-owned."),
 ("10", "Before, build, return", "3",
  "THREE PARAGRAPHS  |  BEFORE — What did not exist.  |  BUILD — What you "
  "created underneath the output.  |  RETURN — What is different now—and how "
  "someone else could tell."),
 ("11", "Primary CTA, Keep the Proof", "1",
  "KEEP THE PROOF  |  A 60-MINUTE CAREER EVIDENCE SYSTEM  |  Reconstruct the "
  "before. Name what you built. Record what it returned.  |  "
  "temidayoafonja.com/keep-the-proof"),
 ("12", "Watch next", "1",
  "WATCH NEXT  |  HOW TO EXPLAIN A NONLINEAR CAREER WITHOUT LOOKING UNFOCUSED "
  " |  Career Portability: Career Pivots, Internal Moves & Growth"),
]

DESCRIPTION = [
 "If you built work from scratch, it can be difficult to show your impact at "
 "work—especially when there was no baseline, blueprint or predecessor for "
 "comparison. This video explains how to make invisible work visible by "
 "identifying the capability, evidence and value your work created.",
 "",
 "Foundational work is hard to prove for a structural reason, not a personal "
 "one. Most workplace evidence is comparative: a baseline moved, a number "
 "improved. When you are the first person to own something, there is no prior "
 "state to compare against, so the instrument that would have recorded your "
 "contribution never existed. The answer is not to speak up more. It is to "
 "reconstruct the record.",
 "",
 "You will learn three moves:",
 "1. Reconstruct the before — document the absence you walked into, not the "
 "story of what it was like.",
 "2. Name what you built — the definitions, relationships, decisions, "
 "judgment and repeatable capability underneath the output.",
 "3. Show what the work returned — adoption, continued use, better decisions, "
 "reduced ambiguity, repeatability, recognition, a clean handoff, a capability "
 "that remained.",
 "",
 "This is not personal branding, networking advice or a suggestion to speak up "
 "more. It is a method for making foundational work legible — without "
 "inventing numbers you cannot defend, and without taking anything "
 "confidential or employer-owned.",
 "",
 "KEEP THE PROOF — A 60-MINUTE CAREER EVIDENCE SYSTEM",
 "A structured way to reconstruct the before, name what you built and record "
 "what the work returned, in a form you can keep and reuse.",
 CTA_URL,
 "",
 "WATCH NEXT",
 "How to Explain a Nonlinear Career Without Looking Unfocused",
 "Playlist — Career Portability: Career Pivots, Internal Moves & Growth",
]

TAGS = ["how to show your impact at work", "how to make invisible work visible",
        "how to document your impact at work", "how to show your value at work",
        "visibility at work", "workplace visibility", "career evidence",
        "career growth", "work built from scratch", "career portability"]

CHAPTERS = [
 ("0:00", "The work that is hardest to describe"),
 ("0:35", "When you have to build the conditions first"),
 ("1:20", "Why foundational work goes quiet"),
 ("2:25", "Move one — reconstruct the before"),
 ("4:20", "Move two — name what you built"),
 ("6:20", "Move three — show what the work returned"),
 ("8:30", "Evidence without invented numbers"),
 ("9:15", "Before, build, return — three paragraphs"),
 ("9:55", "Keep the Proof"),
 ("10:25", "Watch next"),
]

PINNED = (
 "If you built something where there was nothing before, the hardest part is "
 "usually not the work — it is explaining it afterwards, when there was no "
 "baseline to compare against.\n\n"
 "Try the three paragraphs from the end of the video on one piece of work "
 "this week:\n"
 "BEFORE — what did not exist.\n"
 "BUILD — what you created underneath the output.\n"
 "RETURN — what is different now, and how someone else could tell.\n\n"
 "If you want a structured way to do it, Keep the Proof is a 60-minute career "
 "evidence system: " + CTA_URL + "\n\n"
 "What did you build that had no blueprint? Tell me in a reply."
)

LOWER_THIRDS = [
 ("0:08", "Temidayo Afonja"),
 ("0:12", "Capability Formation"),
 ("9:58", "Keep the Proof · temidayoafonja.com/keep-the-proof"),
]

ON_SCREEN = [
 ("1:55", "YOU ARE THE BEFORE", "Small lower-third text, cream on navy. Two "
  "seconds. Do not compete with slide 2."),
 ("5:45", "NAME THE CHOICE · THE ALTERNATIVE · THE REASON", "Matches slide 7 "
  "wording exactly."),
 ("8:45", "A NUMBER YOU CANNOT DEFEND IS WORSE THAN NO NUMBER", "Only if "
  "slide 9 is off screen at that moment."),
 ("9:58", "temidayoafonja.com/keep-the-proof", "Persist to the end of the CTA "
  "block."),
]

BROLL = [
 ("0:35–1:20", "Hold on Temidayo, full screen. No b-roll. The early-proof "
  "block is the only personal passage in the video and it should not be "
  "illustrated."),
 ("2:25–4:20", "Optional: a slow push in on the slide during the absence "
  "checklist. Nothing literal — no notebooks, no office stock, no hands "
  "typing."),
 ("9:15–9:55", "None. The application block should stay on the slide so the "
  "viewer can copy the three headings."),
]

EDITOR_NOTES = [
 "No opening title card. Open full screen on Temidayo; slide 1 lands only "
 "after the hook and the viewer promise, at roughly 0:35.",
 "Reveal builds are duplicate sequential slides, not PowerPoint animations. "
 "Advance in the order given by Video_7_Reveal_Order_Sheet.png — 24 frames.",
 "Slide 12 keeps its right third empty for the YouTube end-screen element. Do "
 "not place graphics there.",
 "The arrow marks on slide 10 and the quotation on slide 9 are drawn or typed "
 "brand elements; do not substitute icons.",
 "Chapter timestamps in this package are planning estimates. Reset them "
 "against the real export before publishing.",
 "One offer only in the CTA block. No Field Kit, no Career Decision Evidence "
 "Check, no book mention anywhere in Video 7.",
]

PRECHECKS = [
 "Record 4K at 30fps; export 1080p at 30fps. Audio at 48 kHz.",
 "More light on Temidayo — the single most important recording improvement "
 "carried across the channel.",
 "Confirm the spoken script matches Video_7_Recording_Script_Clean.txt "
 "word for word before the first take.",
 "Confirm the CTA URL on screen and in the description is "
 "temidayoafonja.com/keep-the-proof.",
 "Confirm the approved Canva thumbnail reads MAKE INVISIBLE WORK VISIBLE and "
 "is 1280 x 720.",
 "Set chapter timestamps from the finished export, not from this package.",
 "Confirm the watch-next route is live; if the single video has not "
 "published, point the end screen at the Career Portability playlist.",
]


# --------------------------------------------------------------------- utils
def shade(p, fill):
    pr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill)
    pr.append(sh)


def build():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.85)
        s.left_margin = s.right_margin = Inches(0.95)

    def para(text, size=11, bold=False, color=None, before=0, after=8,
             italic=False, caps=False, spacing=1.25):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = spacing
        r = p.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        if caps:
            r.font.all_caps = True
        return p

    def h1(text):
        return para(text, size=16, bold=True, color=NAVY, before=22, after=10)

    def table(rows, widths=None, head=True):
        t = doc.add_table(rows=0, cols=len(rows[0]))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for c, val in zip(cells, row):
                c.text = ""
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(str(val))
                r.font.size = Pt(10)
                if head and i == 0:
                    r.bold = True
                    r.font.color.rgb = NAVY
        if widths:
            for row in t.rows:
                for c, w in zip(row.cells, widths):
                    c.width = Inches(w)
        para("", size=6, after=6)
        return t

    # ------------------------------------------------------------- header
    para("CAREER PORTABILITY   |   VIDEO 7", size=10, bold=True, color=GOLD,
         after=4, caps=True)
    para(TITLE, size=22, bold=True, color=NAVY, after=6, spacing=1.1)
    para("First-pass production package  ·  slides, script, publishing and "
         "editor direction", size=11, color=DIM, after=14)

    table([["Field", "Value"],
           ["Final title", TITLE],
           ["Primary target keyword", KEYWORD],
           ["Locked thumbnail copy", THUMB],
           ["Thumbnail status",
            "Created and approved externally in Canva. Not built, "
            "regenerated or edited in this package."],
           ["Primary CTA", CTA_NAME],
           ["CTA URL", CTA_URL],
           ["Target length", "Approximately 9–11 minutes"],
           ["Spoken script length", "%d words" % WORDS],
           ["Slides", "12 main slides, 24 reveal frames"],
           ["Watch next",
            "How to Explain a Nonlinear Career Without Looking Unfocused, or "
            "the Career Portability playlist — whichever is live at publish"]],
          widths=[1.9, 4.7])

    para("LOCKED PRODUCTION DECISIONS", size=11, bold=True, color=NAVY,
         before=10, after=6)
    table([["Decision", "Detail"],
           ["Keyword rationale",
            "Primary keyword weighted score 71/100. Complete title weighted "
            "score 59/100. Complete-title search volume excellent. Competition "
            "fair. Zero exact title matches in the top 20. Production "
            "rationale only — never referenced in the video."],
           ["Personal evidence ceiling",
            "The first-in-role/builder line only. No company, role, date, "
            "conflict, result or causal claim is invented around it."],
           ["Excluded metrics",
            "The 30% retention improvement and the $2M+ estimated turnover "
            "cost avoidance are excluded. See section 2."],
           ["Offer exclusivity",
            "One offer on screen and in the script: Keep the Proof. No Field "
            "Kit, no Career Decision Evidence Check, no book."]],
          widths=[1.9, 4.7])

    # ------------------------------------------------------------- 1
    h1("1. Strategy at a glance")
    table([["Element", "Decision"],
           ["Viewer", "Experienced professional, particularly a senior "
            "corporate woman with substantial career history and real "
            "financial, family and reputational stakes."],
           ["Viewer problem",
            "“I created something the organization did not previously have, "
            "but because there was no before-and-after measurement system, "
            "the work now looks vague, invisible or difficult to explain.”"],
           ["What this video is not",
            "Not personal branding. Not networking advice. Not “speak up "
            "more.” Not a promise that evidence produces promotion."],
           ["Promise", "Leave able to reconstruct the original conditions, "
            "name what you built, explain the judgment involved and preserve "
            "credible, permitted evidence of the value that resulted."],
           ["Teaching spine",
            "Reconstruct the before → name what you built → show what it "
            "returned."]],
          widths=[1.6, 5.0])

    table([["Architecture beat", "Where it lands"],
           ["Specific human situation", "0:35 — first-in-role/builder opening"],
           ["Unanswered meaning or tension", "1:20 — “I made this possible” "
            "sounds like nothing"],
           ["What became visible", "1:45 — you are the before; the instrument "
            "did not exist"],
           ["Practical distinction", "2:25 — the absence, not the story of "
            "what it was like"],
           ["Usable test", "9:15 — before, build, return, in three paragraphs"],
           ["Earned next step", "9:55 — Keep the Proof"]],
          widths=[2.4, 4.2])

    # ------------------------------------------------------------- 2
    h1("2. Editorial thesis and factual boundaries")
    para("Thesis. Foundational work is hard to prove for a structural reason, "
         "not a personal one. Comparative evidence needs a prior state; a "
         "first-in-role builder is the prior state. The remedy is "
         "reconstruction of the record, not more self-promotion.", after=10)
    table([["Boundary", "Ruling"],
           ["Personal evidence",
            "Approved level only: “When I was first in roles, I sometimes had "
            "to create the language, relationships and decision clarity "
            "before the work itself could succeed.” Rendered naturally in the "
            "0:35 block. No company, role, date, conflict, outcome or causal "
            "claim is added."],
           ["30% retention improvement",
            "EXCLUDED. docs/claims-ledger.md records this figure as “Needs "
            "source. No supporting document is on file in this repository,” "
            "and attributes it to an enterprise operating role — not to this "
            "first-in-role story. No supportable attribution exists, so the "
            "figure does not appear."],
           ["$2M+ turnover cost avoidance",
            "EXCLUDED, on the same basis. The ledger records it as “Needs "
            "source,” notes that the estimation model is not on file, and "
            "attributes it to the same enterprise operating role."],
           ["Invented precision",
            "The script states outright that a number you cannot defend is "
            "worse than no number, and gives a defensible qualitative "
            "alternative in its place."],
           ["Permitted evidence",
            "Outcomes, decisions, learning and non-confidential examples in "
            "the viewer's own words. Explicitly not documents, client or "
            "customer data, or anything proprietary or employer-owned. Stated "
            "on screen (slide 9) and out loud."],
           ["Causal claims",
            "None. Evidence is presented as something a reader can verify, "
            "never as a guarantee of promotion or recognition."]],
          widths=[1.7, 4.9])

    para("Open evidence question. If a supporting document is ever placed on "
         "file that connects the 30% retention improvement or the $2M+ "
         "estimated turnover cost avoidance to this specific first-in-role "
         "builder story — with a stated population, baseline and measurement "
         "method — the 0:35 block could carry one of them. Until that "
         "document exists, both stay out. This question is recorded here and "
         "in the QA README, and deliberately not inside the spoken script.",
         italic=True, color=DIM, after=10)

    # ------------------------------------------------------------- 3
    h1("3. Timed script outline")
    table([["Time", "Block", "Slide"]]
          + [[h.split(" | ")[0], h.split(" | ")[1], s] for h, s in
             zip([b[0] for b in SCRIPT],
                 ["—", "1", "2, 3", "4, 5", "6, 7", "8", "9", "10", "11", "12"])],
          widths=[1.2, 3.9, 1.5])

    # ------------------------------------------------------------- 4
    h1("4. Full recording script")
    para("Target 1,450–1,700 spoken words. This draft is %d." % WORDS,
         size=10, italic=True, color=DIM, after=12)
    for header, paras in SCRIPT:
        p = para(header, size=10.5, bold=True, color=GOLD, before=16, after=8,
                 caps=True, spacing=1.1)
        shade(p, "F3F0E8")
        for t in paras:
            para(t, size=11.5, color=INK, after=10, spacing=1.4)

    # ------------------------------------------------------------- 5
    h1("5. Slide deck content and placement")
    table([["#", "Slide", "States", "Copy on screen, verbatim"]]
          + [list(r) for r in SLIDE_COPY], widths=[0.4, 1.5, 0.6, 4.1])
    para("1920 × 1080, 16:9. No stock imagery, gradients, decorative icons or "
         "PowerPoint animations. Reveal builds are duplicate sequential "
         "slides. 24 frames total.", size=10, italic=True, color=DIM)

    # ------------------------------------------------------------- 6
    h1("6. Code handoff — slides and script files")
    table([["File", "What it is"],
           ["Video_7_Main_Slides.pptx", "12 editable slides, final revealed "
            "state. All text is live text; no images."],
           ["Video_7_Reveal_Builds.pptx", "24 slides, one per reveal state, "
            "in advance order."],
           ["Video_7_Slide_Preview.pdf", "12 pages at true 16:9."],
           ["Video_7_Main_Slide_Contact_Sheet.png", "All 12 slides on one "
            "sheet."],
           ["Video_7_Reveal_Order_Sheet.png", "All 24 frames in advance "
            "order."],
           ["Video_7_Phone_Legibility_Sheet.png", "Every slide at 320 × 180."],
           ["Video_7_Teleprompter_Script_with_Slide_Markers.docx",
            "The spoken script with 12 slide markers and the package's own "
            "stage directions kept visually distinct."],
           ["Video_7_Recording_Script_Clean.txt", "Spoken words only. No "
            "timestamps, markers or directions."],
           ["build/slides.py, build/deck.py, build/build.py",
            "Editable sources. deck.py is carried forward from Video 6 "
            "unchanged."]],
          widths=[2.6, 4.0])

    # ------------------------------------------------------------- 7
    h1("7. Thumbnail — externally approved, not rebuilt here")
    table([["Field", "Value"],
           ["Locked copy", THUMB],
           ["Status", "Created and approved in Canva outside this package."],
           ["Instruction", "Do not create, regenerate, reinterpret, redesign "
            "or edit a Video 7 thumbnail. If the approved Canva export is "
            "supplied it is preserved byte-identically and included "
            "unchanged."],
           ["In this package", "The export was not present in the repository "
            "when this package was built. It is recorded as an externally "
            "approved asset still to be added. No substitute was made."]],
          widths=[1.5, 5.1])

    # ------------------------------------------------------------- 8
    h1("8. YouTube publishing package")
    para("Title", size=11, bold=True, color=NAVY, after=4)
    para(TITLE, after=10)
    para("Description", size=11, bold=True, color=NAVY, after=4)
    for line in DESCRIPTION:
        para(line if line else " ", after=6 if line else 2)
    para("Tags", size=11, bold=True, color=NAVY, before=8, after=4)
    para(", ".join(TAGS), after=10)
    para("Chapters", size=11, bold=True, color=NAVY, after=4)
    table([["Time", "Chapter"]] + [list(c) for c in CHAPTERS],
          widths=[1.0, 5.6])
    para("Chapter timestamps are planning estimates and must be reset from the "
         "finished export before publishing.", size=10, italic=True,
         color=DIM, after=10)
    para("Pinned comment", size=11, bold=True, color=NAVY, before=8, after=4)
    for line in PINNED.split("\n"):
        para(line if line else " ", after=5 if line else 2)
    para("End screen", size=11, bold=True, color=NAVY, before=10, after=4)
    table([["Element", "Route"],
           ["Video card",
            "How to Explain a Nonlinear Career Without Looking Unfocused"],
           ["Fallback if unpublished",
            "Career Portability: Career Pivots, Internal Moves & Growth"],
           ["Subscribe element", "Bottom right, standard placement"],
           ["Clear zone", "Slide 12 keeps the right third empty for the card"]],
          widths=[1.9, 4.7])

    # ------------------------------------------------------------- 9
    h1("9. Editor direction")
    para("Lower thirds", size=11, bold=True, color=NAVY, after=4)
    table([["Time", "Copy"]] + [list(l) for l in LOWER_THIRDS],
          widths=[1.0, 5.6])
    para("On-screen text", size=11, bold=True, color=NAVY, before=8, after=4)
    table([["Time", "Text", "Note"]] + [list(o) for o in ON_SCREEN],
          widths=[0.8, 2.6, 3.2])
    para("B-roll", size=11, bold=True, color=NAVY, before=8, after=4)
    table([["Block", "Guidance"]] + [list(b) for b in BROLL],
          widths=[1.4, 5.2])
    para("Editor notes", size=11, bold=True, color=NAVY, before=8, after=4)
    for n in EDITOR_NOTES:
        para("•  " + n, after=6)

    # ------------------------------------------------------------- 10
    h1("10. Final pre-recording and delivery checks")
    for c in PRECHECKS:
        para("☐  " + c, after=6)

    doc.save(OUT)
    print("wrote", os.path.basename(OUT))
    print("spoken paragraphs: %d   words: %d" % (len(SPOKEN), WORDS))
    return OUT


if __name__ == "__main__":
    build()
