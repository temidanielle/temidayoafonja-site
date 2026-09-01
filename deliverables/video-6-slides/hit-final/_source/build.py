# -*- coding: utf-8 -*-
"""Build the Video 6 H.I.T. final recording and Shorts package."""
import os, sys, shutil, zipfile, hashlib
sys.path.insert(0, "/tmp/v6hit")
from script_text import LINES, SPOKEN, MARKERS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0x9B,0x2C,0x10)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

ROOT="/tmp/v6hit/Video_6_HIT_FINAL"
LF=os.path.join(ROOT,"LONG_FORM"); SH=os.path.join(ROOT,"SHORTS")
shutil.rmtree(ROOT, ignore_errors=True)
os.makedirs(LF); os.makedirs(SH)

TITLE="Are You Growing—or Just Being Given More Work?"
THUMB="MORE WORK \u2260 GROWTH"
CTA="Capability Formation Field Kit"
CTA_URL="https://temidayoafonja.com/fieldkit"
NEXT="How to Show Your Impact at Work When You Built It From Scratch"
OLD_NEXT="How to Prove the Value of Work That Had No Blueprint"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO 6",size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t,before=20):
    return keep(P(d,t,size=14,bold=True,color=NAVY,before=before,after=8),True)
def H2(d,t,before=13):
    return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=before,after=5),True)

def compress(d, line_spacing=1.18, after_scale=0.80):
    """Tighten one document's vertical rhythm without touching type size,
    weight, colour or wording. Keeps editor briefs off a near-empty last page."""
    for p in d.paragraphs:
        pf=p.paragraph_format
        if pf.line_spacing and 1.0 < pf.line_spacing < 1.4 and pf.line_spacing > line_spacing:
            pf.line_spacing = line_spacing
        if pf.space_after is not None:
            pf.space_after = Pt(round(pf.space_after.pt*after_scale,1))
        if pf.space_before is not None and pf.space_before.pt:
            pf.space_before = Pt(round(pf.space_before.pt*after_scale,1))
    return d

def pairlist(d, items, indent="—  ", gap="        ", after=4, budget=78):
    """Lay short bullet items several to a line, packed by width.

    Every item is kept verbatim; nothing is removed, reworded or reordered,
    and neither type size nor leading changes. This reclaims the unused
    right-hand part of a short-bullet line, which is what keeps the Shorts
    editor brief to two pages. The width budget is measured against the
    inspection renderer's font, which is wider than Calibri, so Word fits at
    least as much per line and never less."""
    lines, cur = [], []
    for it in items:
        trial = cur + [it]
        width = sum(len(indent) + len(x) for x in trial) + len(gap) * (len(trial) - 1)
        if cur and width > budget:
            lines.append(cur); cur = [it]
        else:
            cur = trial
    if cur: lines.append(cur)
    for group in lines:
        keep(P(d, gap.join(indent + x for x in group), after=after))




# ------------------------------------------------- 1. teleprompter DOCX + TXT
d=newdoc(True)
head(d,TITLE,"Video 6  ·  Teleprompter script with slide markers",
     "Spoken script is the large text. A slide marker in a tinted band tells "
     "the editor which slide to bring up; it is not spoken.")
for line in LINES:
    if line.startswith("[SLIDE:"):
        p=P(d,"SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(),size=11,bold=True,
            color=NAVY,before=14,after=14,spacing=1.1)
        shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
    else:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video6TeleprompterScriptwithslidemarkers_HIT_v2.0.docx"))
tel=[TITLE,"Video 6  ·  Teleprompter script with slide markers",""]
for line in LINES:
    if line.startswith("[SLIDE:"):
        tel += ["", "SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(), ""]
    else: tel += [line, ""]
open(os.path.join(LF,"Video6TeleprompterScriptwithslidemarkers_HIT_v2.0.txt"),
     "w").write("\n".join(tel).strip()+"\n")

# ------------------------------------------------- 2. reading script DOCX+TXT
d=newdoc(True)
head(d,TITLE,"Video 6  ·  Reading script, no markers",
     "Spoken language only. No slide markers, no timestamps, no production "
     "directions.")
for line in SPOKEN:
    keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video6ReadingScriptnomarkers_HIT_v2.0.docx"))
open(os.path.join(LF,"Video6ReadingScriptnomarkers_HIT_v2.0.txt"),
     "w").write("\n\n".join(SPOKEN)+"\n")
print("long-form scripts written")

# --------------------------------------------------- 3. long-form editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 6",size=12,bold=True,color=GOLD,after=2,caps=True)
P(d,TITLE,size=20,bold=True,color=NAVY,after=6,spacing=1.1)
p=P(d,"This document is for the editor. It is NOT Temidayo's teleprompter and "
     "must not be placed on the recording screen.",size=11,italic=True,
     color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"Locked metadata",before=14)
for k,v in (("Title",TITLE),("Thumbnail",THUMB),("Primary CTA",CTA),
            ("CTA URL",CTA_URL),("Watch next",NEXT),
            ("Capability Formation question","What is this work building in me?"),
            ("Core distinction",
             "More responsibility can be part of growth. It is not proof of "
             "it.")):
    keep(P(d,"%-30s %s"%(k+":",v),size=11,after=5))

H1(d,"Slide 12 correction — the only visual change",before=14)
p=P(d,"AUTHORISED AND APPLIED. Slide 12's Watch Next title was stale. Only "
     "that title text changed.",size=11,bold=True,color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
keep(P(d,"FROM:  %s"%OLD_NEXT,size=11,after=4))
keep(P(d,"TO:    %s"%NEXT,size=11,bold=True,after=8))
keep(P(d,"Set as four lines to hold the existing 29pt Montserrat Bold "
       "typography inside the existing text box: HOW TO SHOW YOUR / IMPACT AT "
       "WORK / WHEN YOU BUILT / IT FROM SCRATCH. Every line is narrower than "
       "the widest line of the old title, and the block still clears the gold "
       "rule.",size=10.5,color=DIM,after=8,spacing=1.25))
keep(P(d,"Slides 1 to 11 are byte-identical. The design, typography, layout, "
       "playlist line and end-screen space are unchanged. Reveal frame 23 "
       "carries the same title and received the same single correction; all "
       "other reveal frames are byte-identical.",bold=True,after=10,
       spacing=1.25))

H1(d,"First 30 seconds — H.I.T.",before=14)
P(d,"H = Hook. I = Interest. T = Trust. One audiovisual unit: immediate "
    "conversational contradiction, meaningful visual movement, relevant "
    "personal proof, and a viewer payoff within roughly 20 to 30 seconds. No "
    "generic channel welcome before the promise. No forced statistic. No "
    "invented personal story.",after=8)
P(d,"The visual carries simple scope progression. Temidayo's voice carries "
    "the meaning.",italic=True,color=DIM,after=14)

def beat(t,anchor,layer,body):
    H2(d,t,before=10)
    p=P(d,"Spoken anchor:  “%s”"%anchor,size=10.5,italic=True,color=DIM,after=8)
    shade(p,BAND_CREAM)
    if layer: keep(P(d,layer,size=11,bold=True,color=GOLD,after=6))
    for b in body: keep(P(d,b,after=5))

beat("0:00–0:05","Your workload can grow faster than your career.","H = HOOK",
     ["Visual: direct to camera.",
      "On-screen:  YOUR WORKLOAD GREW.  DID YOUR CAREER?",
      "Do not repeat the thumbnail wording in the first beat."])
beat("0:05–0:16",
     "Across my own career, my scope has expanded beyond the job I was "
     "originally hired to do more than once.","T = TRUST  /  I = INTEREST",
     ["Visual:  ORIGINAL SCOPE  →  EXPANDED SCOPE"])
p=P(d,"FACTUAL BOUNDARY. Do not add an employer, exact role, exact "
     "assignment, exact timeline, metric, outcome or quote.",size=11,bold=True,
     color=RED,before=6,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
beat("0:16–0:20","That taught me not to use volume as the test.",None,
     ["Return to Temidayo.","On-screen:  VOLUME ≠ GROWTH","Let the line breathe."])
beat("0:20–0:30","Before you call more responsibility growth…","PAYOFF",
     ["Progressively reveal:","     COMPLEXITY","     AUTHORITY","     RETURN",
      "Then enter the existing deck."])

H1(d,"Editorial rhythm after the opening",before=14)
P(d,"After the first 30 seconds:",after=6)
for x in ["preserve natural pace;","use slides for distinctions and questions;",
 "avoid constant motion;","avoid decorative B-roll;",
 "avoid constant punch-ins;","allow reflective pauses;",
 "do not make workload and growth visually alarmist."]:
    keep(P(d,"—  "+x,after=4))

H1(d,"Existing slide system",before=14)
P(d,"Slides 1 to 11 are preserved unchanged. Slide 12 carries the authorised "
    "Watch Next title correction and nothing else.",after=8)
for n,job in enumerate(["Core Distinction","More Work / Real Growth",
 "Three Tests","Complexity","Capability Question",
 "Accountability / Authority","Authority Warning","Return","Pattern Read",
 "Before the Scope Expands Again","Capability Formation Field Kit",
 "Watch Next  (title corrected)"],1):
    keep(P(d,"Slide %-3d %s"%(n,job),size=10.5,after=3))
keep(P(d,"Verified in the live files: 12 main slides, 23 reveal-build frames.",
       size=10.5,italic=True,color=DIM,before=6,after=8))

H1(d,"Let the slides carry lists",before=14)
P(d,"Do not add spoken words to repeat every on-screen label.",after=6)
for x in ["the complexity variables;",
 "the responsibility / accountability / authority distinction;",
 "capability / evidence / recognition;",
 "the final manager-conversation questions."]:
    keep(P(d,"—  "+x,after=4))

H1(d,"Fact and proof boundaries",before=14)
keep(P(d,"Confirmed proof: across Temidayo's career, scope has expanded beyond "
       "the original job description more than once.",bold=True,after=6))
P(d,"Do not invent an employer, exact role, assignment, quote, outcome, metric "
    "or causal claim.",after=6)
p=P(d,"Do not attach an approximately 30% retention improvement, a more-than-"
     "$2M avoided-turnover figure, or any other undocumented result.",size=11,
     bold=True,color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
keep(P(d,"The line about a career stalling while employed is conceptual "
       "editorial language, not a prevalence claim. Do not add external "
       "statistics unless Temidayo explicitly authorises them later.",after=8,
       spacing=1.25))
keep(P(d,"Evidence boundary: no confidential information, employer-owned "
       "files, customer data, employee data or proprietary material.",
       bold=True,after=8))

H1(d,"CTA and watch next",before=14)
keep(P(d,"One offer only: %s — %s"%(CTA,CTA_URL),after=5))
keep(P(d,"Do not add Keep the Proof or the Career Decision Evidence Check.",
       bold=True,after=8))
keep(P(d,"Watch next uses the current locked title: %s"%NEXT,bold=True,after=5))
for x in ["use the direct Video 7 end-screen route once Video 7 is public;",
 "before Video 7 is public, use the Career Portability playlist if it is "
 "functioning;",
 "otherwise use YouTube's best-for-viewer option or another appropriate "
 "currently public video temporarily."]:
    keep(P(d,"—  "+x,after=4))
keep(P(d,"Do not leave Subscribe as the only end-screen element.",bold=True,
       color=RED,before=4,after=8))
compress(d)
d.save(os.path.join(LF,"Video_6_EDITOR_ONLY_HIT_Brief_v2.0.docx"))
print("editor brief written")

# The eleven working chapter lines, defined once and reused by the publishing
# package, its reference section and the separate description-only document.
CHAPTERS=[("00:00","Your Workload Can Grow Faster Than Your Career"),
 ("01:55","The 3 Tests for Real Growth"),
 ("02:10","Test 1: Did the Problem Become More Complex?"),
 ("03:30","What Can You Do Now That You Couldn’t Before?"),
 ("04:15","Test 2: Did Your Authority Expand?"),
 ("05:20","Accountability Is Not Authority"),
 ("06:05","Test 3: What Did the Work Return?"),
 ("08:20","Read the Pattern"),
 ("09:45","What to Ask Before Your Scope Expands Again"),
 ("10:55","Capability Formation Field Kit"),
 ("11:20","How to Show Your Impact at Work When You Built It From Scratch")]
CHAPTER_LINES=["%s %s"%(t,c) for t,c in CHAPTERS]

PRIMARY="are you growing or just being given more work"
SUPPORTING=("career growth · more responsibility at work · career development · "
 "workload increase · career stagnation · stretch assignment · professional "
 "growth · career progression · capability formation")
TAGS=("are you growing or just being given more work, more responsibility at "
 "work, career growth, career development, workload increase, career "
 "stagnation, professional growth, stretch assignment, career progression, "
 "work overload, career portability, capability formation, experienced "
 "professionals, Temidayo Afonja")
DESC=[
 "More responsibility at work can look like career growth even when your "
 "workload is simply getting larger.",
 "In this video, I give you three tests for separating genuine development "
 "from a role that has simply learned to rely on you for more:",
 "✨ Complexity — Did the problem become more difficult, or did the volume "
 "simply increase?",
 "✨ Authority — Did your ability to influence decisions expand with the "
 "responsibility?",
 "✨ Return — What did the additional work give back in capability, evidence "
 "or recognition?",
 "Being dependable is valuable. But a career can look successful from the "
 "outside while the work underneath it has stopped adding much new judgment.",
 "The question is not simply:","“Am I doing more?”","It is:",
 "“What is this work building in me?”","",
 "🧭 CAPABILITY FORMATION FIELD KIT",
 "Examine what your current work is building, what appears portable and where "
 "the role may need a boundary or redesign:", CTA_URL,"",
 "⏱️ CHAPTERS"]+CHAPTER_LINES+["",
 "▶️ WATCH NEXT", NEXT, "[ADD VIDEO 7 LINK WHEN LIVE]","",
 "🔗 CONNECT AND EXPLORE",
 "Website:","https://temidayoafonja.com",
 "LinkedIn:","https://www.linkedin.com/in/temidayo-afonja",
 "Substack:","https://temidayoafonja.substack.com","",
 "#CareerGrowth #CareerDevelopment #CapabilityFormation"]
PINNED=["Which part of your role has grown the most recently?",
 "1. Complexity","2. Authority","3. Workload",
 "And what has that additional responsibility actually returned to your career?",
 "You do not need to share confidential details — I am more interested in the "
 "pattern you are seeing.",
 "If you want a fuller private read of what your current work is building, the "
 "Capability Formation Field Kit is here:", CTA_URL]

EMOJI_NOTE=("The restrained emoji system is part of the approved standard: ✨ "
  "teaching points, 🧭 CTA, ⏱️ chapters, ▶️ Watch Next, 🔗 Connect and Explore. "
  "Do not remove them and do not add more.")

def description_block(d, heading_before=14, upload_doc=False):
    """The copy-ready description, its end marker and the internal note.

    upload_doc=True moves the editorial emoji instruction ABOVE an explicit
    COPY-READY ... BEGIN marker, so nothing internal sits inside or immediately
    before the block a person selects when pasting into YouTube."""
    if upload_doc:
        H1(d,"INTERNAL NOTE — DO NOT PASTE INTO YOUTUBE",before=heading_before)
        p=P(d,EMOJI_NOTE,size=10.5,italic=True,color=RED,after=12,spacing=1.25)
        shade(p,BAND_CREAM); keep(p)
        p=keep(P(d,"COPY-READY YOUTUBE DESCRIPTION — BEGIN",size=11,bold=True,
                 color=NAVY,before=14,after=12,spacing=1.2))
        shade(p,BAND_NAVY)
    else:
        H1(d,"Description",before=heading_before)
        keep(P(d,EMOJI_NOTE,size=10.5,italic=True,color=DIM,after=10))
    for para in DESC:
        keep(P(d,para if para else " ",after=7 if para else 3))
    keep(P(d,"— END OF THE COPY-READY DESCRIPTION —",size=10,bold=True,
           color=DIM,before=14,after=12,spacing=1.2))
    H1(d,"Internal note — do not paste into YouTube",before=14)
    p=P(d,"WORKING ESTIMATES — EDITOR MUST REPLACE FROM FINAL CUT",size=11,
        bold=True,color=RED,after=6,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    p=P(d,"The chapter timestamps above are script-derived, not final. Replace "
        "every one of them using the finished cut before publication. Do not "
        "force the edit to match these estimates.",size=10.5,bold=True,
        italic=True,color=RED,after=10,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    H1(d,"Working chapters — reference copy",before=14)
    keep(P(d,"Identical to the eleven chapter lines inside the description "
           "above.",size=10.5,italic=True,color=DIM,after=8))
    for line in CHAPTER_LINES: keep(P(d,line,size=11,after=4))

# ----------------------------------------------------- 4. publishing package
d=newdoc()
head(d,TITLE,"Video 6  ·  Publishing package",
     "Everything needed to upload. Working timestamps must be replaced with "
     "real ones from the finished edit.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14); P(d,PRIMARY,after=10)
H1(d,"Supporting search language",before=14); P(d,SUPPORTING,after=10)
description_block(d)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only. Do not put the full tag field in the "
       "public description.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
H1(d,"Watch next",before=14)
keep(P(d,NEXT,bold=True,after=5))
keep(P(d,"Slide 12 now carries this title. The correction record is in the "
       "EDITOR ONLY brief.",size=10.5,color=DIM,after=8,spacing=1.25))
compress(d)
d.save(os.path.join(LF,"Video_6_Publishing_Package_HIT_v2.0.docx"))

# ------------------------------- 4b. separate description-only document
d=newdoc()
head(d,TITLE,"Video 6  ·  YouTube description",
     "Upload copy only. Everything below the end marker is internal and must "
     "not be pasted into YouTube.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14); P(d,PRIMARY,after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"Watch next",before=14); keep(P(d,NEXT,bold=True,after=8))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
compress(d)
DESC_DOC="/tmp/v6hit/Video_6_YouTube_Description_HIT.docx"
d.save(DESC_DOC)
print("publishing package and description-only document written")

# ---------------------------------------------------------------- 5. Shorts
SHORTS=[
 ("Video_6_Short_1_Workload_Grows_Faster.docx","SHORT 1","Recognition",
  "Your workload can grow faster than your career.",
  ["Your workload can grow faster than your career.",
   "This happens easily when you are dependable.",
   "A project loses its owner.","Someone leaves.",
   "A temporary exception lands with you because people know you will handle "
   "it.",
   "None of those things is automatically a problem.",
   "The question is what happens next.",
   "If the extra work becomes permanent but the problems are not harder, your "
   "authority has not changed and there is no meaningful career return, your "
   "job may be using more of your capacity without developing much more "
   "capability.",
   "Being busy is real.",
   "It just is not proof that you are growing."]),
 ("Video_6_Short_2_Accountability_Not_Authority.docx","SHORT 2","Distinction",
  "Accountability is not the same as authority.",
  ["Accountability is not the same as authority.",
   "Accountability is what you will answer for.",
   "Authority is what you can influence or decide.",
   "Those do not always expand together.",
   "You can be told to own an outcome while someone else still controls the "
   "priorities, resources and decisions that produce it.",
   "That can increase your exposure without increasing your judgment.",
   "You do not need complete control.",
   "But if the responsibility keeps growing, your judgment should have "
   "somewhere to go.",
   "Accountability without authority is a warning, not a development plan."]),
 ("Video_6_Short_3_More_Scope_Was_It_Growth.docx","SHORT 3","Personal proof",
  "My scope has expanded beyond the job I was originally hired to do more "
  "than once.",
  ["My scope has expanded beyond the job I was originally hired to do more "
   "than once.",
   "That taught me not to ask only:",
   "“How much more am I carrying?”",
   "I ask three different questions.",
   "Did the problem become harder?",
   "Did I gain more authority to influence the outcome?",
   "And what did the work return to my career?",
   "Because more volume can make you very busy.",
   "It can even make you very valuable to the organization.",
   "But growth should leave you able to handle something you could not "
   "credibly handle before.",
   "That is the difference I want you to look for."]),
 ("Video_6_Short_4_What_Comes_Off_Your_Plate.docx","SHORT 4",
  "Practical action",
  "Before your scope expands again, ask what comes off your plate.",
  ["Before your scope expands again, ask what comes off your plate.",
   "Then ask two more questions.",
   "Which decisions need to belong to me if I am responsible for these "
   "outcomes?",
   "And how and when will this expanded scope be reviewed?",
   "That is a stronger conversation than simply saying:",
   "“I have too much work.”",
   "You are showing how the role itself has changed.",
   "If something new becomes permanent, something around the role may need to "
   "change too — priorities, authority, support, scope or recognition.",
   "Temporary work without a review point has a way of becoming the new "
   "baseline."])]

for fn,label,role,hook,copy in SHORTS:
    d=newdoc(True)
    P(d,"VIDEO 6 SHORT",size=10,bold=True,color=GOLD,after=4,caps=True)
    P(d,label,size=20,bold=True,color=NAVY,after=8,spacing=1.1)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    keep(P(d,"Verbal hook:  “%s”"%hook,size=11,color=DIM,after=5))
    keep(P(d,"Related long-form:  %s"%TITLE,size=11,color=DIM,after=10))
    H1(d,"RECORDING COPY",before=12)
    for line in copy:
        keep(P(d,line,size=13.5,color=INK,after=10,spacing=1.5))
    d.save(os.path.join(SH,fn))

# ------------------------------------------------------ 6. Shorts editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 6 — FOUR STANDALONE SHORTS",size=18,bold=True,color=NAVY,
  after=8,spacing=1.1)
p=P(d,"This document is for the editor. It is separate from the four Short "
     "recording documents and must not be placed on Temidayo's recording "
     "screen.",size=11,italic=True,color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)
H1(d,"How these are produced",before=14)
keep(P(d,"These are separately recorded 9:16 Shorts. They are NOT excerpts cut "
       "from the long-form video.",bold=True,after=10))
P(d,"Each should have:",after=6)
pairlist(d,["an immediate verbal hook;","a matching on-screen hook;",
 "meaningful visual interest;","accurate mobile-safe captions;",
 "restrained pacing;",
 "Video 6 added as the YouTube Related Video when available."])

def short(label,role,onscreen,body):
    H1(d,label,before=14)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    p=keep(P(d,"On-screen hook:  %s"%onscreen,size=11,bold=True,color=GOLD,after=8))
    shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))
    keep(P(d,"Related Video:  Video 6",size=10.5,color=DIM,before=4,after=6))

short("SHORT 1","Recognition","BUSIER ≠ MORE CAPABLE",
 ["Visual:  MORE CAPACITY USED  vs.  MORE CAPABILITY BUILT",
  "Do not use generic exhausted-office-worker B-roll.",
  "End:  BEING BUSY IS NOT PROOF OF GROWTH"])
short("SHORT 2","Distinction","ACCOUNTABILITY ≠ AUTHORITY",
 ["Use the existing conceptual distinction:",
  "     ACCOUNTABILITY — WHAT YOU ANSWER FOR",
  "     AUTHORITY — WHAT YOU CAN INFLUENCE OR DECIDE",
  "End:  YOUR JUDGMENT NEEDS SOMEWHERE TO GO",
  "No red warning graphics."])
short("SHORT 3","Personal proof","MORE SCOPE. BUT WAS IT GROWTH?",
 ["FACTUAL BOUNDARY: do not add an employer, exact role, assignment, "
  "timeline, result or metric.",
  "Progressively show:  COMPLEXITY  ·  AUTHORITY  ·  RETURN",
  "End:  WHAT CAN YOU HANDLE NOW?"])
short("SHORT 4","Practical action","WHAT COMES OFF YOUR PLATE?",
 ["Reveal progressively:",
  "     WHAT COMES OFF MY PLATE?",
  "     WHICH DECISIONS BELONG TO ME?",
  "     WHEN WILL THIS SCOPE BE REVIEWED?",
  "End:  TEMPORARY NEEDS A REVIEW DATE"])

H1(d,"All Shorts — visual boundaries",before=14)
P(d,"Do not use:",after=5)
pairlist(d,["hyperactive zooming;","fake shock expressions;",
 "red warning graphics;","burnout clichés;","generic office B-roll;",
 "stock résumé footage;","clocks or calendar gimmicks;",
 "AI-style animated icons;","trendy caption templates;",
 "flashy transitions."],after=3)
compress(d, 1.18, 0.62)
d.save(os.path.join(SH,"Video_6_Shorts_EDITOR_ONLY_HIT_Brief.docx"))
print("shorts and shorts editor brief written")

# ---------------------------------------------------------------- 7. README
FILES=(["LONG_FORM/"+f for f in sorted(os.listdir(LF))]
      +["SHORTS/"+f for f in sorted(os.listdir(SH))])
R=["VIDEO 6 — H.I.T. FINAL RECORDING PACKAGE","",
 "Title:             %s"%TITLE,
 "Thumbnail:         %s"%THUMB,
 "CTA:               %s"%CTA,
 "CTA URL:           %s"%CTA_URL,
 "Watch next:        %s"%NEXT,"",
 "Long-form:         Revised under H.I.T.",
 "Slides 1-11:       UNCHANGED.",
 "Slide 12:          WATCH NEXT TITLE UPDATED ONLY, to match the locked",
 "                   Video 7 title. Design, typography, layout, playlist",
 "                   line and end-screen space preserved.",
 "Reveal deck:       Only the corresponding Slide 12 title frame changed",
 "                   (frame 23 of 23). All other frames byte-identical.",
 "Thumbnail:         UNCHANGED.",
 "Shorts:            Four separately recorded vertical scripts.",
 "Editor",
 "instructions:      Separated from all recording copy.",
 "Description-only",
 "document:          Video_6_YouTube_Description_HIT.docx, created separately",
 "                   OUTSIDE this package ZIP.","",
 "Verified in the",
 "live files:        12 main slides, 23 reveal-build frames.","",
 "-"*70,"","WHAT EACH FILE IS","",
 "LONG_FORM/","",
 "  Video6TeleprompterScriptwithslidemarkers_HIT_v2.0.docx",
 "  Video6TeleprompterScriptwithslidemarkers_HIT_v2.0.txt",
 "      Temidayo's recording copy. Spoken script in large text; slide markers",
 "      in tinted bands. The markers are not spoken.","",
 "  Video6ReadingScriptnomarkers_HIT_v2.0.docx",
 "  Video6ReadingScriptnomarkers_HIT_v2.0.txt",
 "      The same spoken words with the slide markers removed.","",
 "  Video_6_EDITOR_ONLY_HIT_Brief_v2.0.docx",
 "      For the editor. The H.I.T. first-30-second plan, the Slide 12",
 "      correction record, editorial rhythm after 0:30, the 12-slide map,",
 "      the let-the-slides-carry principle, and the fact, proof and evidence",
 "      boundaries. Not for the teleprompter.","",
 "  Video_6_Publishing_Package_HIT_v2.0.docx",
 "      Title, thumbnail, search language, the copy-ready description with",
 "      its approved emoji system, working chapter estimates, pinned comment",
 "      and the tag field.","",
 "SHORTS/","",
 "  Four recording documents, one per Short. These contain Temidayo's",
 "  recording copy and no editor directions.","",
 "  Video_6_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "      For the editor. On-screen hooks and visual treatment for all four.","",
 "-"*70,"","ALL FILES IN THIS PACKAGE","",]
for f in FILES: R.append("  "+f)
R+=["  README_FINAL.txt","  SHA256SUMS.txt","",
 "  Video_6_YouTube_Description_HIT.docx is deliberately NOT in this ZIP.","",
 "-"*70,"","WORKING CHAPTER TIMESTAMPS","",
 "The chapter timestamps are WORKING ESTIMATES derived from the script. They",
 "were not measured from an edit. The editor must replace every one of them",
 "from the finished cut before publishing.","",
 "-"*70,"","FACT AND PROOF BOUNDARY","",
 "The confirmed proof is only that, across Temidayo's career, scope has",
 "expanded beyond the original job description more than once. No employer,",
 "exact role, assignment, timeline, quote, outcome or metric is added",
 "anywhere in this package. No retention percentage and no avoided-turnover",
 "figure is attached to anything.","",
 "-"*70,"","CHECKSUMS","",
 "SHA256SUMS.txt covers the other 12 user-facing files in this package. It",
 "does not hash itself, and it carries no ZIP checksum. The archive's own",
 "SHA-256 is in the sibling file:",
 "  Video_6_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256","",
 "-"*70,"","WHAT WAS NOT CHANGED","",
 "Slides 1 to 11, reveal frames 1 to 22, the approved thumbnail, every",
 "website file, every product and every other video are unchanged. The only",
 "visual change in this pass is the authorised Slide 12 Watch Next title,",
 "applied identically to main slide 12 and reveal frame 23.",""]
open(os.path.join(ROOT,"README_FINAL.txt"),"w").write("\n".join(R))

# ------------------------------------------- 8. checksums and the master ZIP
MANIFEST=[
 "LONG_FORM/Video6TeleprompterScriptwithslidemarkers_HIT_v2.0.docx",
 "LONG_FORM/Video6TeleprompterScriptwithslidemarkers_HIT_v2.0.txt",
 "LONG_FORM/Video6ReadingScriptnomarkers_HIT_v2.0.docx",
 "LONG_FORM/Video6ReadingScriptnomarkers_HIT_v2.0.txt",
 "LONG_FORM/Video_6_EDITOR_ONLY_HIT_Brief_v2.0.docx",
 "LONG_FORM/Video_6_Publishing_Package_HIT_v2.0.docx",
 "SHORTS/Video_6_Short_1_Workload_Grows_Faster.docx",
 "SHORTS/Video_6_Short_2_Accountability_Not_Authority.docx",
 "SHORTS/Video_6_Short_3_More_Scope_Was_It_Growth.docx",
 "SHORTS/Video_6_Short_4_What_Comes_Off_Your_Plate.docx",
 "SHORTS/Video_6_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "README_FINAL.txt",
]
SUMS="SHA256SUMS.txt"
ZIP="/tmp/v6hit/Video_6_HIT_FINAL_Recording_and_Shorts_Package.zip"

def sha256(p):
    h=hashlib.sha256()
    with open(p,"rb") as fh:
        for b in iter(lambda: fh.read(1<<20), b""): h.update(b)
    return h.hexdigest()

for m in MANIFEST:
    assert os.path.isfile(os.path.join(ROOT,m)), "missing from build: "+m
on_disk=set()
for dp,dn,fn in os.walk(ROOT):
    dn[:]=[x for x in dn if x!="__pycache__"]
    for f in fn:
        if f.endswith(".pyc"): continue
        on_disk.add(os.path.relpath(os.path.join(dp,f),ROOT).replace(os.sep,"/"))
unexpected=sorted(on_disk-set(MANIFEST)-{SUMS})
assert not unexpected, "unexpected files in package directory: %r"%unexpected

L=["# VIDEO 6 - H.I.T. FINAL RECORDING PACKAGE",
   "# SHA-256 of the 12 user-facing files in this package.",
   "# SHA256SUMS.txt cannot hash itself. The master ZIP cannot contain its own",
   "# checksum either; it is published in the sibling file",
   "# Video_6_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256",
   "# Video_6_YouTube_Description_HIT.docx sits outside this package; its",
   "# SHA-256 is reported in the delivery summary.",""]
for m in MANIFEST: L.append("%s  %s"%(sha256(os.path.join(ROOT,m)),m))
open(os.path.join(ROOT,SUMS),"w").write("\n".join(L)+"\n")

if os.path.exists(ZIP): os.remove(ZIP)
with zipfile.ZipFile(ZIP,"w",zipfile.ZIP_DEFLATED) as z:
    for m in MANIFEST+[SUMS]:
        z.write(os.path.join(ROOT,m), "Video_6_HIT_FINAL/"+m)
zsha=sha256(ZIP)
open(ZIP+".sha256","w").write("%s  %s\n"%(zsha,os.path.basename(ZIP)))

PROV="/tmp/v6hit/_source"
shutil.rmtree(PROV,ignore_errors=True); os.makedirs(PROV)
for f in ("script_text.py","build.py","qa.py","canonical_script.txt"):
    src="/tmp/v6hit/"+f
    if os.path.isfile(src): shutil.copy2(src, os.path.join(PROV,f))
print("ZIP sha256:",zsha)
print("description-only doc sha256:",sha256(DESC_DOC))
