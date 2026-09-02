# -*- coding: utf-8 -*-
"""Build the Video 5 DIFFERENTIATED v3.0 final recording and Shorts package."""
import os, sys, shutil, zipfile, hashlib
sys.path.insert(0, "/tmp/v5v3")
from script_text import LINES, SPOKEN, MARKERS
from shorts_text import SHORTS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0x9B,0x2C,0x10)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

BASE="/tmp/v5v3"
ROOT=os.path.join(BASE,"Video_5_HIT_FINAL")
LF=os.path.join(ROOT,"LONG_FORM"); SH=os.path.join(ROOT,"SHORTS")
shutil.rmtree(ROOT, ignore_errors=True)
os.makedirs(LF); os.makedirs(SH)

TITLE="Should I Make an Internal Move? 3 Questions to Decide"
THUMB="YOU MAY NOT NEED TO LEAVE"
PRIMARY="should I make an internal move"
SUPPORTING=("internal move · internal mobility · internal job transfer · "
    "career growth · career decision · should I leave my company · "
    "internal career move · career portability · internal hiring · "
    "career development")
CTA="Career Decision Evidence Check"
CTA_URL="https://temidayoafonja.com/career-decisions"
NEXT="Are You Growing—or Just Being Given More Work?"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO 5",size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t,before=20):
    return keep(P(d,t,size=14,bold=True,color=NAVY,before=before,after=8),True)
def H2(d,t,before=13):
    return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=before,after=5),True)

def compress(d, line_spacing=1.18, after_scale=0.80):
    """Tighten one document's vertical rhythm without touching type size,
    weight, colour or wording. Keeps editor briefs off a near-empty last page."""
    for p in d.paragraphs:
        pf=p.paragraph_format
        if pf.line_spacing and 1.0 < pf.line_spacing < 1.4 and pf.line_spacing > line_spacing:
            pf.line_spacing = line_spacing
        if pf.space_after is not None:
            pf.space_after = Pt(round(pf.space_after.pt*after_scale,1))
        if pf.space_before is not None and pf.space_before.pt:
            pf.space_before = Pt(round(pf.space_before.pt*after_scale,1))
    return d

def pairlist(d, items, indent="—  ", gap="        ", after=4, budget=78):
    """Lay short bullet items several to a line, packed by width. Every item is
    kept verbatim; type size and leading are untouched."""
    lines, cur = [], []
    for it in items:
        trial = cur + [it]
        width = sum(len(indent) + len(x) for x in trial) + len(gap) * (len(trial) - 1)
        if cur and width > budget:
            lines.append(cur); cur = [it]
        else:
            cur = trial
    if cur: lines.append(cur)
    for group in lines:
        keep(P(d, gap.join(indent + x for x in group), after=after))

TEL="Video5TeleprompterScriptwithslidemarkers_HIT_v3.0"
RDG="Video5ReadingScriptnomarkers_HIT_v3.0"
EDB="Video_5_EDITOR_ONLY_HIT_Brief_v3.0.docx"
PUB="Video_5_Publishing_Package_HIT_v3.0.docx"
SEB="Video_5_Shorts_EDITOR_ONLY_HIT_Brief.docx"

# ------------------------------------------------- 1. teleprompter DOCX + TXT
d=newdoc(True)
head(d,TITLE,"Video 5  ·  Teleprompter script with slide markers",
     "Spoken script is the large text. A slide marker in a tinted band tells "
     "the editor which slide to bring up; it is not spoken.")
for line in LINES:
    if line.startswith("[SLIDE:"):
        name=line[len("[SLIDE:"):-1].strip()
        p=P(d,"SLIDE  —  %s"%name,size=11,bold=True,color=NAVY,
            before=14,after=14,spacing=1.1)
        shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
    else:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,TEL+".docx"))
tel=[TITLE,"Video 5  ·  Teleprompter script with slide markers",""]
for line in LINES:
    if line.startswith("[SLIDE:"):
        tel += ["", "SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(), ""]
    else: tel += [line, ""]
open(os.path.join(LF,TEL+".txt"),"w").write("\n".join(tel).strip()+"\n")

# ------------------------------------------------- 2. reading script DOCX+TXT
d=newdoc(True)
head(d,TITLE,"Video 5  ·  Reading script, no markers",
     "Spoken language only. No slide markers, no timestamps, no production "
     "directions.")
for line in SPOKEN:
    keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,RDG+".docx"))
open(os.path.join(LF,RDG+".txt"),"w").write("\n\n".join(SPOKEN)+"\n")
print("long-form scripts written")

# --------------------------------------------------- 3. long-form editor brief
SLIDE_JOBS=["Core Distinction","The Three Questions","1 — Will the Work Change?",
 "Access Test","2 — Will Your Judgment Expand?","More Tasks / More Judgment",
 "3 — Will the Evidence Travel?","Result / Judgment / Range","Decision Read",
 "Conversation Prompts","Career Decision Evidence Check","Watch Next"]
REVEAL_MAP=[(1,"1–2",2),(2,"3–5",3),(3,"6",1),(4,"7–10",4),(5,"11",1),
 (6,"12–13",2),(7,"14",1),(8,"15–17",3),(9,"18–20",3),(10,"21–23",3),
 (11,"24",1),(12,"25",1)]

d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 5  ·  v3.0",size=12,bold=True,color=GOLD,after=2,caps=True)
P(d,TITLE,size=20,bold=True,color=NAVY,after=6,spacing=1.1)
p=P(d,"This document is for the editor. It is NOT Temidayo's teleprompter and "
     "must not be placed on the recording screen.",size=11,italic=True,
     color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

# 1 -------------------------------------------------------- locked metadata
H1(d,"1.  Locked metadata",before=14)
for k,v in (("Title",TITLE),("Thumbnail",THUMB),("Primary search phrase",PRIMARY),
            ("Format","Searchable decision + organizational mechanics"),
            ("Primary CTA",CTA),("CTA URL",CTA_URL),
            ("Watch next","Video 6 — "+NEXT),
            ("Core distinction","MOVEMENT IS NOT AUTOMATICALLY GROWTH."),
            ("Memory device","The three questions. No acronym, no second "
             "framework.")):
    keep(P(d,"%-24s %s"%(k+":",v),size=11,after=5))
p=P(d,"CTA PRODUCTION GATE: SATISFIED. The Career Decision Evidence Check page "
     "is live and the core production journey has been verified. Video 5 is "
     "NOT blocked on this CTA page. One normal signed-out link check is "
     "retained in the final upload SOP.",size=11,bold=True,color=NAVY,
     before=8,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
keep(P(d,"Do not add CAR — it belongs to Video 6. Do not add the Career "
       "Evidence 3 Cs — they belong to the evidence method, not this decision "
       "video.",bold=True,color=RED,after=8,spacing=1.25))

# 2 --------------------------------------------------------------- H.I.T. map
H1(d,"2.  First 30 seconds — H.I.T. map",before=14)
P(d,"H = Hook. I = Interest. T = Trust. The opening must work as one "
    "audiovisual unit: immediate conversational hook, meaningful visual "
    "interest, relevant lived proof, and a clear viewer payoff by 30 seconds. "
    "No generic welcome before the promise. No title card first.",after=8)

def beat(t,anchor,layer,body):
    H2(d,t,before=10)
    p=P(d,"Spoken anchor:  “%s”"%anchor,size=10.5,italic=True,color=DIM,after=8)
    shade(p,BAND_CREAM)
    if layer: keep(P(d,layer,size=11,bold=True,color=GOLD,after=6))
    for b in body: keep(P(d,b,after=5))

beat("0:00–0:05","You may not need to leave your company.","H = HOOK",
     ["Visual: direct to camera.",
      "On-screen text:  YOU MAY NOT NEED TO LEAVE",
      "No title card before this."])
beat("0:05–0:10","You may need access to work the company has not trusted you "
     "with yet.","I = INTEREST",
     ["On-screen:  SAME COMPANY  /  DIFFERENT WORK",
      "Keep it restrained. No office-building stock footage. No transfer "
      "arrows."])
beat("0:10–0:21","About six months after I returned from maternity leave in "
     "one chapter of my career, my scope expanded beyond the original box of "
     "the role.","T = TRUST",
     ["Visual progression:  MORE WORK?  then  DIFFERENT WORK",
      "The proof is Temidayo saying what actually happened. Keep her visible.",
      "Do not use family or maternity stock imagery. Do not dramatize the "
      "maternity reference."])
beat("0:21–0:30","Before you accept an internal move, ask three questions…",
     "PAYOFF",
     ["Progressively reveal:",
      "     WILL THE WORK CHANGE?",
      "     WILL YOUR JUDGMENT EXPAND?",
      "     WILL THE EVIDENCE TRAVEL?",
      "Viewer payoff is clear by 30 seconds. Then move into the deck."])

# 3 ------------------------------------------- organizational-mechanics purpose
H1(d,"3.  The organizational-mechanics purpose of this video",before=14)
P(d,"This video must not feel like “three tips for getting an internal job.” "
    "It should feel like: “Temidayo is teaching me how to read what an "
    "internal opportunity is actually designed to give me.”",after=8)
P(d,"Use the employee + organization dual perspective throughout. An internal "
    "move is not only a decision the employee makes; it is also a decision "
    "about what work, context, judgment and ownership the organization is "
    "willing to entrust to that person. Do not turn the organizational layer "
    "into HR jargon or enterprise commentary that leaves the viewer behind.",
    after=10)
H2(d,"Lines to let breathe",before=10)
for q in ["“An internal move is not only a decision you make. It is also a "
 "decision the organization is making about what it is willing to trust you "
 "to do next.”",
 "“The same people who know your strengths may also know you too well inside "
 "one box.”",
 "“Sometimes the organizational problem is not your capability. It is "
 "recognition.”",
 "“A stretch opportunity is developmental when the person is trusted with "
 "more judgment, not simply handed more volume.”",
 "“An internal move can fail even when the person is capable.”"]:
    keep(P(d,q,size=10.5,after=6,spacing=1.25))
p=P(d,"Do not convert these into motivational quote cards. Restrained text "
     "emphasis only, where it is genuinely useful.",size=11,bold=True,
     color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

# 4 ---------------------------------------------------- slide-marker mapping
H1(d,"4.  Slide marker → actual slide number",before=14)
P(d,"The teleprompter carries twelve slide markers. They map to the existing "
    "twelve-slide deck in order, marker 1 to slide 1 through marker 12 to "
    "slide 12. Do not add, delete, redesign or reorder slides.",after=8)
for n,job in enumerate(SLIDE_JOBS,1):
    keep(P(d,"Marker %-3d →  Slide %-3d %s"%(n,n,job),size=10.5,after=3))

# 5 -------------------------------------------------------- reveal-frame map
H1(d,"5.  Existing reveal-frame map",before=14)
P(d,"The reveal-build deck contains 25 frames. This is the inspected count "
    "from the actual file, not an assumption. Reveal visuals are unchanged.",
    after=8)
for n,rng,cnt in REVEAL_MAP:
    keep(P(d,"Slide %-3d →  reveal frames %-8s (%d)"%(n,rng,cnt),size=10.5,after=3))
keep(P(d,"Slides 3, 5, 7, 11 and 12 are single-state frames; the rest build.",
       size=10.5,italic=True,color=DIM,before=4,after=8))
H2(d,"Let the visual carry the information",before=10)
P(d,"Do not add spoken wording to compensate for slide copy.",after=6)
for s_,note in [("Slide 2","carries the three-question preview."),
 ("Slide 4","carries new problems / systems / stakeholders / context."),
 ("Slide 6","carries MORE TASKS vs MORE JUDGMENT."),
 ("Slide 8","carries RESULT / JUDGMENT / RANGE."),
 ("Slide 9","carries the 3 / 2 / 0–1 decision read."),
 ("Slide 10","carries the conversation prompts.")]:
    keep(P(d,"%s %s"%(s_,note),after=5))

# 6 ---------------------------------------------------------- factual boundaries
H1(d,"6.  Factual boundaries",before=14)
p=P(d,"FACTUAL BOUNDARY. The confirmed evidence is only this: about six months "
     "after Temidayo returned from maternity leave in one career chapter, her "
     "scope expanded beyond the original box of the role, and the meaningful "
     "part was being trusted with different work.",size=11,bold=True,
     color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
P(d,"Do not add:",after=5)
pairlist(d,["an employer name;","an exact title;","an exact assignment;",
 "an executive conversation;","a quote;","a metric;","a promotion;",
 "a causal claim;","a conflict;","a resistance story;",
 "an unestablished result."],after=3)
keep(P(d,"Do not name the employer publicly in this video.",bold=True,
       color=RED,before=4,after=6))
keep(P(d,"The accurate learning is NOT “my scope got bigger, therefore I "
       "grew.” It is that being trusted with different work was the "
       "meaningful part.",after=8,spacing=1.25))

# 7 ---------------------------------------------------------- safety boundaries
H1(d,"7.  Safety and decision boundaries",before=14)
P(d,"Preserve the script's explicit acknowledgment that:",after=6)
for x in ["a lower-formation move may still be right for pay, flexibility, "
 "benefits, stability or manager fit;","health and safety take priority;",
 "harassment or discrimination is not a situation to optimise around before "
 "seeking appropriate support;",
 "caregiving, immigration status, location, energy and timing may change the "
 "decision;",
 "an internal move will not solve every problem, and a capable person can "
 "still be overlooked."]:
    keep(P(d,"—  "+x,after=4))
keep(P(d,"Never claim that every internal move is growth, and never claim an "
       "external move is automatically better.",bold=True,color=RED,before=4,
       after=8,spacing=1.25))

# 8 ---------------------------------------------------------- CTA + watch next
H1(d,"8.  CTA and watch next",before=14)
keep(P(d,"One resource CTA only: %s — %s"%(CTA,CTA_URL),after=5))
keep(P(d,"Do not add the Capability Formation Field Kit, the Career Evidence "
       "Starter or Keep the Proof to this video.",bold=True,after=6))
keep(P(d,"The production gate is satisfied.",bold=True,color=NAVY,after=8))
keep(P(d,"Watch next: Video 6 — %s"%NEXT,bold=True,after=5))
for x in ["use direct Video 6 end-screen routing once Video 6 is public;",
 "before Video 6 is public, use the Career Portability playlist if it is "
 "functioning;",
 "otherwise use an appropriate currently public video or YouTube's "
 "best-for-viewer option temporarily."]:
    keep(P(d,"—  "+x,after=4))
keep(P(d,"Do not leave Subscribe as the only end-screen element.",bold=True,
       color=RED,before=4,after=8))

# 9 --------------------------------------------------------------- editing rhythm
H1(d,"9.  Editing rhythm",before=14)
P(d,"After the first 30 seconds, let the deck carry the three questions, the "
    "access variables, the task-versus-judgment distinction, the evidence "
    "types, the decision read and the conversation prompts. Temidayo's voice "
    "carries interpretation, organizational mechanics, limits and real-life "
    "tradeoffs.",after=8)
pairlist(d,["preserve natural pacing;","avoid constant motion;",
 "avoid decorative B-roll;","avoid text duplication;",
 "avoid constant punch-ins;","allow reflective pauses."],after=3)

# 10 --------------------------------------------------------- visual do-not-use
H1(d,"10.  Visual “do not use” list",before=14)
pairlist(d,["maternity imagery;","baby imagery;","stock office B-roll;",
 "transfer arrows;","doors;","ladders;","job-change icons;",
 "generic org charts;","fake promotion graphics;","org-chart animation;",
 "employer logos;","résumé graphics;","constant zooms;",
 "red warning graphics;","fake shock expressions;","AI-generated scenery;",
 "social-media template effects."],after=3)

# 11 ------------------------------------------------- speaker-note update record
H1(d,"11.  Speaker-note update record",before=14)
P(d,"Both decks ship with speaker notes rewritten for the v3.0 script. The "
    "superseded v2.0 timings and narration cues have been removed; nothing "
    "on any slide was changed.",after=8)
for x in ["Main deck: 12 notes parts rewritten, one per slide.",
 "Reveal deck: 25 notes parts rewritten, one per reveal frame.",
 "Slide XML, geometry, typography, palette and media: unchanged.",
 "Timings in the notes are script-derived working estimates at 145 words per "
 "minute. Replace them from the finished cut."]:
    keep(P(d,"—  "+x,after=4))
compress(d, 1.12, 0.50)
d.save(os.path.join(LF,EDB))
print("editor brief written")

# --------------------------------------------------- chapters and description
CHAPTERS=[("00:00","You May Not Need to Leave"),
 ("00:38","What an Internal Move Really Gives You"),
 ("01:58","The 3 Questions for an Internal Move"),
 ("02:14","Will the Work Change?"),
 ("03:06","When the Organization Still Sees the Old You"),
 ("04:32","Will Your Judgment Expand?"),
 ("04:39","More Tasks vs. More Judgment"),
 ("06:40","Will the Evidence Travel?"),
 ("08:00","Read the Three Answers"),
 ("09:00","What the Opportunity Says About the Organization"),
 ("10:32","When an Internal Move Will Not Solve the Problem"),
 ("11:17","Career Decision Evidence Check"),
 ("11:45","Are You Growing—or Just Being Given More Work?")]
CHAPTER_LINES=["%s %s"%(t,c) for t,c in CHAPTERS]

EMOJI_NOTE=("The restrained emoji treatment on the section labels and the "
    "three teaching bullets is part of the approved standard: ✨ teaching "
    "points, 🧭 CTA and resource, ⏱️ chapters, ▶️ Watch Next, 🔗 Connect and "
    "Explore. Do not remove it, and do not add more.")

DESC=[
 "You may not need to leave your company to find work that expands your career.",
 "An internal move can give you access to different problems, broader "
 "judgment and evidence you can carry into future roles.",
 "But movement is not automatically growth.",
 "In this video, I use three questions to test an internal role, transfer, "
 "team change or stretch opportunity:",
 "✨ Will the work change?",
 "✨ Will your judgment expand?",
 "✨ Will the evidence travel?",
 "I also look at the organizational side of the decision: what the company is "
 "actually willing to let you learn, decide and own — and why a capable "
 "person can still be overlooked for different work inside the same "
 "organization.",
 "The real question is not simply whether you moved.",
 "It is whether the move increased what you can carry afterward.","",
 "🧭 CAREER DECISION EVIDENCE CHECK",
 "If you are actively deciding whether to stay, move internally or leave, use "
 "the Career Decision Evidence Check to organize the evidence behind the "
 "choice:",
 CTA_URL,"",
 "⏱️ CHAPTERS"]+CHAPTER_LINES+["",
 "▶️ WATCH NEXT", NEXT, "[ADD VIDEO 6 LINK WHEN LIVE]","",
 "🔗 CONNECT AND EXPLORE",
 "Website:","https://temidayoafonja.com",
 "LinkedIn:","https://www.linkedin.com/in/temidayo-afonja",
 "Substack:","https://temidayoafonja.substack.com","",
 "#InternalMobility #CareerGrowth #CareerDecision"]

PINNED=["If you are considering an internal move, which question is least clear?",
 "1. Will the work change?","2. Will my judgment expand?",
 "3. Will the evidence travel?",
 "Then ask one more question:",
 "What is the organization actually willing to trust me to do next?",
 "You do not need to share confidential details.",
 "If you are actively deciding whether to stay, move internally or leave, the "
 "Career Decision Evidence Check is here:", CTA_URL]

TAGS=("should I make an internal move, internal move, internal mobility, "
 "internal job transfer, internal hiring, career growth, career decision, "
 "should I leave my company, career portability, internal career move, "
 "career transition, experienced professionals, Temidayo Afonja, "
 "Capability Formation")

def description_block(d, heading_before=14, upload_doc=False):
    """The copy-ready description, its end marker and the internal note.

    upload_doc=True moves the editorial emoji instruction ABOVE an explicit
    COPY-READY ... BEGIN marker, so nothing internal sits inside or immediately
    before the block a person selects when pasting into YouTube."""
    if upload_doc:
        H1(d,"INTERNAL NOTE — DO NOT PASTE INTO YOUTUBE",before=heading_before)
        p=P(d,EMOJI_NOTE,size=10.5,italic=True,color=RED,after=12,spacing=1.25)
        shade(p,BAND_CREAM); keep(p)
        p=keep(P(d,"COPY-READY YOUTUBE DESCRIPTION — BEGIN",size=11,bold=True,
                 color=NAVY,before=14,after=12,spacing=1.2))
        shade(p,BAND_NAVY)
    else:
        H1(d,"Description",before=heading_before)
        keep(P(d,EMOJI_NOTE,size=10.5,italic=True,color=DIM,after=10))
    for para in DESC:
        keep(P(d,para if para else " ",after=7 if para else 3))
    keep(P(d,"— END OF COPY-READY DESCRIPTION —",size=10,bold=True,
           color=DIM,before=14,after=12,spacing=1.2))
    H1(d,"Internal note — do not paste into YouTube",before=14)
    p=P(d,"WORKING ESTIMATES — EDITOR MUST REPLACE FROM FINAL CUT",size=11,
        bold=True,color=RED,after=6,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    p=P(d,"These timestamps are script-derived and must be replaced using the "
        "finished edit. Do not force the edit to match these estimates.",
        size=10.5,bold=True,italic=True,color=RED,after=10,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    H1(d,"Working chapters — reference copy",before=14)
    keep(P(d,"Identical to the thirteen chapter lines inside the description "
           "above.",size=10.5,italic=True,color=DIM,after=8))
    for line in CHAPTER_LINES: keep(P(d,line,size=11,after=4))

# ----------------------------------------------------- 4. publishing package
d=newdoc()
head(d,TITLE,"Video 5  ·  Publishing package  ·  v3.0",
     "Everything needed to upload. Working timestamps must be replaced with "
     "real ones from the finished edit.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14); P(d,PRIMARY,after=10)
H1(d,"Supporting search language",before=14); P(d,SUPPORTING,after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only. Do not put the full tag field in the "
       "public description.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
H1(d,"Watch next",before=14)
keep(P(d,"%s  (Video 6)"%NEXT,bold=True,after=8))
H1(d,"CTA production gate",before=14)
keep(P(d,"SATISFIED. The Career Decision Evidence Check page is live and the "
       "core production journey has been verified. Video 5 is not blocked on "
       "it. Retain one normal signed-out link check in the final upload SOP.",
       bold=True,after=8,spacing=1.25))
compress(d)
d.save(os.path.join(LF,PUB))

# ------------------------------- 4b. separate description-only document
d=newdoc()
head(d,TITLE,"Video 5  ·  YouTube description",
     "Upload copy only. Everything below the end marker is internal and must "
     "not be pasted into YouTube.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14); P(d,PRIMARY,after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"Watch next",before=14); keep(P(d,"%s  (Video 6)"%NEXT,bold=True,after=8))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
compress(d)
DESC_DOC=os.path.join(BASE,"Video_5_YouTube_Description_HIT.docx")
d.save(DESC_DOC)
print("publishing package and description-only document written")

# ---------------------------------------------------------------- 5. Shorts
LABELS=["SHORT 1","SHORT 2","SHORT 3","SHORT 4"]
for (fn,role,hook,copy),label in zip(SHORTS,LABELS):
    d=newdoc(True)
    P(d,"VIDEO 5 SHORT",size=10,bold=True,color=GOLD,after=4,caps=True)
    P(d,label,size=20,bold=True,color=NAVY,after=8,spacing=1.1)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    keep(P(d,"Verbal hook:  “%s”"%hook,size=11,color=DIM,after=5))
    keep(P(d,"Related long-form:  %s"%TITLE,size=11,color=DIM,after=10))
    H1(d,"RECORDING COPY",before=12)
    for line in copy:
        keep(P(d,line,size=13.5,color=INK,after=10,spacing=1.5))
    d.save(os.path.join(SH,fn))
print("%d Shorts written"%len(SHORTS))

# ------------------------------------------------------ 6. Shorts editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 5 — FOUR STANDALONE SHORTS",size=18,bold=True,color=NAVY,
  after=8,spacing=1.1)
p=P(d,"This document is for the editor. It is separate from the four Short "
     "recording documents and must not be placed on Temidayo's recording "
     "screen.",size=11,italic=True,color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"How these are produced",before=14)
keep(P(d,"These are separately recorded 9:16 Shorts. They are NOT excerpts cut "
       "from the long-form video.",bold=True,after=10))
P(d,"Each Short needs:",after=6)
pairlist(d,["an immediate verbal hook;","a corresponding on-screen hook;",
 "meaningful visual movement;","accurate mobile-safe captions;",
 "premium, restrained editorial pacing;",
 "Video 5 added as the Related Video when available."])

def short(label,role,onscreen,body):
    H1(d,label,before=14)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    p=keep(P(d,"On-screen hook:  %s"%onscreen,size=11,bold=True,color=GOLD,after=8))
    shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))
    keep(P(d,"Related Video:  Video 5",size=10.5,color=DIM,before=4,after=6))

short("SHORT 1","Recognition","YOU MAY NOT NEED TO LEAVE",
 ["Visual:  SAME COMPANY  /  DIFFERENT WORK",
  "End on:  WILL THE WORK CHANGE?"])
short("SHORT 2","Distinction","MORE TASKS ≠ MORE JUDGMENT",
 ["Use the existing restrained contrast:  MORE TASKS  vs  MORE JUDGMENT",
  "No generic office B-roll.",
  "End on:  WHAT WILL YOU BE TRUSTED TO DECIDE?"])
short("SHORT 3","Personal evidence","MORE WORK WASN’T THE POINT",
 ["Then:  TRUSTED WITH DIFFERENT WORK",
  "FACTUAL BOUNDARY: do not use maternity or baby stock imagery, do not name "
  "the employer, and do not add a metric or a result.",
  "End on:  WHAT AM I BECOMING TRUSTED TO DO?"])
short("SHORT 4","Practical test","3 QUESTIONS BEFORE YOU MOVE",
 ["Reveal:",
  "     WILL THE WORK CHANGE?",
  "     WILL YOUR JUDGMENT EXPAND?",
  "     WILL THE EVIDENCE TRAVEL?",
  "Then:  3 YES  /  2 YES  /  0–1 YES",
  "Do not gamify it. Keep the life-factor boundary intact."])

H1(d,"All Shorts — visual boundaries",before=14)
P(d,"Do not use:",after=5)
pairlist(d,["generic office B-roll;","transfer-arrow clichés;",
 "company logos;","org-chart animations;","maternity or baby imagery;",
 "flashy transitions;","constant zooms;","fake shock expressions;",
 "red warning graphics;","AI-generated scenery;",
 "social-media template effects."],after=3)
compress(d, 1.18, 0.62)
d.save(os.path.join(SH,SEB))

# ---------------------------------------------------------------- 7. README
FILES=(["LONG_FORM/"+f for f in sorted(os.listdir(LF))]
      +["SHORTS/"+f for f in sorted(os.listdir(SH))])
R=["VIDEO 5 — DIFFERENTIATED H.I.T. FINAL RECORDING PACKAGE v3.0","",
 "Title:             %s"%TITLE,
 "Thumbnail:         %s"%THUMB,
 "Format:            Searchable decision + organizational mechanics",
 "Core distinction:  Movement is not automatically growth.","",
 "Memory structure:",
 "  Will the work change?",
 "  Will your judgment expand?",
 "  Will the evidence travel?","",
 "Personal proof:    About six months after returning from maternity leave in",
 "                   one career chapter, Temidayo's scope expanded beyond the",
 "                   original box of the role; the meaningful part was being",
 "                   trusted with different work.",
 "Public employer:   NOT NAMED.","",
 "Primary CTA:       %s"%CTA,
 "CTA URL:           %s"%CTA_URL,
 "CTA production",
 "gate:              SATISFIED",
 "Watch next:        %s"%NEXT,"",
 "Slides:            Visual design and on-slide copy unchanged. 12 main",
 "                   slides.",
 "Speaker notes:     Updated for v3.0.",
 "Reveal deck:       Visual design and reveal states unchanged. 25 frames.",
 "Shorts:            Four separately recorded scripts.",
 "Description-only",
 "document:          Separate from this ZIP.",
 "Editor directions: Separated from recording copy.","",
 "-"*70,"","WHAT EACH FILE IS","",
 "LONG_FORM/","",
 "  %s.docx"%TEL,
 "  %s.txt"%TEL,
 "      Temidayo's recording copy. Spoken script in large text; slide markers",
 "      in tinted bands. The markers are not spoken.","",
 "  %s.docx"%RDG,
 "  %s.txt"%RDG,
 "      The same spoken words with the slide markers removed.","",
 "  %s"%EDB,
 "      For the editor. Locked metadata, the H.I.T. first-30-second map, the",
 "      organizational-mechanics purpose, the slide-marker mapping, the",
 "      reveal-frame map, the factual and safety boundaries, the CTA and",
 "      watch-next routing, editing rhythm, the visual do-not-use list and",
 "      the speaker-note update record. Not for the teleprompter.","",
 "  %s"%PUB,
 "      Title, thumbnail, search language, the copy-ready description with",
 "      its approved emoji treatment, working chapter estimates, pinned",
 "      comment and the tag field.","",
 "SHORTS/","",
 "  Four recording documents, one per Short. These contain Temidayo's",
 "  recording copy and no editor directions.","",
 "  %s"%SEB,
 "      For the editor. On-screen hooks and visual treatment for all four.","",
 "-"*70,"","ALL FILES IN THIS PACKAGE","",]
for f in FILES: R.append("  "+f)
R+=["  README_FINAL.txt","  SHA256SUMS.txt","",
 "-"*70,"","WORKING CHAPTER TIMESTAMPS","",
 "The chapter timestamps in the publishing package are WORKING ESTIMATES",
 "derived from the script. They were not measured from an edit. The editor",
 "must replace every one of them from the finished cut before publishing.","",
 "-"*70,"","FACTUAL BOUNDARY ON THE PERSONAL PROOF","",
 "The confirmed evidence is only that about six months after Temidayo returned",
 "from maternity leave in one career chapter, her scope expanded beyond the",
 "original box of the role, and that the meaningful part was being trusted",
 "with different work. No employer, exact role, exact assignment, causal",
 "claim, quote, metric or unsupported result is added anywhere in this",
 "package.","",
 "-"*70,"","CHECKSUMS","",
 "SHA256SUMS.txt covers the other 12 user-facing files in this package. It",
 "does not hash itself, and it carries no ZIP checksum. The archive's own",
 "SHA-256 is in the sibling file:",
 "  Video_5_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256","",
 "-"*70,"","WHAT WAS NOT CHANGED","",
 "The Video 5 PowerPoint deck (12 slides), the reveal-build deck (25 frames),",
 "the approved thumbnail, the Career Decision Evidence Check page, every",
 "website file, every product and every other video are unchanged. Only the",
 "speaker notes in the two decks were updated, to match the v3.0 script.","",
 "The 12-slide deck remains authoritative. The teleprompter's 12 slide markers",
 "map to it in order, slide 1 to slide 12.",""]
open(os.path.join(ROOT,"README_FINAL.txt"),"w").write("\n".join(R))
print("shorts editor brief and README written")

# ------------------------------------------- 8. checksums and the master ZIP
MANIFEST=[
 "LONG_FORM/%s.docx"%TEL,
 "LONG_FORM/%s.txt"%TEL,
 "LONG_FORM/%s.docx"%RDG,
 "LONG_FORM/%s.txt"%RDG,
 "LONG_FORM/%s"%EDB,
 "LONG_FORM/%s"%PUB,
 "SHORTS/Video_5_Short_1_You_May_Not_Need_To_Leave.docx",
 "SHORTS/Video_5_Short_2_More_Tasks_Not_More_Judgment.docx",
 "SHORTS/Video_5_Short_3_Maternity_Return_Scope_Expansion.docx",
 "SHORTS/Video_5_Short_4_Three_Questions_Before_You_Move.docx",
 "SHORTS/%s"%SEB,
 "README_FINAL.txt",
]
SUMS="SHA256SUMS.txt"
ZIP=os.path.join(BASE,"Video_5_HIT_FINAL_Recording_and_Shorts_Package.zip")

def sha256(p):
    h=hashlib.sha256()
    with open(p,"rb") as fh:
        for b in iter(lambda: fh.read(1<<20), b""): h.update(b)
    return h.hexdigest()

for m in MANIFEST:
    assert os.path.isfile(os.path.join(ROOT,m)), "missing from build: "+m
on_disk=set()
for dp,dn,fn in os.walk(ROOT):
    dn[:]=[x for x in dn if x!="__pycache__"]
    for f in fn:
        if f.endswith(".pyc"): continue
        on_disk.add(os.path.relpath(os.path.join(dp,f),ROOT).replace(os.sep,"/"))
unexpected=sorted(on_disk-set(MANIFEST)-{SUMS})
assert not unexpected, "unexpected files in package directory: %r"%unexpected

L=["# VIDEO 5 - DIFFERENTIATED H.I.T. FINAL RECORDING PACKAGE v3.0",
   "# SHA-256 of the 12 user-facing files in this package.",
   "# SHA256SUMS.txt cannot hash itself. The master ZIP cannot contain its own",
   "# checksum either; it is published in the sibling file",
   "# Video_5_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256",""]
for m in MANIFEST: L.append("%s  %s"%(sha256(os.path.join(ROOT,m)),m))
open(os.path.join(ROOT,SUMS),"w").write("\n".join(L)+"\n")

if os.path.exists(ZIP): os.remove(ZIP)
with zipfile.ZipFile(ZIP,"w",zipfile.ZIP_DEFLATED) as z:
    for m in MANIFEST+[SUMS]:
        z.write(os.path.join(ROOT,m), "Video_5_HIT_FINAL/"+m)
zsha=sha256(ZIP)
open(ZIP+".sha256","w").write("%s  %s\n"%(zsha,os.path.basename(ZIP)))
print("ZIP sha256:",zsha)
print("DESC-ONLY sha256:",sha256(DESC_DOC))
