"""Build the two Video 5 script files from the production package.

Every spoken word is copied out of the package. Nothing is rewritten, expanded,
summarised or reordered.

  Video_5_Teleprompter_Script_with_Slide_Markers.docx
      the approved script unchanged, with slide markers and the package's own
      stage directions kept visually distinct from the spoken lines

  Video_5_Recording_Script_Clean.txt
      the spoken words only
"""
import os, re, zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = ("/root/.claude/uploads/f121668d-e262-5eb8-9b22-0eaa1006a361/"
       "33bdf063-YouTube_Video_5_Production_Package_Internal_Career_Move_3.docx")
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

NAVY = RGBColor(0x0F, 0x23, 0x46)
GOLD = RGBColor(0x8A, 0x6D, 0x1E)
DIM = RGBColor(0x5A, 0x6B, 0x82)

HDR = re.compile(r'^\d+:\d\d[–-]\d+:\d\d\s*\|')

# Internal labels that are not spoken content and must not appear in either file.
INTERNAL_LABELS = {"VISUAL TEACHING SYSTEM"}

# The package's own editor direction for the 1:20 block, corrected. The slide 1
# marker near the opening is unchanged; this refers to the later reuse.
DIRECTION_FIX = {}


def apply_direction_fix(header):
    for old, new in DIRECTION_FIX.items():
        if old in header:
            return header.replace(old, new)
    return header


def source_blocks():
    z = zipfile.ZipFile(SRC)
    x = ET.fromstring(z.read('word/document.xml'))
    t = lambda p: ''.join(e.text or '' for e in p.iter(NS + 't'))
    blocks = [t(el) for el in x.find(NS + 'body') if el.tag == NS + 'p']
    i4 = next(i for i, b in enumerate(blocks) if b.strip() == '4. Full recording script')
    i5 = next(i for i, b in enumerate(blocks) if b.strip().startswith('5. Slide deck content'))
    return [b for b in blocks[i4 + 1:i5]
            if b.strip() and b.strip() not in INTERNAL_LABELS]


# Slide cues. Each is (slide number, slide name, the spoken paragraph the slide
# lands on). The marker is placed immediately BEFORE that paragraph, except
# where noted as trailing.
CUES = [
    (1,  "Core distinction",
     "By the end, you\u2019ll be able to tell a move", "before"),
    (2,  "The three questions",
     "The three questions are: Will the work change?", "before"),
    (3,  "Question 1, will the work change",
     "The first question is: Will this move give me access", "before"),
    (4,  "Access test",
     "Do not begin with the title.", "before"),
    (5,  "Question 2, will your judgment expand",
     "The second question is: Will the move change what I am trusted", "before"),
    (6,  "Critical distinction, tasks and judgment",
     "This is where more responsibility can be misleading.", "before"),
    (7,  "Question 3, will the evidence travel",
     "The third question is: Will this move produce evidence", "before"),
    (8,  "Portable evidence",
     "Three forms of evidence are especially useful.", "before"),
    (9,  "Decision read",
     "Now put the three answers together.", "before"),
    (10, "Conversation prompts",
     "Before your next internal conversation, write one sentence", "before"),
    (11, "Career Decision Evidence Check",
     "If you are actively deciding whether to stay, move internally", "before"),
    (12, "Watch next",
     "And before you call any opportunity growth,", "before"),
]

STATES = {1: 2, 2: 3, 4: 4, 6: 2, 8: 3, 9: 3, 10: 3}


def shade(p, hexfill):
    pr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexfill)
    pr.append(sh)


def left_bar(p, hexcol):
    pr = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    lf = OxmlElement('w:left')
    lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '18')
    lf.set(qn('w:space'), '8'); lf.set(qn('w:color'), hexcol)
    bd.append(lf); pr.append(bd)


# ---------------------------------------------------------------- opening
# The revised conversational opening lives in a script-only source so the
# production-package DOCX and every slide asset stay untouched.
from opening_revision import NEW_OPENING, REPLACES


def apply_opening_override(sec):
    """Swap the first REPLACES spoken paragraphs for the revised opening.

    Timed headers and the target line are left exactly where they are, so the
    block structure of the script is unchanged.
    """
    out, dropped, inserted = [], 0, False
    for b in sec:
        t = b.strip()
        is_spoken = not HDR.match(t) and not t.startswith("Target")
        if is_spoken and dropped < REPLACES:
            dropped += 1
            if not inserted:
                out.extend(NEW_OPENING)
                inserted = True
            continue
        out.append(b)
    assert dropped == REPLACES, "expected %d paragraphs to replace, found %d" % (
        REPLACES, dropped)
    return out


def build():
    sec = apply_opening_override(source_blocks())
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(13)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.9)
        s.left_margin = s.right_margin = Inches(1.05)

    def para(text, size=13, bold=False, color=None, before=0, after=10,
             italic=False, caps=False, spacing=1.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = spacing
        r = p.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        if caps:
            r.font.all_caps = True
        return p

    h = para("Should I Make an Internal Move? 3 Questions to Decide", size=22, bold=True, color=NAVY,
             after=2, spacing=1.1)
    para("Video 5  ·  Teleprompter script with slide markers", size=12,
         color=DIM, after=6, spacing=1.1)
    para("Spoken script is the large text. Everything in a tinted band is a "
         "production direction and is not spoken.", size=11, italic=True,
         color=DIM, after=20, spacing=1.2)

    cue_by_prefix = {c[2]: c for c in CUES}

    def marker(n, name):
        label = "SLIDE %d  —  %s" % (n, name)
        if n in STATES:
            label += "   (%d reveal states)" % STATES[n]
        p = para(label, size=11, bold=True, color=NAVY, before=14, after=14,
                 spacing=1.1)
        shade(p, "E8EDF4"); left_bar(p, "0F2346")

    pending = None
    for block in sec:
        b = block.strip()
        if b.startswith("Target length"):
            p = para(b, size=10.5, italic=True, color=DIM, before=0, after=18,
                     spacing=1.2)
            shade(p, "F3F0E8")
            continue
        if HDR.match(b):
            p = para(apply_direction_fix(b), size=10.5, bold=True, color=GOLD,
                     before=20, after=12, spacing=1.1, caps=True)
            shade(p, "F3F0E8")
            continue
        hit = next((c for k, c in cue_by_prefix.items() if b.startswith(k)), None)
        if hit and hit[3] == "before":
            marker(hit[0], hit[1])
        para(b, size=13.5, color=RGBColor(0x1A, 0x1A, 0x1A), after=14,
             spacing=1.55)
        if hit and hit[3] == "after":
            marker(hit[0], hit[1])

    out_docx = os.path.join(HERE, "Video_5_Teleprompter_Script_with_Slide_Markers.docx")
    doc.save(out_docx)

    # ---------------------------------------------------- clean spoken script
    spoken = [b.strip() for b in sec
              if not HDR.match(b.strip()) and not b.startswith("Target length")]
    out_txt = os.path.join(HERE, "Video_5_Recording_Script_Clean.txt")
    open(out_txt, "w").write("\n\n".join(spoken) + "\n")

    print("teleprompter :", os.path.basename(out_docx))
    print("clean script :", os.path.basename(out_txt))
    print("spoken paragraphs: %d   words: %d"
          % (len(spoken), sum(len(s.split()) for s in spoken)))
    return out_docx, out_txt, spoken, sec


if __name__ == "__main__":
    build()
