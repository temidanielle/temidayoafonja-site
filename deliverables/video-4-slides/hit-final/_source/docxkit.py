# -*- coding: utf-8 -*-
"""Shared DOCX construction helpers for the direct-address revision pass.
Identical typography, palette and pagination behaviour to the locked packages."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0x9B,0x2C,0x10)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,vid,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO %d"%vid,size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t,before=20):
    return keep(P(d,t,size=14,bold=True,color=NAVY,before=before,after=8),True)
def H2(d,t,before=13):
    return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=before,after=5),True)

def compress(d, line_spacing=1.18, after_scale=0.80):
    for p in d.paragraphs:
        pf=p.paragraph_format
        if pf.line_spacing and 1.0 < pf.line_spacing < 1.4 and pf.line_spacing > line_spacing:
            pf.line_spacing = line_spacing
        if pf.space_after is not None:
            pf.space_after = Pt(round(pf.space_after.pt*after_scale,1))
        if pf.space_before is not None and pf.space_before.pt:
            pf.space_before = Pt(round(pf.space_before.pt*after_scale,1))
    return d

def pairlist(d, items, indent="—  ", gap="        ", after=4, budget=78):
    lines, cur = [], []
    for it in items:
        trial = cur + [it]
        width = sum(len(indent)+len(x) for x in trial)+len(gap)*(len(trial)-1)
        if cur and width > budget: lines.append(cur); cur=[it]
        else: cur = trial
    if cur: lines.append(cur)
    for g in lines: keep(P(d, gap.join(indent+x for x in g), after=after))

def scripts(vid, title, LINES, SPOKEN, LF, tel_stem, rdg_stem):
    """Teleprompter and reading script, DOCX and TXT, from the canonical blocks."""
    d=newdoc(True)
    head(d,vid,title,"Video %d  ·  Teleprompter script with slide markers"%vid,
         "Spoken script is the large text. A slide marker in a tinted band tells "
         "the editor which slide to bring up; it is not spoken.")
    for line in LINES:
        if line.startswith("[SLIDE:"):
            p=P(d,"SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(),size=11,
                bold=True,color=NAVY,before=14,after=14,spacing=1.1)
            shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
        else:
            keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
    d.save(os.path.join(LF,tel_stem+".docx"))
    tel=[title,"Video %d  ·  Teleprompter script with slide markers"%vid,""]
    for line in LINES:
        if line.startswith("[SLIDE:"):
            tel += ["","SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(),""]
        else: tel += [line,""]
    open(os.path.join(LF,tel_stem+".txt"),"w").write("\n".join(tel).strip()+"\n")

    d=newdoc(True)
    head(d,vid,title,"Video %d  ·  Reading script, no markers"%vid,
         "Spoken language only. No slide markers, no timestamps, no production "
         "directions.")
    for line in SPOKEN:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
    d.save(os.path.join(LF,rdg_stem+".docx"))
    open(os.path.join(LF,rdg_stem+".txt"),"w").write("\n\n".join(SPOKEN)+"\n")

def direct_address_section(d, heading, beats):
    """The DIRECT ADDRESS IS PART OF THE CREATIVE section required in every
    editor brief in this pass."""
    H1(d,heading,before=14)
    p=P(d,"Temidayo is speaking to ONE experienced professional sitting across "
         "from her. Not a crowd, not a conference room, not an abstract "
         "audience.",size=11,bold=True,color=NAVY,after=8,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    H2(d,"Editing should support that relationship",before=10)
    pairlist(d,["keep direct questions connected to Temidayo's face where useful;",
     "avoid making this feel like a keynote;","avoid overusing quote cards;",
     "let the slides carry structure;",
     "let Temidayo carry relationship, interpretation and trust;",
     "preserve the pauses around direct questions;",
     "avoid generic “career guru” visual language."],after=3)
    if beats:
        H2(d,"Relational beats — do not cut these as filler",before=10)
        for q in beats: keep(P(d,q,size=10.5,after=5,spacing=1.25))
        p=P(d,"They are part of the viewer relationship, not padding. Do not trim "
             "them for pace, and do not let a graphic cover Temidayo while she "
             "says one.",size=11,bold=True,color=RED,after=8,spacing=1.25)
        shade(p,BAND_CREAM); keep(p)
    keep(P(d,"Do not reintroduce detached phrasing — “some people”, "
           "“professionals often”, “people may”, “a person should”, “employees "
           "often”, “workers may”, “many professionals” — anywhere in the edit, "
           "the captions or the on-screen text.",after=8,spacing=1.25))

def sha256(p):
    import hashlib
    h=hashlib.sha256()
    with open(p,"rb") as fh:
        for b in iter(lambda: fh.read(1<<20), b""): h.update(b)
    return h.hexdigest()

def package(ROOT, MANIFEST, zippath, inner, header):
    """Checksums and the 13-file ZIP, built from the explicit allowlist."""
    import os, zipfile
    for m in MANIFEST:
        assert os.path.isfile(os.path.join(ROOT,m)), "missing from build: "+m
    on_disk=set()
    for dp,dn,fn in os.walk(ROOT):
        dn[:]=[x for x in dn if x!="__pycache__"]
        for f in fn:
            if f.endswith(".pyc"): continue
            on_disk.add(os.path.relpath(os.path.join(dp,f),ROOT).replace(os.sep,"/"))
    unexpected=sorted(on_disk-set(MANIFEST)-{"SHA256SUMS.txt"})
    assert not unexpected, "unexpected files in package directory: %r"%unexpected
    L=header+[""]
    for m in MANIFEST: L.append("%s  %s"%(sha256(os.path.join(ROOT,m)),m))
    open(os.path.join(ROOT,"SHA256SUMS.txt"),"w").write("\n".join(L)+"\n")
    if os.path.exists(zippath): os.remove(zippath)
    with zipfile.ZipFile(zippath,"w",zipfile.ZIP_DEFLATED) as z:
        for m in MANIFEST+["SHA256SUMS.txt"]:
            z.write(os.path.join(ROOT,m), inner+"/"+m)
    zs=sha256(zippath)
    open(zippath+".sha256","w").write("%s  %s\n"%(zs,os.path.basename(zippath)))
    return zs
