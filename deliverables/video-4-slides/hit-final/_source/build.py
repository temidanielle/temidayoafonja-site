# -*- coding: utf-8 -*-
"""Build the Video 4 H.I.T. final recording and Shorts package."""
import os, sys, shutil, zipfile, hashlib
sys.path.insert(0, "/tmp/v4p")
from script_text import LINES, SPOKEN, MARKERS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0x9B,0x2C,0x10)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

ROOT="/tmp/v4p/Video_4_HIT_FINAL"
LF=os.path.join(ROOT,"LONG_FORM"); SH=os.path.join(ROOT,"SHORTS")
shutil.rmtree(ROOT, ignore_errors=True)
os.makedirs(LF); os.makedirs(SH)

TITLE="How to Explain Your Career Change"
THUMB="YOUR CAREER MAKES SENSE"
CTA="Free Career Evidence Starter"
CTA_SHORT="Career Evidence Starter"
CTA_URL="https://temidayoafonja.com/career-evidence-starter"
CTA_DESCRIPTOR="FREE CAREER ACCOMPLISHMENT TRACKER"
CTA_PROMISE=("Turn one accomplishment into proof you can use in a "
             "performance review, interview, internal move or career pivot.")
OLD_CTA="Keep the Proof"
NEXT="Should I Make an Internal Move? 3 Questions to Decide"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO 4",size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t,before=20):
    return keep(P(d,t,size=14,bold=True,color=NAVY,before=before,after=8),True)
def H2(d,t,before=13):
    return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=before,after=5),True)

def compress(d, line_spacing=1.18, after_scale=0.80):
    """Tighten one document's vertical rhythm without touching type size,
    weight, colour or wording. Keeps editor briefs off a near-empty last page."""
    for p in d.paragraphs:
        pf=p.paragraph_format
        if pf.line_spacing and 1.0 < pf.line_spacing < 1.4 and pf.line_spacing > line_spacing:
            pf.line_spacing = line_spacing
        if pf.space_after is not None:
            pf.space_after = Pt(round(pf.space_after.pt*after_scale,1))
        if pf.space_before is not None and pf.space_before.pt:
            pf.space_before = Pt(round(pf.space_before.pt*after_scale,1))
    return d

def pairlist(d, items, indent="—  ", gap="        ", after=4, budget=78):
    """Lay short bullet items several to a line, packed by width.

    Every item is kept verbatim; nothing is removed, reworded or reordered,
    and neither type size nor leading changes. This reclaims the unused
    right-hand part of a short-bullet line, which is what keeps the Shorts
    editor brief to two pages. The width budget is measured against the
    inspection renderer's font, which is wider than Calibri, so Word fits at
    least as much per line and never less."""
    lines, cur = [], []
    for it in items:
        trial = cur + [it]
        width = sum(len(indent) + len(x) for x in trial) + len(gap) * (len(trial) - 1)
        if cur and width > budget:
            lines.append(cur); cur = [it]
        else:
            cur = trial
    if cur: lines.append(cur)
    for group in lines:
        keep(P(d, gap.join(indent + x for x in group), after=after))


# ------------------------------------------------- 1. teleprompter DOCX + TXT
d=newdoc(True)
head(d,TITLE,"Video 4  ·  Teleprompter script with slide markers",
     "Spoken script is the large text. A slide marker in a tinted band tells "
     "the editor which slide to bring up; it is not spoken.")
for line in LINES:
    if line.startswith("[SLIDE:"):
        name=line[len("[SLIDE:"):-1].strip()
        p=P(d,"SLIDE  —  %s"%name,size=11,bold=True,color=NAVY,
            before=14,after=14,spacing=1.1)
        shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
    else:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video4TeleprompterScriptwithslidemarkers_HIT_v2.1.docx"))

tel=[TITLE,"Video 4  ·  Teleprompter script with slide markers",""]
for line in LINES:
    if line.startswith("[SLIDE:"):
        tel += ["", "SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(), ""]
    else:
        tel += [line, ""]
open(os.path.join(LF,"Video4TeleprompterScriptwithslidemarkers_HIT_v2.1.txt"),
     "w").write("\n".join(tel).strip()+"\n")

# ------------------------------------------------- 2. reading script DOCX+TXT
d=newdoc(True)
head(d,TITLE,"Video 4  ·  Reading script, no markers",
     "Spoken language only. No slide markers, no timestamps, no production "
     "directions.")
for line in SPOKEN:
    keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video4ReadingScriptnomarkers_HIT_v2.1.docx"))
open(os.path.join(LF,"Video4ReadingScriptnomarkers_HIT_v2.1.txt"),
     "w").write("\n\n".join(SPOKEN)+"\n")
print("long-form scripts written")

# --------------------------------------------------- 3. long-form editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 4",size=12,bold=True,color=GOLD,after=2,caps=True)
P(d,TITLE,size=20,bold=True,color=NAVY,after=6,spacing=1.1)
p=P(d,"This document is for the editor. It is NOT Temidayo's teleprompter and "
     "must not be placed on the recording screen.",size=11,italic=True,
     color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"Locked metadata",before=14)
for k,v in (("Title",TITLE),("Thumbnail",THUMB),("Primary CTA",CTA_SHORT),
            ("CTA URL",CTA_URL),("Watch next",NEXT),
            ("Core distinction",
             "A chronology tells people where you have been. A portability "
             "explanation tells them what traveled with you.")):
    keep(P(d,"%-18s %s"%(k+":",v),size=11,after=5))

H1(d,"Fact boundary — the “cat with nine lives” comment",before=14)
p=P(d,"ESTABLISHED: the recurring description. One of Temidayo's "
     "senior-manager friends at EY used to joke that she was a “cat with nine "
     "lives.”",size=11,bold=True,color=NAVY,after=6,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
p=P(d,"NOT ESTABLISHED: the original conversation that first prompted the "
     "joke. Do not invent the original meeting, the exact first conversation, "
     "any additional quotation, a reaction from Temidayo or another person, "
     "or a causal story around when or why the phrase first appeared.",
     size=11,bold=True,color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
keep(P(d,"Do not use cat imagery, cat animation, cat sound effects, "
       "reenactments or memes. Do not put “cat with nine lives” in large "
       "on-screen text.",bold=True,after=8,spacing=1.25))

H1(d,"First 30 seconds — H.I.T.",before=14)
P(d,"H = Hook. I = Interest. T = Trust. The opening must work as one "
    "audiovisual unit: immediate human hook, meaningful visual movement, "
    "relevant lived credibility, and a clear payoff by roughly 20 to 30 "
    "seconds. No résumé recital. No title card before the hook. No forced "
    "statistic. No artificial suspense.",after=8)
P(d,"The visual carries the chronology. Temidayo's voice carries the meaning.",
  italic=True,color=DIM,after=14)

def beat(t,anchor,layer,body):
    H2(d,t,before=10)
    p=P(d,"Spoken anchor:  “%s”"%anchor,size=10.5,italic=True,color=DIM,after=8)
    shade(p,BAND_CREAM)
    if layer: keep(P(d,layer,size=11,bold=True,color=GOLD,after=6))
    for b in body: keep(P(d,b,after=5))

beat("0:00–0:05","A senior colleague once called me a cat with nine lives.",
     "H = HOOK",
     ["Visual: begin direct to camera, medium/tight Temidayo.",
      "On-screen text:  MY CAREER LOOKED DISCONNECTED",
      "Do NOT put “cat with nine lives” in large on-screen text.",
      "Do NOT use cat imagery, animation, sound effect or reenactment."])
beat("0:05–0:12",
     "She meant my career kept moving into work that looked unrelated.",
     "I = INTEREST  /  T = TRUST",
     ["Briefly show the existing career progression:",
      "     ACCOUNTING & AUDIT  →  CYBERSECURITY  →  PEOPLE STRATEGY  →  "
      "ENTERPRISE TRANSFORMATION",
      "Temidayo deliberately does not read the labels aloud.",
      "The visual carries chronology. The spoken copy carries meaning."])
beat("0:12–0:19",
     "When I explained it as a list of jobs, the path sounded more "
     "disconnected than it was.",None,
     ["Transition from the full path toward a restrained CHRONOLOGY treatment.",
      "Do not use employer logos. Do not create a résumé animation."])
beat("0:19–0:23","The career wasn’t the problem. The explanation was.",None,
     ["Return cleanly to Temidayo.",
      "On-screen text:  THE EXPLANATION WAS",
      "Let this line breathe. No unnecessary motion."])
beat("0:23–0:30","Here are three steps…","PAYOFF",
     ["Progressively preview:",
      "     NAME THE CHAPTERS",
      "     FIND THE REPEATED WORK",
      "     EXPLAIN THE DIRECTION",
      "Then move into the existing slide system."])

H1(d,"Editorial rhythm after the opening",before=14)
P(d,"After the first 30 seconds:",after=6)
for x in ["preserve natural pacing;",
 "use slides when they help explain distinctions or structures;",
 "avoid constant movement;","avoid decorative B-roll;",
 "avoid unnecessary text duplication;","avoid constant punch-ins;",
 "preserve reflective pauses."]:
    keep(P(d,"—  "+x,after=4))

H1(d,"Existing slides — unchanged",before=14)
P(d,"Use exactly the existing 11-slide sequence. The reveal-build deck is "
    "unchanged. Do not add, remove, redesign or reorder slides.",after=8)
for n,job in enumerate(["Career Path","Chronology / Portability",
 "1 — Name the Chapters Briefly","2 — Find the Repeated Work",
 "Look Beneath the Nouns","3 — Explain the Direction",
 "Three-Sentence Structure","Do Not Invent a Perfect Plan",
 "Explanation Test","Career Evidence Starter  (CTA corrected)",
 "Watch Next"],1):
    keep(P(d,"Slide %-3d %s"%(n,job),size=10.5,after=3))

H1(d,"Let the visual carry information",before=14)
P(d,"Do not add spoken wording to compensate for slide copy.",after=6)
for s_,note in [("Slide 1","carries the full career chronology."),
 ("Slide 2","carries CHRONOLOGY vs PORTABILITY."),
 ("Slides 4 and 5","carry the repeated-work prompts and the noun/verb "
  "contrast."),
 ("Slide 7","carries the three sentence stems."),
 ("Slide 9","carries the three explanation-test questions.")]:
    keep(P(d,"%s %s"%(s_,note),after=5))

H1(d,"Fact and evidence boundaries",before=14)
for x in ["the recurring “cat with nine lives” description is established;",
 "the original first conversation is not established;",
 "the 2008 financial crisis is legitimate context;",
 "do not imply every move was planned;",
 "do not imply all experience transfers;",
 "acknowledge relearning and context;","no unsupported metrics;",
 "no confidential or proprietary detail."]:
    keep(P(d,"—  "+x,after=4))
keep(P(d,"Evidence retention means Temidayo's own recollection and "
       "information she is permitted to retain. Nothing in this video "
       "encourages taking confidential, proprietary, customer, employee or "
       "employer-owned material.",before=8,after=8,spacing=1.25))

H1(d,"CTA slide correction — applied",before=14)
p=P(d,"AUTHORISED AND APPLIED. The %s CTA is SUPERSEDED for this video by the "
     "free Career Evidence Starter. Slide 10 and reveal frame 25 carry the new "
     "CTA. Text only; slide and reveal counts, design, typography family, "
     "colours, layout and the Slide 11 Watch Next card are unchanged."%OLD_CTA,
     size=11,bold=True,color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
for frm,to in (("KEEP THE PROOF","FREE CAREER EVIDENCE STARTER"),
               ("A 60-Minute Career Evidence System",
                "ONE ACCOMPLISHMENT → ONE PORTABLE PROOF LINE"),
               ("temidayoafonja.com/keep-the-proof",
                "temidayoafonja.com/career-evidence-starter")):
    keep(P(d,"FROM:  %s"%frm,size=10.5,after=3))
    keep(P(d,"TO:    %s"%to,size=10.5,bold=True,after=6))
keep(P(d,"One type size came down. The headline measures 11.47in at the "
       "existing 46pt in a 10.42in box, so it goes to 41pt and holds one line "
       "at 10.21in, ending 0.21in clear of the gold rule; two lines at 46pt "
       "would run 1.53in and cross it. The sub-copy (8.66in) and the URL "
       "(10.07in) both fit at their existing sizes and were not resized. No "
       "product mockup was added to the deck.",size=10.5,color=DIM,
       before=4,after=8,spacing=1.25))
p=P(d,"CTA VISUAL. The preferred proof visual is the REAL Career Evidence "
     "Starter artifact: Starter cover in front, Portable Proof Line page "
     "visible behind, warm cream background. Do not use generic AI graphics, "
     "an unrelated portrait, a decorative quote card or a fake worksheet, and "
     "never show the direct PDF URL. The artifact itself is the proof.",
     size=11,bold=True,color=RED,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"CTA and watch next",before=14)
keep(P(d,"One product CTA only: %s — %s"%(CTA,CTA_URL),after=5))
keep(P(d,"Public descriptor: %s. Promise: %s Outcome: one portable Proof "
       "Line. Expected time: about 10 to 15 focused minutes."
       %(CTA_DESCRIPTOR,CTA_PROMISE),size=10.5,color=DIM,after=6,spacing=1.25))
keep(P(d,"Use the landing-page URL publicly. Never expose the direct PDF URL.",
       bold=True,color=RED,after=6))
keep(P(d,"Do not add Keep the Proof, the Capability Formation Field Kit or the "
       "Career Decision Evidence Check.",bold=True,after=8))
keep(P(d,"Watch next: %s"%NEXT,bold=True,after=5))
for x in ["use the direct Video 5 end-screen route when Video 5 is public;",
 "before Video 5 is public, use the Career Portability playlist if the "
 "playlist element is functioning;",
 "otherwise use a suitable currently public video or YouTube's "
 "best-for-viewer option temporarily."]:
    keep(P(d,"—  "+x,after=4))
keep(P(d,"Do not leave Subscribe as the only end-screen element.",bold=True,
       color=RED,before=4,after=8))
compress(d)
d.save(os.path.join(LF,"Video_4_EDITOR_ONLY_HIT_Brief_v2.1.docx"))
print("editor brief written")

# The eleven working chapter lines, defined once so the copy-ready description
# and the reference section cannot drift apart.
CHAPTERS=[("00:00","When a Career Looks Disconnected"),
 ("01:13","Chronology vs. Portability"),
 ("01:57","Name the Chapters Briefly"),
 ("02:56","Find the Repeated Work"),
 ("03:22","Look Beneath the Job Titles"),
 ("04:30","Explain Why the Direction Follows"),
 ("04:48","The Three-Sentence Career Explanation"),
 ("05:46","Do Not Invent a Perfect Plan"),
 ("06:32","Test Your Career Explanation"),
 ("07:54","Free Career Evidence Starter"),
 ("08:09","Should You Make an Internal Move?")]
CHAPTER_LINES=["%s %s"%(t,c) for t,c in CHAPTERS]

EMOJI_NOTE=("The restrained emoji system is part of the approved standard: "
  "🧭 marks the CTA/resource section. Do not remove it and do not add more.")

def description_block(d, heading_before=14, upload_doc=False):
    """The copy-ready description, its end marker and the internal notes.

    Shared verbatim by the publishing package and the separate description-only
    document so their public copy cannot drift apart.
    """
    if upload_doc:
        H1(d,"INTERNAL NOTE — DO NOT PASTE INTO YOUTUBE",before=heading_before)
        p=P(d,EMOJI_NOTE,size=10.5,italic=True,color=RED,after=12,spacing=1.25)
        shade(p,BAND_CREAM); keep(p)
        p=keep(P(d,"COPY-READY YOUTUBE DESCRIPTION — BEGIN",size=11,bold=True,
                 color=NAVY,before=14,after=12,spacing=1.2))
        shade(p,BAND_NAVY)
    for para in DESC:
        keep(P(d,para if para else " ",after=7 if para else 3))
    keep(P(d,"— END OF THE COPY-READY DESCRIPTION —",size=10,bold=True,
           color=DIM,before=14,after=12,spacing=1.2))
    H1(d,"Internal note — do not paste into YouTube",before=14)
    p=P(d,"WORKING ESTIMATES — EDITOR MUST REPLACE FROM FINAL CUT",size=11,
        bold=True,color=RED,after=6,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    p=P(d,"These timestamps were estimated from the patched script, not "
        "measured from the finished edit. Replace every timestamp using the "
        "finished cut before publication. Do not force the edit to match "
        "these estimates.",size=10.5,bold=True,italic=True,color=RED,
        after=10,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    H1(d,"Working chapters — reference copy",before=14)
    keep(P(d,"Identical to the eleven chapter lines inside the description "
           "above.",size=10.5,italic=True,color=DIM,after=8))
    for line in CHAPTER_LINES: keep(P(d,line,size=11,after=4))

TAGS=("how to explain your career change, how to explain a nonlinear "
 "career, nonlinear career path, career pivot, career story, transferable "
 "skills, career portability, career transition, career change, career "
 "change advice, how to explain your experience, transferable experience, "
 "experienced professionals, Temidayo Afonja, Capability Formation")

PINNED=["What part of your career looks disconnected from the outside—but "
 "taught you something you still use?",
 "Try completing this sentence:",
 "“Across my career, I keep being asked to…”",
 "Then ask yourself: what evidence sits behind that verb?",
 "If you want to try this on one accomplishment, the Career Evidence Starter "
 "is free.",
 "It takes about 10 to 15 focused minutes and helps you turn one piece of "
 "work into a portable Proof Line.",
 CTA_URL]



# ----------------------------------------------------- 4. publishing package
d=newdoc()
head(d,TITLE,"Video 4  ·  Publishing package",
     "Everything needed to upload. Working timestamps must be replaced with "
     "real ones from the finished edit.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14)
P(d,"how to explain your career change",after=10)
H1(d,"Supporting search language",before=14)
P(d,"how to explain a nonlinear career · nonlinear career path · career "
    "pivot · career story · transferable skills · career transition · career "
    "portability",after=10)

H1(d,"Description",before=14)
DESC=[
 "Does your career path look disconnected on paper?",
 "In this video, I show you how to explain your career change without "
 "pretending every move was part of a perfect plan.",
 "One of my senior-manager friends at EY used to joke that I was a “cat with "
 "nine lives” because my career kept moving into work that looked unrelated. "
 "Over time, I realized the problem was not necessarily the career. It was "
 "how I was explaining what had traveled between the chapters.",
 "I use a simple three-part method:",
 "1. Name the major career chapters briefly.",
 "2. Find the repeated work beneath the titles.",
 "3. Explain why your next direction follows from what you have already built.",
 "A chronology tells people where you have been. A portability explanation "
 "tells them what traveled with you.",
 "Not everything transfers. Different roles, functions and industries can "
 "require real relearning. The goal is not to invent a perfect career story. "
 "It is to make the continuity you can actually support easier to hear.","",
 "🧭 FREE CAREER EVIDENCE STARTER",
 "Turn one accomplishment into a portable Proof Line you can use in a "
 "performance review, interview, internal move or career pivot:", CTA_URL,"",
 "CHAPTERS"]+CHAPTER_LINES+["",
 "WATCH NEXT", NEXT, "[ADD VIDEO 5 LINK WHEN LIVE]","",
 "PLAYLIST","Career Portability: Career Pivots, Internal Moves & Growth",
 "[ADD PLAYLIST LINK]","",
 "Temidayo Afonja helps experienced professionals understand what they can "
 "carry across roles, functions, employers and industries so they can make "
 "career pivots and internal moves without starting from zero."]
description_block(d, upload_doc=True)

H1(d,"Pinned comment",before=14)
for para in PINNED:
    keep(P(d,para,after=6))

H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only. Do not place the full tag field in "
       "the public description.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))

H1(d,"Fact boundary for the description",before=14)
keep(P(d,"The recurring “cat with nine lives” description is established. The "
       "original conversation that first prompted it is not established and "
       "is not described anywhere in this package.",bold=True,after=8,
       spacing=1.25))
compress(d, 1.14, 0.56)
d.save(os.path.join(LF,"Video_4_Publishing_Package_HIT_v2.1.docx"))

# ------------------------------- 4b. separate description-only document
# Outside the 13-file ZIP. Same architecture approved for Videos 1 and 6 to 8.
d=newdoc()
head(d,TITLE,"Video 4  ·  YouTube description",
     "Upload copy only. Everything below the end marker is internal and must "
     "not be pasted into YouTube.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14)
P(d,"how to explain your career change",after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"Watch next",before=14); keep(P(d,NEXT,bold=True,after=8))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
compress(d)
DESC_DOC="/tmp/v4p/Video_4_YouTube_Description_HIT.docx"
d.save(DESC_DOC)
print("publishing package and description-only document written")

# ---------------------------------------------------------------- 5. Shorts
SHORTS=[
 ("Video_4_Short_1_Cat_With_Nine_Lives.docx","SHORT 1","Recognition / story",
  "A senior colleague once called me a cat with nine lives.",
  ["A senior colleague once called me a cat with nine lives.",
   "She meant my career kept moving into work that looked unrelated.",
   "And for a while, the way I explained it made it sound even more "
   "disconnected.",
   "I would list the jobs.","The functions.","The industries.",
   "But I was leaving the listener to figure out what connected them.",
   "That is what changed for me.",
   "The career was not necessarily the problem.",
   "The explanation was.",
   "A nonlinear career can make sense without pretending it was perfectly "
   "planned.",
   "You have to make what traveled between the chapters easier to hear."]),
 ("Video_4_Short_2_Chronology_Not_Explanation.docx","SHORT 2",
  "Distinction / myth",
  "A chronology is not the same as a career explanation.",
  ["A chronology is not the same as a career explanation.",
   "Chronology tells me where you have been.",
   "A useful career explanation tells me what traveled with you.",
   "If you only list titles, companies and industries, I know the sequence.",
   "I still may not know what you became able to do.",
   "So after the chronology, ask:",
   "What kind of work kept finding me?",
   "What judgment did people repeatedly trust me to use?",
   "And what remained useful when the context changed?",
   "That is often where a nonlinear career starts to make sense."]),
 ("Video_4_Short_3_Not_A_Perfect_Plan.docx","SHORT 3","Proof / honesty",
  "My career was not a carefully designed portfolio career.",
  ["My career was not a carefully designed portfolio career.",
   "I graduated with an accounting degree in December 2008, during the "
   "financial crisis.",
   "So the first turn was already shaped by a market I did not control.",
   "That matters because career coherence is not the same as pretending every "
   "move was strategic.",
   "Some moves happen because of markets, caregiving, health, compensation, "
   "restructuring or an unexpected opportunity.",
   "You do not need to rewrite history.",
   "Name the context honestly.",
   "Then ask what the chapter taught you to do that remained useful later.",
   "Context can explain a transition without becoming an apology."]),
 ("Video_4_Short_4_Three_Sentences.docx","SHORT 4",
  "Practical test / action",
  "Use these three sentences to explain a career change.",
  ["Use these three sentences to explain a career change.",
   "“My career has moved across…”",
   "“Across those chapters, I kept being asked to…”",
   "“That is why I am now focused on…”",
   "The first sentence gives the chronology.",
   "The second makes the repeated work visible.",
   "The third explains the direction.",
   "But there is one rule:",
   "Do not put a verb in the second sentence unless you have evidence behind "
   "it.",
   "If you say you translate complexity, build alignment or solve ambiguous "
   "problems, be ready to show where your work actually required you to do "
   "that.",
   "That is what makes the explanation credible."])]

for fn,label,role,hook,copy in SHORTS:
    d=newdoc(True)
    P(d,"VIDEO 4 SHORT",size=10,bold=True,color=GOLD,after=4,caps=True)
    P(d,label,size=20,bold=True,color=NAVY,after=8,spacing=1.1)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    keep(P(d,"Verbal hook:  “%s”"%hook,size=11,color=DIM,after=5))
    keep(P(d,"Related long-form:  %s"%TITLE,size=11,color=DIM,after=10))
    H1(d,"RECORDING COPY",before=12)
    for line in copy:
        keep(P(d,line,size=13.5,color=INK,after=10,spacing=1.5))
    d.save(os.path.join(SH,fn))
print("publishing package and %d Shorts written"%len(SHORTS))

# ------------------------------------------------------ 6. Shorts editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 4 — FOUR STANDALONE SHORTS",size=18,bold=True,color=NAVY,
  after=8,spacing=1.1)
p=P(d,"This document is for the editor. It is separate from the four Short "
     "recording documents and must not be placed on Temidayo's recording "
     "screen.",size=11,italic=True,color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"How these are produced",before=14)
keep(P(d,"These are separately recorded 9:16 Shorts. They are NOT excerpts cut "
       "from the long-form video.",bold=True,after=10))
P(d,"Each Short needs:",after=6)
pairlist(d,["an immediate verbal hook;","a corresponding on-screen hook;",
 "meaningful visual interest;","accurate mobile-safe captions;",
 "restrained editorial pacing;",
 "Video 4 added as the Related Video when available."])

def short(label,role,onscreen,body):
    H1(d,label,before=14)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    p=keep(P(d,"On-screen hook:  %s"%onscreen,size=11,bold=True,color=GOLD,after=8))
    shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))
    keep(P(d,"Related Video:  Video 4",size=10.5,color=DIM,before=4,after=6))

short("SHORT 1","Recognition / story","MY CAREER LOOKED DISCONNECTED",
 ["Begin Temidayo direct to camera.",
  "When she says the career looked unrelated, briefly show:",
  "     ACCOUNTING & AUDIT  →  CYBERSECURITY  →  PEOPLE STRATEGY  →  "
  "ENTERPRISE TRANSFORMATION",
  "Do NOT use cat imagery or sounds.",
  "End visually on:  THE CAREER WASN’T THE PROBLEM"])
short("SHORT 2","Distinction / myth","CHRONOLOGY ≠ EXPLANATION",
 ["Use a restrained two-column treatment:",
  "     CHRONOLOGY — WHERE YOU HAVE BEEN",
  "     PORTABILITY — WHAT TRAVELED WITH YOU",
  "End on:  WHAT BECAME USEFUL AGAIN?"])
short("SHORT 3","Proof / honesty","IT WASN’T A PERFECT PLAN",
 ["Keep Temidayo visible for the 2008 financial-crisis context.",
  "Do not use dramatic stock-market crash footage. A simple DECEMBER 2008  ·  "
  "FINANCIAL CRISIS is enough if visual context is needed.",
  "End on:  CONTEXT ≠ APOLOGY"])
short("SHORT 4","Practical test / action","3 SENTENCES FOR A CAREER CHANGE",
 ["Reveal progressively:",
  "     MY CAREER HAS MOVED ACROSS…",
  "     ACROSS THOSE CHAPTERS, I KEPT BEING ASKED TO…",
  "     THAT IS WHY I AM NOW FOCUSED ON…",
  "Then end on:  EVERY VERB NEEDS EVIDENCE"])

H1(d,"All Shorts — visual boundaries",before=14)
P(d,"Do not use:",after=5)
pairlist(d,["cat imagery;","cat sounds;","employer logos;",
 "generic office B-roll;","résumé icons;","career ladders;",
 "flashy transitions;","constant zooms;","fake shock expressions;",
 "red warning graphics;","AI-generated scenery;",
 "social-media template effects."],after=3)
compress(d, 1.18, 0.62)
d.save(os.path.join(SH,"Video_4_Shorts_EDITOR_ONLY_HIT_Brief.docx"))

# ---------------------------------------------------------------- 7. README
FILES=(["LONG_FORM/"+f for f in sorted(os.listdir(LF))]
      +["SHORTS/"+f for f in sorted(os.listdir(SH))])
R=["VIDEO 4 — H.I.T. FINAL RECORDING PACKAGE","",
 "Title:             %s"%TITLE,
 "Thumbnail:         %s"%THUMB,
 "CTA:               %s"%CTA,
 "CTA URL:           %s"%CTA_URL,
 "                   The earlier %s CTA is SUPERSEDED for this"%OLD_CTA,
 "                   video by the free Career Evidence Starter. Slide 10 and",
 "                   reveal frame 25 carry the new CTA; every other slide and",
 "                   frame is byte-identical.",
 "Watch next:        %s"%NEXT,"",
 "Long-form:         Revised under H.I.T.",
 "Slides:            UNCHANGED. 11 main slides.",
 "Reveal deck:       UNCHANGED.",
 "Thumbnail:         UNCHANGED.",
 "Shorts:            Four separately recorded vertical scripts.",
 "Editor directions: Separated from recording copy.","",
 "Fact boundary:     The recurring “cat with nine lives” description is",
 "                   established. The original first conversation is not",
 "                   established and is not invented.","",
 "-"*70,"","WHAT EACH FILE IS","",
 "LONG_FORM/","",
 "  Video4TeleprompterScriptwithslidemarkers_HIT_v2.1.docx",
 "  Video4TeleprompterScriptwithslidemarkers_HIT_v2.1.txt",
 "      Temidayo's recording copy. Spoken script in large text; slide markers",
 "      in tinted bands. The markers are not spoken.","",
 "  Video4ReadingScriptnomarkers_HIT_v2.1.docx",
 "  Video4ReadingScriptnomarkers_HIT_v2.1.txt",
 "      The same spoken words with the slide markers removed.","",
 "  Video_4_EDITOR_ONLY_HIT_Brief_v2.1.docx",
 "      For the editor. The H.I.T. first-30-second plan, editorial rhythm",
 "      after 0:30, the existing 11-slide map, the let-the-visual-carry",
 "      principle, and the fact and evidence boundaries.",
 "      Not for the teleprompter.","",
 "  Video_4_Publishing_Package_HIT_v2.1.docx",
 "      Title, thumbnail, search language, the copy-ready description,",
 "      working chapter estimates, pinned comment and the tag field.","",
 "SHORTS/","",
 "  Four recording documents, one per Short. These contain Temidayo's",
 "  recording copy and no editor directions.","",
 "  Video_4_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "      For the editor. On-screen hooks and visual treatment for all four.","",
 "-"*70,"","ALL FILES IN THIS PACKAGE","",]
for f in FILES: R.append("  "+f)
R+=["  README_FINAL.txt","  SHA256SUMS.txt","",
 "-"*70,"","WORKING CHAPTER TIMESTAMPS","",
 "The chapter timestamps in the publishing package are WORKING ESTIMATES",
 "derived from the script. They were not measured from an edit. The editor",
 "must replace every one of them from the finished cut before publishing.","",
 "-"*70,"","CHECKSUMS","",
 "SHA256SUMS.txt covers the other 12 user-facing files in this package. It",
 "does not hash itself, and it carries no ZIP checksum. The archive's own",
 "SHA-256 is in the sibling file:",
 "  Video_4_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256","",
 "-"*70,"","WHAT WAS NOT CHANGED","",
 "The existing Video 4 PowerPoint deck (11 slides), the reveal-build deck,",
 "the approved thumbnail, every website file, every product and every other",
 "video are unchanged. This revision is spoken script, editor instruction and",
 "publishing copy only.","",
 "The 11-slide deck remains authoritative. The teleprompter's 11 slide",
 "markers map to it in order, slide 1 to slide 11.",""]
open(os.path.join(ROOT,"README_FINAL.txt"),"w").write("\n".join(R))
print("shorts editor brief and README written")

# ------------------------------------------- 8. checksums and the master ZIP
MANIFEST=[
 "LONG_FORM/Video4TeleprompterScriptwithslidemarkers_HIT_v2.1.docx",
 "LONG_FORM/Video4TeleprompterScriptwithslidemarkers_HIT_v2.1.txt",
 "LONG_FORM/Video4ReadingScriptnomarkers_HIT_v2.1.docx",
 "LONG_FORM/Video4ReadingScriptnomarkers_HIT_v2.1.txt",
 "LONG_FORM/Video_4_EDITOR_ONLY_HIT_Brief_v2.1.docx",
 "LONG_FORM/Video_4_Publishing_Package_HIT_v2.1.docx",
 "SHORTS/Video_4_Short_1_Cat_With_Nine_Lives.docx",
 "SHORTS/Video_4_Short_2_Chronology_Not_Explanation.docx",
 "SHORTS/Video_4_Short_3_Not_A_Perfect_Plan.docx",
 "SHORTS/Video_4_Short_4_Three_Sentences.docx",
 "SHORTS/Video_4_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "README_FINAL.txt",
]
SUMS="SHA256SUMS.txt"
ZIP="/tmp/v4p/Video_4_HIT_FINAL_Recording_and_Shorts_Package.zip"

def sha256(p):
    h=hashlib.sha256()
    with open(p,"rb") as fh:
        for b in iter(lambda: fh.read(1<<20), b""): h.update(b)
    return h.hexdigest()

for m in MANIFEST:
    assert os.path.isfile(os.path.join(ROOT,m)), "missing from build: "+m
on_disk=set()
for dp,dn,fn in os.walk(ROOT):
    dn[:]=[x for x in dn if x!="__pycache__"]
    for f in fn:
        if f.endswith(".pyc"): continue
        on_disk.add(os.path.relpath(os.path.join(dp,f),ROOT).replace(os.sep,"/"))
unexpected=sorted(on_disk-set(MANIFEST)-{SUMS})
assert not unexpected, "unexpected files in package directory: %r"%unexpected

L=["# VIDEO 4 - H.I.T. FINAL RECORDING PACKAGE",
   "# SHA-256 of the 12 user-facing files in this package.",
   "# SHA256SUMS.txt cannot hash itself. The master ZIP cannot contain its own",
   "# checksum either; it is published in the sibling file",
   "# Video_4_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256",""]
for m in MANIFEST: L.append("%s  %s"%(sha256(os.path.join(ROOT,m)),m))
open(os.path.join(ROOT,SUMS),"w").write("\n".join(L)+"\n")

if os.path.exists(ZIP): os.remove(ZIP)
with zipfile.ZipFile(ZIP,"w",zipfile.ZIP_DEFLATED) as z:
    for m in MANIFEST+[SUMS]:
        z.write(os.path.join(ROOT,m), "Video_4_HIT_FINAL/"+m)
zsha=sha256(ZIP)
open(ZIP+".sha256","w").write("%s  %s\n"%(zsha,os.path.basename(ZIP)))

PROV="/tmp/v4p/_source"
shutil.rmtree(PROV,ignore_errors=True); os.makedirs(PROV)
for f in ("script_text.py","build.py","qa.py"):
    src="/tmp/v4p/"+f
    if os.path.isfile(src): shutil.copy2(src, os.path.join(PROV,f))
print("ZIP sha256:",zsha)
print("description-only doc sha256:",sha256(DESC_DOC))
