# -*- coding: utf-8 -*-
"""Build the Video 2 H.I.T. final recording and Shorts package."""
import os, sys, shutil, zipfile, hashlib
sys.path.insert(0, "/tmp/v2hit")
from script_text import LINES, SPOKEN, MARKERS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

ROOT="/tmp/v2hit/Video_2_HIT_FINAL"
LF=os.path.join(ROOT,"LONG_FORM"); SH=os.path.join(ROOT,"SHORTS")
shutil.rmtree(ROOT, ignore_errors=True)
os.makedirs(LF); os.makedirs(SH)

TITLE="Is Your Job Making You Less Marketable?"
THUMB="YOUR SKILLS ARE STALLING"
CTA="Capability Formation Field Kit"
CTA_URL="https://temidayoafonja.com/fieldkit"
NEXT="3 Things to Do Before Quitting Your Job"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO 2",size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t): return keep(P(d,t,size=14,bold=True,color=NAVY,before=20,after=8),True)
def H2(d,t): return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=13,after=5),True)

# ------------------------------------------------- 1. teleprompter DOCX + TXT
d=newdoc(True)
head(d,TITLE,"Video 2  ·  Teleprompter script with slide markers",
     "Spoken script is the large text. A slide marker in a tinted band tells "
     "the editor which slide to bring up; it is not spoken.")
for line in LINES:
    if line.startswith("[SLIDE:"):
        name=line[len("[SLIDE:"):-1].strip()
        p=P(d,"SLIDE  —  %s"%name,size=11,bold=True,color=NAVY,
            before=14,after=14,spacing=1.1)
        shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
    else:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video2TeleprompterScriptwithslidemarkers_HIT_v2.0.docx"))

tel_txt=[TITLE,"Video 2  ·  Teleprompter script with slide markers","",]
for line in LINES:
    if line.startswith("[SLIDE:"):
        tel_txt += ["", "SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(), ""]
    else:
        tel_txt += [line, ""]
open(os.path.join(LF,"Video2TeleprompterScriptwithslidemarkers_HIT_v2.0.txt"),
     "w").write("\n".join(tel_txt).strip()+"\n")

# ------------------------------------------------- 2. reading script DOCX+TXT
d=newdoc(True)
head(d,TITLE,"Video 2  ·  Reading script, no markers",
     "Spoken language only. No slide markers, no timestamps, no production "
     "directions.")
for line in SPOKEN:
    keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video2ReadingScriptnomarkers_HIT_v2.0.docx"))
open(os.path.join(LF,"Video2ReadingScriptnomarkers_HIT_v2.0.txt"),
     "w").write("\n\n".join(SPOKEN)+"\n")
print("long-form scripts written")

# --------------------------------------------------- 3. long-form editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RGBColor(0x9B,0x2C,0x10),after=2)
P(d,"VIDEO 2",size=12,bold=True,color=GOLD,after=2,caps=True)
P(d,TITLE,size=20,bold=True,color=NAVY,after=6,spacing=1.1)
p=P(d,"This document is for the editor. It is NOT Temidayo's teleprompter and "
     "must not be placed on the recording screen.",size=11,italic=True,
     color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"Locked metadata")
for k,v in (("Title",TITLE),("Thumbnail",THUMB),("Primary CTA",CTA),
            ("CTA URL",CTA_URL),("Watch next",NEXT),
            ("Core distinction",
             "VALUABLE HERE is not automatically the same as LEGIBLE ELSEWHERE.")):
    keep(P(d,"%-18s %s"%(k+":",v),size=11,after=5))

H1(d,"First 30 seconds — H.I.T.")
P(d,"H = Hook.  I = Interest.  T = Trust. The opening must work as an "
    "audiovisual unit: immediate conversational hook, meaningful visual "
    "interest, a concrete reason to trust Temidayo, and a clear viewer payoff "
    "by roughly 20 to 30 seconds. Do not force a statistic into the opening.",
  after=12)
P(d,"This supersedes the older Video 2 opening and any earlier instruction "
    "that kept Temidayo visually static or full-screen for approximately the "
    "first 55 seconds.", italic=True, color=DIM, after=14)

def beat(t,anchor,layer,body):
    H2(d,t)
    p=P(d,"Spoken anchor:  “%s”"%anchor,size=10.5,italic=True,color=DIM,after=8)
    shade(p,BAND_CREAM)
    if layer: keep(P(d,layer,size=11,bold=True,color=GOLD,after=6))
    for b in body: keep(P(d,b,after=5))

beat("0:00–0:06",
 "You can be doing very well at work and still be narrowing your next options.",
 "H = HOOK",
 ["Visual: begin on Temidayo, medium/tight, direct to camera.",
  "On-screen text: DOING WELL. FEWER OPTIONS?",
  "Do not use a title card before this."])
beat("0:06–0:14",
 "My own career has crossed very different functions and industries…",
 "T = TRUST   ·   I = INTEREST",
 ["Show a restrained visual proof of the career path. Short labels only:",
  "     ACCOUNTING & AUDIT  →  CYBER / PRIVACY  →  PEOPLE / EMPLOYEE "
  "EXPERIENCE  →  TRANSFORMATION",
  "IMPORTANT: Temidayo deliberately does NOT recite these chapters aloud. The "
  "visual carries the chronology; her spoken words carry the meaning.",
  "Do not change the PowerPoint deck to accomplish this. Use a simple "
  "editorial overlay or approved visual treatment."])
beat("0:14–0:23",
 "Being valuable here is not the same as being legible somewhere else.",
 None,
 ["Bring forward:  VALUABLE HERE  vs.  LEGIBLE ELSEWHERE",
  "This should visually establish the central idea of the video."])
beat("0:23–0:30",
 "Here are three tests…", None,
 ["Return cleanly to Temidayo.",
  "Optional small text: 3 MARKETABILITY TESTS",
  "The viewer should understand the payoff before the teaching begins."])

H1(d,"Editorial rhythm after 0:30")
P(d,"The opening is intentionally more visually active than the remainder of "
    "the video. Do NOT read H.I.T. as permission for constant visual movement.",
  after=10)
P(d,"Once the viewer is inside the teaching:",after=6)
for x in ["preserve Temidayo's natural pacing;",
          "use the existing slides as teaching support;",
          "avoid decorative motion;","avoid constant punch-ins;",
          "avoid stock office footage;","avoid generic résumé B-roll;",
          "avoid red warning graphics;","avoid fake urgency;",
          "avoid flashy transitions."]:
    keep(P(d,"—  "+x,after=4))
P(d,"The edit should feel premium, calm and editorial.",before=8,after=10)

H1(d,"Existing slides")
P(d,"The existing 13-slide Video 2 deck remains UNCHANGED. Do not add or "
    "redesign slides. Use each slide for its existing job:",after=8)
for n,job in enumerate(["Title","Valuable Here / Legible Elsewhere",
 "01 Remove the Company Nouns","Test One","02 Find Outside-Context Evidence",
 "Test Two","03 Read the Last 90 Days","Test Three","The Three Tests",
 "Read the Pattern","Before You Leave","Capability Formation Field Kit",
 "Watch Next"],1):
    keep(P(d,"Slide %-3d %s"%(n,job),size=10.5,after=3))

H1(d,"Overlay principle")
P(d,"When a visual can prove information more efficiently than Temidayo "
    "speaking it, let the visual carry it. Do not duplicate every spoken "
    "sentence on screen. The first-30-second career-path visual is a "
    "deliberate example of this principle.",after=10)

H1(d,"CTA and watch next")
keep(P(d,"One product CTA only: %s — %s"%(CTA,CTA_URL),after=5))
keep(P(d,"Do not add Keep the Proof.",bold=True,after=8))
keep(P(d,"Watch next: route to %s."%NEXT,after=6))
d.save(os.path.join(LF,"Video_2_EDITOR_ONLY_HIT_Brief_v2.0.docx"))
print("editor brief written")

# ----------------------------------------------------- 4. publishing package
d=newdoc()
head(d,TITLE,"Video 2  ·  Publishing package",
     "Everything needed to upload. Working timestamps must be replaced with "
     "real ones from the finished edit.")
H1(d,"Title"); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail"); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase"); P(d,"is your job making you less marketable",after=10)
H1(d,"Supporting search language")
P(d,"career stagnation · career marketability · transferable skills · career "
    "growth · transferable experience · internal mobility · career "
    "portability",after=10)

H1(d,"Description")
for para in [
 "Can you be successful in your current job and still become less marketable?",
 "In this video, I give you three practical tests for separating the value you "
 "have inside your current organization from the experience and judgment "
 "another employer, function or industry can recognize and use.",
 "You will learn how to remove company-specific language from the way you "
 "explain your work, find evidence that your judgment remains useful outside "
 "its original context, and read what the last 90 days have actually added to "
 "your capability.",
 "Being valuable where you are matters. But it answers a different question "
 "from whether your value can travel.","",
 "CAPABILITY FORMATION FIELD KIT",
 "A private, evidence-led assessment of what your current work is building, "
 "what appears portable and what needs attention next.",
 CTA_URL,"",
 "WATCH NEXT", NEXT, "[ADD VIDEO LINK WHEN LIVE]","",
 "Temidayo Afonja helps experienced professionals understand what they can "
 "carry across roles, functions, employers and industries so they can make "
 "career pivots and internal moves without starting from zero."]:
    keep(P(d,para if para else " ",after=7 if para else 3))

H1(d,"Working chapters")
p=P(d,"WORKING TIMESTAMPS. The editor must replace every one of these with the "
    "actual timestamps from the finished edit before publishing.",size=10.5,
    bold=True,italic=True,color=RGBColor(0x9B,0x2C,0x10),after=10)
shade(p,BAND_CREAM)
for t,c in [("00:00","Doing well can still narrow your options"),
 ("00:30","Valuable here vs. legible elsewhere"),
 ("01:30","Test 1: Remove the company nouns"),
 ("03:00","Test 2: Find outside-context evidence"),
 ("04:30","Test 3: Read the last 90 days"),
 ("05:45","Read the pattern"),
 ("06:50","What can you change before leaving?"),
 ("07:55","Capability Formation Field Kit"),
 ("08:25","Before you quit")]:
    keep(P(d,"%s   %s"%(t,c),size=11,after=4))

H1(d,"Pinned comment")
for para in ["Which of the three tests is hardest for you right now?",
 "1. Remove the company nouns","2. Find outside-context evidence",
 "3. Read the last 90 days",
 "And if you removed your employer, systems and internal language from the way "
 "you describe your work, what would still be left?"]:
    keep(P(d,para,after=6))
d.save(os.path.join(LF,"Video_2_Publishing_Package_HIT_v2.0.docx"))

# ---------------------------------------------------------------- 5. Shorts
SHORTS=[
 ("Video_2_Short_1_Doing_Well_Fewer_Options.docx","SHORT 1","Recognition",
  "You can be doing well at work and still be narrowing your options.",
  ["You can be doing well at work and still be narrowing your options.",
   "There are two questions I want experienced professionals to separate.",
   "Am I valuable here?",
   "And can someone somewhere else understand and use what I know how to do?",
   "Those are not always the same.",
   "You can be the person everyone relies on because you know the company, the "
   "relationships, the history and the workarounds.",
   "That is real value.",
   "But if most of it only makes sense inside that environment, your "
   "marketability may not be growing at the same rate.",
   "In the full video, I give you three tests to check the difference."]),
 ("Video_2_Short_2_Indispensable_Not_Marketable.docx","SHORT 2",
  "Distinction / myth",
  "Being indispensable is not the same as being marketable.",
  ["Being indispensable is not the same as being marketable.",
   "Sometimes an organization depends on you because the undocumented history "
   "sits in your head.",
   "Certain relationships run through you.",
   "Or you are the person who knows how to rescue a fragile process.",
   "That can make you extremely valuable there.",
   "But now remove the company.",
   "Could another organization still recognize the judgment you bring and the "
   "problem you know how to solve?",
   "If the answer is hard to explain, that does not mean your experience has "
   "no value.",
   "It means you may need to separate what belongs to you from what belongs to "
   "the context."]),
 ("Video_2_Short_3_Not_Everything_Travels.docx","SHORT 3",
  "Proof / personal evidence",
  "Changing contexts taught me that not everything travels.",
  ["Changing contexts taught me that not everything travels.",
   "My own career has crossed functions and industries that looked very "
   "different on paper.",
   "And the useful question was never, “Can I take everything with me?”",
   "I could not.",
   "Some knowledge belongs to the company, industry or moment where you "
   "learned it.",
   "The better question is:",
   "What remains useful when the context changes?",
   "Look for one example where your judgment was useful beyond your immediate "
   "team, role or employer.",
   "That is a much stronger clue to portability than simply being very good at "
   "one company’s way of working."]),
 ("Video_2_Short_4_Remove_Company_Nouns.docx","SHORT 4",
  "Practical test / action",
  "Try this on one sentence from your résumé.",
  ["Try this on one sentence from your résumé or LinkedIn profile.",
   "Remove the employer.","Remove the internal program.",
   "Remove the system name, product name and acronym.",
   "Then read what is left.",
   "If your sentence collapses, that is useful information.",
   "For example, “I own the QBR process for this business unit” may "
   "mean a lot inside one company.",
   "Outside it, I still do not know what judgment you bring.",
   "Rewrite the sentence around the problem you solve, the decisions you help "
   "make and what becomes possible because of your work.",
   "Now your experience is easier for another context to understand."])]

for fn,label,role,hook,copy in SHORTS:
    d=newdoc(True)
    P(d,"VIDEO 2 SHORT",size=10,bold=True,color=GOLD,after=4,caps=True)
    P(d,label,size=20,bold=True,color=NAVY,after=8,spacing=1.1)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    keep(P(d,"Verbal hook:  “%s”"%hook,size=11,color=DIM,after=5))
    keep(P(d,"Related long-form:  %s"%TITLE,size=11,color=DIM,after=16))
    H1(d,"RECORDING COPY")
    for line in copy:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
    d.save(os.path.join(SH,fn))
print("publishing package and %d Shorts written"%len(SHORTS))

# ------------------------------------------------------ 6. Shorts editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RGBColor(0x9B,0x2C,0x10),after=2)
P(d,"VIDEO 2 — FOUR STANDALONE SHORTS",size=18,bold=True,color=NAVY,
  after=8,spacing=1.1)
p=P(d,"This document is for the editor. It is separate from the four Short "
     "recording documents and must not be placed on Temidayo's recording "
     "screen.",size=11,italic=True,color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"How these are produced")
keep(P(d,"These are separately recorded 9:16 Shorts. They are NOT excerpts cut "
       "from the long-form video.",bold=True,after=10))
P(d,"Each Short should have:",after=6)
for x in ["an immediate verbal hook;","a corresponding on-screen hook;",
          "meaningful visual movement;","accurate burned-in captions;",
          "restrained editorial pacing;",
          "Video 2 added as the YouTube Related Video when available."]:
    keep(P(d,"—  "+x,after=4))

def short(label,role,onscreen,body):
    H1(d,label)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    p=keep(P(d,"On-screen hook:  %s"%onscreen,size=11,bold=True,color=GOLD,after=8))
    shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))

short("SHORT 1","Recognition","DOING WELL. FEWER OPTIONS?",
 ["Begin direct to camera.",
  "Briefly introduce:  VALUABLE HERE  /  LEGIBLE ELSEWHERE",
  "No generic office B-roll.",
  "End: FULL VIDEO ON YOUTUBE, or use the YouTube Related Video route without "
  "adding a forced verbal CTA."])
short("SHORT 2","Distinction / myth","INDISPENSABLE ≠ MARKETABLE",
 ["Use a restrained contrast between YOU and THE CONTEXT.",
  "When Temidayo mentions undocumented history, relationships and fragile "
  "processes, do not illustrate each item with literal stock footage. Keep the "
  "emphasis conceptual.",
  "End with: WHAT ACTUALLY TRAVELS?"])
short("SHORT 3","Proof / personal evidence","NOT EVERYTHING TRAVELS",
 ["When Temidayo says “My own career has crossed functions and "
  "industries…”, briefly show:",
  "     ACCOUNTING & AUDIT  →  CYBER / PRIVACY  →  PEOPLE / EMPLOYEE "
  "EXPERIENCE  →  TRANSFORMATION",
  "Temidayo intentionally does not narrate the list.",
  "End visually on: WHAT REMAINS USEFUL?"])
short("SHORT 4","Practical action","REMOVE THE COMPANY NOUNS",
 ["Visually demonstrate removing: EMPLOYER · INTERNAL PROGRAM · SYSTEM NAME · "
  "PRODUCT NAME · ACRONYM",
  "Then show: “I own the QBR process for this business unit.”",
  "Transition to a clearer capability-focused expression.",
  "Do not overcrowd the screen with the entire replacement sentence at once. "
  "Use progressive text."])

H1(d,"All Shorts")
for x in ["No hyperactive zooming.","No fake shock expressions.",
          "No red warning graphics.","No generic career B-roll.",
          "No stock résumé footage.","No AI-style animated icons.",
          "Keep captions mobile-safe and accurate."]:
    keep(P(d,"—  "+x,after=4))
d.save(os.path.join(SH,"Video_2_Shorts_EDITOR_ONLY_HIT_Brief.docx"))

# ---------------------------------------------------------------- 7. README
FILES=(["LONG_FORM/"+f for f in sorted(os.listdir(LF))]
      +["SHORTS/"+f for f in sorted(os.listdir(SH))])
R=["VIDEO 2 — H.I.T. FINAL RECORDING PACKAGE","",
 "Title:        %s"%TITLE,
 "Thumbnail:    %s"%THUMB,
 "CTA:          %s"%CTA,
 "              %s"%CTA_URL,
 "Watch next:   %s"%NEXT,"",
 "Long-form:    Revised under the H.I.T. first-30-second standard.",
 "Slides:       UNCHANGED.",
 "Shorts:       Four separately recorded vertical scripts.",
 "Editor",
 "instructions: Separated from all recording copy.","",
 "-"*70,"","WHAT EACH FILE IS","",
 "LONG_FORM/","",
 "  Video2TeleprompterScriptwithslidemarkers_HIT_v2.0.docx",
 "  Video2TeleprompterScriptwithslidemarkers_HIT_v2.0.txt",
 "      Temidayo's recording copy. Spoken script in large text; slide markers",
 "      in tinted bands. The markers are not spoken.","",
 "  Video2ReadingScriptnomarkers_HIT_v2.0.docx",
 "  Video2ReadingScriptnomarkers_HIT_v2.0.txt",
 "      The same spoken words with the slide markers removed.","",
 "  Video_2_EDITOR_ONLY_HIT_Brief_v2.0.docx",
 "      For the editor. The H.I.T. first-30-second plan, editorial rhythm",
 "      after 0:30, the existing slide map and the overlay principle.",
 "      Not for the teleprompter.","",
 "  Video_2_Publishing_Package_HIT_v2.0.docx",
 "      Title, thumbnail, search language, description, working chapters and",
 "      pinned comment.","",
 "SHORTS/","",
 "  Four recording documents, one per Short. These contain Temidayo's",
 "  recording copy and no editor directions.","",
 "  Video_2_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "      For the editor. On-screen hooks and visual treatment for all four.","",
 "-"*70,"","ALL FILES IN THIS PACKAGE","",]
for f in FILES: R.append("  "+f)
R+=["  README_FINAL.txt","  SHA256SUMS.txt","",
 "-"*70,"","WHAT WAS NOT CHANGED","",
 "The existing Video 2 PowerPoint deck, the reveal deck, the thumbnail, every",
 "website file, every product and every other video are unchanged. This",
 "revision is spoken script and editor instruction only.","",
 "The 13-slide deck remains authoritative. The teleprompter's 13 slide markers",
 "map to it in order, slide 1 to slide 13.",""]
open(os.path.join(ROOT,"README_FINAL.txt"),"w").write("\n".join(R))
print("shorts editor brief and README written")
