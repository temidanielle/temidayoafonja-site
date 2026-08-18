"""Build the Video 2 teleprompter scripts from the approved un-script v1.0.

Two versions, each as .docx and .txt:
  Video-2-Teleprompter-Script-with-slide-markers_v1.0
  Video-2-Reading-Script-no-markers_v1.0

Plain continuous speech in running order. No tables, no labels, no side notes.
Same structure and formatting as the Video 1 scripts.
"""
import os, re

HERE = os.path.dirname(os.path.abspath(__file__))

SCRIPT = """
A job can make you more valuable inside one company while making you harder to understand everywhere else.

That can happen while you are being praised, well paid, and the person everyone depends on.

In this video, I am going to give you three tests for whether your current success is expanding your marketability or quietly narrowing it.

By the end, you will know what to strip away from the way you describe your work, what outside-context evidence to look for, and what the last 90 days have actually added to your judgment.

[SLIDE: Title]

I am Temidayo Afonja. On this channel, I help experienced professionals make clearer career decisions by looking at what their work is actually building in them.

[SLIDE: Valuable Here, Legible Elsewhere]

Let me start with why a good job can hide this risk.

Praise can increase while the work quietly becomes more context-bound. Compensation and responsibility can rise because you know one environment extremely well. And people can depend on you because knowledge, relationships, a proprietary system or a fragile process all run through you.

All of that raises your importance inside the building. None of it tells you whether your capability is easier for another organization to recognize.

Being valuable here answers only one question. The second question is whether another context can recognize and use what you have built.

Praise can tell you that you matter here. It cannot tell you how easily your value travels.

I want to be clear about something. This is not a reason to panic, and it is not a reason to leave. A strong role can still deserve a closer read. It is simply a reason to test what your success is actually producing.

So here are the three tests.

[SLIDE: Remove the Company Nouns]

The first test is to remove the company nouns.

Take the way you currently describe your work, and cross out the employer name. Cross out the internal programs, the proprietary systems, the product names and the acronyms.

What remains should still tell another person what problem you solve and what judgment the work requires.

Here is what I mean. Someone might say, I own the QBR process for this business unit.

Inside that company, everyone knows what that means. Outside it, almost no one does.

Now here is the same work with the company nouns removed. I combine incomplete operating data, surface the decision leaders are avoiding, and create a shared view of what needs to happen next.

That is a capability. Another organization can recognize that problem, even if they have never heard of your business unit or your process.

So take one sentence from your resume, your LinkedIn profile, or the way you introduce your work at a dinner. Remove the company nouns. If the sentence collapses, rewrite it around the problem you solve and the judgment you bring.

If another organization can recognize the problem, the capability becomes easier to see.

[SLIDE: Find Outside-Context Evidence]

The second test is to find evidence that your capability has been useful outside its original context.

I want to be precise here, because this is not about job offers, and it is not about public visibility.

The question is simpler. Has anyone beyond your immediate environment valued this capability? You are looking for evidence that the usefulness survived some distance from the conditions that originally formed it.

That evidence can be small. It only has to show that the capability remained useful when the context changed.

It might be a cross-functional project, where another group used or sought your judgment.

It might be an external client or customer, where the capability created value beyond your immediate team.

It might be a former colleague who came back to you for your judgment after the original context had changed.

Or it might be a contribution in another setting or industry, where a similar problem or a similar approach remained useful.

What does not count, on its own, is visibility. Being known is not the same as being useful somewhere else. The question is whether the judgment was actually used beyond the place that produced it.

This is also where indispensability can mislead you. An organization may depend on you because knowledge, relationships or a fragile process run through you. That increases your internal importance. It does not necessarily increase what you can do somewhere else.

So write down one example where your judgment was useful beyond your immediate role, team or employer. And if you cannot find one yet, that is not a failure. That is useful information, and it tells you what to investigate.

You are looking for proof that the usefulness survived some distance from the original context.

[SLIDE: Read the Last 90 Days]

The third test is to read the last 90 days.

Do not ask only whether you were busy, or whether you performed well. Ask what you can now do, see or decide that you could not do 90 days ago.

New judgment might come from handling a more ambiguous problem. From making a decision with incomplete information. From influencing a different group. From working across a new constraint. Or from building judgment in a place where you previously needed help.

There are four questions I use.

What unfamiliar problem did I have to solve?

What decision can I now make with less help, or with better judgment?

What constraint, audience or context did I learn to work across?

And the honest one. Is the main change new judgment, or am I mostly doing the same work faster?

Speed has value. Getting faster at familiar work is real, and it is worth something. But new judgment tells you something different, and it is the thing that travels.

[SLIDE: Recap]

So those are the three tests. Remove the company nouns. Find outside-context evidence. And read the last 90 days.

Together, they help you separate internal success from external marketability, because internal value and external marketability are two different questions.

[SLIDE: Two Questions]

Now, how do you read what you find?

Please do not diagnose yourself from one familiar sign. Look at the pattern, across two things: whether your judgment is growing, and whether another context can use it.

If your judgment is deepening and the capability is useful across contexts, the work is compounding.

If your expertise is deepening but its usefulness is tightly trapped inside one environment, you may be in what I call a depth trap.

If visibility or opportunities are growing but the work itself is not building enough depth, the position may be more fragile than it looks.

And if neither the capability nor the options are growing, the role may be stagnant, even while your performance is strong.

Hold all of that lightly. This is a reading, not a verdict.

[SLIDE: What Can I Change Before I Leave]

A concern in one test does not automatically mean the answer is to resign.

Before you change the employer, there are usually things you can change about the work.

You can ask for work that creates new judgment, not only more volume.

You can make the capability visible outside your immediate team or function.

You can document knowledge or results that currently exist mainly in your head.

Or you can take on an adjacent problem, one that forces the capability to work in a different context.

Choose one thing you can test over the next 30 days. Sometimes the first move is to change the work before changing the employer.

[SLIDE: Field Kit]

If these three tests show that you need a fuller read of what your current work is building and how portable it is, the Capability Formation Field Kit gives you a private, evidence-led assessment using the last 90 days of your actual work. It helps you read the position before you decide what comes next.

You can find it at temidayoafonja.com/fieldkit.

[SLIDE: Watch Next]

Before you decide that a concern means you should resign, there are three things I want you to check.

That is the next video. Before You Quit Your Job, Check These 3 Things. Watch that one next.
"""


def paragraphs(text):
    return [p.strip() for p in text.strip().split("\n\n") if p.strip()]


def strip_markers(paras):
    return [p for p in paras if not re.fullmatch(r"\[SLIDE:.*\]", p)]


def write_txt(path, paras):
    with open(path, "w") as f:
        f.write("\n\n".join(paras) + "\n")


def write_docx(path, title, paras):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_LINE_SPACING
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(16)
    normal.paragraph_format.space_after = Pt(14)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    head = doc.add_paragraph()
    r = head.add_run(title)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x5A, 0x6B, 0x82)
    head.paragraph_format.space_after = Pt(26)

    for p in paras:
        para = doc.add_paragraph()
        run = para.add_run(p)
        if re.fullmatch(r"\[SLIDE:.*\]", p):
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)
            para.paragraph_format.space_before = Pt(18)
    doc.save(path)


def main():
    marked = paragraphs(SCRIPT)
    clean = strip_markers(marked)
    words = len(" ".join(clean).split())

    write_txt(os.path.join(HERE, "Video-2-Teleprompter-Script-with-slide-markers_v1.0.txt"), marked)
    write_txt(os.path.join(HERE, "Video-2-Reading-Script-no-markers_v1.0.txt"), clean)
    write_docx(os.path.join(HERE, "Video-2-Teleprompter-Script-with-slide-markers_v1.0.docx"),
               "Video 2 teleprompter script, with slide markers", marked)
    write_docx(os.path.join(HERE, "Video-2-Reading-Script-no-markers_v1.0.docx"),
               "Video 2 reading script", clean)

    print("spoken paragraphs: %d | slide markers: %d | words: %d"
          % (len(clean), len(marked) - len(clean), words))
    print("at 130 wpm: %.1f min | at 145 wpm: %.1f min" % (words / 130, words / 145))
    print("em/en dashes: %d" % (SCRIPT.count("—") + SCRIPT.count("–")))


if __name__ == "__main__":
    main()
