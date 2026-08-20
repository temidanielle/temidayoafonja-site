"""Video 3 v1.1 script package. Final title: 3 Things to Do Before Quitting Your Job."""
import os, re

HERE = os.path.dirname(os.path.abspath(__file__))

SCRIPT = """
If you are seriously thinking about quitting your job, there are three things I want you to check before you go.

I am not going to try to talk you into staying.

I want to help you leave, if leaving is the right decision, with a clear record of what this work built in you and a better read on what the next move needs to do.

Because once you leave, access changes.

Records become harder to reach. Systems close. People who saw the work move on to other priorities. And something that feels obvious while you are still inside the role can become surprisingly difficult to reconstruct six months later.

By the end of this video, you will know what evidence to preserve, how to name what the work actually built in you, and how to test whether the next move uses something you have already proved while requiring you to build something genuinely new.

I’ve worked inside the systems where performance, talent decisions and employee transitions are documented, so I know how quickly the evidence behind someone’s work becomes harder to reconstruct once they are no longer inside the role.

That is why the first thing I want you to check is what evidence you should preserve.

[SLIDE: Title]

I'm Temidayo Afonja. On this channel, I help experienced professionals make clearer career decisions by looking at what their work is actually building in them.

[SLIDE: Once You Leave, Access Changes]

Before anything else, one thing needs saying clearly.

If your health or safety is at risk, or you are facing harassment, discrimination or another urgent threat, nothing in this video is a reason to delay leaving.

Please act on that first.

Everything I am about to say assumes you have the time to think. If you do not, that is a different situation and it deserves a different response.

For everyone else, here is why the timing matters.

While you are still in the role, you can look up a date. You can check what a project actually involved. You can read your own review history. You can ask someone what they remember. The day after you give notice, most of that is harder, and some of it is gone.

None of this is about hesitating. It is about not losing the record of your own work on the way out.

[SLIDE: 01 — Preserve the Evidence]

The first check is to preserve the evidence.

[SLIDE: What to Keep / What Not to Take]

I want to be careful here, because this is not permission to take things.

What you can keep is what is genuinely yours. Your own performance reviews, where you are permitted to keep them. Recognition you received. Nonconfidential metrics that were already shared with you. Project dates. The scope of what you were responsible for. And permitted notes about decisions you influenced or problems you helped resolve.

What you must not take is anything that belongs to the employer. Confidential information. Customer data. Employee data. Proprietary documents. Anything you are not entitled to retain.

The rule is simple. If you are not entitled to keep it, do not take it.

Preserving your record does not mean taking their material.

Now, the shape of the record itself. For each thing worth keeping, write down what changed. What the starting condition was. What you decided or influenced. Who was affected. And what permitted evidence supports the result.

That is a paragraph, not a filing cabinet.

And a lot of what makes experience valuable is surprisingly easy to forget once the context disappears.

[SLIDE: 02 — Name What the Work Built]

The second check is to name what the work built.

[SLIDE: Problem / Constraint / Judgment / Outcome]

A resume bullet can tell another employer what happened. It does not automatically tell them what you became able to do.

So for each achievement that matters, capture four things.

The problem. What problem was being solved?

The constraint. What made the situation difficult?

The judgment. What did you notice, decide, interpret or influence?

And the outcome. What changed, or what was prevented?

Then ask the question that makes it portable. Where else could that combination matter?

Let me make that concrete. Someone might say they reduced the time an internal process took. That is the bullet.

But the value that travels is usually not the number. It is that they could identify where the work was getting stuck. That they could align people who owned different parts of the same system. That they could redesign the handoff. And that they could do it without creating a new control failure.

That combination is the capability. Another organization can recognize it, even if they have never heard of the process you fixed.

[SLIDE: 03 — Test the Next Move]

The third check is to test the next move.

[SLIDE: Uses Something Proven / Builds Something New]

A strong move does two things at once. It uses something proven and builds something new.

Three questions will tell you whether the move in front of you does both.

What will this next role allow me to carry?

What new judgment, exposure or responsibility will it force me to develop?

And what will I be able to do after a year that I cannot do now?

Two things worth watching for.

A move that uses nothing you have already developed may create an unnecessary reset. Sometimes that is exactly the right trade. It is worth making it knowingly rather than by accident.

And a move that repeats the same work at another employer may change the setting without materially changing what the career is building. The title can improve while the position stays where it was.

[SLIDE: The Three Checks]

So those are the three checks. Preserve the evidence. Name what the work built. Test the next move.

In that order, before you resign, not after.

[SLIDE: Decision Reading]

Once you have done that, the evidence tends to point in one of three directions.

It may say that leaving is right and you are ready.

It may say that repositioning inside is worth trying first. Another role, another project, or a change in scope could restore growth without an immediate exit.

Or it may say that you need to build a bridge. That something is missing before the move: a credential, one piece of outside-context evidence, financial runway, or a clearer way of translating what your experience can do somewhere else.

I am not going to tell you which one you are looking at. That is not something I can read from here, and it is not something a single sign can tell you either.

The point is not to make the decision slow. The point is to make it legible.

And if the boundary I mentioned at the start applies to you, none of this does. A harmful situation does not need a bridge plan.

[SLIDE: Before You Resign]

So before you resign, answer three questions.

What evidence do I need to preserve now?

What does my strongest evidence show I can do?

What must the next move use, and what must it build?

Pause the video here if you want to write those down. They are worth more on paper than in your head.

[SLIDE: Career Decision Evidence Check]

If you want a structured read of the evidence behind the decision you are weighing, that is what the Career Decision Evidence Check is for.

You can find it at temidayoafonja.com/career-decisions. I have also linked it below.

[SLIDE: Watch Next]

And if the move you are considering changes your function or your industry, there is one more question underneath all of this. What actually carries across, and what does not?

That is the video I would watch next. How to Change Jobs Without Starting Your Career Over.
"""

SHORT_3A = """
Before you resign, preserve the evidence.

Once you leave, access changes. Records and the people who can confirm your work get harder to reach.

While you are still there, keep what is genuinely yours. Your own performance reviews where permitted. Recognition you received. Nonconfidential metrics already shared with you. Project dates, scope, and permitted notes on decisions you influenced.

What you must not take is anything the employer owns. Confidential information. Customer data. Employee data. Proprietary documents.

The rule is simple. If you are not entitled to keep it, do not take it.

And if your health or safety is at risk, or you are facing harassment or discrimination, none of this is a reason to wait.

The full version is in my video, 3 Things to Do Before Quitting Your Job.
"""

SHORT_3B = """
A better next move does two things.

It uses something you have already proved, and it requires you to build something genuinely new.

Ask three questions about the role in front of you.

What will it allow me to carry?

What new judgment or responsibility will it force me to develop?

And what will I be able to do after a year that I cannot do now?

Here is why both halves matter. A move that uses nothing you have already developed can create an unnecessary reset. And a move that repeats the same work somewhere else may change the setting without changing what your career is building.

If you are weighing a move right now, the full version is in my video, 3 Things to Do Before Quitting Your Job.
"""


def paragraphs(t):
    return [p.strip() for p in t.strip().split("\n\n") if p.strip()]


def strip_markers(ps):
    return [p for p in ps if not re.fullmatch(r"\[SLIDE:.*\]", p)]


def write_txt(path, ps):
    open(path, "w").write("\n\n".join(ps) + "\n")


def write_docx(path, title, ps):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_LINE_SPACING
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(1.0)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(16)
    n.paragraph_format.space_after = Pt(14)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h = doc.add_paragraph(); r = h.add_run(title)
    r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x5A, 0x6B, 0x82)
    h.paragraph_format.space_after = Pt(26)
    for p in ps:
        par = doc.add_paragraph(); run = par.add_run(p)
        if re.fullmatch(r"\[SLIDE:.*\]", p):
            run.font.size = Pt(12); run.font.bold = True
            run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)
            par.paragraph_format.space_before = Pt(18)
    doc.save(path)


def main():
    marked = paragraphs(SCRIPT)
    clean = strip_markers(marked)
    write_txt(os.path.join(HERE, "Video-3-Teleprompter-Script-with-slide-markers_v1.2.txt"), marked)
    write_txt(os.path.join(HERE, "Video-3-Reading-Script-no-markers_v1.2.txt"), clean)
    write_docx(os.path.join(HERE, "Video-3-Teleprompter-Script-with-slide-markers_v1.2.docx"),
               "3 Things to Do Before Quitting Your Job, teleprompter script with slide markers, v1.2", marked)
    write_docx(os.path.join(HERE, "Video-3-Reading-Script-no-markers_v1.2.docx"),
               "3 Things to Do Before Quitting Your Job, reading script, v1.2", clean)
    words = len(" ".join(clean).split())
    markers = [p for p in marked if re.fullmatch(r"\[SLIDE:.*\]", p)]
    print("main script: %d spoken paragraphs, %d markers, %d words"
          % (len(clean), len(markers), words))
    print("  at 130 wpm: %.1f min | at 145 wpm: %.1f min" % (words / 130, words / 145))
    spoken = "\n".join(clean)
    print("  em/en dashes in spoken copy: %d"
          % (spoken.count("—") + spoken.count("–")))
    for m in markers:
        print("   ", m)


if __name__ == "__main__":
    main()
