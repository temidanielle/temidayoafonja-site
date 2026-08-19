"""Build the Video 3 script package.

  Video-3-Unscript-Working-Sheet_v1.0_Temidayo_Afonja.docx
  Video-3-Teleprompter-Script-with-slide-markers_v1.0.docx / .txt
  Video-3-Reading-Script-no-markers_v1.0.docx / .txt
  Short-3A-Before-You-Resign-Preserve-the-Evidence_v1.0.docx / .txt
  Short-3B-A-Better-Next-Move-Does-Two-Things_v1.0.docx / .txt
"""
import os, re

HERE = os.path.dirname(os.path.abspath(__file__))

SCRIPT = """
If you are seriously thinking about resigning, this video is not going to talk you out of it.

It is going to help you do three things before you go, so the decision you make rests on evidence rather than on how the last few months have felt.

Because once you leave, access changes. Records, systems and the people who can confirm your work all become harder to reach.

By the end of this video, you will know what to preserve before that happens, how to name what your work actually built, and how to test whether the next move uses something you have already proved while requiring you to build something genuinely new.

[SLIDE: Title]

I am Temidayo Afonja. On this channel, I help experienced professionals make clearer career decisions by looking at what their work is actually building in them.

[SLIDE: Once You Leave, Access Changes]

Before anything else, I want to be clear about one thing.

If your health or safety is at risk, or you are facing harassment, discrimination or another urgent threat, nothing in this video is a reason to delay leaving. Please act on that first.

Everything I am about to say assumes you have the time to think. If you do not, that is a different situation and it deserves a different response.

For everyone else, here is why the timing matters.

While you are still in the role, you can see your own review history. You can look up dates. You can check what a project actually involved. You can ask a colleague what they remember. The day after you resign, most of that becomes harder, and some of it becomes impossible.

None of this is about hesitating. It is about not losing the record of your own work on the way out.

[SLIDE: Preserve the Evidence]

The first check is to preserve the evidence.

I want to be careful here, because this is not permission to take things.

What you can keep is what is genuinely yours. Your own performance reviews. Recognition you received. Nonconfidential metrics that were already shared with you. Project dates. The scope of what you were responsible for. And permitted notes about decisions you influenced or problems you helped resolve.

What you must not take is anything that belongs to the employer. Confidential information. Customer data. Employee data. Proprietary documents. Anything you do not have the right to keep.

Preserving your record does not mean taking their material. If you would not be comfortable explaining it to your next employer's legal team, leave it.

Now, the shape of the record itself. For each thing worth keeping, write down what changed, what the starting condition was, what you decided or influenced, who was affected, and what permitted evidence supports the result.

That is a paragraph. It is not a filing cabinet. And it is the thing you will wish you had written down when someone asks you about this work eighteen months from now.

[SLIDE: Name What the Work Built]

The second check is to name what the work built.

A resume bullet is not automatically portable evidence. It tells someone what you were assigned. It does not tell them what you can do.

So for each achievement that matters, capture four things.

The problem. What problem was being solved?

The constraint. What made the situation difficult?

The judgment. What did you notice, decide, interpret or influence?

And the outcome. What changed, or what was prevented?

Then ask the question that makes it portable. Where else could that combination matter?

Let me make that concrete. Someone might say they reduced the time an internal process took. That is the bullet.

But the portable value is usually not the number. It is that they could identify where the work was getting stuck. That they could align people who owned different parts of the same system. That they could redesign the handoff. And that they could do all of that without creating a new control failure.

That combination is the capability. Another organization can recognize that, even if they have never heard of the process you fixed.

[SLIDE: Test the Next Move]

The third check is to test the next move.

A strong move usually does two things at once. It uses something you have already proved, and it requires you to build something genuinely new.

Three questions will tell you whether the move in front of you does both.

What will this next role allow me to carry?

What new judgment, exposure or responsibility will it force me to develop?

And what will I be able to do after a year that I cannot do now?

Two cautions, from watching a lot of these decisions.

A move that uses nothing you have already developed may impose an unnecessary reset. You can absolutely choose that, but choose it knowingly.

And a move that simply repeats the same work at another employer may change the scenery without changing your position. The title improves, the commute changes, and a year later you are in the same place.

[SLIDE: Recap]

So those are the three checks. Preserve the evidence. Name what the work built. Test the next move.

Take them in that order, before you resign, not after.

[SLIDE: Three Directions]

Once you have done that, the evidence usually points in one of three directions.

It may say that leaving is right and you are ready. The evidence supports the move and the move is available.

It may say that repositioning inside is worth trying first. Another role, another project, or a change in scope could restore growth without an immediate exit.

Or it may say that you need to build a bridge. That there is something missing before the move: a credential, one piece of outside-context evidence, financial runway, or a clearer way of translating what your experience can do somewhere else.

I am not going to tell you which one you are looking at. That is not something I can read from here, and it is not something one sign can tell you either.

The point is not to make the decision slow. The point is to make it legible.

And if the safety boundary I mentioned earlier applies to you, none of this applies. A harmful situation does not need a bridge plan.

[SLIDE: Before You Resign]

So before you resign, answer three questions.

What evidence do I need to preserve now?

What does my strongest evidence show I can do?

What must the next move use, and what must it build?

Pause the video here if you want to write those down. They are worth more on paper than in your head.

[SLIDE: Career Decision Evidence Check]

If you want a structured read of the evidence behind the decision you are weighing, that is what the Career Decision Evidence Check is for.

You can find it at temidayoafonja dot com slash career dash decisions. I have also linked it below.

[SLIDE: Watch Next]

And if the move you are considering changes your function or your industry, there is one more question underneath all of this. What actually carries across, and what does not?

That is the video I would watch next. How to Change Jobs Without Starting Your Career Over.
"""

SHORT_3A = """
Before you resign, preserve the evidence.

Once you leave, access changes. Records and the people who can confirm your work get harder to reach.

While you are still there, keep what is genuinely yours. Your own performance reviews. Recognition you received. Nonconfidential metrics already shared with you. Project dates, scope of responsibility, and permitted notes on decisions you influenced.

What you must not take is anything the employer owns. Confidential information. Customer data. Employee data. Proprietary documents. Anything you have no right to keep.

Preserving your record does not mean taking their material.

And if your health or safety is at risk, or you are facing harassment or discrimination, none of this is a reason to wait.

The full version is in my video on the three things to check before you quit.
"""

SHORT_3B = """
A better next move does two things.

It uses something you have already proved, and it requires you to build something genuinely new.

Ask three questions about the role in front of you.

What will it allow me to carry?

What new judgment or responsibility will it force me to develop?

And what will I be able to do after a year that I cannot do now?

Here is why both halves matter. A move that uses nothing you have already developed can impose an unnecessary reset. And a move that repeats the same work at another employer changes the scenery without changing your position.

If you are weighing a move right now, the full version is in my video on the three things to check before you quit.
"""


def paragraphs(text):
    return [p.strip() for p in text.strip().split("\n\n") if p.strip()]


def strip_markers(paras):
    return [p for p in paras if not re.fullmatch(r"\[SLIDE:.*\]", p)]


def write_txt(path, paras):
    with open(path, "w") as f:
        f.write("\n\n".join(paras) + "\n")


def write_docx(path, title, paras):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_LINE_SPACING
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(1.0)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(16)
    n.paragraph_format.space_after = Pt(14)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x5A, 0x6B, 0x82)
    h.paragraph_format.space_after = Pt(26)
    for p in paras:
        para = doc.add_paragraph()
        run = para.add_run(p)
        if re.fullmatch(r"\[SLIDE:.*\]", p):
            run.font.size = Pt(12); run.font.bold = True
            run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)
            para.paragraph_format.space_before = Pt(18)
    doc.save(path)


def main():
    marked = paragraphs(SCRIPT)
    clean = strip_markers(marked)
    write_txt(os.path.join(HERE, "Video-3-Teleprompter-Script-with-slide-markers_v1.0.txt"), marked)
    write_txt(os.path.join(HERE, "Video-3-Reading-Script-no-markers_v1.0.txt"), clean)
    write_docx(os.path.join(HERE, "Video-3-Teleprompter-Script-with-slide-markers_v1.0.docx"),
               "Video 3 teleprompter script, with slide markers", marked)
    write_docx(os.path.join(HERE, "Video-3-Reading-Script-no-markers_v1.0.docx"),
               "Video 3 reading script", clean)

    for name, body, title in (
        ("Short-3A-Before-You-Resign-Preserve-the-Evidence_v1.0", SHORT_3A,
         "Short 3A, Before You Resign, Preserve the Evidence"),
        ("Short-3B-A-Better-Next-Move-Does-Two-Things_v1.0", SHORT_3B,
         "Short 3B, A Better Next Move Does Two Things"),
    ):
        ps = paragraphs(body)
        write_txt(os.path.join(HERE, name + ".txt"), ps)
        write_docx(os.path.join(HERE, name + ".docx"), title, ps)
        w = len(" ".join(ps).split())
        print("%-52s %3d words, about %d to %d seconds"
              % (name, w, w / 3.0, w / 2.2))

    words = len(" ".join(clean).split())
    print("main script: %d paragraphs, %d markers, %d words"
          % (len(clean), len(marked) - len(clean), words))
    print("  at 130 wpm: %.1f min | at 145 wpm: %.1f min" % (words / 130, words / 145))
    print("  em/en dashes: %d" % (SCRIPT.count("—") + SCRIPT.count("–")))


if __name__ == "__main__":
    main()
