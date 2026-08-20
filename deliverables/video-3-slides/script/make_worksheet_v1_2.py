"""Build the Video 3 un-script working sheet, in the format of Videos 1 and 2."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
NAVY, GOLD, DIM = RGBColor(0x0F,0x23,0x46), RGBColor(0xA8,0x86,0x28), RGBColor(0x5A,0x6B,0x82)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.8)
    s.left_margin = s.right_margin = Inches(0.9)
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)
n.paragraph_format.space_after = Pt(8)


def h1(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)


def h2(t, sub=None):
    p = doc.add_paragraph(); r = p.add_run(t.upper())
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(4)
    if sub:
        q = doc.add_paragraph(); s = q.add_run(sub)
        s.font.size = Pt(10); s.font.italic = True; s.font.color.rgb = DIM
        q.paragraph_format.space_after = Pt(8)


def para(t, bold=False, size=11, color=None):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(size); r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p


def bullets(items):
    for i in items:
        p = doc.add_paragraph(i, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def table(rows, widths=None, header=True):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if header and ri == 0:
                r.font.bold = True
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------- title page
h1("VIDEO 3 - FILMING WORKING SHEET  (v1.2)")
para("3 Things to Do Before Quitting Your Job", bold=True, size=15, color=NAVY)
para("Natural-delivery un-script matched to the v1.0 slide architecture: viewer-first "
     "opening, an explicit safety boundary, three separate checks, recap, a provisional "
     "decision reading, practical questions, the Career Decision Evidence Check "
     "invitation, and the next-video bridge.")
para("Prepared for Temidayo Afonja", color=DIM)
para("Revision: v1.2, Wednesday, August 19, 2026, America/Chicago", color=DIM)

h2("Viewer transformation", "Read this before anything else.")
para("By the end, the viewer knows:")
bullets([
    "What evidence to preserve before access changes.",
    "How to name what the work actually built in them.",
    "How to test whether the next move uses something already proven while requiring genuinely new growth.",
    "How to read whether the evidence points toward leaving, repositioning inside, or building a bridge first.",
])

h2("Production snapshot")
table([
    ["Field", "Locked decision"],
    ["YouTube title", "3 Things to Do Before Quitting Your Job (approved, TubeBuddy comparison)"],
    ["On-screen deck title", "3 Things to Do Before Quitting Your Job"],
    ["Thumbnail words", "WAIT BEFORE YOU QUIT"],
    ["Viewer promise", "Three checks that turn an urgent exit decision into an evidence-led one."],
    ["Primary CTA", "Career Decision Evidence Check, https://temidayoafonja.com/career-decisions"],
    ["End-screen video", "How to Change Jobs Without Starting Your Career Over"],
    ["Run time", "Approximately 9 to 9.5 minutes"],
    ["Delivery model", "Memorize the opening, the safety boundary, the CTA and the closing bridge. Speak the three checks from ideas."],
], widths=[1.6, 5.2])

h2("Delivery legend")
bullets([
    "MEMORIZE EXACTLY: opening promise, safety boundary, CTA, closing bridge.",
    "SPEAK FROM IDEAS: the three checks, the example, the decision reading.",
    "PAUSE AND SHOW: let the section slides and the evidence slides carry the screen.",
])

h2("Run of show")
table([
    ["Time", "Section", "Delivery"],
    ["0:00-0:50", "Hook, payoff, early proof beat, title, introduction", "Memorize exactly"],
    ["0:50-1:35", "Safety boundary, then why access matters", "Memorize the boundary"],
    ["1:35-1:43", "01 section slide, preserve the evidence", "Section slide"],
    ["1:43-3:15", "Check 1, what to keep and what not to take", "Speak from ideas"],
    ["3:15-3:23", "02 section slide, name what the work built", "Section slide"],
    ["3:23-4:55", "Check 2, problem, constraint, judgment, outcome", "Speak from ideas"],
    ["4:55-5:03", "03 section slide, test the next move", "Section slide"],
    ["5:03-6:30", "Check 3, uses something proven, builds something new", "Speak from ideas"],
    ["6:30-6:45", "Recap, all three together", "Synthesis"],
    ["6:45-7:45", "Decision reading, three directions", "Provisional, not diagnostic"],
    ["7:45-8:35", "Before you resign, three questions", "Slow down and pause"],
    ["8:35-9:00", "Career Decision Evidence Check", "Memorize exactly"],
    ["9:00-9:25", "Next-video bridge", "Memorize exactly"],
], widths=[1.1, 4.2, 1.5])
para("Timing note. These times come from the v1.1 read and are carried unchanged so "
     "they match the speaker notes in the deck. The early proof beat added in v1.2 runs "
     "roughly twenty to thirty seconds, so the finished video will land closer to nine "
     "minutes fifty than nine twenty-five. Time one full rehearsal and adjust the deck "
     "notes and the description chapters together, in one pass, rather than guessing "
     "now.", bold=True)

# ------------------------------------------------------------------ sections
h1("01  OPENING          0:00-0:50, see the timing note")
h2("Memorize exactly", "Begin full-screen on Temidayo. The title slide arrives after the payoff.")
para("If you are seriously thinking about quitting your job, there are three things I "
     "want you to check before you go.")
para("I am not going to try to talk you into staying.")
para("I want to help you leave, if leaving is the right decision, with a clear record "
     "of what this work built in you and a better read on what the next move needs to do.")
para("Because once you leave, access changes.")
para("Records become harder to reach. Systems close. People who saw the work move on to "
     "other priorities. And something that feels obvious while you are still inside the "
     "role can become surprisingly difficult to reconstruct six months later.")
para("By the end of this video, you will know what evidence to preserve, how to name "
     "what the work actually built in you, and how to test whether the next move uses "
     "something you have already proved while requiring you to build something "
     "genuinely new.")
h2("Early proof beat", "Locked standard. Say this before the title slide appears.")
para("I’ve worked inside the systems where performance, talent decisions and employee "
     "transitions are documented, so I know how quickly the evidence behind someone’s "
     "work becomes harder to reconstruct once they are no longer inside the role.")
para("That is why the first thing I want you to check is what evidence you should preserve.")
h2("Brief introduction", "After the title slide.")
para("I'm Temidayo Afonja. On this channel, I help experienced professionals make "
     "clearer career decisions by looking at what their work is actually building in them.")
h2("Performance notes")
bullets([
    "Opening order is fixed: hook, viewer payoff, early proof beat, title, brief introduction, teaching.",
    "Do not open with a welcome, a greeting or any creator-style intro.",
    "Do not argue against leaving. The viewer has usually been thinking about this for months.",
    "Let the title slide appear only after the three-check promise is clear.",
])
h2("Presentation and camera cue")
para("SLIDE: 1, title. CAMERA: full frame through the promise, then the title.")

h1("02  SAFETY BOUNDARY, THEN WHY ACCESS MATTERS          0:50-1:35")
h2("Memorize exactly", "This boundary is not optional and must not be softened.")
para("If your health or safety is at risk, or you are facing harassment, discrimination "
     "or another urgent threat, nothing in this video is a reason to delay leaving. "
     "Please act on that first. Everything I am about to say assumes you have the time "
     "to think. If you do not, that is a different situation and it deserves a different "
     "response.", bold=True)
h2("Then, speak from ideas")
bullets([
    "While you are still in the role you can see your own review history, look up dates, check what a project involved, and ask a colleague what they remember.",
    "The day after you resign, most of that is harder and some of it is impossible.",
    "This is not about hesitating. It is about not losing the record of your own work on the way out.",
])
h2("Presentation and camera cue")
para("SLIDE: 2, once you leave access changes. The boundary line is on the slide "
     "beneath a small rust rule. Say it as well as showing it.")

h1("03  CHECK 1: PRESERVE THE EVIDENCE          1:35-3:15")
h2("Section-intro beat")
para('On the standalone 01 slide, say: "The first check is to preserve the evidence." '
     "Hold for roughly four to seven seconds, then move into the teaching slide.")
h2("What the viewer may keep")
bullets([
    "Their own performance reviews.",
    "Recognition they received.",
    "Nonconfidential metrics already shared with them.",
    "Project dates.",
    "Scope of responsibility.",
    "Permitted notes about decisions they influenced or problems they helped resolve.",
])
h2("What the viewer must not take", "State this plainly, every time.")
bullets([
    "Confidential information.",
    "Customer data.",
    "Employee data.",
    "Proprietary documents.",
    "Anything employer-owned they do not have the right to keep.",
])
para('Anchor line: "Preserving your record does not mean taking their material. If you '
     'are not entitled to keep it, do not take it."')
h2("The record itself", "Spoken, not on the slide.")
table([
    ["Line", "Question"],
    ["What changed", "What changed as a result?"],
    ["Starting condition", "What was the situation before?"],
    ["Decision or influence", "What did I decide or influence?"],
    ["Who was affected", "Who did the change reach?"],
    ["Permitted evidence", "What evidence supports it that I am entitled to keep?"],
], widths=[1.8, 5.0])
para("That is a paragraph, not a filing cabinet.")
h2("Presentation and camera cue")
para("SLIDE: 3, standalone 01; then 4, what to keep and what not to take. Reveal the "
     "left column two items at a time, then the boundary column. No employer documents "
     "or screenshots at any point.")

h1("04  CHECK 2: NAME WHAT THE WORK BUILT          3:15-4:55")
h2("Section-intro beat")
para('On the standalone 02 slide, say: "The second check is to name what the work built."')
h2("Speak from ideas")
para("A resume bullet can tell another employer what happened. It does not "
     "automatically tell them what you became able to do.")
table([
    ["Capture", "Question"],
    ["Problem", "What problem was being solved?"],
    ["Constraint", "What made the situation difficult?"],
    ["Judgment", "What did I notice, decide, interpret or influence?"],
    ["Outcome", "What changed or was prevented?"],
], widths=[1.6, 5.2])
para("Then ask: where else could that combination matter?")
h2("The generic example", "Use this one. Do not fabricate statistics.")
para("Someone may have reduced the time an internal process took. That is the bullet. "
     "The portable value is usually the ability to identify where the work was getting "
     "stuck, align people who owned different parts of the same system, redesign the "
     "handoff, and do it without creating a new control failure.")
h2("Presentation and camera cue")
para("SLIDE: 5, standalone 02; then 6, the four rows revealed one at a time, then the "
     "closing question.")

h1("05  CHECK 3: TEST THE NEXT MOVE          4:55-6:30")
h2("Section-intro beat")
para('On the standalone 03 slide, say: "The third check is to test the next move."')
h2("The locked distinction")
para("A strong move uses something already proven and builds something genuinely new.",
     bold=True)
bullets([
    "What will this next role allow me to carry?",
    "What new judgment, exposure or responsibility will it force me to develop?",
    "What will I be able to do after a year that I cannot do now?",
])
h2("Two cautions", "Spoken, not on the slide.")
bullets([
    "A move that uses nothing already developed may impose an unnecessary reset.",
    "A move that repeats the same work at another employer may change the setting without materially changing what the career is building.",
])
h2("Presentation and camera cue")
para("SLIDE: 7, standalone 03; then 8, the contrast, then the three questions one at a time.")

h1("06  RECAP AND DECISION READING          6:30-7:45")
h2("Recap talk track")
para("Those are the three checks. Preserve the evidence. Name what the work built. "
     "Test the next move. Take them in that order, before you resign, not after.")
h2("Three directions", "Provisional. Do not tell the viewer which one is theirs.")
table([
    ["Direction", "What the evidence is saying"],
    ["Leave", "Leaving is right and the person is ready."],
    ["Reposition inside", "Another role, project or scope change could restore growth without an immediate exit."],
    ["Build a bridge", "Something is missing first: a credential, outside-context evidence, financial runway, or a clearer translation of what the experience can do elsewhere."],
], widths=[1.6, 5.2])
para("Exact line: The point is not to make the decision slow. The point is to make it "
     "legible.", bold=True)
para("If the safety boundary applies to someone watching, repeat it briefly here.")
para("Exact line: A harmful situation does not need a bridge plan.", bold=True)

h1("07  BEFORE YOU RESIGN          7:45-8:35")
h2("Pause and show", "Ask one question at a time. Leave five to seven seconds after the third.")
bullets([
    "What evidence do I need to preserve now?",
    "What does my strongest evidence show I can do?",
    "What must the next move use, and what must it build?",
])

h1("08  CTA AND NEXT-VIDEO BRIDGE          8:35-9:25")
h2("Memorize exactly", "The Career Decision Evidence Check is the only invitation in this video. Do not substitute the Field Kit.")
para("If you want a structured read of the evidence behind the decision you are "
     "weighing, that is what the Career Decision Evidence Check is for. You can find it "
     "at temidayoafonja.com/career-decisions. I have also linked it below.")
h2("Final bridge", "Do not summarize the video again.")
para("And if the move you are considering changes your function or your industry, there "
     "is one more question underneath all of this. What actually carries across, and "
     "what does not? That is the video I would watch next. How to Change Jobs Without "
     "Starting Your Career Over.")
para("Before publishing, verify that temidayoafonja.com/career-decisions is live and "
     "reaches the intended destination.", bold=True)

h1("09  PRESENTATION MAP          13 slides")
table([
    ["#", "Job", "On-screen content", "Builds"],
    ["1", "Title", "3 Things to Do Before Quitting Your Job", "None"],
    ["2", "Recognition", "Once you leave, access changes, plus the safety boundary", "None"],
    ["3", "Section 01", "Preserve the evidence", "None"],
    ["4", "Check 1", "Yours to keep, and not yours to take", "4"],
    ["5", "Section 02", "Name what the work built", "None"],
    ["6", "Check 2", "Problem, constraint, judgment, outcome", "5"],
    ["7", "Section 03", "Test the next move", "None"],
    ["8", "Check 3", "Uses something proven, builds something new", "4"],
    ["9", "Recap", "All three checks together, once", "None"],
    ["10", "Decision reading", "Leave, reposition inside, build a bridge", "3"],
    ["11", "Before you resign", "Three questions", "3"],
    ["12", "CTA", "Career Decision Evidence Check", "None"],
    ["13", "Watch next", "How to Change Jobs Without Starting Your Career Over", "None"],
], widths=[0.4, 1.3, 3.8, 0.8])

h1("10  REHEARSAL CHECKLIST")
bullets([
    "Read the whole script aloud once for meaning, not speed.",
    "Memorize the opening, the safety boundary, the CTA and the closing bridge.",
    "Practice saying the safety boundary without hedging or rushing it.",
    "Practice each check from the standalone slide title alone.",
    "Keep the confidentiality boundary explicit every time evidence is discussed.",
    "Present the three directions as provisional, never as a diagnosis of the viewer.",
    "Confirm temidayoafonja.com/career-decisions is live before publishing.",
    "Record a 60-second test and check eye line, audio, lighting and presenter-box placement.",
    "Verify every slide on a phone-sized preview.",
    "Add the Video 1 end-screen link before publishing and review automatic captions.",
])

h1("11  UPLOAD PACKAGE")
h2("Title, thumbnail, CTA")
bullets([
    "TITLE: 3 Things to Do Before Quitting Your Job",
    "THUMBNAIL: WAIT BEFORE YOU QUIT",
    "PRIMARY CTA: Career Decision Evidence Check, https://temidayoafonja.com/career-decisions",
    "END SCREEN: How to Change Jobs Without Starting Your Career Over",
])
h2("Description", "Adjust chapter times against the final export.")
para("Read the evidence before you decide: https://temidayoafonja.com/career-decisions")
para("Resignation decisions are often made when growth has stopped, conditions have "
     "changed, or another opportunity has appeared. This video does not argue for "
     "staying. It gives you three checks to run before you go, while you still have "
     "access to the record of your own work.")
bullets([
    "Preserve the evidence you are entitled to keep, and nothing you are not.",
    "Name what the work built: problem, constraint, judgment, outcome.",
    "Test whether the next move uses something proven and builds something new.",
])
para("If your health or safety is at risk, or you are facing harassment or "
     "discrimination, this framework is not a reason to delay leaving.")
para("CHAPTERS  0:00 Before you resign  0:50 Once you leave, access changes  "
     "1:35 Preserve the evidence  3:15 Name what the work built  4:55 Test the next move  "
     "6:30 The three checks  6:45 What the evidence can point to  7:45 Three questions "
     "before you resign  8:35 Career Decision Evidence Check  9:00 Watch next")
h2("Pinned comment")
para("What evidence would you want in hand before you resign? If you want a structured "
     "read of the evidence behind the decision, it is here: "
     "https://temidayoafonja.com/career-decisions")

doc.save(os.path.join(HERE, "Video-3-Unscript-Working-Sheet_v1.2_Temidayo_Afonja.docx"))
print("working sheet written")
