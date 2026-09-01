# -*- coding: utf-8 -*-
"""Build the Video 2 H.I.T. final recording and Shorts package."""
import os, sys, shutil, zipfile, hashlib
sys.path.insert(0, "/tmp/v3hit")
from script_text import LINES, SPOKEN, MARKERS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

ROOT="/tmp/v3hit/Video_3_HIT_FINAL"
LF=os.path.join(ROOT,"LONG_FORM"); SH=os.path.join(ROOT,"SHORTS")
shutil.rmtree(ROOT, ignore_errors=True)
os.makedirs(LF); os.makedirs(SH)

TITLE="3 Things to Do Before Quitting Your Job"
THUMB="WAIT BEFORE YOU QUIT"
CTA="Career Decision Evidence Check"
CTA_URL="https://temidayoafonja.com/career-decisions"
NEXT="How to Change Jobs Without Starting Your Career Over"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO 3",size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t,before=20):
    return keep(P(d,t,size=14,bold=True,color=NAVY,before=before,after=8),True)
def H2(d,t,before=13):
    return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=before,after=5),True)

def compress(d, line_spacing=1.22, after_scale=0.8):
    """Tighten one document's vertical rhythm without touching type size,
    weight, colour or wording. Used on the long-form editor brief so its last
    section cannot spill onto a near-empty extra page in Word."""
    for p in d.paragraphs:
        pf = p.paragraph_format
        if pf.line_spacing and pf.line_spacing > line_spacing:
            if pf.line_spacing < 1.4:          # leave recording copy alone
                pf.line_spacing = line_spacing
        if pf.space_after is not None:
            pf.space_after = Pt(round(pf.space_after.pt * after_scale, 1))
        if pf.space_before is not None and pf.space_before.pt:
            pf.space_before = Pt(round(pf.space_before.pt * after_scale, 1))
    return d

# ------------------------------------------------- 1. teleprompter DOCX + TXT
d=newdoc(True)
head(d,TITLE,"Video 3  ·  Teleprompter script with slide markers",
     "Spoken script is the large text. A slide marker in a tinted band tells "
     "the editor which slide to bring up; it is not spoken.")
for line in LINES:
    if line.startswith("[SLIDE:"):
        name=line[len("[SLIDE:"):-1].strip()
        p=P(d,"SLIDE  —  %s"%name,size=11,bold=True,color=NAVY,
            before=14,after=14,spacing=1.1)
        shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
    else:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video3TeleprompterScriptwithslidemarkers_HIT_v2.0.docx"))

tel_txt=[TITLE,"Video 3  ·  Teleprompter script with slide markers","",]
for line in LINES:
    if line.startswith("[SLIDE:"):
        tel_txt += ["", "SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(), ""]
    else:
        tel_txt += [line, ""]
open(os.path.join(LF,"Video3TeleprompterScriptwithslidemarkers_HIT_v2.0.txt"),
     "w").write("\n".join(tel_txt).strip()+"\n")

# ------------------------------------------------- 2. reading script DOCX+TXT
d=newdoc(True)
head(d,TITLE,"Video 3  ·  Reading script, no markers",
     "Spoken language only. No slide markers, no timestamps, no production "
     "directions.")
for line in SPOKEN:
    keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video3ReadingScriptnomarkers_HIT_v2.0.docx"))
open(os.path.join(LF,"Video3ReadingScriptnomarkers_HIT_v2.0.txt"),
     "w").write("\n\n".join(SPOKEN)+"\n")
print("long-form scripts written")


# --------------------------------------------------- 3. long-form editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RGBColor(0x9B,0x2C,0x10),after=2)
P(d,"VIDEO 3",size=12,bold=True,color=GOLD,after=2,caps=True)
P(d,TITLE,size=20,bold=True,color=NAVY,after=6,spacing=1.1)
p=P(d,"This document is for the editor. It is NOT Temidayo's teleprompter and "
     "must not be placed on the recording screen.",size=11,italic=True,
     color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"Locked metadata",before=14)
for k,v in (("Title",TITLE),("Thumbnail",THUMB),("Primary CTA",CTA),
            ("CTA URL",CTA_URL),("Watch next",NEXT),
            ("Core distinction",
             "Before leaving, preserve the evidence, understand what your work "
             "built in you, and test what the next move must use and build.")):
    keep(P(d,"%-18s %s"%(k+":",v),size=11,after=5))
p=P(d,"PUBLICATION GATE. Do not treat Video 3 as publication-ready until "
     "https://temidayoafonja.com/career-decisions is production-live and "
     "usable. The page is confirmed live, so the gate is currently "
     "satisfied. One signed-out production check of the page is still "
     "required before Video 3 is uploaded or scheduled.",
     size=11,bold=True,color=RGBColor(0x9B,0x2C,0x10),
     before=8,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"First 30 seconds — H.I.T.",before=14)
P(d,"H = Hook.  I = Interest.  T = Trust. The opening must work as an "
    "audiovisual unit: immediate conversational hook, meaningful visual "
    "interest, a relevant concrete reason to trust Temidayo, a clear practical "
    "payoff inside roughly the first 20 seconds, and an immediate safety and "
    "integrity boundary. Do not force a statistic into the opening.",after=12)
P(d,"This supersedes the older Video 3 opening and any earlier instruction "
    "that kept Temidayo visually static or full-screen through the extended "
    "opening.",italic=True,color=DIM,after=14)

def beat(t,anchor,layer,body):
    H2(d,t)
    p=P(d,"Spoken anchor:  “%s”"%anchor,size=10.5,italic=True,color=DIM,after=8)
    shade(p,BAND_CREAM)
    if layer: keep(P(d,layer,size=11,bold=True,color=GOLD,after=6))
    for b in body: keep(P(d,b,after=5))

beat("0:00–0:08",
 "If you’re going to quit your job, don’t wait until after you leave to figure "
 "out what your work actually built in you.","H = HOOK",
 ["Visual: begin on Temidayo, medium/tight, direct to camera.",
  "On-screen text: BEFORE YOU RESIGN…",
  "No title card before this. No “I QUIT” graphic. No resignation "
  "stock footage."])
beat("0:08–0:12","Access changes. Systems close. People move on.","I = INTEREST",
 ["Three restrained editorial text states:  ACCESS  ·  SYSTEMS  ·  PEOPLE",
  "Then briefly resolve to:  HARDER TO RECONSTRUCT",
  "Do not use fake software screens, login pages, employee records or literal "
  "employer systems."])
beat("0:12–0:21","Before you resign, check three things…","PAYOFF",
 ["Brief progressive preview:",
  "     PRESERVE THE EVIDENCE",
  "     NAME WHAT YOUR WORK BUILT",
  "     TEST THE NEXT MOVE",
  "This is a micro-preview only. Do not replace or redesign the existing "
  "recap slide."])
beat("0:21–0:29",
 "I’ve worked inside systems where performance and talent decisions are "
 "documented…","T = TRUST",
 ["Return cleanly to Temidayo. Keep this visual treatment simple.",
  "Do not add résumé graphics, employer logos or a career montage.",
  "This proof is specifically about Temidayo's relevant organizational "
  "experience.",
  "Do not imply that Temidayo personally watched a named employee lose "
  "evidence, that she witnessed a specific person regret leaving, or any "
  "undocumented anecdote."])
beat("~0:29 onward",
 "And if your health or safety is at risk, this is not a reason to wait.",
 "INTEGRITY / SAFETY",
 ["Stay on Temidayo.",
  "No dramatic music or visual warning effect.",
  "The fuller safety boundary follows in the body."])

H1(d,"Editorial rhythm after the opening",before=14)
P(d,"The first roughly 30 seconds may be more visually active than the "
    "remainder. Do NOT read H.I.T. as permission for constant motion.",after=10)
P(d,"Once the viewer is inside the teaching:",after=6)
for x in ["preserve Temidayo's natural pacing;",
 "use the existing slides as teaching support;","allow reflective pauses;",
 "avoid constant punch-ins;","avoid generic corporate B-roll;",
 "avoid stock resignation scenes;","avoid firing or layoff imagery;",
 "avoid red warning graphics;","avoid countdown clocks;",
 "avoid fake HR or employer screens;",
 "avoid boxes being carried out of an office;","avoid flashy transitions."]:
    keep(P(d,"—  "+x,after=4))
P(d,"The edit should feel calm, premium, specific and editorial.",
  before=8,after=10)

H1(d,"Existing slides",before=14)
P(d,"The existing 13-slide Video 3 deck remains UNCHANGED. Do not add, delete, "
    "redesign or reorder slides.",after=8)
for n,job in enumerate(["Title","Once You Leave, Access Changes",
 "01 Preserve the Evidence","What to Keep / What Not to Take",
 "02 Name What Your Work Built","Problem / Constraint / Judgment / Outcome",
 "03 Test the Next Move","Uses Something Proven / Builds Something New",
 "The Three Checks","Decision Reading","Before You Resign",
 "Career Decision Evidence Check","Watch Next"],1):
    keep(P(d,"Slide %-3d %s"%(n,job),size=10.5,after=3))
p=P(d,"REVIEWED AND INTENTIONALLY RETAINED. Slides 5 and 6 use the "
    "conceptual heading “Name what the work built.” The spoken script "
    "addresses the viewer directly with “Name what your work built.” No "
    "slide change is required.",size=10.5,bold=True,color=NAVY,
    before=10,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"Let the slide carry lists",before=14)
P(d,"When a visual can carry factual or list information efficiently, "
    "Temidayo does not need to narrate every item. Do not add spoken lines to "
    "compensate for slide copy.",after=8)
keep(P(d,"Slide 4 — What to Keep / What Not to Take. Temidayo states the "
       "governing rule; the slide carries the supporting examples.",after=5))
keep(P(d,"Slide 6 — Problem / Constraint / Judgment / Outcome. Temidayo does "
       "not read every label mechanically. The slide provides the structure "
       "while the spoken script carries judgment and meaning.",after=8))

H1(d,"Safety and evidence boundaries",before=14)
keep(P(d,"Preserve the safety boundary exactly as written. Do not visually or "
       "editorially imply that viewers should delay leaving unsafe work, "
       "harassment, discrimination or another urgent threat.",bold=True,after=8))
P(d,"Evidence preservation is NOT permission to remove employer property. Do "
    "not suggest downloading confidential files, forwarding employer "
    "documents, retaining customer or employee data, taking proprietary "
    "material, or bypassing employer policies or controls.",after=8)
keep(P(d,"Keep the framing: a permitted record of the viewer's own "
       "contribution, in their own words, using information they are entitled "
       "to retain.",italic=True,after=10))

H1(d,"CTA and watch next",before=14)
keep(P(d,"One product or resource CTA only: %s — %s"%(CTA,CTA_URL),after=5))
keep(P(d,"Do not add the Capability Formation Field Kit, Keep the Proof, the "
       "book, the newsletter or any other product.",bold=True,after=8))
keep(P(d,"Watch next: route to %s. Use Video 1 as the intended clickable "
       "direct-video route where available."%NEXT,after=6))
compress(d,1.18,0.80)
d.save(os.path.join(LF,"Video_3_EDITOR_ONLY_HIT_Brief_v2.0.docx"))
print("editor brief written")

# ----------------------------------------------------- 4. publishing package
d=newdoc()
head(d,TITLE,"Video 3  ·  Publishing package",
     "Everything needed to upload. Working timestamps must be replaced with "
     "real ones from the finished edit.")
H1(d,"Title"); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail"); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase")
P(d,"things to do before quitting your job",after=10)
H1(d,"Supporting search language")
P(d,"before quitting your job · before you resign · career change · should I "
    "quit my job · career transition · career decision · transferable "
    "experience · career evidence · career portability",after=10)

H1(d,"Description")
for para in [
 "If you are seriously thinking about quitting your job, there are three "
 "things worth checking before you resign.",
 "In this video, I show you how to preserve permitted evidence of your work, "
 "identify what your work actually built in you, and test whether your next "
 "move uses something you have already proved while asking you to build "
 "something genuinely new.",
 "The three checks are:",
 "1. Preserve the evidence.",
 "2. Name what your work built in you.",
 "3. Test the next move.",
 "This is not advice to delay leaving an unsafe or harmful situation. If your "
 "health or safety is at risk, or you are facing harassment, discrimination "
 "or another urgent threat, act on that first.",
 "And preserving career evidence does not mean taking confidential, "
 "proprietary, customer, employee or employer-owned material. Keep only what "
 "you are entitled to retain.","",
 "CAREER DECISION EVIDENCE CHECK",
 "If you want a structured way to read the evidence behind a stay, move or "
 "leave decision:", CTA_URL,"",
 "WATCH NEXT", NEXT, "[ADD VIDEO 1 LINK]","",
 "Temidayo Afonja helps experienced professionals understand what they can "
 "carry across roles, functions, employers and industries so they can make "
 "career pivots and internal moves without starting from zero."]:
    keep(P(d,para if para else " ",after=7 if para else 3))

H1(d,"Working chapters")
p=P(d,"WORKING ESTIMATES — EDITOR MUST REPLACE FROM FINAL CUT. These "
    "timestamps were estimated from the script, not measured from an edit. "
    "Every one must be replaced with the actual timestamp from the finished "
    "cut before publishing.",size=10.5,bold=True,italic=True,
    color=RGBColor(0x9B,0x2C,0x10),after=10,spacing=1.25)
shade(p,BAND_CREAM)
for t,c in [("00:00","Do not wait until after you leave"),
 ("00:35","The safety boundary"),
 ("00:50","Once you leave, access changes"),
 ("01:50","Check 1: Preserve the evidence"),
 ("02:55","Check 2: Name what your work built in you"),
 ("04:10","Check 3: Test the next move"),
 ("05:15","The three checks together"),
 ("05:35","Reading your decision"),
 ("06:50","Before you resign"),
 ("07:30","Career Decision Evidence Check"),
 ("07:55","Watch next")]:
    keep(P(d,"%s   %s"%(t,c),size=11,after=4))

H1(d,"Pinned comment")
for para in [
 "Before you resign, which of these three questions is hardest to answer?",
 "1. What evidence do I need to preserve now?",
 "2. What does my strongest evidence prove I can do?",
 "3. What must the next move use — and what must it build?",
 "If your health or safety is at risk, you do not need to delay leaving in "
 "order to complete a career exercise."]:
    keep(P(d,para,after=6))

H1(d,"Publication gate")
p=keep(P(d,"Do not publish Video 3 until %s is production-live and usable."%CTA_URL,
       size=11,bold=True,color=NAVY,after=6,spacing=1.25))
shade(p,BAND_CREAM)
keep(P(d,"The page is confirmed live. The gate is currently satisfied. One signed-out production check of the page is still required before Video 3 is uploaded or scheduled: confirm it loads for a visitor who is not signed in and is not holding a preview link.",size=11,after=8,spacing=1.25))
d.save(os.path.join(LF,"Video_3_Publishing_Package_HIT_v2.0.docx"))

# ---------------------------------------------------------------- 5. Shorts
SHORTS=[
 ("Video_3_Short_1_Before_You_Lose_Context.docx","SHORT 1","Recognition",
  "The day after you quit is a bad time to start reconstructing your career.",
  ["The day after you quit is a bad time to start reconstructing your career.",
   "While you are still inside the role, you remember why a project mattered, "
   "what changed, what you actually decided and who saw the work.",
   "Later, some of that context gets surprisingly hard to retrieve.",
   "So before you leave, write down the evidence you are permitted to keep "
   "and enough context to explain your contribution in your own words.",
   "You are not taking company material.",
   "You are making sure your own experience does not become harder to explain "
   "simply because you changed employers."]),
 ("Video_3_Short_2_Keep_Proof_Not_Files.docx","SHORT 2","Distinction / myth",
  "Preserving your career evidence does not mean taking company files.",
  ["Preserving your career evidence does not mean taking company files.",
   "If you do not have the right to keep something, do not take it.",
   "The better habit is to create a permitted record in your own words.",
   "What was the situation?",
   "What did you decide or influence?",
   "What changed?",
   "What evidence are you actually entitled to retain?",
   "That is usually what you will need later to explain your value.",
   "Not a folder full of somebody else’s material."]),
 ("Video_3_Short_3_Context_Disappears_Fast.docx","SHORT 3",
  "Proof / perspective",
  "Working around performance and talent systems taught me something about "
  "leaving a job.",
  ["Working around performance and talent systems taught me something about "
   "leaving a job.",
   "While you are inside an organization, the context around your work is "
   "everywhere.",
   "Reviews.","Recognition.","Project history.",
   "People who remember why something mattered.",
   "Once the context changes, your experience did not disappear — but "
   "explaining it can become much harder.",
   "That is why I want you to capture more than what you did.",
   "Capture what your work proved you can do.",
   "That is the part you may need in the next context."]),
 ("Video_3_Short_4_Three_Questions_Before_You_Quit.docx","SHORT 4",
  "Practical test / action",
  "Before you resign, answer these three questions.",
  ["Before you resign, answer these three questions.",
   "What evidence do I need to preserve now?",
   "What does my strongest evidence prove I can do?",
   "And what must the next move use — and what must it build?",
   "Those questions do not tell you whether quitting is right or wrong.",
   "They tell you whether you understand what you are leaving with and what "
   "you are moving toward.",
   "And if your health or safety is at risk, you do not need to delay leaving "
   "in order to complete a career exercise."])]

for fn,label,role,hook,copy in SHORTS:
    d=newdoc(True)
    P(d,"VIDEO 3 SHORT",size=10,bold=True,color=GOLD,after=4,caps=True)
    P(d,label,size=20,bold=True,color=NAVY,after=8,spacing=1.1)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    keep(P(d,"Verbal hook:  “%s”"%hook,size=11,color=DIM,after=5))
    keep(P(d,"Related long-form:  %s"%TITLE,size=11,color=DIM,after=16))
    H1(d,"RECORDING COPY")
    for line in copy:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
    d.save(os.path.join(SH,fn))
print("publishing package and %d Shorts written"%len(SHORTS))

# ------------------------------------------------------ 6. Shorts editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RGBColor(0x9B,0x2C,0x10),after=2)
P(d,"VIDEO 3 — FOUR STANDALONE SHORTS",size=18,bold=True,color=NAVY,
  after=8,spacing=1.1)
p=P(d,"This document is for the editor. It is separate from the four Short "
     "recording documents and must not be placed on Temidayo's recording "
     "screen.",size=11,italic=True,color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"How these are produced")
keep(P(d,"These are separately recorded 9:16 Shorts. They are NOT excerpts cut "
       "from the long-form video.",bold=True,after=10))
P(d,"Each Short should have:",after=6)
for x in ["an immediate verbal hook;","a corresponding on-screen hook;",
          "meaningful visual movement;","accurate burned-in captions;",
          "restrained editorial pacing;",
          "Video 3 added as the YouTube Related Video when available."]:
    keep(P(d,"—  "+x,after=4))

def short(label,role,onscreen,body):
    H1(d,label)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    p=keep(P(d,"On-screen hook:  %s"%onscreen,size=11,bold=True,color=GOLD,after=8))
    shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))
    keep(P(d,"Related Video:  Video 3",size=10.5,color=DIM,before=4,after=6))

short("SHORT 1","Recognition","BEFORE YOU LOSE THE CONTEXT",
 ["Begin direct to camera.",
  "Use restrained word states:  CONTEXT  ·  EVIDENCE  ·  WHAT YOU DECIDED  ·  "
  "WHY IT MATTERED",
  "Do not use fake employer systems.",
  "Do not show downloading files.",
  "Do not show resignation stock footage.",
  "End visually on:  PRESERVE WHAT YOU CAN EXPLAIN"])
short("SHORT 2","Distinction / myth","KEEP THE PROOF. NOT THEIR FILES.",
 ["Use a restrained contrast:  YOUR RECORD  vs.  THEIR MATERIAL",
  "When Temidayo says “If you do not have the right to keep something, do not "
  "take it,” keep her full-screen.",
  "Do not show confidential documents, file downloading, email forwarding, "
  "USB drives, or screenshots of company systems.",
  "End visually on:  YOUR CONTRIBUTION.  YOUR WORDS.  PERMITTED EVIDENCE."])
short("SHORT 3","Proof / perspective","CONTEXT DISAPPEARS FAST",
 ["When Temidayo says “the context around your work is everywhere,” allow the "
  "words to appear around her:",
  "     REVIEWS  ·  RECOGNITION  ·  PROJECT HISTORY  ·  PEOPLE",
  "Then let them recede.",
  "IMPORTANT: Do not imply Temidayo personally witnessed a specific person "
  "lose records or regret leaving. Her proof is her experience working around "
  "performance and talent systems.",
  "End visually on:  WHAT DID YOUR WORK PROVE?"])
short("SHORT 4","Practical test / action","3 QUESTIONS BEFORE YOU QUIT",
 ["Reveal progressively:",
  "     1   WHAT DO I NEED TO PRESERVE?",
  "     2   WHAT DOES MY EVIDENCE PROVE?",
  "     3   WHAT MUST THE NEXT MOVE USE + BUILD?",
  "Do not show all three crowded on screen at once.",
  "When Temidayo gives the health and safety boundary, return cleanly to "
  "full-screen Temidayo.",
  "No warning graphic or alarm effect."])

H1(d,"All Shorts — visual boundaries")
for x in ["No hyperactive zooming.","No fake shock expressions.",
          "No red warning graphics.","No resignation memes.",
          "No firing footage.","No cardboard-box office exit footage.",
          "No fake HR software.","No stock résumé footage.",
          "No AI-style animated icons.","No confidential document visuals."]:
    keep(P(d,"—  "+x,after=4))
keep(P(d,"Keep captions accurate and mobile-safe. The tone should remain calm, "
       "intelligent, human and specific.",before=8,after=8,spacing=1.25))
d.save(os.path.join(SH,"Video_3_Shorts_EDITOR_ONLY_HIT_Brief.docx"))

# ---------------------------------------------------------------- 7. README
FILES=(["LONG_FORM/"+f for f in sorted(os.listdir(LF))]
      +["SHORTS/"+f for f in sorted(os.listdir(SH))])
R=["VIDEO 3 — H.I.T. FINAL RECORDING PACKAGE","",
 "Title:        %s"%TITLE,
 "Thumbnail:    %s"%THUMB,
 "CTA:          %s"%CTA,
 "CTA URL:      %s"%CTA_URL,
 "Publication",
 "gate:         Do not publish until the Career Decision Evidence Check page",
 "              is production-live and usable. The page is confirmed live,",
 "              so the gate is currently satisfied. One signed-out",
 "              production check is still required before Video 3 is",
 "              uploaded or scheduled.",
 "Watch next:   %s"%NEXT,"",
 "Long-form:    Revised under the H.I.T. first-30-second standard.",
 "Slides:       UNCHANGED.",
 "Shorts:       Four separately recorded vertical scripts.",
 "Editor",
 "instructions: Separated from all recording copy.","",
 "-"*70,"","WHAT EACH FILE IS","",
 "LONG_FORM/","",
 "  Video3TeleprompterScriptwithslidemarkers_HIT_v2.0.docx",
 "  Video3TeleprompterScriptwithslidemarkers_HIT_v2.0.txt",
 "      Temidayo's recording copy. Spoken script in large text; slide markers",
 "      in tinted bands. The markers are not spoken.","",
 "  Video3ReadingScriptnomarkers_HIT_v2.0.docx",
 "  Video3ReadingScriptnomarkers_HIT_v2.0.txt",
 "      The same spoken words with the slide markers removed.","",
 "  Video_3_EDITOR_ONLY_HIT_Brief_v2.0.docx",
 "      For the editor. The H.I.T. first-30-second plan, editorial rhythm",
 "      after 0:30, the existing 13-slide map, the reviewed slide 5 and 6",
 "      wording note and the safety and evidence boundaries.",
 "      Not for the teleprompter.","",
 "  Video_3_Publishing_Package_HIT_v2.0.docx",
 "      Title, thumbnail, search language, description, working chapter",
 "      estimates, pinned comment and the publication gate.","",
 "SHORTS/","",
 "  Four recording documents, one per Short. These contain Temidayo's",
 "  recording copy and no editor directions.","",
 "  Video_3_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "      For the editor. On-screen hooks and visual treatment for all four.","",
 "-"*70,"","ALL FILES IN THIS PACKAGE","",]
for f in FILES: R.append("  "+f)
R+=["  README_FINAL.txt","  SHA256SUMS.txt","",
 "-"*70,"","WORKING CHAPTER TIMESTAMPS","",
 "The chapter timestamps in the publishing package are WORKING ESTIMATES",
 "derived from the script. They were not measured from an edit. The editor",
 "must replace every one of them from the finished cut before publishing.","",
 "-"*70,"","SLIDE WORDING — REVIEWED AND INTENTIONALLY RETAINED","",
 "REVIEWED AND INTENTIONALLY RETAINED. Slides 5 and 6 use the conceptual",
 "heading “Name what the work built.” The spoken script addresses the",
 "viewer directly with “Name what your work built.” No slide change is",
 "required.","",
 "-"*70,"","WHAT WAS NOT CHANGED","",
 "The existing Video 3 PowerPoint deck, the reveal deck, the thumbnail, the",
 "Career Decision Evidence Check page, every website file, every product and",
 "every other video are unchanged. This revision is spoken script and editor",
 "instruction only.","",
 "The 13-slide deck remains authoritative. The teleprompter's 13 slide markers",
 "map to it in order, slide 1 to slide 13.",""]
open(os.path.join(ROOT,"README_FINAL.txt"),"w").write("\n".join(R))
print("shorts editor brief and README written")


# ------------------------------------------- 8. checksums and the master ZIP
# The archive is built from an explicit allowlist, never from a directory walk.
MANIFEST=[
 "LONG_FORM/Video3TeleprompterScriptwithslidemarkers_HIT_v2.0.docx",
 "LONG_FORM/Video3TeleprompterScriptwithslidemarkers_HIT_v2.0.txt",
 "LONG_FORM/Video3ReadingScriptnomarkers_HIT_v2.0.docx",
 "LONG_FORM/Video3ReadingScriptnomarkers_HIT_v2.0.txt",
 "LONG_FORM/Video_3_EDITOR_ONLY_HIT_Brief_v2.0.docx",
 "LONG_FORM/Video_3_Publishing_Package_HIT_v2.0.docx",
 "SHORTS/Video_3_Short_1_Before_You_Lose_Context.docx",
 "SHORTS/Video_3_Short_2_Keep_Proof_Not_Files.docx",
 "SHORTS/Video_3_Short_3_Context_Disappears_Fast.docx",
 "SHORTS/Video_3_Short_4_Three_Questions_Before_You_Quit.docx",
 "SHORTS/Video_3_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "README_FINAL.txt",
]
SUMS="SHA256SUMS.txt"
ZIP="/tmp/v3hit/Video_3_HIT_FINAL_Recording_and_Shorts_Package.zip"

def sha256(p):
    h=hashlib.sha256()
    with open(p,"rb") as fh:
        for b in iter(lambda: fh.read(1<<20), b""): h.update(b)
    return h.hexdigest()

# every allowlisted file must exist, and nothing on disk may be unaccounted for
for m in MANIFEST:
    assert os.path.isfile(os.path.join(ROOT,m)), "missing from build: "+m
on_disk=set()
for dp,dn,fn in os.walk(ROOT):
    dn[:]=[x for x in dn if x!="__pycache__"]
    for f in fn:
        if f.endswith(".pyc"): continue
        on_disk.add(os.path.relpath(os.path.join(dp,f),ROOT).replace(os.sep,"/"))
unexpected=sorted(on_disk-set(MANIFEST)-{SUMS})
assert not unexpected, "unexpected files in package directory: %r"%unexpected

L=["# VIDEO 3 - H.I.T. FINAL RECORDING PACKAGE",
   "# SHA-256 of the 12 user-facing files in this package.",
   "# SHA256SUMS.txt cannot hash itself. The master ZIP cannot contain its own",
   "# checksum either; it is published in the sibling file",
   "# Video_3_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256",""]
for m in MANIFEST: L.append("%s  %s"%(sha256(os.path.join(ROOT,m)),m))
open(os.path.join(ROOT,SUMS),"w").write("\n".join(L)+"\n")

if os.path.exists(ZIP): os.remove(ZIP)
with zipfile.ZipFile(ZIP,"w",zipfile.ZIP_DEFLATED) as z:
    for m in MANIFEST+[SUMS]:
        z.write(os.path.join(ROOT,m), "Video_3_HIT_FINAL/"+m)
zsha=sha256(ZIP)
open(ZIP+".sha256","w").write("%s  %s\n"%(zsha,os.path.basename(ZIP)))

# build provenance is kept beside the package, never inside it
PROV="/tmp/v3hit/_source"
shutil.rmtree(PROV,ignore_errors=True); os.makedirs(PROV)
for f in ("script_text.py","build.py","qa.py"):
    shutil.copy2("/tmp/v3hit/"+f, os.path.join(PROV,f))
print("ZIP sha256:",zsha)
