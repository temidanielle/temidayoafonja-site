# -*- coding: utf-8 -*-
"""Video_8_FINAL_Content_and_Publishing_Summary.docx — human-readable reference
sheet. Deliberately OUTSIDE the 13-file recording/Shorts ZIP.

Every field is derived from the final v2.2 build, not retyped from memory."""
import json, importlib.util, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
BAND_CREAM="F3F0E8"; BAND_NAVY="E8EDF4"

s=importlib.util.spec_from_file_location("st","script_text.py")
st=importlib.util.module_from_spec(s); s.loader.exec_module(st)
sh=importlib.util.spec_from_file_location("sh","shorts_text.py")
shm=importlib.util.module_from_spec(sh); sh.loader.exec_module(shm)
DESC=json.load(open("desc_block.json"))
PUB=[p.text for p in Document("Video_8_HIT_FINAL/LONG_FORM/Video_8_Publishing_Package_HIT_v2.2.docx").paragraphs]
WORDS=sum(len(x.split()) for x in st.SPOKEN)
RUNTIME="%d:%02d"%(WORDS/145*60//60, WORDS/145*60%60)

d=Document()
stl=d.styles['Normal']; stl.font.name='Calibri'; stl.font.size=Pt(10.5)
ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0'); d.settings.element.append(ah)
for sec in d.sections:
    sec.top_margin=sec.bottom_margin=Inches(0.75)
    sec.left_margin=sec.right_margin=Inches(0.9)

def P(txt="",size=10.5,bold=False,italic=False,color=INK,after=5,before=0,spacing=1.2,caps=False):
    p=d.add_paragraph(); r=p.add_run(txt)
    r.font.size=Pt(size); r.bold=bold; r.italic=italic; r.font.color.rgb=color
    if caps: r.font.all_caps=True
    pf=p.paragraph_format; pf.space_after=Pt(after); pf.space_before=Pt(before)
    pf.line_spacing=spacing; return p
def shade(p,fill):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),fill)
    p._p.get_or_add_pPr().append(sh); return p
def H(t):
    p=P(t,size=11,bold=True,color=NAVY,before=14,after=6,caps=True)
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'6'); bt.set(qn('w:color'),'8A6D1E')
    bt.set(qn('w:space'),'3'); b.append(bt); pr.append(b); return p
def KV(k,v):
    p=d.add_paragraph(); r1=p.add_run(k+"  "); r1.bold=True; r1.font.size=Pt(10.5); r1.font.color.rgb=NAVY
    r2=p.add_run(v); r2.font.size=Pt(10.5); r2.font.color.rgb=INK
    p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.2; return p
def B(t,after=3):
    p=P("—  "+t,after=after); p.paragraph_format.left_indent=Inches(0.16); return p

P("CAPABILITY FORMATION",size=9,bold=True,color=GOLD,after=2,caps=True)
P("Video 8 — Final Content and Publishing Summary",size=19,bold=True,color=NAVY,after=2,spacing=1.05)
P("v2.2.1  ·  research-alignment precision pass",size=11,bold=True,color=GOLD,after=4)
p=P("Reference sheet for Temidayo and the editing/publishing workflow. Not a recording "
    "document, and deliberately not inside the 13-file package ZIP.",size=10,italic=True,
    color=DIM,after=12,spacing=1.25); shade(p,BAND_CREAM)

H("Identity")
KV("Video number:","8")
KV("Final public title:","How to Switch Industries Without Starting Over")
KV("Final thumbnail copy:","YOUR EXPERIENCE STILL COUNTS")
KV("Status:","FINAL + LOCKED  ·  v2.2.1 research-alignment precision pass complete")

H("Video job / core question")
P("What from my experience still counts when I move into a different industry, and what "
  "do I genuinely have to learn?",spacing=1.3)

H("Target viewer")
P("An experienced professional considering or actively making an industry change, who "
  "keeps meeting “direct industry experience” requirements in the roles they want.",spacing=1.3)

H("One-line promise")
p=P("Separate what travels, what the new context changes and what you actually have to "
    "earn — so an industry move is neither a restart nor a bluff.",bold=True,color=NAVY,
    spacing=1.3,after=6); shade(p,BAND_CREAM)
P("Drawn from the script's own promise line: “what you can carry, what the new context "
  "changes, and what you actually have to earn.” No new framework or marketing claim.",
  size=9.5,italic=True,color=DIM,spacing=1.25)

H("Core tension")
P("“I have years of experience, but the next context may still ask me to prove I belong "
  "there.”",spacing=1.3)

H("Primary teaching structure — the three Cs")
KV("Capability","What remains useful when the setting changes: the judgment, decisions "
   "and patterns that are not owned by one industry — but only where they meet a problem "
   "that actually exists in the destination.")
KV("Context","What must be learned because the environment is genuinely different — "
   "language, stakeholders, incentives, regulation, operating rhythm and risk. Some of it "
   "can be researched; some has to be observed or practiced.")
KV("Credential","What has to be formally earned. Sometimes genuinely required, in which "
   "case experience does not substitute for it; sometimes a signal that shortens the "
   "conversation; sometimes neither.")

H("Permanent audit — answerable by the end of the video")
for q,a in (("What travels?","The judgment, decisions and patterns that stay useful when the setting changes."),
            ("What does not?","Knowledge tied to the old context, and anything that has to be learned, practiced or formally earned."),
            ("What can I prove?","Judgment with evidence attached — a comparable decision, comparable complexity, a relevant outcome."),
            ("What must I relearn?","The context you have not earned yet, and any credential the destination genuinely requires.")):
    KV(q,a)

H("Key recognition / opening")
p=P("“You may have fifteen or twenty years of experience, open a job description in "
    "another industry, and see the same line: direct industry experience required. That "
    "one phrase can make a career you are proud of suddenly feel like it does not count.”",
    spacing=1.3,after=6); shade(p,BAND_NAVY)
P("Held full screen on Temidayo. The first slide is the answer to it: changing industries "
  "does not make you entry-level at everything — it makes you new to a context.",
  size=9.5,italic=True,color=DIM,spacing=1.25)

H("Key lived evidence — only what the final script uses")
B("Nearly two decades across eight industries and sectors.")
B("The move toward cybersecurity and privacy: prepared for the CISM exam and did not pass "
  "the first time.")
P("Nothing else is claimed. No employer, client, metric or result appears in this video.",
  size=9.5,italic=True,color=DIM,before=3,spacing=1.25)

H("Employer-legibility point")
P("“Direct industry experience required” does not always mean the same thing. Sometimes it "
  "points at something real — a regulation to understand properly, domain knowledge that "
  "takes time, a license, decisions where being wrong carries real consequences. Sometimes "
  "it is shorthand for hiring risk: the employer is not sure how quickly somebody from "
  "outside becomes useful here.",spacing=1.3)
P("The video does not tell the viewer to ignore the requirement either way. It asks: what "
  "is this employer trying not to get wrong by asking for it? Then it sorts the gap into "
  "four parts — what I can already do with evidence behind it; what context I could "
  "reasonably learn; what is a genuine gap I cannot talk my way around; and what this "
  "employer would need to see before they could trust the overlap.",spacing=1.3)
P("The two are not mutually exclusive. A single “direct industry experience” requirement can contain both a real gap to close and uncertainty the employer needs reduced, and the script says so rather than forcing a choice between them.",spacing=1.3)

H("Honest limits — what the video does NOT promise")
for x in ("That all experience transfers. It says plainly that some knowledge becomes much "
          "less useful the moment the context changes.",
          "That framing or translation can overcome a genuine credential requirement.",
          "That every context gap is an information gap. Some has to be observed or practiced.",
          "That the Field Kit will tell the viewer which industry to choose, or remove the "
          "need to learn the field they are entering.",
          "That the third column should be empty. The script explicitly says do not try to "
          "make it empty — and do not make the first one empty either."):
    B(x,after=4)

H("Primary CTA")
KV("Offer:","Capability Formation Field Kit")
KV("URL:","https://temidayoafonja.com/fieldkit")
P("Why it fits this viewer's stage: they are mid-decision about an industry move and need "
  "to work out what their work has actually built, what looks portable and where they "
  "still need development, evidence or context. That is the Field Kit's job. One offer "
  "only — no Keep the Proof, no Career Decision Evidence Check, no book.",spacing=1.3)

H("Watch Next")
KV("Route:","Video 9 — What to Do Before a Layoff Happens")
P("Spoken bridge: if part of the reason for considering an industry move is that the "
  "current situation feels less secure, the best time to understand what you can carry is "
  "before somebody else decides when you have to move.",spacing=1.3)

H("Final YouTube description — as approved")
p=P("Paste exactly. Everything below the chapter block carries working timestamps that the "
    "editor must replace from the final cut.",size=9.5,italic=True,color=DIM,after=8,
    spacing=1.25); shade(p,BAND_CREAM)
for line in DESC:
    P(line if line.strip() else " ",size=10,after=4 if line.strip() else 2,spacing=1.25)

H("Working chapters")
p=P("WORKING TIMESTAMPS — script-derived at 145 words per minute, NOT measured from a "
    "final edit. Replace every one before publication.",size=10,bold=True,color=RGBColor(0x9B,0x2C,0x10),
    after=8,spacing=1.25); shade(p,BAND_CREAM)
import re
ch=[l for l in DESC if re.match(r"^\d\d:\d\d ",l)]
for line in ch: P(line,size=10,after=3)

H("Pinned comment — as approved")
for line in PUB[77:85]: P(line,size=10,after=4,spacing=1.25)

H("Search phrase and tags")
KV("Primary search phrase:",PUB[9])
P("Tag field — paste into the tag field only, never into the public description:",
  size=9.5,italic=True,color=DIM,before=4,after=4)
P(PUB[87],size=10,spacing=1.25)

H("Shorts — four standalone recordings")
REL=("Standalone 9:16 recording, not an excerpt cut from the long-form video. "
     "Related long-form: How to Switch Industries Without Starting Over.")
CORE=["Names the audience pain directly and helps the viewer separate what they "
      "genuinely need to learn from what the employer needs more evidence to trust — "
      "without presenting the two as mutually exclusive. Deliberately not an "
      "“ignore the requirement and apply anyway” Short.",
      "Separates the three Cs so the viewer stops solving all three problems with another "
      "certificate or an adjective.",
      "Lived proof that not everything travels, using the eight industries and the CISM "
      "first-attempt non-pass.",
      "The practical exercise: one page, three columns — what travels, what changes, what I "
      "must earn."]
for i,((fn,role,hook,copy),core) in enumerate(zip(shm.SHORTS,CORE),1):
    P("SHORT %d  ·  %s"%(i,role),size=10.5,bold=True,color=GOLD,before=8,after=3)
    KV("File:",fn)
    KV("Opening hook:","“%s”"%hook)
    KV("Core point:",core)
    KV("CTA:","None. These Shorts carry no product CTA; the long-form video holds the "
       "single Field Kit CTA." if i!=4 else "None. Route viewers to the long-form video.")
    KV("Relationship:",REL)

H("Slide summary")
KV("Main slides:","12")
KV("Reveal frames:","24")
KV("Visible slide content changed in this pass:","YES — two authorized precision "
   "corrections, main slides 2 and 4, mirrored on reveal frames 4 and 9. Nothing else.")
P("Slide 2 payoff — old: “NONE OF THAT IS YOUR COMPETENCE. / ALL OF IT IS CONTEXT.” with “And context is learnable.”  New: “NOT EVERY DIFFERENCE IS A CAPABILITY GAP. / MUCH OF IT IS CONTEXT.” with “Some can be researched. Some must be learned through exposure.” The list and visual design are unchanged.",spacing=1.3)
P("Slide 4 closing line — old: “None of these belong to an industry.”  New: “THESE FORMS OF JUDGMENT CAN TRAVEL — / WHEN THE NEW CONTEXT NEEDS THEM.” Same position, same type, two lines instead of one.",spacing=1.3)
KV("Everything else on the decks:","Speaker/editor notes only, with NON-NOTES PARTS "
   "CHANGED: [] on both files apart from the two authorized slide parts each.")
p=P("VISIBLE SLIDE CHANGES LIMITED TO THE TWO AUTHORIZED PRECISION CORRECTIONS "
    "(SLIDES 2 AND 4). GEOMETRY, LAYOUT, THEME, MEDIA AND FONT SIZES UNCHANGED.",
    bold=True,color=NAVY,before=6,after=6,spacing=1.25); shade(p,BAND_NAVY)

H("Production metadata")
KV("Final spoken word count:","{:,}".format(WORDS))
KV("Estimated runtime:","%s at 145 words per minute"%RUNTIME)
KV("H.I.T. payoff timing:","0:24 — “what you can carry, what the new context changes, and "
   "what you actually have to earn”")
KV("U.S. English:","Applied across scripts, Shorts, descriptions, decks and notes — "
   "traveled, practiced, recognized, license.")
KV("Recognition beat:","0:00 — “direct industry experience required”")
KV("Lived evidence from:","0:34 (eight industries) and 0:50 (CISM non-pass)")
KV("Framework arrives:","2:36 — the three Cs, after recognition")
KV("Identity-exit timing:","11:47 — “a way of moving rather than a plan for one change”")

H("Final status")
for t in ("VIDEO 8 — FINAL + LOCKED",
          "v2.2.1 RESEARCH-ALIGNMENT PRECISION PASS COMPLETE",
          "TITLE:  HOW TO SWITCH INDUSTRIES WITHOUT STARTING OVER",
          "THUMBNAIL:  YOUR EXPERIENCE STILL COUNTS"):
    p=P(t,size=11,bold=True,color=NAVY,after=5,spacing=1.25); shade(p,BAND_CREAM)

d.save("Video_8_FINAL_Content_and_Publishing_Summary.docx")
print("summary written")
