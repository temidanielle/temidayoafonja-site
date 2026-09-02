# -*- coding: utf-8 -*-
"""Build the Video 8 H.I.T. final recording and Shorts package.

Formatting helpers are the approved house system carried over unchanged
from the Videos 1-7 packages; only wording, metadata and the manifest
are Video 8 specific.
"""
import os, sys, shutil, zipfile, hashlib
sys.path.insert(0, "/tmp/v8p")
from script_text import LINES, SPOKEN, MARKERS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0x9B,0x2C,0x10)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

ROOT="/tmp/v8p/Video_8_HIT_FINAL"
LF=os.path.join(ROOT,"LONG_FORM"); SH=os.path.join(ROOT,"SHORTS")
shutil.rmtree(ROOT, ignore_errors=True)
os.makedirs(LF); os.makedirs(SH)

TITLE="How to Switch Industries Without Starting Over"
THUMB="YOUR EXPERIENCE STILL COUNTS"
CTA="Capability Formation Field Kit"
CTA_SUB="Capability Formation Field Kit"
CTA_URL="https://temidayoafonja.com/fieldkit"
NEXT="What to Do Before a Layoff Happens"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO 8",size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t,before=20):
    return keep(P(d,t,size=14,bold=True,color=NAVY,before=before,after=8),True)
def H2(d,t,before=13):
    return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=before,after=5),True)

def compress(d, line_spacing=1.18, after_scale=0.80):
    """Tighten one document's vertical rhythm without touching type size,
    weight, colour or wording. Keeps editor briefs off a near-empty last page."""
    for p in d.paragraphs:
        pf=p.paragraph_format
        if pf.line_spacing and 1.0 < pf.line_spacing < 1.4 and pf.line_spacing > line_spacing:
            pf.line_spacing = line_spacing
        if pf.space_after is not None:
            pf.space_after = Pt(round(pf.space_after.pt*after_scale,1))
        if pf.space_before is not None and pf.space_before.pt:
            pf.space_before = Pt(round(pf.space_before.pt*after_scale,1))
    return d

def pairlist(d, items, indent="—  ", gap="        ", after=4, budget=78):
    """Lay short bullet items several to a line, packed by width.

    Every item is kept verbatim; nothing is removed, reworded or reordered,
    and neither type size nor leading changes. This reclaims the unused
    right-hand part of a short-bullet line, which is what keeps the Shorts
    editor brief to two pages. The width budget is measured against the
    inspection renderer's font, which is wider than Calibri, so Word fits at
    least as much per line and never less."""
    lines, cur = [], []
    for it in items:
        trial = cur + [it]
        width = sum(len(indent) + len(x) for x in trial) + len(gap) * (len(trial) - 1)
        if cur and width > budget:
            lines.append(cur); cur = [it]
        else:
            cur = trial
    if cur: lines.append(cur)
    for group in lines:
        keep(P(d, gap.join(indent + x for x in group), after=after))






# ------------------------------------------------- 1. teleprompter DOCX + TXT
TEL_DOCX="Video8TeleprompterScriptwithslidemarkers_HIT_v2.1.docx"
TEL_TXT ="Video8TeleprompterScriptwithslidemarkers_HIT_v2.1.txt"
RD_DOCX ="Video8ReadingScriptnomarkers_HIT_v2.1.docx"
RD_TXT  ="Video8ReadingScriptnomarkers_HIT_v2.1.txt"

d=newdoc(True)
head(d,TITLE,"Video 8  ·  Teleprompter script with slide markers",
     "Spoken script is the large text. A slide marker in a tinted band tells "
     "the editor which slide to bring up; it is not spoken.")
for line in LINES:
    if line.startswith("[SLIDE:"):
        p=P(d,"SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(),size=11,bold=True,
            color=NAVY,before=14,after=14,spacing=1.1)
        shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
    else:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,TEL_DOCX))
tel=[TITLE,"Video 8  ·  Teleprompter script with slide markers",""]
for line in LINES:
    if line.startswith("[SLIDE:"):
        tel += ["", "SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(), ""]
    else: tel += [line, ""]
open(os.path.join(LF,TEL_TXT),"w").write("\n".join(tel).strip()+"\n")

d=newdoc(True)
head(d,TITLE,"Video 8  ·  Reading script, no markers",
     "Spoken language only. No slide markers, no timestamps, no production "
     "directions.")
for line in SPOKEN:
    keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,RD_DOCX))
open(os.path.join(LF,RD_TXT),"w").write("\n\n".join(SPOKEN)+"\n")
print("long-form scripts written")

# --------------------------------------------------- 2. long-form editor brief
# SCRIPT MARKER -> SLIDE NUMBER -> ACTUAL LIVE SLIDE TITLE AFTER THE AUTHORISED
# CORRECTIONS. Mapping is by position; marker 12 is named Watch Next but the
# live slide is the CONTINUE THE SERIES card, which is why this table exists.
MAPPING=[
 (1,"New to a Context",1,
  "CHANGING INDUSTRIES DOES NOT MAKE YOU ENTRY-LEVEL AT EVERYTHING. / "
  "IT MAKES YOU NEW TO A CONTEXT."),
 (2,"What Actually Changes",2,"WHAT ACTUALLY CHANGES"),
 (3,"Capability Context Credential",3,
  "MOVE ONE / CAPABILITY · CONTEXT · CREDENTIAL"),
 (4,"What Travels",4,"CAPABILITY / WHAT TRAVELS"),
 (5,"What Must Be Relearned",5,"CONTEXT / WHAT MUST BE RELEARNED"),
 (6,"What Must Be Earned",6,"CREDENTIAL / WHAT MUST BE EARNED"),
 (7,"Start From the Destination",7,
  "MOVE TWO / START FROM THE DESTINATION, NOT FROM YOUR CV."),
 (8,"Translate Do Not Recite",8,"TRANSLATE, DO NOT RECITE"),
 (9,"Bridge Evidence",9,"MOVE THREE / BRIDGE EVIDENCE"),
 (10,"One Page Three Columns",10,"ONE PAGE, THREE COLUMNS"),
 (11,"Field Kit",11,"CAPABILITY FORMATION / FIELD KIT"),
 (12,"Watch Next",12,
  "CONTINUE THE SERIES / CAREER PORTABILITY / "
  "CAREER PIVOTS · INTERNAL MOVES · GROWTH"),
]

d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 8",size=12,bold=True,color=GOLD,after=2,caps=True)
P(d,TITLE,size=20,bold=True,color=NAVY,after=6,spacing=1.1)
p=P(d,"This document is for the editor. It is NOT Temidayo's teleprompter and "
     "must not be placed on the recording screen.",size=11,italic=True,
     color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"Locked metadata",before=14)
for k,v in (("Title",TITLE),("Thumbnail",THUMB),("CTA",CTA),
            ("CTA URL",CTA_URL),("Watch next",NEXT),
            ("Core teaching","WHAT TRAVELS · WHAT CHANGES · WHAT I MUST EARN"),
            ("Core distinction",
             "Changing industries does not make you entry-level at "
             "everything. It makes you new to a context."),
            ("Credibility boundary","Not everything transfers.")):
    keep(P(d,"%-24s %s"%(k+":",v),size=11,after=5))

H1(d,"The 3 Cs — memorable teaching device",before=14)
p=P(d,"This video has ONE memorable device: THE 3 Cs OF AN INDUSTRY CHANGE — "
     "Capability, Context, Credential. Do not add a second acronym or named "
     "framework.",size=11,bold=True,color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
keep(P(d,"Mapping:",after=5))
for x in ["Capability  →  What travels","Context     →  What changes",
          "Credential  →  What I must earn"]:
    keep(P(d,x,size=11,bold=True,color=GOLD,after=3))
keep(P(d,"The three-column exercise is the APPLICATION of the same mnemonic, "
       "not a second framework.",bold=True,before=4,after=6,spacing=1.25))
keep(P(d,"The spoken script names the three Cs three times: when the method is "
       "named, when it is applied, and when it closes. They always mean "
       "Capability, Context and Credential.",after=6,spacing=1.25))
keep(P(d,"NO POWERPOINT OR REVEAL-DECK CHANGE. The cue below is an editor "
       "overlay on the existing deck, not a slide edit.",bold=True,after=6))
keep(P(d,"On the existing Capability / Context / Credential slide, add one "
       "restrained cue:",after=5))
for x in ["THE 3 Cs OF AN INDUSTRY CHANGE","     C — CAPABILITY",
          "     C — CONTEXT","     C — CREDENTIAL"]:
    keep(P(d,x,size=11,bold=True,color=GOLD,after=3))
keep(P(d,"Use one clean visual card or initial-letter emphasis. Do not animate "
       "each C excessively, add another mnemonic, obscure the existing slide, "
       "change Slide 5's already-approved context correction, or change the "
       "intentional Slide 12 playlist end card.",before=4,after=8,spacing=1.25))

H1(d,"Script marker → slide number → live slide title",before=14)
keep(P(d,"The twelve [SLIDE: ...] markers map to the twelve deck slides BY "
       "POSITION, in order, one to one. Marker 12 is named Watch Next, but the "
       "live slide 12 is the CONTINUE THE SERIES / CAREER PORTABILITY card and "
       "carries no video title. Follow the slide number in this table, not the "
       "marker name.",size=10.5,color=DIM,after=10,spacing=1.25))
tbl=d.add_table(rows=1, cols=4)
tbl.style="Table Grid"; tbl.autofit=False
COLW=(0.45,2.30,0.55,3.10)
for cell,wdt,label in zip(tbl.rows[0].cells,COLW,
                          ("MK","SCRIPT MARKER","SLIDE","LIVE SLIDE TITLE")):
    cell.width=Inches(wdt)
    cp=cell.paragraphs[0]; cp.paragraph_format.space_after=Pt(2)
    r=cp.add_run(label); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=NAVY
    shade(cp,BAND_NAVY)
for mk,name,sl,title in MAPPING:
    cells=tbl.add_row().cells
    for cell,wdt,val,bold in zip(cells,COLW,(str(mk),name,str(sl),title),
                                 (False,False,False,True)):
        cell.width=Inches(wdt)
        cp=cell.paragraphs[0]
        cp.paragraph_format.space_after=Pt(2); cp.paragraph_format.line_spacing=1.1
        r=cp.add_run(val); r.font.size=Pt(8.5); r.bold=bold
        if bold: r.font.color.rgb=NAVY
P(d,"",size=6,after=2)
keep(P(d,"Validated in the live files: 12 markers, 12 slides, one-to-one and "
       "in order, no duplicate mapping, no missing slide, no reordered slide.",
       size=10.5,italic=True,color=DIM,before=4,after=8,spacing=1.25))

H1(d,"Authorised slide correction — applied",before=14)
p=P(d,"AUTHORISED AND APPLIED. Slide 5 carried stale conceptual framing. Text "
     "only changed; visual system, typography family, colours, composition and "
     "hierarchy are unchanged.",size=11,bold=True,color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
keep(P(d,"Slide 5 (gold emphasis line)",size=11,bold=True,color=NAVY,before=6,after=3))
keep(P(d,"FROM:  It feels like a competence gap. It is an information gap.",
       size=10.5,after=3))
keep(P(d,"TO:    IT CAN FEEL LIKE A COMPETENCE GAP. / SOME CONTEXT CAN BE "
       "RESEARCHED. / SOME MUST BE LEARNED THROUGH EXPOSURE.",size=10.5,
       bold=True,after=5,spacing=1.25))
keep(P(d,"Set as three lines inside the existing text box at the existing 20pt "
       "Montserrat Bold in the existing gold. Measured against the real "
       "embedded font first: the widest line is 7.06in inside an 11.11in box, "
       "so nothing wraps, and the block ends at 6.71in of a 7.5in slide. The "
       "replaced line was 7.98in on one line, so the block is narrower than "
       "what it replaces.",size=10.5,color=DIM,after=8,spacing=1.25))
keep(P(d,"Reveal frames corrected: 11 only. Frame 10 is the title-only build "
       "of the same slide and never carried the line. Main slides 1 to 4 and 6 "
       "to 12 and the other twenty-three reveal frames are byte-identical.",
       bold=True,after=8,spacing=1.25))
p=P(d,"SLIDE 12 IS INTENTIONAL AND STAYS UNCHANGED. Confirmed by Temidayo. "
     "The live slide 12 is CONTINUE THE SERIES / CAREER PORTABILITY / CAREER "
     "PIVOTS · INTERNAL MOVES · GROWTH — a deliberate series and playlist end "
     "card, not a stale Watch Next error. It carries no video title by design. "
     "The spoken script and the end screen carry the route to Video 9. Do not "
     "correct this slide.",size=11,bold=True,color=RED,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"First 30 seconds — H.I.T.",before=14)
P(d,"H = Hook. I = Interest. T = Trust. The opening must match the title and "
    "thumbnail promise, create immediate recognition, use visual movement, "
    "establish concrete trust quickly, make clear that Temidayo is not "
    "promising that all experience transfers, and deliver the viewer payoff by "
    "roughly 20 to 30 seconds.",after=14,spacing=1.25)

def beat(t,anchor,body):
    H2(d,t,before=10)
    if anchor:
        p=P(d,"Spoken:  “%s”"%anchor,size=10.5,italic=True,color=DIM,after=8)
        shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))

beat("0:00–0:05  ·  HOOK",
     "Changing industries does not make you entry-level at everything. It "
     "makes you new to a context.",
     ["On-screen:  NEW INDUSTRY ≠ STARTING OVER"])
beat("0:05–0:16  ·  TRUST + INTEREST",
     "Across nearly two decades, I’ve worked across eight industries and "
     "sectors…",
     ["On-screen:  8 INDUSTRIES + SECTORS","Then:  WHAT TRAVELS?  WHAT CHANGES?",
      "No employer-logo montage."])
beat("0:16–0:24  ·  TRUST BOUNDARY",
     "When I moved toward cybersecurity and privacy, I prepared for the CISM "
     "exam and didn’t pass the first time.",
     ["On-screen:  NOT EVERYTHING TRAVELS",
      "Visual tone: calm, factual, non-dramatic."])
p=P(d,"Do not show a fake exam result, a score, a failed stamp or any "
     "sensational failure imagery. Do not dramatize the non-pass.",size=11,
     bold=True,color=RED,before=6,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
beat("0:24–0:30  ·  PAYOFF",
     "So I’m going to help you separate three things: what you can carry, what "
     "the new context changes, and what you actually have to earn.",
     ["Progressive reveal:","     WHAT TRAVELS","     WHAT CHANGES",
      "     WHAT I MUST EARN"])

H1(d,"Editing principle",before=14)
keep(P(d,"Let visual structure carry the three-part distinction. Temidayo's "
       "voice carries meaning.",after=8,spacing=1.25))
P(d,"Avoid:",after=5)
pairlist(d,["airport/travel metaphors;","suitcase imagery;",
 "passport imagery;","generic career ladder;",
 "“new industry = total restart” stock visuals;","hyperactive transitions."])

H1(d,"Factual boundary",before=14)
keep(P(d,"Approved personal evidence:",after=6))
for x in ["eight industries and sectors across roughly eighteen years;",
 "a career crossing accounting/audit, cyber/privacy, people/employee "
 "experience and enterprise transformation;",
 "CISM exam preparation and a first-attempt non-pass."]:
    keep(P(d,"—  "+x,after=5,spacing=1.25))
p=P(d,"Do not invent an exam score, an exam date, any further attempts, a "
     "later passing result, employer-specific details or confidential industry "
     "information. Do not imply that all prior experience transfers.",size=11,
     bold=True,color=RED,before=6,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
keep(P(d,"Some judgment travels. Context does not travel for free. Some "
       "knowledge, exposure or credentials must genuinely be earned.",
       bold=True,after=10,spacing=1.25))

H1(d,"CTA and watch next",before=14)
keep(P(d,"One offer only: %s — %s"%(CTA,CTA_URL),after=5,spacing=1.25))
keep(P(d,"Do not add Keep the Proof or the Career Decision Evidence Check.",
       bold=True,after=8))
keep(P(d,"Watch next: %s  (Video 9)"%NEXT,bold=True,after=5))
keep(P(d,"Do not leave Subscribe as the only end-screen element.",bold=True,
       color=RED,before=4,after=8))
compress(d, 1.14, 0.56)
d.save(os.path.join(LF,"Video_8_EDITOR_ONLY_HIT_Brief_v2.1.docx"))
print("editor brief written")

# The twelve working chapter lines, defined once and reused by the publishing
# package, its reference section and the separate description-only document.
# Offsets are script-derived at 145 wpm from the canonical spoken text; the
# first sits at 00:00 and the rest at markers 2 to 12. They are estimates.
CHAPTERS=[("00:00","New Industry Does Not Mean Starting Over"),
 ("01:27","What Actually Changes"),
 ("02:18","The 3 Cs: Capability, Context and Credential"),
 ("02:30","What Travels"),
 ("03:18","What Must Be Relearned"),
 ("04:12","What Must Be Earned"),
 ("04:58","Start From the Destination"),
 ("05:54","Translate, Do Not Recite"),
 ("06:54","Build Bridge Evidence"),
 ("08:20","The Three-Column Test"),
 ("09:55","Capability Formation Field Kit"),
 ("10:31","What to Do Before a Layoff Happens")]
CHAPTER_LINES=["%s %s"%(t,c) for t,c in CHAPTERS]

PRIMARY="how to switch industries without starting over"
SUPPORTING=("how to switch industries · change industries · career change · "
 "how to change careers without starting over · transferable skills · career "
 "transition · new industry · career portability · experienced professionals · "
 "career change after 40")
TAGS=("how to switch industries without starting over, how to switch "
 "industries, change industries, career change, how to change careers without "
 "starting over, transferable skills, career transition, new industry, career "
 "portability, experienced professionals, career change after 40, capability "
 "formation, Temidayo Afonja")
DESC=[
 "Changing industries does not mean every part of your experience disappears "
 "— but it also does not mean everything transfers.",
 "In this video, I use the three Cs of an industry change—Capability, "
 "Context and Credential—to separate three different problems before you "
 "move:",
 "✨ What travels — the judgment, decisions and patterns that remain useful "
 "when the setting changes.",
 "✨ What changes — the language, stakeholders, incentives, regulation, "
 "operating rhythm and risks you need to learn in the new context.",
 "✨ What I must earn — the credential, direct exposure or practice that "
 "previous experience cannot replace.",
 "Across nearly two decades, I have worked across eight industries and "
 "sectors. That taught me that a credible industry move requires both "
 "confidence and humility: knowing what you can carry and being precise about "
 "what you still need to learn.",
 "I also share one example from my move toward cybersecurity and privacy, "
 "where preparing for the CISM exam — and not passing the first time — "
 "reminded me that not everything travels simply because you are experienced.",
 "The goal is not to prove that your whole career is transferable.",
 "It is to answer:",
 "“What can I credibly carry into this context, what must I learn, and what "
 "still has to be earned?”","",
 "🧭 CAPABILITY FORMATION FIELD KIT",
 "Examine what your current work has built, what appears portable and where "
 "you still need development, evidence or context:",
 CTA_URL,"",
 "⏱️ CHAPTERS"]+CHAPTER_LINES+["",
 "▶️ WATCH NEXT", NEXT, "[ADD VIDEO 9 LINK WHEN LIVE]","",
 "🔗 CONNECT AND EXPLORE",
 "Website:","https://temidayoafonja.com",
 "LinkedIn:","https://www.linkedin.com/in/temidayo-afonja",
 "Substack:","https://temidayoafonja.substack.com","",
 "#CareerChange #CareerGrowth #CapabilityFormation"]
PINNED=["Which of the three Cs is the biggest gap in the industry you want "
 "to enter?",
 "1. Capability — what travels",
 "2. Context — what changes",
 "3. Credential — what I must earn",
 "You do not need to share confidential details.",
 "The useful exercise is to separate what your experience genuinely supports "
 "from what the new context still requires you to learn.",
 "If you want a fuller private read of what your current work has built, the "
 "Capability Formation Field Kit is here:", CTA_URL]

EMOJI_NOTE=("The restrained emoji system is part of the approved standard: ✨ "
  "teaching points, 🧭 CTA/resource, ⏱️ chapters, ▶️ Watch Next, 🔗 Connect "
  "and Explore. Do not remove them and do not add more.")

def description_block(d, heading_before=14, upload_doc=False):
    """The copy-ready description, its end marker and the internal note.

    upload_doc=True moves the editorial emoji instruction ABOVE an explicit
    COPY-READY ... BEGIN marker, so nothing internal sits inside or immediately
    before the block a person selects when pasting into YouTube."""
    if upload_doc:
        H1(d,"INTERNAL NOTE — DO NOT PASTE INTO YOUTUBE",before=heading_before)
        p=P(d,EMOJI_NOTE,size=10.5,italic=True,color=RED,after=12,spacing=1.25)
        shade(p,BAND_CREAM); keep(p)
        p=keep(P(d,"COPY-READY YOUTUBE DESCRIPTION — BEGIN",size=11,bold=True,
                 color=NAVY,before=14,after=12,spacing=1.2))
        shade(p,BAND_NAVY)
    else:
        H1(d,"Description",before=heading_before)
        keep(P(d,EMOJI_NOTE,size=10.5,italic=True,color=DIM,after=10))
    for para in DESC:
        keep(P(d,para if para else " ",after=7 if para else 3))
    keep(P(d,"— END OF THE COPY-READY DESCRIPTION —",size=10,bold=True,
           color=DIM,before=14,after=12,spacing=1.2))
    H1(d,"Internal note — do not paste into YouTube",before=14)
    p=P(d,"WORKING ESTIMATES — EDITOR MUST REPLACE FROM FINAL CUT",size=11,
        bold=True,color=RED,after=6,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    p=P(d,"The chapter timestamps above are script-derived, not final. Replace "
        "every one of them using the finished cut before publication. Do not "
        "force the edit to match these estimates.",size=10.5,bold=True,
        italic=True,color=RED,after=10,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    H1(d,"Working chapters — reference copy",before=14)
    keep(P(d,"Identical to the twelve chapter lines inside the description "
           "above.",size=10.5,italic=True,color=DIM,after=8))
    for line in CHAPTER_LINES: keep(P(d,line,size=11,after=4))



# ----------------------------------------------------- 3. publishing package
d=newdoc()
head(d,TITLE,"Video 8  ·  Publishing package",
     "Everything needed to upload. Working timestamps must be replaced with "
     "real ones from the finished edit.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14); P(d,PRIMARY,after=10)
H1(d,"Supporting search language",before=14); P(d,SUPPORTING,after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only. Do not put the full tag field in the "
       "public description.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
H1(d,"Watch next",before=14)
keep(P(d,"%s  (Video 9)"%NEXT,bold=True,after=5))
keep(P(d,"Live slide 12 is the CONTINUE THE SERIES card and carries no video "
       "title. See the EDITOR ONLY brief.",size=10.5,color=DIM,after=8,
       spacing=1.25))
compress(d)
d.save(os.path.join(LF,"Video_8_Publishing_Package_HIT_v2.1.docx"))

# ------------------------------- 3b. separate description-only document
d=newdoc()
head(d,TITLE,"Video 8  ·  YouTube description",
     "Upload copy only. Everything below the end marker is internal and must "
     "not be pasted into YouTube.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=10)
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14); P(d,PRIMARY,after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"Watch next",before=14); keep(P(d,"%s  (Video 9)"%NEXT,bold=True,after=8))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
compress(d)
DESC_DOC="/tmp/v8p/Video_8_YouTube_Description_HIT.docx"
d.save(DESC_DOC)
print("publishing package and description-only document written")

# ---------------------------------------------------------------- 4. Shorts
from shorts_text import SHORTS
LABELS=["SHORT 1","SHORT 2","SHORT 3","SHORT 4"]
for (fn,role,hook,copy),label in zip(SHORTS,LABELS):
    d=newdoc(True)
    P(d,"VIDEO 8 SHORT",size=10,bold=True,color=GOLD,after=4,caps=True)
    P(d,label,size=20,bold=True,color=NAVY,after=8,spacing=1.1)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    keep(P(d,"Verbal hook:  “%s”"%hook,size=11,color=DIM,after=5))
    keep(P(d,"Related long-form:  %s"%TITLE,size=11,color=DIM,after=10))
    H1(d,"RECORDING COPY",before=12)
    for line in copy:
        keep(P(d,line,size=13.5,color=INK,after=10,spacing=1.5))
    d.save(os.path.join(SH,fn))

# ------------------------------------------------------ 5. Shorts editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 8 — FOUR STANDALONE SHORTS",size=18,bold=True,color=NAVY,
  after=8,spacing=1.1)
p=P(d,"This document is for the editor. It is separate from the four Short "
     "recording documents and must not be placed on Temidayo's recording "
     "screen.",size=11,italic=True,color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)
H1(d,"How these are produced",before=14)
keep(P(d,"These are separately recorded 9:16 Shorts. They are NOT excerpts cut "
       "from the long-form video.",bold=True,after=10))
P(d,"Each should have:",after=6)
pairlist(d,["an immediate verbal hook;","a matching on-screen hook;",
 "accurate mobile-safe captions;","restrained editorial pacing;",
 "Video 8 added as the YouTube Related Video when available."])

def short(label,role,onscreen,body):
    H1(d,label,before=14)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    p=keep(P(d,"On-screen:  %s"%onscreen,size=11,bold=True,color=GOLD,after=8))
    shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))
    keep(P(d,"Related Video:  Video 8",size=10.5,color=DIM,before=4,after=6))

short("SHORT 1","Recognition","NEW INDUSTRY ≠ ENTRY LEVEL AT EVERYTHING",
 ["Visual:  CURRENT CONTEXT  →  NEW CONTEXT",
  "Avoid the career-ladder cliché."])
short("SHORT 2","Distinction","CAPABILITY / CONTEXT / CREDENTIAL",
 ["Visual: three distinct columns."])
short("SHORT 3","Personal proof","NOT EVERYTHING TRAVELS",
 ["Secondary:","     8 INDUSTRIES + SECTORS",
  "     CISM: FIRST ATTEMPT NON-PASS",
  "Do not show a fake test score, a failed stamp, or a credential badge "
  "Temidayo does not hold."])
short("SHORT 4","Practical action",
 "WHAT TRAVELS / WHAT CHANGES / WHAT I MUST EARN",
 ["Visual: clean three-column worksheet reveal."])

H1(d,"All Shorts — visual boundaries",before=14)
P(d,"Do not use:",after=5)
pairlist(d,["airport or suitcase metaphors;","passport imagery;",
 "stock “career switch” montage;","generic career ladder;",
 "fake shock expressions;","fake exam results or scores;",
 "a failed stamp;","credential badges not owned;",
 "hyperactive transitions;","trendy caption templates."],after=3)
compress(d, 1.18, 0.62)
d.save(os.path.join(SH,"Video_8_Shorts_EDITOR_ONLY_HIT_Brief.docx"))
print("shorts and shorts editor brief written")

# ---------------------------------------------------------------- 6. README
ZIPNAME="Video_8_HIT_FINAL_Recording_and_Shorts_Package.zip"
FILES=(["LONG_FORM/"+f for f in sorted(os.listdir(LF))]
      +["SHORTS/"+f for f in sorted(os.listdir(SH))])
R=["VIDEO 8 — H.I.T. FINAL RECORDING PACKAGE","",
 "Title:             %s"%TITLE,
 "Thumbnail:         %s"%THUMB,
 "CTA:               %s"%CTA,
 "CTA URL:           %s"%CTA_URL,
 "Watch next:        %s"%NEXT,"",
 "Long-form:         Revised under H.I.T.",
 "Memorable device:  THE 3 Cs OF AN INDUSTRY CHANGE — Capability, Context,",
 "                   Credential. Capability maps to what travels, Context to",
 "                   what changes, Credential to what must be earned. Named",
 "                   three times in the long-form script and once in Short 2.",
 "                   The three-column exercise is the application of the same",
 "                   mnemonic, not a second framework. NO PowerPoint or",
 "                   reveal-deck change: the 3 Cs cue is an editor overlay.","",
 "Approved proof:",
 "  - roughly eighteen-year cross-context career",
 "  - eight industries/sectors",
 "  - CISM preparation and first-attempt non-pass","",
 "Core distinction:",
 "  Some judgment travels.",
 "  Context does not travel for free.",
 "  Some knowledge, exposure or credentials must genuinely be earned.","",
 "Shorts:            Four separately recorded vertical scripts.",
 "Editor",
 "instructions:      Separated from all recording copy.",
 "Description-only",
 "document:          Video_8_YouTube_Description_HIT.docx, created separately",
 "                   OUTSIDE this package ZIP.","",
 "Verified in the",
 "live files:        12 main slides, 24 reveal-build frames.","",
 "Slide correction:  Slide 5 carried the stale conceptual framing",
 "                   \"It feels like a competence gap. It is an information",
 "                   gap.\" It was replaced under explicit authorisation with",
 "                   the three approved statements. Text only; visual system,",
 "                   typography family, colours, composition and hierarchy",
 "                   unchanged. Reveal frame 11 carried the same line and",
 "                   received the same correction. Main slides 1-4 and 6-12",
 "                   and the other twenty-three reveal frames are",
 "                   byte-identical.",
 "Slide 12:          INTENTIONAL. UNCHANGED. Confirmed by Temidayo. The live",
 "                   slide 12 is CONTINUE THE SERIES / CAREER PORTABILITY, a",
 "                   deliberate series and playlist end card, not a stale",
 "                   Watch Next error. It carries no video title by design.",
 "                   The spoken script and the end screen carry the route to",
 "                   Video 9. Do not correct this slide.",
 "Thumbnail:         UNCHANGED.","",
 "-"*70,"","WHAT EACH FILE IS","",
 "LONG_FORM/","",
 "  "+TEL_DOCX,
 "  "+TEL_TXT,
 "      Temidayo's recording copy. Spoken script in large text; slide markers",
 "      in tinted bands. The markers are not spoken.","",
 "  "+RD_DOCX,
 "  "+RD_TXT,
 "      The same spoken words with the slide markers removed.","",
 "  Video_8_EDITOR_ONLY_HIT_Brief_v2.1.docx",
 "      For the editor. Locked metadata, the script-marker to slide-number",
 "      mapping table, the Slide 5 correction record and the Slide 12",
 "      question, the H.I.T. first-30-second plan, the editing principle and",
 "      the factual boundary. Not for the teleprompter.","",
 "  Video_8_Publishing_Package_HIT_v2.1.docx",
 "      Title, thumbnail, search language, the copy-ready description with",
 "      its approved emoji system, working chapter estimates, pinned comment",
 "      and the tag field.","",
 "SHORTS/","",
 "  Four recording documents, one per Short. These contain Temidayo's",
 "  recording copy and no editor directions.","",
 "  Video_8_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "      For the editor. On-screen hooks and visual treatment for all four.","",
 "-"*70,"","ALL FILES IN THIS PACKAGE","",]
for f in FILES: R.append("  "+f)
R+=["  README_FINAL.txt","  SHA256SUMS.txt","",
 "  Video_8_YouTube_Description_HIT.docx is deliberately NOT in this ZIP.","",
 "-"*70,"","WORKING CHAPTER TIMESTAMPS","",
 "The twelve chapter timestamps are WORKING ESTIMATES derived from the script",
 "at 145 words per minute. They were not measured from an edit. The editor",
 "must replace every one of them from the finished cut before publishing.","",
 "-"*70,"","FACT AND PROOF BOUNDARY","",
 "No exam score, exam date, further attempt, later passing result,",
 "employer-specific detail or confidential industry information is invented",
 "anywhere in this package. Nothing claims that all prior experience",
 "transfers.","",
 "-"*70,"","CHECKSUMS","",
 "SHA256SUMS.txt covers the other 12 user-facing files in this package. It",
 "does not hash itself, and it carries no ZIP checksum. The archive's own",
 "SHA-256 is in the sibling file:",
 "  "+ZIPNAME+".sha256","",
 "-"*70,"","WHAT WAS NOT CHANGED","",
 "Every website file, every product, the approved thumbnail and every other",
 "video are unchanged. The only visual change in this pass is the authorised",
 "Slide 5 text correction, applied identically to main slide 5 and reveal",
 "frame 11.",""]
open(os.path.join(ROOT,"README_FINAL.txt"),"w").write("\n".join(R))

# ------------------------------------------- 7. checksums and the master ZIP
MANIFEST=[
 "LONG_FORM/"+TEL_DOCX,
 "LONG_FORM/"+TEL_TXT,
 "LONG_FORM/"+RD_DOCX,
 "LONG_FORM/"+RD_TXT,
 "LONG_FORM/Video_8_EDITOR_ONLY_HIT_Brief_v2.1.docx",
 "LONG_FORM/Video_8_Publishing_Package_HIT_v2.1.docx",
 "SHORTS/Video_8_Short_1_New_Industry_Not_Starting_Over.docx",
 "SHORTS/Video_8_Short_2_Stop_Calling_Everything_Transferable.docx",
 "SHORTS/Video_8_Short_3_Not_Everything_Travels.docx",
 "SHORTS/Video_8_Short_4_Three_Columns.docx",
 "SHORTS/Video_8_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "README_FINAL.txt",
]
SUMS="SHA256SUMS.txt"
ZIP="/tmp/v8p/"+ZIPNAME

def sha256(p):
    h=hashlib.sha256()
    with open(p,"rb") as fh:
        for b in iter(lambda: fh.read(1<<20), b""): h.update(b)
    return h.hexdigest()

for m in MANIFEST:
    assert os.path.isfile(os.path.join(ROOT,m)), "missing from build: "+m
on_disk=set()
for dp,dn,fn in os.walk(ROOT):
    dn[:]=[x for x in dn if x!="__pycache__"]
    for f in fn:
        if f.endswith(".pyc"): continue
        on_disk.add(os.path.relpath(os.path.join(dp,f),ROOT).replace(os.sep,"/"))
unexpected=sorted(on_disk-set(MANIFEST)-{SUMS})
assert not unexpected, "unexpected files in package directory: %r"%unexpected

L=["# VIDEO 8 - H.I.T. FINAL RECORDING PACKAGE",
   "# SHA-256 of the 12 user-facing files in this package.",
   "# SHA256SUMS.txt cannot hash itself. The master ZIP cannot contain its own",
   "# checksum either; it is published in the sibling file",
   "# "+ZIPNAME+".sha256",
   "# Video_8_YouTube_Description_HIT.docx sits outside this package; its",
   "# SHA-256 is reported in the delivery summary.",""]
for m in MANIFEST: L.append("%s  %s"%(sha256(os.path.join(ROOT,m)),m))
open(os.path.join(ROOT,SUMS),"w").write("\n".join(L)+"\n")

if os.path.exists(ZIP): os.remove(ZIP)
with zipfile.ZipFile(ZIP,"w",zipfile.ZIP_DEFLATED) as z:
    for m in MANIFEST+[SUMS]:
        z.write(os.path.join(ROOT,m), "Video_8_HIT_FINAL/"+m)
zsha=sha256(ZIP)
open(ZIP+".sha256","w").write("%s  %s\n"%(zsha,os.path.basename(ZIP)))

PROV="/tmp/v8p/_source"
shutil.rmtree(PROV,ignore_errors=True); os.makedirs(PROV)
for f in ("script_text.py","shorts_text.py","build.py","qa.py",
          "verify_canonical.py","fix_slide5.py"):
    src="/tmp/v8p/"+f
    if os.path.isfile(src): shutil.copy2(src, os.path.join(PROV,f))
CANON_SRC=("/root/.claude/uploads/f121668d-e262-5eb8-9b22-0eaa1006a361/"
           "7c400758-Video_8_Code_Prompt_HIT_Final.txt")
if os.path.isfile(CANON_SRC):
    shutil.copy2(CANON_SRC, os.path.join(PROV,"Video_8_Code_Prompt_HIT_Final.txt"))
print("ZIP sha256:",zsha)
print("description-only doc sha256:",sha256(DESC_DOC))
