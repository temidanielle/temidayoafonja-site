# -*- coding: utf-8 -*-
"""Author the Video 8 production package DOCX.

Single source of truth for Video 8. The teleprompter DOCX and the clean
recording TXT are generated from the DOCX this writes, so the three cannot
drift. Structure mirrors the approved Video 7 package section for section.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
OUT = os.path.join(ROOT, "YouTube_Video_8_Production_Package_New_Industry.docx")

NAVY = RGBColor(0x0F, 0x23, 0x46)
GOLD = RGBColor(0x8A, 0x6D, 0x1E)
DIM  = RGBColor(0x5A, 0x6B, 0x82)
INK  = RGBColor(0x1A, 0x1A, 0x1A)

TITLE = "How to Move Into a New Industry Without Starting Over"
KEYWORD = "how to move into a new industry"
CTA_NAME = "Capability Formation Field Kit"
CTA_URL = "https://temidayoafonja.com/fieldkit"
THUMB = "YOUR EXPERIENCE STILL COUNTS"

SCRIPT = [
 ("0:00-0:30 | HOOK + PROMISE", [
  "Changing industries does not make you entry-level at everything. It makes "
  "you new to a context.",
  "Those are two different problems with two different solutions, and most of "
  "the advice you will be given collapses them into one. Either you are told "
  "to start over, or you are told that everything you have done will carry "
  "across on its own. Neither prepares you for the conversation you are "
  "actually going to have.",
  "So over the next ten minutes or so, I want to help you sort your own "
  "experience into three piles: what travels with you, what changes when the "
  "context changes, and what you will genuinely have to earn in the new "
  "field. Then how to show the first pile to someone who has no particular "
  "reason yet to believe you.",
 ]),
 ("0:30-1:15 | EARLY PROOF + INTRODUCTION", [
  "I have made this move more than once. Over nearly two decades I have "
  "worked across eight industries and sectors, including at Deloitte, EY and "
  "PwC, and moved from Consultant to Head of Function.",
  "What I want you to take from that is not that it was smooth, because it "
  "was not. Some of what I brought with me was useful immediately. Some of it "
  "was no help at all in the room I had just walked into. When I moved toward "
  "cybersecurity and privacy, I prepared for the CISM exam and did not pass "
  "the first time. I mention that because it is the honest shape of this "
  "problem. Judgment travels. Context does not travel for free.",
  "I am Temidayo Afonja. I help experienced professionals work out what they "
  "can carry across roles, functions, employers and industries, so that the "
  "next move does not feel like starting from zero.",
 ]),
 ("1:15-2:15 | WHAT ACTUALLY CHANGES", [
  "Let me be precise about what changes when you move industries, because "
  "“everything” is not true and neither is “nothing.”",
  "The language changes. The same idea goes by a different name, and using "
  "the old name marks you as an outsider faster than any gap on your CV. The "
  "stakeholders change: who has to be convinced, who can quietly stop "
  "something, who is actually accountable when it goes wrong.",
  "The incentives change, and this is the one people underestimate. In one "
  "industry the pressure is quarterly. In another it is a regulator, or a "
  "safety record, or a clinical outcome. That pressure shapes every decision "
  "being made around you, including the ones nobody explains to you.",
  "The regulation changes. The operating rhythm changes — what counts as "
  "a fast decision, what a normal cycle looks like, when the year actually "
  "begins. And the risks change, so what everyone is quietly worried about is "
  "different, and until you know what it is you will keep proposing sensible "
  "things that land badly.",
  "None of that is your competence. All of it is context. And context is "
  "learnable, which makes it a far smaller problem than the one you thought "
  "you had.",
 ]),
 ("2:15-4:10 | MOVE ONE — SEPARATE CAPABILITY, CONTEXT AND CREDENTIAL", [
  "That brings me to the first move, which is to stop carrying your "
  "experience around as one undifferentiated lump. Separate it into "
  "capability, context and credential.",
  "Capability is judgment or ability that stays useful when the setting "
  "changes. Finding the real decision in a room full of stated positions. "
  "Structuring a problem nobody has framed yet. Knowing when a number is too "
  "clean to be true. Holding a position when the room would rather you did "
  "not. Those do not belong to an industry.",
  "Context is what the new field knows and you do not — its language, "
  "its stakeholders, its incentives, its regulation, its rhythm and its "
  "risks. From the inside it feels like a competence gap. It is an "
  "information gap, and it can close when you are deliberate about learning "
  "it.",
  "Credential is the formal evidence or permission the destination may "
  "require. Sometimes it is genuinely required and there is no way around it, "
  "so you plan for it. Sometimes it is not required but it shortens the "
  "conversation, which is a real benefit worth weighing. And sometimes it is "
  "neither, and people pursue it anyway because it feels like progress when "
  "the rest of the move feels stuck. That is an understandable thing to do "
  "and an expensive one.",
  "The reason to separate the three is that they have different remedies. A "
  "certificate will not fix a context gap. Enthusiasm will not fix a missing "
  "licence. And you do not need to rebuild a capability you already have "
  "simply because the room is unfamiliar.",
 ]),
 ("4:10-6:15 | MOVE TWO — MATCH JUDGMENT TO DESTINATION PROBLEMS", [
  "The second move is to match the judgment you already have to the problems "
  "the destination industry actually needs solved.",
  "This is where most industry-change pitches come apart, because people "
  "reach for adjectives. Adaptable. Strategic. A fast learner. Those describe "
  "how you would like to be seen. They give the person across from you "
  "very little to evaluate, and experienced interviewers have heard them many "
  "times.",
  "Work in the other direction. Start with the destination rather than with "
  "your CV. What does that industry keep failing to solve? What is expensive "
  "there, or slow, or contested? You can find a surprising amount of that in "
  "trade press, earnings calls, regulator commentary, and in how job "
  "descriptions read when a team is frustrated.",
  "Then come back to your own history and ask what you already know how to do "
  "about problems shaped like that. Which of those decisions have you made "
  "before, under comparable pressure? Which patterns do you recognise early "
  "because you have watched them play out? Which stakeholders have you had to "
  "align when they wanted different things? Which constraints have you worked "
  "inside — a budget, a regulator, a safety requirement, a board that "
  "had already decided?",
  "The unit that transfers is not your job title and it is not your task "
  "list. It is judgment, with evidence attached.",
  "Be equally clear about where that judgment still needs translation. Saying "
  "you have managed a regulated change programme is transferable. Saying you "
  "know how their regulator behaves is not, unless you do. Claiming the "
  "second when you only have the first is the quickest way to lose a room you "
  "had already won.",
 ]),
 ("6:15-8:30 | MOVE THREE — BUILD BRIDGE EVIDENCE AND A LEARNING PLAN", [
  "The third move is to close the credibility gap using evidence you are "
  "actually permitted to have.",
  "The person considering you is not being unreasonable. They are being asked "
  "to take a risk on somebody who has not done this in their world. Your job "
  "is to make that risk smaller and more specific, not to argue that it is "
  "not there.",
  "A relevant project or assignment counts, even if it sat inside your old "
  "industry, and especially if it touched the same kind of problem. "
  "Cross-functional work counts, because it shows you can operate usefully in "
  "a room where you are not the expert.",
  "Research counts once you have turned it into a point of view. Not "
  "“I have read about the sector,” but “here is what I think "
  "is happening in it, here is what I would want to check first, and here is "
  "where I might be wrong.” That is a different quality of statement, "
  "and the people you are talking to can tell the difference immediately.",
  "A case study or work sample built from non-confidential information "
  "counts. A course or credential counts where it is genuinely relevant, and "
  "is dead weight where it is not. Conversations with practitioners count, "
  "and they do two jobs at once: they give you the language, and they tell "
  "you which of your assumptions are wrong before you say them out loud in an "
  "interview.",
  "A clear first-ninety-days learning plan can also reduce the credibility "
  "gap. It shows you understand there is something here to learn, which is "
  "the opposite of the overconfidence they are quietly screening for.",
  "What none of this permits is inventing experience or implying more "
  "familiarity than you have. The gap is real. You are making it smaller and "
  "legible. You are not pretending it is not there.",
 ]),
 ("8:30-9:30 | THE THREE COLUMNS", [
  "So here is the exercise, and it fits on one page. Three columns.",
  "What travels: the judgment, the decisions, the patterns and the ways of "
  "working that stay useful when the setting changes.",
  "What changes: the language, the stakeholders, the incentives, the "
  "regulation, the rhythm and the risks you will have to learn.",
  "What I must earn: the credential, the licence or the specific exposure "
  "that no amount of adjacent experience substitutes for.",
  "You may find the first column is longer than you feared and the third is "
  "shorter than you assumed. That is the reason to write it down instead of "
  "carrying it around as a feeling.",
  "Then build one sentence out of it, in your own words. Something close to: "
  "I am new to this industry, but I am not new to this problem. Here is the "
  "evidence. Here is what I am learning right now.",
  "That sentence declines the false modesty of starting over without "
  "pretending the context comes free. And it gives the other person something "
  "they can actually check, which is what they wanted from the beginning.",
 ]),
 ("9:30-10:05 | PRIMARY CTA", [
  "If you want a structured way to work through this, the Capability "
  "Formation Field Kit is built for exactly it. It helps you look at what "
  "your work has formed in you, how portable that is, and where you still "
  "need development or evidence. It will not tell you which industry is "
  "hiring, and it does not replace researching the field you are moving "
  "toward — that part stays yours. You will find it at "
  "temidayoafonja.com/fieldkit.",
 ]),
 ("10:05-10:20 | CONTINUE THE SERIES", [
  "When you are ready for the next step, continue with the Career Portability "
  "playlist.",
 ]),
]

SPOKEN = [p for _, ps in SCRIPT for p in ps]
WORDS = sum(len(p.split()) for p in SPOKEN)

SLIDE_COPY = [
 ("1", "New to a context", "2",
  "CHANGING INDUSTRIES DOES NOT MAKE YOU ENTRY-LEVEL AT EVERYTHING.  |  IT "
  "MAKES YOU NEW TO A CONTEXT."),
 ("2", "What actually changes", "2",
  "WHAT ACTUALLY CHANGES  |  Language  |  Stakeholders  |  Incentives  |  "
  "Regulation  |  Operating rhythm  |  Risks  |  NONE OF THAT IS YOUR "
  "COMPETENCE.  |  ALL OF IT IS CONTEXT.  |  And context is learnable."),
 ("3", "Capability, context, credential", "3",
  "MOVE ONE  |  1 CAPABILITY — Judgment that stays useful when the "
  "setting changes.  |  2 CONTEXT — What the new field knows and you do "
  "not — yet.  |  3 CREDENTIAL — Formal evidence or permission the "
  "destination requires."),
 ("4", "What travels", "2",
  "CAPABILITY  |  WHAT TRAVELS  |  Finding the real decision in a room of "
  "stated positions.  |  Structuring a problem nobody has framed yet.  |  "
  "Knowing when a number is too clean to be true.  |  Holding a position when "
  "the room would rather you did not.  |  None of these belong to an "
  "industry."),
 ("5", "What must be relearned", "2",
  "CONTEXT  |  WHAT MUST BE RELEARNED  |  The names the same idea goes by "
  "here.  |  Who must be convinced, and who can quietly stop it.  |  What the "
  "pressure is, and where it comes from.  |  What everyone here is quietly "
  "afraid of.  |  It feels like a competence gap. It is an information gap."),
 ("6", "What must be earned", "1",
  "CREDENTIAL  |  WHAT MUST BE EARNED  |  Genuinely required — There is "
  "no way around it. Plan for it.  |  A signal — Not required, but it "
  "shortens the conversation.  |  Neither — Chased because it feels like "
  "progress."),
 ("7", "Start from the destination", "2",
  "MOVE TWO  |  START FROM THE DESTINATION, NOT FROM YOUR CV.  |  What does "
  "that industry keep failing to solve?  |  Which of those decisions have you "
  "made before?  |  Which patterns do you already recognise early?  |  Which "
  "constraints have you worked inside?"),
 ("8", "Translate, do not recite", "2",
  "TRANSLATE, DO NOT RECITE  |  “ADAPTABLE.” “STRATEGIC.” "
  "“A FAST LEARNER.”  |  These describe how you would like to be "
  "seen.  |  THE UNIT THAT TRANSFERS IS JUDGMENT, WITH EVIDENCE ATTACHED.  |  "
  "Not a job title. Not a task list."),
 ("9", "Bridge evidence", "3",
  "MOVE THREE  |  BRIDGE EVIDENCE  |  A RELEVANT PROJECT · CROSS-FUNCTIONAL "
  "WORK  |  RESEARCH TURNED INTO A POINT OF VIEW · A WORK SAMPLE  |  "
  "PRACTITIONER CONVERSATIONS · A FIRST-90-DAYS PLAN"),
 ("10", "The three columns", "3",
  "ONE PAGE, THREE COLUMNS  |  WHAT TRAVELS — Judgment, decisions, "
  "patterns, ways of working.  |  WHAT CHANGES — Language, stakeholders, "
  "incentives, rhythm, risk.  |  WHAT I MUST EARN — The credential or "
  "exposure nothing substitutes for."),
 ("11", "Primary CTA, Capability Formation Field Kit", "1",
  "CAPABILITY FORMATION  |  FIELD KIT  |  What has your work built in you? "
  "How portable is it? What is still missing?  |  "
  "temidayoafonja.com/fieldkit"),
 ("12", "Continue the series", "1",
  "CONTINUE THE SERIES  |  CAREER PORTABILITY  |  CAREER PIVOTS · INTERNAL "
  "MOVES · GROWTH"),
]

DESCRIPTION = [
 "You are considering a move into a different industry, and the question "
 "underneath it is whether your experience will still count once you get "
 "there. This video is about how to move into a new industry without "
 "starting over: how to work out what actually transfers, what you will have "
 "to relearn because the context is different, and how to build credible "
 "evidence before you have direct experience in the destination field.",
 "",
 "Changing industries does not make you entry-level at everything. It makes "
 "you new to a context. Those are different problems, and advice that "
 "collapses them into one leaves you either underselling years of judgment or "
 "overclaiming familiarity you have not earned yet.",
 "",
 "Three moves:",
 "1. Separate capability, context and credential — they feel like one "
 "thing and they have completely different remedies.",
 "2. Match your judgment to the problems the destination industry needs "
 "solved, starting from that industry rather than from your CV.",
 "3. Build bridge evidence and a first-90-days learning plan, using only what "
 "you are permitted to have.",
 "",
 "You will also get a one-page exercise — what travels, what changes, "
 "what I must earn — and a way to say it out loud that does not sound "
 "like a slogan.",
 "",
 "This is not a promise that you can avoid every tradeoff of moving "
 "industries, and it is not permission to overstate what you know. It is a "
 "method for being accurate about your own portability.",
 "",
 "CAPABILITY FORMATION FIELD KIT",
 "A structured way to examine what your work has built in you, how portable "
 "it is, and where development or evidence is still needed.",
 CTA_URL,
 "",
 "CONTINUE THE SERIES",
 "Career Portability: Career Pivots, Internal Moves & Growth",
]

TAGS = ["how to move into a new industry",
        "how to switch industries without starting over",
        "transferable skills for a career change",
        "how to change industries without starting over",
        "career change", "changing industries", "career pivot",
        "transferable skills", "career portability", "mid career change",
        "experienced professionals"]

CHAPTERS = [
 ("0:00", "New to a context, not new to everything"),
 ("0:30", "Why I have made this move more than once"),
 ("1:15", "What actually changes in an industry move"),
 ("2:15", "Move one - capability, context and credential"),
 ("4:10", "Move two - start from the destination"),
 ("6:15", "Move three - bridge evidence and a learning plan"),
 ("8:30", "The three columns"),
 ("9:30", "Capability Formation Field Kit"),
 ("10:05", "Continue the series"),
]

PINNED = (
 "The line I keep coming back to: changing industries does not make you "
 "entry-level at everything. It makes you new to a context.\n\n"
 "Try the one-page exercise from the end of the video:\n"
 "WHAT TRAVELS - the judgment, decisions and patterns that stay useful.\n"
 "WHAT CHANGES - language, stakeholders, incentives, rhythm, risk.\n"
 "WHAT I MUST EARN - the credential or exposure nothing substitutes for.\n\n"
 "You may find the first column is longer than you feared and the third is "
 "shorter than you assumed.\n\n"
 "The Field Kit is a structured way to work through it: " + CTA_URL + "\n\n"
 "Which column is hardest for you to fill in? Tell me in a reply."
)

LOWER_THIRDS = [
 ("0:35", "Temidayo Afonja"),
 ("0:40", "Capability Formation"),
 ("9:33", "Capability Formation Field Kit / temidayoafonja.com/fieldkit"),
]

ON_SCREEN = [
 ("2:05", "CONTEXT IS LEARNABLE", "Small lower third, cream on navy. Two "
  "seconds. Do not compete with slide 2."),
 ("5:55", "JUDGMENT, WITH EVIDENCE ATTACHED", "Matches slide 8 wording "
  "exactly."),
 ("9:33", "temidayoafonja.com/fieldkit", "Persist to the end of the CTA "
  "block."),
]

BROLL = [
 ("0:30-1:15", "Hold on Temidayo, full screen. No b-roll. The early-proof "
  "block is the only personal passage in the video and should not be "
  "illustrated."),
 ("2:15-4:10", "Optional: a slow push in on slide 3 as each row lands. "
  "Nothing literal - no office stock, no handshakes, no laptops."),
 ("8:30-9:30", "None. The three columns should stay on screen so the viewer "
  "can copy the headings."),
]

EDITOR_NOTES = [
 "No opening title card. Open full screen on Temidayo; slide 1 lands only "
 "after the hook and the viewer promise, at roughly 0:30.",
 "Reveal builds are duplicate sequential slides, not PowerPoint animations. "
 "Advance in the order given by Video_8_Reveal_Order_Sheet.png - 24 frames.",
 "Slide 12 keeps its right third empty for the YouTube end-screen element. Do "
 "not place graphics there.",
 "No employer logos anywhere, on screen or in lower thirds, even though "
 "Deloitte, EY and PwC are named in the spoken script.",
 "Chapter timestamps in this package are planning estimates. Reset them "
 "against the real export before publishing.",
 "One offer only. No Keep the Proof, no Career Decision Evidence Check, no "
 "Maven session, no book, anywhere in Video 8.",
]

PRECHECKS = [
 "Record 4K at 30fps; export 1080p at 30fps. Audio at 48 kHz.",
 "More light on Temidayo - the single most important recording improvement "
 "carried across the channel.",
 "Confirm the spoken script matches Video_8_Recording_Script_Clean.txt word "
 "for word before the first take.",
 "Confirm the CTA URL on screen and in the description is "
 "temidayoafonja.com/fieldkit.",
 "Confirm the approved Canva thumbnail reads YOUR EXPERIENCE STILL COUNTS and "
 "is 1280 x 720.",
 "Set chapter timestamps from the finished export, not from this package.",
 "Neither the spoken script nor slide 12 names an unpublished video. Set the "
 "end-screen video element per the routing rules in section 8, and verify the "
 "playlist opens signed out and holds another public video.",
]


def shade(p, fill):
    pr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill)
    pr.append(sh)


def keep_together(p, with_next=False):
    """Keep a paragraph's lines on one page, optionally with the next block."""
    pr = p._p.get_or_add_pPr()
    for tag in (('w:keepLines',) + (('w:keepNext',) if with_next else ())):
        el = OxmlElement(tag)
        el.set(qn('w:val'), '1')
        pr.append(el)
    return p


def no_split(tbl):
    """Stop every row of a table from breaking across a page."""
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit')
        trPr.append(cs)
        for cell in row.cells:
            for p in cell.paragraphs:
                keep_together(p)
    return tbl


def no_hyphenation(doc):
    """Word may otherwise break a URL or a filename mid-token."""
    st = doc.settings.element
    el = OxmlElement('w:autoHyphenation')
    el.set(qn('w:val'), '0')
    st.append(el)


def build():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(11)
    no_hyphenation(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.85)
        s.left_margin = s.right_margin = Inches(0.95)

    def para(text, size=11, bold=False, color=None, before=0, after=8,
             italic=False, caps=False, spacing=1.25):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = spacing
        r = p.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        if caps:
            r.font.all_caps = True
        return p

    def h1(text):
        # A heading never strands at the foot of a page on its own.
        return keep_together(
            para(text, size=16, bold=True, color=NAVY, before=22, after=10),
            with_next=True)

    def table(rows, widths=None):
        t = doc.add_table(rows=0, cols=len(rows[0]))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for c, val in zip(cells, row):
                c.text = ""
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(str(val))
                r.font.size = Pt(10)
                if i == 0:
                    r.bold = True; r.font.color.rgb = NAVY
        if widths:
            for row in t.rows:
                for c, w in zip(row.cells, widths):
                    c.width = Inches(w)
        no_split(t)
        para("", size=6, after=6)
        return t

    para("CAREER PORTABILITY   |   VIDEO 8", size=10, bold=True, color=GOLD,
         after=4, caps=True)
    para(TITLE, size=22, bold=True, color=NAVY, after=6, spacing=1.1)
    para("First-pass production package  ·  slides, script, publishing and "
         "editor direction", size=11, color=DIM, after=14)

    table([["Field", "Value"],
           ["Final title", TITLE],
           ["Primary target keyword", KEYWORD],
           ["Secondary search language",
            "how to switch industries without starting over; transferable "
            "skills for a career change; how to change industries without "
            "starting over"],
           ["Viewer question", "What can I carry into a field where I am new?"],
           ["Locked thumbnail copy", THUMB],
           ["Thumbnail status",
            "To be created and approved externally in Canva. Not built, "
            "regenerated or substituted in this package."],
           ["Primary CTA", CTA_NAME],
           ["CTA URL", CTA_URL],
           ["Target length", "Approximately 9-12 minutes"],
           ["Spoken script length", "%d words" % WORDS],
           ["Slides", "12 main slides, 24 reveal frames"],
           ["Watch next",
            "Career Portability playlist. The spoken script and slide 12 name "
            "no unpublished video; see section 8 for the end-screen routing "
            "rules."]],
          widths=[1.9, 4.7])

    para("LOCKED PRODUCTION DECISIONS", size=11, bold=True, color=NAVY,
         before=10, after=6)
    table([["Decision", "Detail"],
           ["Keyword rationale",
            "TubeBuddy, recorded 30 August 2026. “how to move into a new "
            "industry” 73/100 weighted, Very Good. “how to switch "
            "industries without starting over” 71/100. “transferable "
            "skills for a career change” 69/100. The roadmap's earlier "
            "wording, “how to change industries without starting at entry "
            "level,” scored 58/100. Production rationale only - never "
            "referenced in the video, and not a promise of performance."],
           ["Title relationship to the roadmap",
            "The roadmap's working title is “How to Change Industries "
            "Without Starting at Entry Level.” The final title is an "
            "intentional SEO refinement of it. Viewer question, thumbnail "
            "copy, CTA and the three teaching moves are unchanged from the "
            "roadmap."],
           ["Personal evidence ceiling",
            "Career span, eight industries and sectors, Deloitte/EY/PwC, "
            "Consultant to Head of Function, CISM preparation and an "
            "unsuccessful first attempt. Nothing beyond this. See section 2."],
           ["Excluded metrics",
            "The 30% retention improvement and the $2M+ estimated turnover "
            "cost avoidance are excluded. See section 2."],
           ["Offer exclusivity",
            "One offer on screen and in the script: the Capability Formation "
            "Field Kit. No Keep the Proof, no Career Decision Evidence Check, "
            "no Maven offering, no book."]],
          widths=[1.9, 4.7])

    h1("1. Strategy at a glance")
    table([["Element", "Decision"],
           ["Viewer",
            "Experienced professionals — especially mid-career and senior professionals, managers, directors and experienced specialists — considering a move into a new industry without wanting to erase their prior experience."],
           ["Viewer question", "What can I carry into a field where I am new?"],
           ["Central distinction",
            "Changing industries does not make someone entry-level at "
            "everything. It makes them new to a context."],
           ["Both extremes rejected",
            "“None of my previous experience counts” and "
            "“everything I did before will transfer automatically.”"],
           ["Teaching spine",
            "Separate capability, context and credential; match judgment to "
            "destination problems; build bridge evidence and a learning "
            "plan."],
           ["What this video is not",
            "Not a promise that the viewer can avoid every tradeoff of "
            "changing industries. Not permission to overstate familiarity."]],
          widths=[1.6, 5.0])

    table([["Architecture beat", "Where it lands"],
           ["Hook", "0:00 - the locked opening distinction"],
           ["Viewer payoff", "0:15 - three piles, and how to show the first"],
           ["Early proof", "0:30 - career span, sectors, firms, CISM"],
           ["Brief introduction", "1:05 - one positioning line, after the "
            "proof"],
           ["Teaching", "1:15 onward - what changes, then the three moves"],
           ["Word-count method",
            "Python str.split() on whitespace, over the spoken paragraphs of "
            "section 4 only. Timed block headers and the target line are "
            "excluded. The same method is used in every file that reports a "
            "count."],
           ["Usable test", "8:30 - what travels / what changes / what I must "
            "earn"],
           ["Earned next step", "9:30 - Capability Formation Field Kit"]],
          widths=[2.4, 4.2])

    h1("2. Editorial thesis and factual boundaries")
    para("Thesis. Moving industries is a context problem wearing the costume "
         "of a competence problem. Capability is portable, context is "
         "learnable, and credential is the only category that may genuinely "
         "have to be earned - but the three feel identical from the inside, "
         "which is why the move is mishandled in both directions.", after=10)
    table([["Boundary", "Ruling"],
           ["Career span",
            "Spoken as “nearly two decades.” NOT as “18 "
            "years.” docs/claims-ledger.md section 6 records that the "
            "published wording was changed from “eighteen years” to "
            "“nearly two decades” by operator decision in August "
            "2026, and that “eighteen years” now appears nowhere in "
            "the repository. Using the retired wording in a public video "
            "would reintroduce it. Same fact, approved wording. Flagged in "
            "the QA README as a conflict with the brief and the roadmap."],
           ["Industries and firms",
            "“eight industries and sectors” and “Deloitte, EY "
            "and PwC” are roadmap-sourced and used as written. No "
            "employer logo appears on screen."],
           ["Progression",
            "“Consultant to Head of Function.” The roadmap boundary "
            "“do not claim promotion in every industry” is "
            "observed: no promotion is attributed to any specific industry."],
           ["CISM",
            "Only that Temidayo prepared for the CISM exam and did not pass "
            "the first time. No score, date, reason or consequence is stated, "
            "and no career result is attributed to it. It is used to show "
            "that portability still requires relearning."],
           ["30% retention improvement",
            "EXCLUDED. docs/claims-ledger.md records it as “Needs "
            "source” with no supporting document on file, and the "
            "roadmap instructs: do not attach either metric to a role or "
            "intervention until the relationship is documented. No connection "
            "to this story exists."],
           ["$2M+ turnover cost avoidance", "EXCLUDED, on the same basis."],
           ["Causal claims",
            "None. No employer is said to have caused a later outcome, no "
            "move is described as seamless, and no prior skill is claimed to "
            "have transferred automatically."],
           ["Confidentiality",
            "Bridge evidence is limited to non-confidential material. The "
            "script states outright that inventing experience or implying "
            "more familiarity than you have is not permitted."]],
          widths=[1.7, 4.9])

    h1("3. Timed script outline")
    table([["Time", "Block", "Slide"]]
          + [[h.split(" | ")[0], h.split(" | ")[1], s] for h, s in
             zip([b[0] for b in SCRIPT],
                 ["1", "-", "2", "3, 4, 5, 6", "7, 8", "9", "10", "11", "12"])],
          widths=[1.2, 3.9, 1.5])

    h1("4. Full recording script")
    para("Target 1,450-1,700 spoken words. This draft is %d." % WORDS,
         size=10, italic=True, color=DIM, after=12)
    for header, paras in SCRIPT:
        p = para(header, size=10.5, bold=True, color=GOLD, before=16, after=8,
                 caps=True, spacing=1.1)
        shade(p, "F3F0E8")
        for t in paras:
            para(t, size=11.5, color=INK, after=10, spacing=1.4)

    h1("5. Slide deck content and placement")
    table([["#", "Slide", "States", "Copy on screen, verbatim"]]
          + [list(r) for r in SLIDE_COPY], widths=[0.4, 1.5, 0.6, 4.1])
    para("1920 x 1080, 16:9. No stock imagery, gradients, decorative icons, "
         "employer logos or PowerPoint animations. All text and shapes remain "
         "editable; nothing is rasterised. Reveal builds are duplicate "
         "sequential slides. 24 frames total.", size=10, italic=True,
         color=DIM)

    h1("6. Code handoff - slides and script files")
    table([["File", "What it is"],
           ["Video_8_Main_Slides.pptx", "12 editable slides, final revealed "
            "state. All text is live text; no images."],
           ["Video_8_Reveal_Builds.pptx", "24 slides, one per reveal state, "
            "in advance order."],
           ["Video_8_Slide_Preview.pdf", "12 pages at true 16:9."],
           ["Video_8_Main_Slide_Contact_Sheet.png", "All 12 slides on one "
            "sheet."],
           ["Video_8_Reveal_Order_Sheet.png", "All 24 frames in advance "
            "order."],
           ["Video_8_Phone_Legibility_Sheet.png", "Every slide at 320 x 180."],
           ["Video_8_Teleprompter_Script_with_Slide_Markers.docx",
            "The spoken script with 12 slide markers and the package's own "
            "stage directions kept visually distinct."],
           ["Video_8_Recording_Script_Clean.txt", "Spoken words only. No "
            "timestamps, markers or directions."],
           ["build/slides.py, build/deck.py, build/build.py",
            "Editable sources. deck.py is carried forward from Video 7 "
            "unchanged."]],
          widths=[3.1, 3.5])

    h1("7. Thumbnail - to be created externally, not built here")
    table([["Field", "Value"],
           ["Locked copy", THUMB],
           ["Status", "Temidayo will create the final thumbnail in Canva."],
           ["Instruction", "Do not create, redesign or substitute a Video 8 "
            "thumbnail. None was made in this package and no placeholder was "
            "produced."],
           ["Standards to apply", "1280 x 720 upload, exact 16:9. Palette "
            "cream #F5F0E8, navy #0F2346, gold #C9A84C. Legible at 200 px "
            "wide. See CAPABILITY_FORMATION_YOUTUBE_STANDARDS.md."]],
          widths=[1.5, 5.1])

    h1("8. YouTube publishing package")
    keep_together(para("Title", size=11, bold=True, color=NAVY, after=4), with_next=True)
    para(TITLE, after=10)
    keep_together(para("Description", size=11, bold=True, color=NAVY, after=4), with_next=True)
    for line in DESCRIPTION:
        keep_together(para(line if line else " ", after=6 if line else 2))
    para("Tags - the exact target keyword first", size=11, bold=True,
         color=NAVY, before=8, after=4)
    para(", ".join(TAGS), after=10)
    keep_together(para("Chapters", size=11, bold=True, color=NAVY, after=4), with_next=True)
    table([["Time", "Chapter"]] + [list(c) for c in CHAPTERS],
          widths=[1.0, 5.6])
    para("ESTIMATES ONLY. These timestamps are planning figures and must be "
         "reset from the finished export before publishing.", size=10,
         italic=True, color=DIM, after=10)
    keep_together(para("Pinned comment", size=11, bold=True, color=NAVY, before=8, after=4), with_next=True)
    for line in PINNED.split("\n"):
        para(line if line else " ", after=5 if line else 2)
    keep_together(para("End screen", size=11, bold=True, color=NAVY, before=10, after=4), with_next=True)
    table([["Element", "Route"],
           ["Video element, first choice",
            "If the next video in the sequence is public when Video 8 "
            "publishes, the end-screen video element may route directly to "
            "it."],
           ["Otherwise",
            "Route to the public Career Portability playlist."],
           ["Verify before publication",
            "The selected playlist contains at least one other public video "
            "and opens successfully in a signed-out browser."],
           ["Do not use",
            "A playlist containing only Video 8 as the fallback."],
           ["Spoken and on-slide wording",
            "Neither names an unpublished video. Both point at the Career "
            "Portability playlist, so no re-render or re-record is needed "
            "whichever route is chosen."],
           ["Subscribe element", "Bottom right, standard placement"],
           ["Clear zone", "Slide 12 keeps the right third empty for the "
            "card"]],
          widths=[1.9, 4.7])

    h1("9. Editor direction")
    keep_together(para("Lower thirds", size=11, bold=True, color=NAVY, after=4), with_next=True)
    table([["Time", "Copy"]] + [list(l) for l in LOWER_THIRDS],
          widths=[1.0, 5.6])
    keep_together(para("On-screen text", size=11, bold=True, color=NAVY, before=8, after=4), with_next=True)
    table([["Time", "Text", "Note"]] + [list(o) for o in ON_SCREEN],
          widths=[0.8, 2.6, 3.2])
    keep_together(para("B-roll", size=11, bold=True, color=NAVY, before=8, after=4), with_next=True)
    table([["Block", "Guidance"]] + [list(b) for b in BROLL],
          widths=[1.4, 5.2])
    keep_together(para("Editor notes", size=11, bold=True, color=NAVY, before=8, after=4), with_next=True)
    for n in EDITOR_NOTES:
        para("-  " + n, after=6)

    h1("10. Final pre-recording and delivery checks")
    for c in PRECHECKS:
        para("[ ]  " + c, after=6)

    doc.save(OUT)
    print("wrote", os.path.basename(OUT))
    print("spoken paragraphs: %d   words: %d" % (len(SPOKEN), WORDS))
    return OUT


if __name__ == "__main__":
    build()
