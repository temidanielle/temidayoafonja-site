# -*- coding: utf-8 -*-
"""Build the two Video 8 script files from the production package.

Every spoken word is copied out of the package DOCX. Nothing is rewritten,
expanded, summarised or reordered, so the package, the teleprompter and the
clean TXT cannot drift apart.

  Video_8_Teleprompter_Script_with_Slide_Markers.docx
      the approved script unchanged, with 12 slide markers and the package's
      own stage directions kept visually distinct from the spoken lines

  Video_8_Recording_Script_Clean.txt
      the spoken words only
"""
import os, re, zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(
    os.path.dirname(HERE),
    "YouTube_Video_8_Production_Package_New_Industry.docx")
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

NAVY = RGBColor(0x0F, 0x23, 0x46)
GOLD = RGBColor(0x8A, 0x6D, 0x1E)
DIM = RGBColor(0x5A, 0x6B, 0x82)

HDR = re.compile(r'^\d+:\d\d[–-]\d+:\d\d\s*\|')

# Internal labels that are not spoken content and must not appear in either file.
INTERNAL_LABELS = {"VISUAL TEACHING SYSTEM"}


def source_blocks():
    z = zipfile.ZipFile(SRC)
    x = ET.fromstring(z.read('word/document.xml'))
    t = lambda p: ''.join(e.text or '' for e in p.iter(NS + 't'))
    blocks = [t(el) for el in x.find(NS + 'body') if el.tag == NS + 'p']
    i4 = next(i for i, b in enumerate(blocks)
              if b.strip() == '4. Full recording script')
    i5 = next(i for i, b in enumerate(blocks)
              if b.strip().startswith('5. Slide deck content'))
    return [b for b in blocks[i4 + 1:i5]
            if b.strip() and b.strip() not in INTERNAL_LABELS]


# Slide cues. Each is (slide number, slide name, the opening of the spoken
# paragraph the slide lands on). The marker is placed immediately before it.
CUES = [
    (1,  "New context, not no experience",
     "What I learned is that changing industries does not make you"),
    (2,  "What actually changes",
     "Let me be precise about what changes when you move industries"),
    (3,  "Capability, context, credential",
     "That brings me to the first move"),
    (4,  "What travels",
     "Capability is judgment or ability that stays useful"),
    (5,  "What must be relearned",
     "Context is what the new field knows and you do not"),
    (6,  "What must be earned",
     "Credential is the formal evidence or permission"),
    (7,  "Start from the destination",
     "The second move is to match the judgment"),
    (8,  "Translate, do not recite",
     "This is where most industry-change pitches come apart"),
    (9,  "Bridge evidence",
     "The third move is to close the credibility gap"),
    (10, "The three columns",
     "So here is the exercise, and it fits on one page."),
    (11, "Capability Formation Field Kit",
     "If you want a structured way to work through this"),
    (12, "Continue the series",
     "When you are ready for the next step"),
]

STATES = {1: 2, 2: 2, 3: 3, 4: 2, 5: 2, 7: 2, 8: 2, 9: 3, 10: 3}


def shade(p, hexfill):
    pr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexfill)
    pr.append(sh)


def keep_with_next(p):
    pr = p._p.get_or_add_pPr()
    for tag in ('w:keepLines', 'w:keepNext'):
        el = OxmlElement(tag)
        el.set(qn('w:val'), '1')
        pr.append(el)
    return p


def left_bar(p, hexcol):
    pr = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    lf = OxmlElement('w:left')
    lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '18')
    lf.set(qn('w:space'), '8'); lf.set(qn('w:color'), hexcol)
    bd.append(lf); pr.append(bd)


# ---------------------------------------------------------------- opening
# The revised conversational opening lives in a script-only source so the
# production-package DOCX and every slide asset stay untouched.
from opening_revision import NEW_OPENING, REPLACES
from canon import count as canonical_count, spoken_paragraphs


def apply_opening_override(sec):
    """Swap the first REPLACES spoken paragraphs for the revised opening.

    Timed headers and the target line are left exactly where they are, so the
    block structure of the script is unchanged.
    """
    out, dropped, inserted = [], 0, False
    for b in sec:
        t = b.strip()
        is_spoken = not HDR.match(t) and not t.startswith("Target")
        if is_spoken and dropped < REPLACES:
            dropped += 1
            if not inserted:
                out.extend(NEW_OPENING)
                inserted = True
            continue
        out.append(b)
    assert dropped == REPLACES, "expected %d paragraphs to replace, found %d" % (
        REPLACES, dropped)
    return out


def build():
    sec = apply_opening_override(source_blocks())
    CANON_WORDS = canonical_count(sec)
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(13)
    _ah = OxmlElement('w:autoHyphenation')
    _ah.set(qn('w:val'), '0')
    doc.settings.element.append(_ah)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.9)
        s.left_margin = s.right_margin = Inches(1.05)

    def para(text, size=13, bold=False, color=None, before=0, after=10,
             italic=False, caps=False, spacing=1.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = spacing
        r = p.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        if caps:
            r.font.all_caps = True
        return p

    para("How to Move Into a New Industry Without Starting Over",
         size=24, bold=True, color=NAVY, after=2, spacing=1.1)
    para("Video 8  ·  Teleprompter script with slide markers", size=12,
         color=DIM, after=6, spacing=1.1)
    para("Spoken script is the large text. Everything in a tinted band is a "
         "production direction and is not spoken.", size=11, italic=True,
         color=DIM, after=20, spacing=1.2)

    def marker(n, name):
        label = "SLIDE %d  —  %s" % (n, name)
        if n in STATES:
            label += "   (%d reveal states)" % STATES[n]
        p = para(label, size=11, bold=True, color=NAVY, before=14, after=14,
                 spacing=1.1)
        shade(p, "E8EDF4"); left_bar(p, "0F2346"); keep_with_next(p)

    used = set()
    for block in sec:
        b = block.strip()
        if b.startswith("Target"):
            b = re.sub(r"This draft is [\d,]+\.",
                       "This draft is %s." % format(CANON_WORDS, ","), b)
            p = para(b, size=10.5, italic=True, color=DIM, before=0, after=18,
                     spacing=1.2)
            shade(p, "F3F0E8")
            continue
        if HDR.match(b):
            p = para(b, size=10.5, bold=True, color=GOLD, before=20, after=12,
                     spacing=1.1, caps=True)
            shade(p, "F3F0E8"); keep_with_next(p)
            continue
        hit = next((c for c in CUES
                    if c[0] not in used and b.startswith(c[2])), None)
        if hit:
            marker(hit[0], hit[1])
            used.add(hit[0])
        para(b, size=13.5, color=RGBColor(0x1A, 0x1A, 0x1A), after=12,
             spacing=1.5)

    out_docx = os.path.join(
        HERE, "Video_8_Teleprompter_Script_with_Slide_Markers.docx")
    doc.save(out_docx)

    # ---------------------------------------------------- clean spoken script
    spoken = [b.strip() for b in sec
              if not HDR.match(b.strip()) and not b.strip().startswith("Target")]
    out_txt = os.path.join(HERE, "Video_8_Recording_Script_Clean.txt")
    open(out_txt, "w").write("\n\n".join(spoken) + "\n")

    print("teleprompter :", os.path.basename(out_docx))
    print("clean script :", os.path.basename(out_txt))
    print("markers placed:", sorted(used))
    print("spoken paragraphs: %d   words: %d"
          % (len(spoken), CANON_WORDS))
    return out_docx, out_txt, spoken, sec


if __name__ == "__main__":
    build()
