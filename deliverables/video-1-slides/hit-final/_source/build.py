# -*- coding: utf-8 -*-
"""Build the Video 1 H.I.T. final recording and Shorts package."""
import os, sys, shutil, zipfile, hashlib
sys.path.insert(0, "/tmp/v1hit")
from script_text import LINES, SPOKEN, MARKERS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x23,0x46); GOLD=RGBColor(0x8A,0x6D,0x1E)
DIM=RGBColor(0x5A,0x6B,0x82); INK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0x9B,0x2C,0x10)
BAND_NAVY="E8EDF4"; BAND_CREAM="F3F0E8"

ROOT="/tmp/v1hit/Video_1_HIT_FINAL"
LF=os.path.join(ROOT,"LONG_FORM"); SH=os.path.join(ROOT,"SHORTS")
shutil.rmtree(ROOT, ignore_errors=True)
os.makedirs(LF); os.makedirs(SH)

TITLE="How to Change Jobs Without Starting Your Career Over"
DECK_TITLE="How I Changed Jobs Without Starting My Career Over"
THUMB="DON’T START FROM ZERO"
CTA="Free Career Evidence Starter"
CTA_URL="https://temidayoafonja.com/career-evidence-starter"
CTA_DESCRIPTOR="FREE CAREER ACCOMPLISHMENT TRACKER"
CTA_PROMISE=("Turn one accomplishment into proof you can use in a "
             "performance review, interview, internal move or career pivot.")
OLD_CTA="Capability Formation Field Kit"
NEXT="Is Your Job Making You Less Marketable?"

def newdoc(teleprompter=False):
    d=Document(); st=d.styles['Normal']
    st.font.name='Calibri'; st.font.size=Pt(13 if teleprompter else 11)
    ah=OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'),'0')
    d.settings.element.append(ah)
    for s in d.sections:
        s.top_margin=s.bottom_margin=Inches(0.9)
        s.left_margin=s.right_margin=Inches(1.05)
    return d

def keep(p,nxt=False):
    pr=p._p.get_or_add_pPr()
    for t in ('w:keepLines',)+(('w:keepNext',) if nxt else ()):
        e=OxmlElement(t); e.set(qn('w:val'),'1'); pr.append(e)
    return p

def shade(p,fill):
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),fill); pr.append(s)

def bar(p,col):
    pr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18')
    l.set(qn('w:space'),'8'); l.set(qn('w:color'),col); b.append(l); pr.append(b)

def P(d,t,size=11,bold=False,color=None,before=0,after=8,italic=False,
      caps=False,spacing=1.3):
    p=d.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=spacing
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if caps: r.font.all_caps=True
    return p

def head(d,title,sub,note=None):
    P(d,"CAPABILITY FORMATION   |   VIDEO 1",size=10,bold=True,color=GOLD,
      after=4,caps=True)
    P(d,title,size=20,bold=True,color=NAVY,after=4,spacing=1.1)
    P(d,sub,size=11,color=DIM,after=6,spacing=1.1)
    if note: P(d,note,size=10.5,italic=True,color=DIM,after=18,spacing=1.2)

def H1(d,t,before=20):
    return keep(P(d,t,size=14,bold=True,color=NAVY,before=before,after=8),True)
def H2(d,t,before=13):
    return keep(P(d,t,size=11.5,bold=True,color=NAVY,before=before,after=5),True)

def compress(d, line_spacing=1.18, after_scale=0.80):
    """Tighten one document's vertical rhythm without touching type size,
    weight, colour or wording. Keeps editor briefs off a near-empty last page."""
    for p in d.paragraphs:
        pf=p.paragraph_format
        if pf.line_spacing and 1.0 < pf.line_spacing < 1.4 and pf.line_spacing > line_spacing:
            pf.line_spacing = line_spacing
        if pf.space_after is not None:
            pf.space_after = Pt(round(pf.space_after.pt*after_scale,1))
        if pf.space_before is not None and pf.space_before.pt:
            pf.space_before = Pt(round(pf.space_before.pt*after_scale,1))
    return d

def pairlist(d, items, indent="—  ", gap="        ", after=4, budget=78):
    """Lay short bullet items several to a line, packed by width.

    Every item is kept verbatim; nothing is removed, reworded or reordered,
    and neither type size nor leading changes. This reclaims the unused
    right-hand part of a short-bullet line, which is what keeps the Shorts
    editor brief to two pages. The width budget is measured against the
    inspection renderer's font, which is wider than Calibri, so Word fits at
    least as much per line and never less."""
    lines, cur = [], []
    for it in items:
        trial = cur + [it]
        width = sum(len(indent) + len(x) for x in trial) + len(gap) * (len(trial) - 1)
        if cur and width > budget:
            lines.append(cur); cur = [it]
        else:
            cur = trial
    if cur: lines.append(cur)
    for group in lines:
        keep(P(d, gap.join(indent + x for x in group), after=after))

# ------------------------------------------------- 1. teleprompter DOCX + TXT
d=newdoc(True)
head(d,DECK_TITLE,"Video 1  ·  Teleprompter script with slide markers",
     "Spoken script is the large text. A slide marker in a tinted band tells "
     "the editor which slide to bring up; it is not spoken.")
for line in LINES:
    if line.startswith("[SLIDE:"):
        name=line[len("[SLIDE:"):-1].strip()
        p=P(d,"SLIDE  —  %s"%name,size=11,bold=True,color=NAVY,
            before=14,after=14,spacing=1.1)
        shade(p,BAND_NAVY); bar(p,"0F2346"); keep(p,True)
    else:
        keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video1TeleprompterScriptwithslidemarkers_HIT_v3.1.docx"))

tel=[DECK_TITLE,"Video 1  ·  Teleprompter script with slide markers",""]
for line in LINES:
    if line.startswith("[SLIDE:"):
        tel += ["", "SLIDE  —  %s"%line[len("[SLIDE:"):-1].strip(), ""]
    else:
        tel += [line, ""]
open(os.path.join(LF,"Video1TeleprompterScriptwithslidemarkers_HIT_v3.1.txt"),
     "w").write("\n".join(tel).strip()+"\n")

# ------------------------------------------------- 2. reading script DOCX+TXT
d=newdoc(True)
head(d,DECK_TITLE,"Video 1  ·  Reading script, no markers",
     "Spoken language only. No slide markers, no timestamps, no production "
     "directions.")
for line in SPOKEN:
    keep(P(d,line,size=13.5,color=INK,after=12,spacing=1.5))
d.save(os.path.join(LF,"Video1ReadingScriptnomarkers_HIT_v3.1.docx"))
open(os.path.join(LF,"Video1ReadingScriptnomarkers_HIT_v3.1.txt"),
     "w").write("\n\n".join(SPOKEN)+"\n")
print("long-form scripts written")

# --------------------------------------------------- 3. long-form editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 1",size=12,bold=True,color=GOLD,after=2,caps=True)
P(d,TITLE,size=20,bold=True,color=NAVY,after=6,spacing=1.1)
p=P(d,"This document is for the editor. It is NOT Temidayo's teleprompter and "
     "must not be placed on the recording screen.",size=11,italic=True,
     color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"Locked metadata",before=14)
for k,v in (("Public title",TITLE),("On-screen deck title",DECK_TITLE),
            ("Thumbnail",THUMB),("Primary CTA",CTA),("CTA URL",CTA_URL),
            ("Watch next",NEXT),
            ("Core distinction",
             "A new context does not make an experienced professional new to "
             "everything. The task is to identify what travels, support it "
             "with evidence and name what must still be learned.")):
    keep(P(d,"%-22s %s"%(k+":",v),size=11,after=5))
p=P(d,"The public title and the on-screen deck title differ INTENTIONALLY. Do "
     "not change the title slide to match the public metadata title.",size=11,
     bold=True,color=RED,before=8,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
p=P(d,"STATUS. This is a replacement recording package. The currently "
     "published Video 1, its published description, the slides, the reveal "
     "deck and the thumbnail all remain untouched until the replacement has "
     "been recorded, edited, reviewed and uploaded.",size=11,bold=True,
     color=NAVY,before=4,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"First 30 seconds — H.I.T.",before=14)
P(d,"H = Hook. I = Interest. T = Trust. The opening must work as one "
    "audiovisual unit: immediate conversational tension, meaningful visual "
    "interest, a relevant concrete reason to trust Temidayo, and a clear "
    "viewer payoff by roughly 20 to 30 seconds. No generic channel welcome "
    "before the promise. No forced statistic. No résumé recital.",after=8)
P(d,"The visual carries the chronology. Temidayo's voice carries the meaning.",
  italic=True,color=DIM,after=8)
P(d,"This supersedes the older Video 1 instruction that kept Temidayo "
    "visually static or full-screen for roughly the first 45 seconds. The "
    "existing slides themselves are unchanged.",italic=True,color=DIM,after=14)

def beat(t,anchor,layer,body):
    H2(d,t,before=10)
    p=P(d,"Spoken anchor:  “%s”"%anchor,size=10.5,italic=True,color=DIM,after=8)
    shade(p,BAND_CREAM)
    if layer: keep(P(d,layer,size=11,bold=True,color=GOLD,after=6))
    for b in body: keep(P(d,b,after=5))

beat("0:00–0:07",
     "Changing jobs can make years of experience feel as though they belong "
     "to the place you’re leaving.","H = HOOK",
     ["Visual: begin on Temidayo, medium/tight, direct to camera.",
      "On-screen hook:  DOES YOUR EXPERIENCE STILL COUNT?",
      "No title card before this. No résumé animation, stock career footage "
      "or opening logo."])
beat("0:07–0:12","But a new context does not make you new to everything.",
     "I = INTEREST",
     ["Bring in:  NEW CONTEXT ≠ ZERO EXPERIENCE",
      "Keep the text restrained, large and mobile-readable.",
      "No beginner icon, zero graphic or literal reset animation."])
beat("0:12–0:21",
     "Over roughly eighteen years, my career has crossed very different "
     "functions and industries…","T = TRUST  /  I = INTEREST",
     ["Briefly show the existing career-path progression:",
      "     ACCOUNTING & AUDIT  →  CYBERSECURITY  →  PEOPLE STRATEGY  →  "
      "CAPABILITY FORMATION",
      "IMPORTANT: Temidayo deliberately does not narrate the individual "
      "chapters aloud. The visual provides the chronology; her spoken words "
      "provide the meaning.",
      "Do not alter Slide 2. Use a restrained overlay, crop or adapted "
      "treatment based on the existing slide system. No employer logos."])
beat("0:21–0:30","I’ll show you three things…","PAYOFF",
     ["Progressively preview:",
      "     LOOK UNDERNEATH THE TITLE",
      "     EXPLAIN WHAT CHANGED",
      "     KEEP THE EVIDENCE",
      "Return cleanly to Temidayo before moving into the existing title slide.",
      "This is a micro-preview only. Do not replace or redesign Slide 10."])

H1(d,"Editorial rhythm after the opening",before=14)
P(d,"The first roughly 30 seconds may be more visually active than the "
    "remainder. Do NOT read H.I.T. as permission for constant motion.",after=6)
P(d,"After the opening:",after=6)
for x in ["preserve Temidayo's natural pace;","allow reflective pauses;",
 "use the existing slides as teaching support;",
 "let the slide carry lists and chronology;","avoid constant punch-ins;",
 "avoid decorative B-roll;","avoid unnecessary text duplication;",
 "avoid flashy transitions."]:
    keep(P(d,"—  "+x,after=4))
P(d,"The edit should remain premium, calm, intelligent and editorial.",
  before=6,after=8)

H1(d,"Existing slide system",before=14)
P(d,"The existing Video 1 v2.4 visual system is authoritative and UNCHANGED: "
    "exactly 13 main slides and exactly 22 reveal-build frames. Do not add, "
    "delete, redesign or reorder slides or reveal frames.",after=8)
for n,job in enumerate(["Title","My Career Path",
 "01: Look Underneath the Title","Move One",
 "02: Explain What the Work Changed","Move Two","One Result: 47 to 75",
 "03: Keep Evidence Before You Need It","Move Three",
 "Three Things I Learned to Do","Before Your Next Move",
 "Career Evidence Starter  (CTA corrected)","Watch Next"],1):
    keep(P(d,"Slide %-3d %s"%(n,job),size=10.5,after=3))

H1(d,"Let the visual carry information",before=14)
P(d,"The production principle established in Videos 2 and 3 applies here. "
    "When the visual can carry information efficiently, Temidayo does not "
    "need to narrate every item. Do not add spoken lines to compensate for "
    "slide copy.",after=8)
for s,note in [("Slide 2","The career-path labels. Temidayo does not read the "
 "full path aloud."),
 ("Slide 4","The three questions that reveal the work beneath a title. "
  "Reveal them progressively."),
 ("Slide 7","The 47-to-75 evidence. Let the visual breathe."),
 ("Slide 9","Situation / My role / What changed / What this shows. Do not "
  "require Temidayo to repeat every line mechanically."),
 ("Slide 11","The three before-your-next-move questions. Reveal one at a "
  "time and allow a real pause.")]:
    keep(P(d,"%s — %s"%(s,note),after=5))

H1(d,"Evidence boundaries",before=14)
p=keep(P(d,"The onboarding result must remain precisely bounded.",bold=True,
       color=NAVY,after=6))
for x in ["Temidayo led an onboarding redesign WITH HER TEAM.",
 "One measure of how well new hires felt integrated moved from 47 to 75.",
 "It was one measure of new-hire integration.",
 "It was NOT a claim about every effect of the redesign.",
 "It was team-based work that Temidayo led. It was NOT solo work."]:
    keep(P(d,"—  "+x,after=4))
p=P(d,"Do not attach an approximately 30% retention improvement, a "
     "more-than-$2-million avoided-turnover figure, or any other undocumented "
     "result to this onboarding story.",size=11,bold=True,color=RED,
     before=8,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
keep(P(d,"Do not imply that every part of Temidayo's experience transferred "
       "across contexts. Preserve the explicit boundary: some knowledge "
       "belongs to the company, industry, regulation, relationships or "
       "context; a move may require real relearning; being new to a context "
       "is not the same as being new to every underlying problem.",
       before=4,after=8,spacing=1.25))
keep(P(d,"Evidence preservation means a permitted, high-level record, in the "
       "person's own words, using information they are entitled to retain. Do "
       "not suggest taking confidential information, customer or employee "
       "data, proprietary documents, employer-owned files, or anything the "
       "person does not have the right to keep.",after=8,spacing=1.25))

H1(d,"Visual boundaries",before=14)
P(d,"Do not add:",after=5)
for x in ["generic office B-roll;","résumé icons;",
 "literal luggage or “carrying skills” imagery;","ladder graphics;",
 "dramatic career-change montages;","clocks or career-anniversary graphics;",
 "employer logos;","constant zooms;","artificial sound effects;",
 "red warning graphics;","AI-generated scenery;",
 "synthetic or altered images of Temidayo;","fake shock expressions;",
 "social-media template transitions."]:
    keep(P(d,"—  "+x,after=3))
keep(P(d,"Preserve Temidayo's natural appearance.",bold=True,before=6,after=8))

H1(d,"Slide 12 CTA correction — applied",before=14)
p=P(d,"AUTHORISED AND APPLIED. The %s CTA is SUPERSEDED for this video by the "
     "live Free Career Evidence Starter. Slide 12 and reveal frame 21 carry "
     "the new CTA. Text only; design system, typography family, palette, "
     "layout, box positions and composition preserved."%OLD_CTA,size=11,
     bold=True,color=RED,after=8,spacing=1.25)
shade(p,BAND_CREAM); keep(p)
for frm,to in (("THE CAPABILITY FORMATION FIELD KIT","FREE CAREER EVIDENCE STARTER"),
               ("Is your job still / building you?",
                "ONE ACCOMPLISHMENT → / ONE PORTABLE PROOF LINE"),
               ("Complete a private, evidence-led career position assessment "
                "using the last 90 days of your actual work.", CTA_PROMISE),
               ("temidayoafonja.com/fieldkit",
                "temidayoafonja.com/career-evidence-starter")):
    keep(P(d,"FROM:  %s"%frm,size=10.5,after=3))
    keep(P(d,"TO:    %s"%to,size=10.5,bold=True,after=6))
keep(P(d,"Two type sizes came down because the new strings are materially "
       "longer and neither box can widen without colliding with the artifact "
       "images. The headline goes 32pt → 29pt: at 32pt no two-line split fits "
       "the 6.11in box, and a three-line block would collide with the body "
       "copy. The URL goes 20pt → 14pt: at 20pt it measures 6.50in and bursts "
       "the 5.00in button; at 14pt it is 4.51in and stays centred in it. "
       "Everything else on the slide is untouched.",size=10.5,color=DIM,
       before=4,after=8,spacing=1.25))
p=P(d,"CTA ARTWORK — CORRECTED. Slide 12 and reveal frame 21 now carry the "
     "REAL Career Evidence Starter artifact: the Starter cover in front and "
     "the Portable Proof Line page visible behind it, on the warm cream "
     "ground. Both images are the artifact's own pages, placed at the existing "
     "picture positions and sizes. No Field Kit imagery remains anywhere in "
     "either deck. Do not substitute generic AI graphics, an unrelated "
     "portrait, a decorative quote card or a fake worksheet, and never show "
     "the direct PDF URL. The artifact itself is the proof.",size=11,bold=True,
     color=NAVY,after=10,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"CTA and watch next",before=14)
keep(P(d,"One offer only: %s — %s"%(CTA,CTA_URL),after=5))
keep(P(d,"Public descriptor: %s. Promise: %s Outcome: one portable Proof "
       "Line. Expected time: about 10 to 15 focused minutes."
       %(CTA_DESCRIPTOR,CTA_PROMISE),size=10.5,color=DIM,after=6,spacing=1.25))
keep(P(d,"Use the landing-page URL publicly. Never expose the direct PDF URL.",
       bold=True,color=RED,after=6))
keep(P(d,"Do not add the Capability Formation Field Kit, Keep the Proof, the "
       "Career Decision Evidence Check, The Capability Audit, Maven, a "
       "newsletter CTA or another product.",bold=True,after=6))
keep(P(d,"The general Website, LinkedIn and Substack links may appear in the "
       "description's Connect and Explore section. Those are profile and "
       "navigation links, not competing product CTAs.",after=8))
keep(P(d,"Spoken watch-next route: %s"%NEXT,bold=True,after=5))
for x in ["add the direct Video 2 link to the description when Video 2 is "
 "public;","use Video 2 as the direct clickable end-screen route when it is "
 "public;","before Video 2 is public, use the Career Portability playlist if "
 "the playlist element is functioning;","if the playlist element is "
 "unavailable, use YouTube's “best for viewer” option or another appropriate "
 "currently public video temporarily;","update the end screen to point "
 "directly to Video 2 once Video 2 is public."]:
    keep(P(d,"—  "+x,after=4))
keep(P(d,"Do not leave the end screen with Subscribe only.",bold=True,
       color=RED,before=4,after=8))
compress(d)
d.save(os.path.join(LF,"Video_1_EDITOR_ONLY_HIT_Brief_v3.1.docx"))
print("editor brief written")

# The nine working chapter lines. Defined once so the copy-ready description
# and the reference section can never drift apart.
CHAPTERS=[("00:00","Change Jobs Without Starting Over"),
 ("01:00","What My Career Shifts Taught Me"),
 ("02:00","Look Beyond Job Titles"),
 ("02:25","Three Questions That Reveal Transferable Skills"),
 ("03:45","Translate Your Impact"),
 ("05:25","Preserve Career Evidence Early"),
 ("07:20","Three Questions Before Your Next Move"),
 ("08:10","Free Career Evidence Starter"),
 ("08:40","Is Your Job Making You Less Marketable?")]
CHAPTER_LINES=["%s %s"%(t,c) for t,c in CHAPTERS]

DESC=[
 "I’m Temidayo Afonja, founder of The Density Group and creator of Capability "
 "Formation. I help experienced professionals understand what they can carry "
 "across roles, functions, employers and industries so they can make career "
 "pivots and internal moves without starting from zero.",
 "In this video, I share three things that helped me change jobs without "
 "treating every transition as a return to zero:",
 "✨ Look underneath the job title and identify what your work trained you to "
 "notice, decide and solve.",
 "✨ Explain what your work changed in language another employer can "
 "understand.",
 "✨ Preserve permitted evidence before you need a résumé, interview story or "
 "promotion case.",
 "Not everything transfers. A new role, employer or industry may require real "
 "relearning. But a new context does not make you new to everything.",
 "I also share one carefully bounded example from my own work: an onboarding "
 "redesign I led with my team, where one measure of how well new hires felt "
 "integrated moved from 47 to 75.",
 "What has your work built in you—and what could you still do if the title, "
 "employer or industry changed?","",
 "🧭 FREE CAREER EVIDENCE STARTER",
 "Turn one accomplishment into a portable Proof Line you can use in a "
 "performance review, interview, internal move or career pivot:", CTA_URL,"",
 "⏱️ CHAPTERS"]+CHAPTER_LINES+["",
 "▶️ WATCH NEXT", NEXT, "[ADD VIDEO 2 LINK WHEN LIVE]","",
 "🔗 CONNECT AND EXPLORE",
 "Website:","https://temidayoafonja.com",
 "LinkedIn:","https://www.linkedin.com/in/temidayo-afonja",
 "Substack:","https://temidayoafonja.substack.com","",
 "#CareerGrowth #CareerChange #TransferableSkills"]

PINNED=["Before your next move, answer these three questions:",
 "1. What can I solve now that I could not solve two years ago?",
 "2. What result can I describe in language another employer would understand?",
 "3. What could I still do if the title, employer or industry changed?",
 "Which answer is clearest for you, and which one still needs evidence?",
 "If you want to try this on one accomplishment, the Career Evidence Starter "
 "is free.",
 "It takes about 10 to 15 focused minutes and helps you turn one piece of "
 "work into a portable Proof Line.",
 CTA_URL]

TAGS=("how to change jobs without starting over, how to change careers "
 "without starting over, change jobs without starting over, transferable "
 "skills for career change, transferable skills, changing jobs, career "
 "transition, career change, career change advice, how to identify "
 "transferable skills, how to explain your experience, career portability, "
 "transferable experience, career growth, Temidayo Afonja, Capability "
 "Formation")

EMOJI_NOTE=("The restrained emoji system is part of the approved standard: ✨ "
  "teaching points, 🧭 CTA/resource, ⏱️ chapters, ▶️ Watch Next, 🔗 Connect "
  "and Explore. Do not remove them and do not add more.")

def description_block(d, heading_before=14, upload_doc=False):
    """The copy-ready description, its end marker and the internal notes.

    Shared verbatim by the publishing package and the separate description-only
    document, so their public copy cannot drift apart. upload_doc=True places
    the editorial emoji instruction ABOVE an explicit BEGIN marker, so nothing
    internal sits inside or immediately before the block a person selects when
    pasting into YouTube.
    """
    if upload_doc:
        H1(d,"INTERNAL NOTE — DO NOT PASTE INTO YOUTUBE",before=heading_before)
        p=P(d,EMOJI_NOTE,size=10.5,italic=True,color=RED,after=12,spacing=1.25)
        shade(p,BAND_CREAM); keep(p)
        p=keep(P(d,"COPY-READY YOUTUBE DESCRIPTION — BEGIN",size=11,bold=True,
                 color=NAVY,before=14,after=12,spacing=1.2))
        shade(p,BAND_NAVY)
    else:
        H1(d,"Description",before=heading_before)
        keep(P(d,EMOJI_NOTE,size=10.5,italic=True,color=DIM,after=10))
    for para in DESC:
        keep(P(d,para if para else " ",after=7 if para else 3))
    keep(P(d,"— END OF THE COPY-READY DESCRIPTION —",size=10,bold=True,
           color=DIM,before=14,after=12,spacing=1.2))
    H1(d,"Internal note — do not paste into YouTube",before=14)
    p=P(d,"WORKING ESTIMATES — EDITOR MUST REPLACE FROM FINAL CUT",size=11,
        bold=True,color=RED,after=6,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    p=P(d,"These timestamps were estimated from the script, not measured from "
        "the finished edit. Replace every timestamp using the finished cut "
        "before publication. Do not force the edit to match these estimates.",
        size=10.5,bold=True,italic=True,color=RED,after=10,spacing=1.25)
    shade(p,BAND_CREAM); keep(p)
    H1(d,"Working chapters — reference copy",before=14)
    keep(P(d,"Identical to the nine chapter lines inside the description "
           "above.",size=10.5,italic=True,color=DIM,after=8))
    for line in CHAPTER_LINES: keep(P(d,line,size=11,after=4))


# ----------------------------------------------------- 4. publishing package
d=newdoc()
head(d,TITLE,"Video 1  ·  Publishing package",
     "Everything needed to upload. Working timestamps must be replaced with "
     "real ones from the finished edit.")
p=P(d,"The currently published Video 1 and its published description remain "
     "untouched. This package is for the replacement upload.",size=10.5,
     bold=True,color=NAVY,after=12,spacing=1.25)
shade(p,BAND_CREAM); keep(p)

H1(d,"Title",before=14); P(d,TITLE,size=12,after=6)
keep(P(d,"On-screen deck title (intentionally different): %s"%DECK_TITLE,
       size=10.5,italic=True,color=DIM,after=10))
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14)
P(d,"how to change jobs without starting over",after=6)
keep(P(d,"Supporting search phrase: transferable skills for career change",
       size=10.5,color=DIM,after=10))
H1(d,"Supporting search language",before=14)
P(d,"transferable skills for career change · how to identify transferable "
    "skills · career portability · changing jobs · career transition · "
    "transferable experience · how to explain your experience · career change "
    "without starting over · career growth · career pivot",after=10)

description_block(d, upload_doc=True)

H1(d,"Pinned comment",before=14)
for para in PINNED:
    keep(P(d,para,after=6))

H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only. Do not place the full tag field in "
       "the public description.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))

H1(d,"Evidence boundary for the description and pinned comment",before=14)
keep(P(d,"The onboarding example stays bounded: one measure of new-hire "
       "integration, moved from 47 to 75, team-based work that Temidayo led. "
       "Do not attach a retention percentage, an avoided-turnover figure or "
       "any other result to this story.",bold=True,after=8,spacing=1.25))
compress(d)
d.save(os.path.join(LF,"Video_1_Publishing_Package_HIT_v3.1.docx"))

# ------------------------------- 4b. separate description-only document
# Outside the 13-file ZIP. Same architecture approved for Videos 6 to 8.
d=newdoc()
head(d,TITLE,"Video 1  ·  YouTube description",
     "Upload copy only. Everything below the end marker is internal and must "
     "not be pasted into YouTube.")
H1(d,"Title",before=14); P(d,TITLE,size=12,after=6)
keep(P(d,"On-screen deck title (intentionally different): %s"%DECK_TITLE,
       size=10.5,italic=True,color=DIM,after=10))
H1(d,"Thumbnail",before=14); P(d,THUMB,size=12,bold=True,after=10)
H1(d,"Primary search phrase",before=14)
P(d,"how to change jobs without starting over",after=10)
description_block(d, upload_doc=True)
H1(d,"Pinned comment",before=14)
for para in PINNED: keep(P(d,para,after=6))
H1(d,"Watch next",before=14); keep(P(d,NEXT,bold=True,after=8))
H1(d,"YouTube tag field",before=14)
keep(P(d,"Paste into the tag field only.",size=10.5,italic=True,color=DIM,after=6))
keep(P(d,TAGS,size=10.5,after=10))
compress(d)
DESC_DOC="/tmp/v1hit/Video_1_YouTube_Description_HIT.docx"
d.save(DESC_DOC)
print("publishing package and description-only document written")

# ---------------------------------------------------------------- 5. Shorts
SHORTS=[
 ("Video_1_Short_1_New_Context_Not_Zero.docx","SHORT 1","Recognition",
  "A new job does not make you new to everything.",
  ["A new job does not make you new to everything.",
   "It may make you new to a company, an industry, a set of regulations or a "
   "particular way of working.",
   "That is a real learning curve.",
   "But it does not automatically erase the problems you have learned to "
   "diagnose, the decisions you can make or the situations people already "
   "trust you to handle.",
   "Before you call yourself a beginner, ask:",
   "What has my work made me able to do, and where else could that be useful?",
   "Portability is not pretending that everything transfers.",
   "It is knowing what you can carry and what you still need to learn."]),
 ("Video_1_Short_2_Experience_Needs_Evidence.docx","SHORT 2",
  "Distinction / myth",
  "Years of experience do not automatically become portable.",
  ["Years of experience do not automatically become portable just because you "
   "have them.",
   "Some knowledge depends on company systems, internal relationships or "
   "industry context.",
   "What travels is usually more specific:",
   "Judgment you have demonstrated.",
   "Problems you know how to solve.",
   "And evidence another setting can understand.",
   "So do not say, “I have twenty years of experience,” and stop there.",
   "Ask:",
   "What can I now do because of those years?",
   "Where have I demonstrated it?",
   "And what would still be useful if the company language disappeared?",
   "Experience matters.",
   "Evidence is what helps another context understand it."]),
 ("Video_1_Short_3_Result_Needs_Context.docx","SHORT 3",
  "Proof / personal evidence",
  "A project name is not proof of your value.",
  ["A project name is not proof of your value.",
   "I could say, “I led an onboarding redesign.”",
   "That is true, but somebody outside that organization still does not know "
   "what changed.",
   "A clearer version is:",
   "I led an onboarding redesign with my team, and one measure of how well "
   "new hires felt integrated moved from 47 to 75.",
   "The point is not that the number sounds impressive.",
   "The point is that I can explain what it measured, what I led and what "
   "changed.",
   "I am also careful to say it was one measure and team-based work.",
   "Evidence travels better when it is specific and honest."]),
 ("Video_1_Short_4_Look_Under_The_Title.docx","SHORT 4",
  "Practical test / action",
  "Put your job title aside for one minute.",
  ["Put your job title aside for one minute.",
   "Underneath it, answer three questions.",
   "What problems do people repeatedly trust me to solve?",
   "What decisions can I now make with better judgment?",
   "And what could I do in another setting because of what I learned here?",
   "Then look for the verbs that repeat across your examples.",
   "Maybe you diagnose, stabilize, translate, design, align or build.",
   "Do not choose a word simply because it sounds good.",
   "Attach it to evidence.",
   "That is how you begin to see what your work built in you, rather than "
   "only where the work happened."])]

for fn,label,role,hook,copy in SHORTS:
    d=newdoc(True)
    P(d,"VIDEO 1 SHORT",size=10,bold=True,color=GOLD,after=4,caps=True)
    P(d,label,size=20,bold=True,color=NAVY,after=8,spacing=1.1)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    keep(P(d,"Verbal hook:  “%s”"%hook,size=11,color=DIM,after=5))
    keep(P(d,"Related long-form:  %s"%TITLE,size=11,color=DIM,after=10))
    # 12pt lead-in and a 10pt paragraph gap keep the longest Short (Short 2)
    # on a single page. Line spacing stays at 1.5 for on-screen readability.
    H1(d,"RECORDING COPY",before=12)
    for line in copy:
        keep(P(d,line,size=13.5,color=INK,after=10,spacing=1.5))
    d.save(os.path.join(SH,fn))
print("publishing package and %d Shorts written"%len(SHORTS))

# ------------------------------------------------------ 6. Shorts editor brief
d=newdoc()
P(d,"EDITOR ONLY",size=22,bold=True,color=RED,after=2)
P(d,"VIDEO 1 — FOUR STANDALONE SHORTS",size=18,bold=True,color=NAVY,
  after=8,spacing=1.1)
p=P(d,"This document is for the editor. It is separate from the four Short "
     "recording documents and must not be placed on Temidayo's recording "
     "screen.",size=11,italic=True,color=DIM,after=16,spacing=1.25)
shade(p,BAND_CREAM)

H1(d,"How these are produced",before=14)
keep(P(d,"These are separately recorded 9:16 Shorts. They are NOT excerpts cut "
       "from the long-form video.",bold=True,after=10))
P(d,"Each Short must have:",after=6)
pairlist(d,["an immediate verbal hook;","a corresponding on-screen hook;",
            "meaningful visual movement;","accurate burned-in captions;",
            "restrained editorial pacing;",
            "Video 1 added as the YouTube Related Video when available."])

def short(label,role,onscreen,body):
    H1(d,label,before=14)
    keep(P(d,"Role:  %s"%role,size=11,color=DIM,after=5))
    p=keep(P(d,"On-screen hook:  %s"%onscreen,size=11,bold=True,color=GOLD,after=8))
    shade(p,BAND_CREAM)
    for b in body: keep(P(d,b,after=5))
    keep(P(d,"Related Video:  Video 1",size=10.5,color=DIM,before=4,after=6))

short("SHORT 1","Recognition","NEW CONTEXT ≠ ZERO EXPERIENCE",
 ["Begin direct to camera.",
  "Allow these words to appear briefly:  NEW COMPANY  ·  NEW CONTEXT  ·  "
  "NEW LANGUAGE",
  "Then resolve to:  NOT NEW TO EVERYTHING",
  "Do not use beginner icons, a reset button or literal zero imagery.",
  "End visually on:  WHAT CAN YOU CARRY?"])
short("SHORT 2","Distinction / myth","EXPERIENCE NEEDS EVIDENCE",
 ["Use a restrained visual contrast:  YEARS OF EXPERIENCE  versus  "
  "JUDGMENT + EVIDENCE",
  "When Temidayo says “Do not say, ‘I have twenty years of experience,’ and "
  "stop there,” keep her visible.",
  "Do not use clocks, calendars, anniversary graphics or résumé B-roll.",
  "End visually on:  WHAT CAN YOU NOW DO?"])
short("SHORT 3","Proof / personal evidence","THE RESULT NEEDS CONTEXT",
 ["First show:  LED AN ONBOARDING REDESIGN",
  "Then reveal:  ONE MEASURE OF NEW-HIRE INTEGRATION",
  "     47  →  75",
  "Keep this visible:  WITH MY TEAM",
  "The edit must preserve the evidence boundary: one measure; new-hire "
  "integration; team-based work; Temidayo led the redesign; not presented as "
  "solo work; no other metrics attached.",
  "Do not add the 30% retention or more-than-$2-million figures.",
  "End visually on:  SPECIFIC + HONEST"])
short("SHORT 4","Practical test / action","LOOK UNDER THE TITLE",
 ["Reveal the three questions progressively:",
  "     1   WHAT PROBLEMS DO PEOPLE TRUST ME TO SOLVE?",
  "     2   WHAT DECISIONS CAN I NOW MAKE?",
  "     3   WHAT COULD I DO SOMEWHERE ELSE?",
  "Then show a restrained verb sequence:",
  "     DIAGNOSE  ·  STABILIZE  ·  TRANSLATE  ·  DESIGN  ·  ALIGN  ·  BUILD",
  "End visually on:  WHICH VERB HAS EVIDENCE?",
  "Do not overcrowd the screen."])

H1(d,"All Shorts — visual boundaries",before=14)
P(d,"Do not use:",after=5)
pairlist(d,["hyperactive zooming;","fake shock expressions;",
 "red warning graphics;","generic office B-roll;","stock résumé footage;",
 "clocks or anniversary imagery;","literal luggage;",
 "animated career ladders;","employer logos;","AI-style animated icons;",
 "trendy caption templates;","individual words bouncing constantly;",
 "flashy transitions."],after=3)
keep(P(d,"Keep captions accurate, large, mobile-safe, restrained and "
       "consistent with the deep navy, warm cream and muted-gold system.",
       before=6,after=8,spacing=1.25))
compress(d, 1.18, 0.62)
d.save(os.path.join(SH,"Video_1_Shorts_EDITOR_ONLY_HIT_Brief.docx"))

# ---------------------------------------------------------------- 7. README
FILES=(["LONG_FORM/"+f for f in sorted(os.listdir(LF))]
      +["SHORTS/"+f for f in sorted(os.listdir(SH))])
R=["VIDEO 1 — H.I.T. FINAL RECORDING PACKAGE","",
 "Public title:      %s"%TITLE,
 "On-screen",
 "deck title:        %s"%DECK_TITLE,
 "                   The two differ intentionally. Do not change the title",
 "                   slide to match the public metadata title.",
 "Thumbnail:         %s"%THUMB,
 "CTA:               %s"%CTA,
 "CTA URL:           %s"%CTA_URL,
 "                   The earlier %s CTA is SUPERSEDED"%OLD_CTA,
 "                   for this video by the live Free Career Evidence Starter.",
 "                   Slide 12 and reveal frame 21 carry the new CTA. The",
 "                   The supporting artwork on that slide is the REAL Career",
 "                   Evidence Starter artifact: cover in front, Portable Proof",
 "                   Line page behind. No Field Kit imagery remains.",
 "Watch next:        %s"%NEXT,"",
 "Status:            Replacement recording package. The existing published",
 "                   Video 1 remains untouched until the replacement",
 "                   recording and edit are approved.","",
 "Long-form:         Revised under the H.I.T. first-30-second standard.",
 "Slides:            UNCHANGED. 13 main slides.",
 "Reveal deck:       UNCHANGED. 22 reveal-build frames.",
 "Thumbnail:         UNCHANGED.",
 "Shorts:            Four separately recorded vertical scripts.",
 "Editor",
 "instructions:      Separated from all recording copy.",
 "Publishing",
 "description:       Uses the current published description as its structural",
 "                   and tonal base, with targeted updates for the revised",
 "                   script.","",
 "-"*70,"","WHAT EACH FILE IS","",
 "LONG_FORM/","",
 "  Video1TeleprompterScriptwithslidemarkers_HIT_v3.1.docx",
 "  Video1TeleprompterScriptwithslidemarkers_HIT_v3.1.txt",
 "      Temidayo's recording copy. Spoken script in large text; slide markers",
 "      in tinted bands. The markers are not spoken.","",
 "  Video1ReadingScriptnomarkers_HIT_v3.1.docx",
 "  Video1ReadingScriptnomarkers_HIT_v3.1.txt",
 "      The same spoken words with the slide markers removed.","",
 "  Video_1_EDITOR_ONLY_HIT_Brief_v3.1.docx",
 "      For the editor. The H.I.T. first-30-second plan, editorial rhythm",
 "      after 0:30, the existing 13-slide map, the let-the-visual-carry",
 "      principle, and the evidence and visual boundaries.",
 "      Not for the teleprompter.","",
 "  Video_1_Publishing_Package_HIT_v3.1.docx",
 "      Title, thumbnail, search language, the full description, working",
 "      chapter estimates, pinned comment and the YouTube tag field.","",
 "SHORTS/","",
 "  Four recording documents, one per Short. These contain Temidayo's",
 "  recording copy and no editor directions.","",
 "  Video_1_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "      For the editor. On-screen hooks and visual treatment for all four.","",
 "-"*70,"","ALL FILES IN THIS PACKAGE","",]
for f in FILES: R.append("  "+f)
R+=["  README_FINAL.txt","  SHA256SUMS.txt","",
 "-"*70,"","WORKING CHAPTER TIMESTAMPS","",
 "The chapter timestamps in the publishing package are WORKING ESTIMATES",
 "derived from the script. They were not measured from an edit. The editor",
 "must replace every one of them from the finished cut before publishing,",
 "and the finished edit must not be forced to match the estimates.","",
 "-"*70,"","EVIDENCE BOUNDARY","",
 "The onboarding example stays bounded: Temidayo led an onboarding redesign",
 "with her team, and one measure of how well new hires felt integrated moved",
 "from 47 to 75. It was one measure of new-hire integration, not a claim",
 "about every effect of the redesign, and it was team-based work she led,",
 "not solo work. No retention percentage and no avoided-turnover figure is",
 "attached to this story anywhere in this package.","",
 "-"*70,"","CHECKSUMS","",
 "SHA256SUMS.txt covers the other 12 user-facing files in this package. It",
 "does not hash itself, and it carries no ZIP checksum. The archive's own",
 "SHA-256 is in the sibling file:",
 "  Video_1_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256","",
 "-"*70,"","WHAT WAS NOT CHANGED","",
 "The existing Video 1 PowerPoint deck (13 slides), the reveal-build deck",
 "(22 frames), the approved thumbnail, the currently published YouTube video",
 "and its published description, the Capability Formation Field Kit product",
 "itself, every website file, every product and every other video are",
 "unchanged. Beyond the authorised Slide 12 CTA correction, this revision is",
 "spoken script, editor instruction and publishing copy only.","",
 "Slide 12 and reveal frame 21 carry the authorised CTA text correction.",
 "Slides 1 to 11 and 13, and the other 21 reveal frames, are byte-identical.","",
 "The 13-slide deck remains authoritative. The teleprompter's 13 slide",
 "markers map to it in order, slide 1 to slide 13.",""]
open(os.path.join(ROOT,"README_FINAL.txt"),"w").write("\n".join(R))
print("shorts editor brief and README written")

# ------------------------------------------- 8. checksums and the master ZIP
# The archive is built from an explicit allowlist, never from a directory walk.
MANIFEST=[
 "LONG_FORM/Video1TeleprompterScriptwithslidemarkers_HIT_v3.1.docx",
 "LONG_FORM/Video1TeleprompterScriptwithslidemarkers_HIT_v3.1.txt",
 "LONG_FORM/Video1ReadingScriptnomarkers_HIT_v3.1.docx",
 "LONG_FORM/Video1ReadingScriptnomarkers_HIT_v3.1.txt",
 "LONG_FORM/Video_1_EDITOR_ONLY_HIT_Brief_v3.1.docx",
 "LONG_FORM/Video_1_Publishing_Package_HIT_v3.1.docx",
 "SHORTS/Video_1_Short_1_New_Context_Not_Zero.docx",
 "SHORTS/Video_1_Short_2_Experience_Needs_Evidence.docx",
 "SHORTS/Video_1_Short_3_Result_Needs_Context.docx",
 "SHORTS/Video_1_Short_4_Look_Under_The_Title.docx",
 "SHORTS/Video_1_Shorts_EDITOR_ONLY_HIT_Brief.docx",
 "README_FINAL.txt",
]
SUMS="SHA256SUMS.txt"
ZIP="/tmp/v1hit/Video_1_HIT_FINAL_Recording_and_Shorts_Package.zip"

def sha256(p):
    h=hashlib.sha256()
    with open(p,"rb") as fh:
        for b in iter(lambda: fh.read(1<<20), b""): h.update(b)
    return h.hexdigest()

for m in MANIFEST:
    assert os.path.isfile(os.path.join(ROOT,m)), "missing from build: "+m
on_disk=set()
for dp,dn,fn in os.walk(ROOT):
    dn[:]=[x for x in dn if x!="__pycache__"]
    for f in fn:
        if f.endswith(".pyc"): continue
        on_disk.add(os.path.relpath(os.path.join(dp,f),ROOT).replace(os.sep,"/"))
unexpected=sorted(on_disk-set(MANIFEST)-{SUMS})
assert not unexpected, "unexpected files in package directory: %r"%unexpected

L=["# VIDEO 1 - H.I.T. FINAL RECORDING PACKAGE",
   "# SHA-256 of the 12 user-facing files in this package.",
   "# SHA256SUMS.txt cannot hash itself. The master ZIP cannot contain its own",
   "# checksum either; it is published in the sibling file",
   "# Video_1_HIT_FINAL_Recording_and_Shorts_Package.zip.sha256",""]
for m in MANIFEST: L.append("%s  %s"%(sha256(os.path.join(ROOT,m)),m))
open(os.path.join(ROOT,SUMS),"w").write("\n".join(L)+"\n")

if os.path.exists(ZIP): os.remove(ZIP)
with zipfile.ZipFile(ZIP,"w",zipfile.ZIP_DEFLATED) as z:
    for m in MANIFEST+[SUMS]:
        z.write(os.path.join(ROOT,m), "Video_1_HIT_FINAL/"+m)
zsha=sha256(ZIP)
open(ZIP+".sha256","w").write("%s  %s\n"%(zsha,os.path.basename(ZIP)))

PROV="/tmp/v1hit/_source"
shutil.rmtree(PROV,ignore_errors=True); os.makedirs(PROV)
for f in ("script_text.py","build.py","qa.py"):
    src="/tmp/v1hit/"+f
    if os.path.isfile(src): shutil.copy2(src, os.path.join(PROV,f))
print("ZIP sha256:",zsha)
print("description-only doc sha256:",sha256(DESC_DOC))
