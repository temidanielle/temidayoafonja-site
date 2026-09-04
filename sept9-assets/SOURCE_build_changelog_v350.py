# -*- coding: utf-8 -*-
"""Lightning Lesson v3.5.0 — Change Log and QA Report.

The lesson family had no change-log or QA document; its record lived in chat and
in two scripts. This builds the document the family was missing, and it does not
retype any result: the verification tables are populated from qa_sept9.py and
audit_family.py at build time, so the document cannot claim a pass the checks did
not produce.

Styling is not reinvented. The document is built on the flagship report's own
file — its section properties, margins, fonts, banner and table looks are reused
by deep-copying its elements and rewriting their text, so the two documents in
this business look like one family rather than two.
"""
import copy, hashlib, importlib, os, sys
import docx

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
SHELL = ("sept16-v204-assets/Free_Flagship_60MIN_v2.0.7_Change_Log_and_QA_Report.docx")
OUT = "scratchpad/sept9/out"
DST = f"{OUT}/Lightning_Lesson_v3.5.0_Change_Log_and_QA_Report.docx"
STAMP = "Friday, September 4, 2026 at 6:40 AM CT"

DECK = "sept9-assets/How_to_Tell_If_Your_Career_Is_Stalling_Lightning_Lesson_v3.5.0_FINAL.pptx"
DECK_PDF = "sept9-assets/How_to_Tell_If_Your_Career_Is_Stalling_Lightning_Lesson_v3.5.0_FINAL.pdf"
STALL = "lightning-lesson-v3/Career_Stall_Check_v1.0.docx"
STALL_PDF = "lightning-lesson-v3/Career_Stall_Check_v1.0.pdf"

sha = lambda p: hashlib.sha256(open(p, "rb").read()).hexdigest()

CHANGES = [
    ("Session date",
     "Wednesday, September 2 becomes Wednesday, September 9, 2026, 6:00 PM CT. The "
     "lesson stays 45 minutes and its architecture is unchanged."),
    ("Continuation date",
     "The continuation session moves from September 16 to Wednesday, September 23, "
     "2026. It appears on slide faces 13, 14 and 15 and in six speaker notes."),
    ("Slide 6 speaker note",
     "The existing sentence “Some things travel with you and some have to be relearned "
     "in the new context” is extended rather than duplicated: “Some things travel with "
     "you and some have to be relearned. Some of this may come back quickly, but it "
     "still has to be learned in the new context. That is different from having nothing "
     "to carry.” The slide face is untouched and keeps “Starting as a learner is not "
     "the same as starting from zero.”"),
    ("Slide 8 speaker note",
     "The outsider-legibility test on the face is unchanged. The note gains the spoken "
     "follow-up — “and would they have enough evidence to trust that the experience is "
     "useful in their context?” — and an explicit instruction not to imply that "
     "translating your value guarantees an employer will accept it."),
    ("career-evidence-starter.html",
     "A stale September 2 block survived the earlier reschedule on this page alone. "
     "Corrected in this pass: the expiry instant to 2026-09-09T18:45:00-05:00, the body "
     "copy to “Wednesday, September 9, 2026, 6:00 PM CT” naming the start time only, "
     "and three occurrences inside the explanatory comment. It now matches the "
     "equivalent block on career-decisions.html."),
    ("Method",
     "A scripted patch of the approved v3.4.1 FINAL file, which is how this deck has "
     "always been maintained. Run text is rewritten in place so every run keeps its own "
     "formatting, and the build refuses to save unless each differing face and note "
     "matches the approved text with only a date substitution or a named sanctioned "
     "edit applied."),
]

UNCHANGED = [
    "The three signs, on slides 7, 8 and 9. Byte-identical to v3.4.1.",
    "The Career Stall Check on slide 10, the last-90-days discipline and the "
    "non-diagnostic boundary on slide 11.",
    "The 45-minute architecture: fifteen timed slides running 0:00 to 45:00 with no gap "
    "or overlap, and every TIMING cue identical to v3.4.1.",
    "The slide 2 worked example distinguishing what travelled from audit into "
    "cybersecurity and privacy from what had to be learned in the new context. It is "
    "lived evidence and is deliberately not expanded.",
    "The two continuation routes and their prices. No third route, no bundle, no "
    "scarcity language, and no reference to the retired paid group workshop.",
    "The hidden AI appendix at slide 16, which carries no timing cue and sits outside "
    "the session.",
    "The Career Stall Check leave-behind, which carries no date, price or version-bound "
    "claim and therefore needs no edit when the session moves.",
    "Every other speaker note. Thirteen of sixteen are byte-identical.",
]

INVENTORY = [
    ("Presenter deck — SHIPPED",
     "How_to_Tell_If_Your_Career_Is_Stalling_Lightning_Lesson_v3.5.0_FINAL.pptx. "
     "16 slides, 15 active, 1 hidden. sha256 " + sha(DECK)),
    ("Exported PDF — SHIPPED",
     "How_to_Tell_If_Your_Career_Is_Stalling_Lightning_Lesson_v3.5.0_FINAL.pdf. "
     "15 pages, the hidden appendix correctly excluded. sha256 " + sha(DECK_PDF)),
    ("Leave-behind — SHIPPED, UNCHANGED",
     "Career_Stall_Check_v1.0.docx and .pdf. Its four questions match slide 10 verbatim "
     "and it carries no date or price, so it is reissued at v1.0 rather than renumbered "
     "for consistency. sha256 " + sha(STALL)[:16] + "… / " + sha(STALL_PDF)[:16] + "…"),
    ("This report — SHIPPED, NEW",
     "Lightning_Lesson_v3.5.0_Change_Log_and_QA_Report.docx, with a LibreOffice preview. "
     "The family had no change-log or QA document; this is it."),
    ("Facilitator Guide v3.2 — NOT SHIPPED, SUPERSEDED",
     "It describes a 15-slide deck and a $249 paid group workshop that has since been "
     "retired from active public sale, and names a secondary Substack route. Bringing it "
     "current is a rewrite, not a date edit, so under the instruction not to rewrite "
     "evergreen material for version consistency it is left untouched and excluded. It "
     "must not be handed to a facilitator for the September 9 delivery."),
    ("Facilitator SOP for v3.5.0 — DOES NOT EXIST",
     "Stated, not reconstructed. No facilitator or delivery SOP matching the current "
     "deck exists in the repository. The delivery instrument for this lesson is the "
     "speaker notes, which carry TIMING, FACILITATION, WHAT TO SAY, WHAT NOT TO SAY and "
     "PARTICIPANT ACTION on every slide."),
    ("Change Log and Timing Map v3.2 — NOT SHIPPED, HISTORICAL",
     "Records of a 15-slide build superseded by v3.3.8, v3.4.0, v3.4.1 and v3.5.0. They "
     "are retained in the repository as the audit trail and are not part of the "
     "delivery package."),
    ("Day-of Control Sheet and Rehearsal Protocol v1.0 — NOT SHIPPED, DIFFERENT SESSION",
     "Both are dated Thursday, August 20, 2026 and belong to the retired “Is Your Work "
     "Still Building You?” lesson, as do the promo and Maven listing copy in "
     "lightning-lesson/. A different session, not this family."),
    ("Attendee instructions and reminder or follow-up copy — DO NOT EXIST",
     "Stated, not reconstructed. No attendee-facing instruction, reminder or follow-up "
     "document for this lesson exists in the repository. The registration and joining "
     "path is the Maven event page, which is configured on the platform and is not "
     "stored here."),
    ("Build and QA scripts — IN THE REPOSITORY, NOT IN THE PACKAGE",
     "SOURCE_build_v350.py, SOURCE_qa_sept9.py and SOURCE_audit_family.py are committed "
     "to sept9-assets/ so a container rebuild cannot lose them as it lost the flagship "
     "generators. They are production tooling rather than delivery material, so they are "
     "kept out of the ZIP."),
]

VISUAL = (
    "The v3.5.0 deck was exported through LibreOffice Impress and inspected page by "
    "page. It exports 15 pages — the hidden appendix is correctly excluded — with no "
    "text overflow, no collision and no element off the canvas. The three changed slide "
    "faces, 13, 14 and 15, were read at full resolution rather than by string search, "
    "because a mixed-case sweep on this deck once missed a panel header typeset as "
    "SEPTEMBER 16 and only the rendered slide showed it. LibreOffice substitutes fonts "
    "it does not have, so line breaking in the PDF is indicative rather than identical "
    "to PowerPoint.")

PROVENANCE = (
    "The generators for this deck lived in a gitignored scratchpad and did not survive a "
    "container rebuild, and are not reconstructed from memory. This is not a fallback "
    "for this file: the lesson deck has been maintained by scripted patching of the "
    "approved PPTX since v3.3.x, so the method is the same one that produced v3.4.0 and "
    "v3.4.1. What changed is that the build and QA scripts are now committed to the "
    "repository rather than left in scratch space.")


def carrier(para):
    runs = para.runs
    if not runs:
        return para.add_run("")
    return next((r for r in runs if r.text.strip()), runs[0])


def set_para(para, text):
    keep = carrier(para)
    keep.text = text
    for r in para.runs:
        if r._element is not keep._element:
            r._element.getparent().remove(r._element)
    return para


def set_cell(cell, text):
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    set_para(p, text)


def build():
    os.makedirs(OUT, exist_ok=True)
    d = docx.Document(SHELL)
    body = d.element.body

    # Templates lifted from the flagship report before its body is emptied.
    banner = copy.deepcopy(d.tables[0]._element)
    tbl2 = copy.deepcopy(d.tables[6]._element)          # 2-column change table
    tbl4 = copy.deepcopy(d.tables[19]._element)         # 4-column QA table
    stamp_p = copy.deepcopy(d.paragraphs[1]._element)   # dark-red revision stamp
    head_p = copy.deepcopy(d.paragraphs[26]._element)   # 16pt navy section heading
    sub_p = copy.deepcopy(d.paragraphs[39]._element)    # 12.5pt navy group heading
    body_p = copy.deepcopy(d.paragraphs[27]._element)   # 9.5pt body
    bullet_p = copy.deepcopy(d.paragraphs[28]._element)  # list bullet
    total_p = copy.deepcopy(d.paragraphs[63]._element)  # bold navy total line
    foot_p = copy.deepcopy(d.paragraphs[91]._element)   # grey footer stamp
    blank_p = copy.deepcopy(d.paragraphs[33]._element)

    sectPr = body.find(docx.oxml.ns.qn("w:sectPr"))
    for child in list(body):
        if child is not sectPr:
            body.remove(child)

    def add(el):
        clone = copy.deepcopy(el)
        body.insert(list(body).index(sectPr), clone)
        return clone

    def para(tmpl, text):
        return set_para(docx.text.paragraph.Paragraph(add(tmpl), d), text)

    def table(tmpl, rows):
        t = docx.table.Table(add(tmpl), d)
        while len(t.rows) > 2:
            t._element.remove(t.rows[-1]._element)
        template_row = copy.deepcopy(t.rows[1]._element)
        for ci, val in enumerate(rows[0]):
            set_cell(t.rows[0].cells[ci], val)
        t._element.remove(t.rows[1]._element)
        for row in rows[1:]:
            t._element.append(copy.deepcopy(template_row))
            cells = t.rows[-1].cells
            assert len(cells) == len(row), "row width does not match the template"
            for ci, val in enumerate(row):
                set_cell(cells[ci], val)
        return t

    # ── banner ──────────────────────────────────────────────────────────────
    b = docx.table.Table(add(banner), d)
    lines = ["INTERNAL OPERATIONS — NOT PARTICIPANT-FACING",
             "Lightning Lesson — Change Log and QA Report",
             "How to Tell If Your Career Is Stalling  ·  v3.5.0 FINAL  ·  Revised " + STAMP]
    cell = b.rows[0].cells[0]
    while len(cell.paragraphs) > len(lines):
        cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)
    for p, text in zip(cell.paragraphs, lines):
        set_para(p, text)

    para(stamp_p, "PRODUCTION REVISION STAMP — Revised " + STAMP)
    para(body_p,
         "v3.5.0 is a controlled revision of the approved v3.4.1 FINAL lesson, not a "
         "rebuild. The session moves to Wednesday, September 9, 2026 and the guided read "
         "it hands off to now falls on Wednesday, September 23, 2026; two speaker notes are "
         "refined. Nothing else in the lesson changed, and this document records both "
         "what moved and what was verified to have stayed still.")

    para(head_p, "What changed in v3.5.0")
    table(tbl2, [("Where", "Change")] + CHANGES)
    add(blank_p)

    para(head_p, "What was deliberately not changed")
    para(body_p,
         "Evergreen teaching language was not rewritten for version consistency. Each "
         "item below is asserted against the approved v3.4.1 file on disk rather than "
         "assumed.")
    for line in UNCHANGED:
        para(bullet_p, line)
    add(blank_p)

    para(head_p, "Family inventory — what exists, and what does not")
    para(body_p,
         "Where an asset no longer exists, or never existed, that is stated here rather "
         "than filled in from memory.")
    table(tbl2, [("Asset", "Status")] + INVENTORY)
    add(blank_p)

    # ── verification, populated from the checks themselves ──────────────────
    qa = importlib.import_module("qa_sept9")
    fam = importlib.import_module("audit_family")
    total = len(qa.R) + len(fam.R)
    passed = sum(1 for r in qa.R if r[3] == "PASS") + sum(1 for r in fam.R if r[2] == "PASS")

    para(head_p, f"Verification — {total} items")
    para(body_p,
         "Executed against the delivered files, not the build directory. The rows below "
         "are written by the checks at build time, so this document cannot report a pass "
         "the checks did not produce.")

    GROUPS = {"A": "Group A — Stale dates and timezone, zero tolerance",
              "B": "Group B — The new dates are present",
              "C": "Group C — The two speaker-note refinements",
              "D": "Group D — Nothing else moved",
              "E": "Group E — Deck structure and links"}
    for letter, title in GROUPS.items():
        rows = [(str(n), label, status, note)
                for n, g, label, status, note in qa.R if g == letter]
        if not rows:
            continue
        para(sub_p, title)
        table(tbl4, [("#", "Category", "Status", "Notes")] + rows)

    para(sub_p, "Group F — Family audit: supporting assets, boundary and provenance")
    table(tbl4, [("#", "Category", "Status", "Notes")]
          + [(str(n), label, status, note) for n, label, status, note in fam.R])
    para(total_p, f"{passed} of {total} PASS.")
    add(blank_p)

    para(head_p, "Visual QA")
    para(body_p, VISUAL)
    add(blank_p)

    para(head_p, "Provenance")
    para(body_p, PROVENANCE)
    add(blank_p)

    para(foot_p,
         "LIGHTNING LESSON v3.5.0 FINAL  ·  CHANGE LOG AND QA REPORT  ·  INTERNAL ONLY  "
         "·  Revised " + STAMP)

    d.save(DST)
    assert passed == total, f"{total - passed} checks did not pass; the report is not FINAL"
    print("built", os.path.basename(DST), f"({passed}/{total})")
    print("  sha256", sha(DST))
    return DST


if __name__ == "__main__":
    build()
